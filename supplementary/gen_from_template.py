"""Clone template.docx, keep formatting/shapes/dark pages, replace text for TrustDesk"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import copy, os

ns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
      'mc':'http://schemas.openxmlformats.org/markup-compatibility/2006',
      'a':'http://schemas.openxmlformats.org/drawingml/2006/main',
      'wps':'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'}

doc = Document(r'D:\TrustDesk\template.docx')

# === HELPER: Replace text in all w:t elements within an XML element ===
def replace_box_text(alt_elem, replacements):
    """Replace text in text box. replacements = {old: new}"""
    for t in alt_elem.findall('.//w:t', ns):
        if t.text:
            for old, new in replacements.items():
                if old in t.text:
                    t.text = t.text.replace(old, new)

# === STEP 1: Fix Page de Garde text (paragraphs 0-31) ===
# P0-P1: Republic header (keep as-is from template)
# P2: Institut line
replacements_para = {
    2: None,  # Keep
    3: None,  # Keep (Echahid...)
    # Memoire title lines
}

# Replace student names
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    if 'Machat Wassim' in txt:
        for r in p.runs:
            if 'Machat Wassim' in r.text:
                r.text = r.text.replace('Machat Wassim', 'Khalef Abdelmadjid')
    if 'Benbehaz Houssam' in txt:
        for r in p.runs:
            if 'Benbehaz' in r.text:
                r.text = r.text.replace('Benbehaz Houssam Elddine', 'Dai Khaled Wassim')
                r.text = r.text.replace('Benbehaz Houssam', 'Dai Khaled Wassim')
    # Replace "Genisurv" references in page de garde area (first 35 paras)
    if i < 35:
        for r in p.runs:
            if r.text and 'Gestation Des Caisse' in r.text:
                r.text = r.text.replace('Gestation Des Caisse R\u00e9gionale', 'Gestion des Incidents de S\u00e9curit\u00e9')
            if r.text and 'GENISURV' in r.text:
                r.text = r.text.replace('GENISURV', 'TrustDesk')
            if r.text and 'Genisurv' in r.text:
                r.text = r.text.replace('Genisurv', 'TrustDesk')

# === STEP 2: Fix text boxes ===
body = doc.element.body
alts = body.findall('.//mc:AlternateContent', ns)

# Box 3 = CHAPITRE I title box
for alt in alts:
    texts = alt.findall('.//w:t', ns)
    combined = ''.join(t.text or '' for t in texts)
    # Fix chapter titles
    for t in texts:
        if t.text:
            t.text = t.text.replace('Genisurv', 'TrustDesk')
            t.text = t.text.replace('GENISURV', 'TrustDesk')
            t.text = t.text.replace('Gestation Des Caisse R\u00e9gionale', 'Gestion des Incidents de S\u00e9curit\u00e9')

# === STEP 3: Replace ALL body content after Chapter 1 heading ===
# Find where Chapter 1 content starts (after the CHAPITRE I divider page)
# The template has: P154 = "Presentation de Genisurv :"
# We need to replace from P154 onwards up to the Chapter 2 divider

# First, find key paragraph indices
ch1_start = None
ch2_divider = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'Pr\u00e9sentation de Genisurv' in txt or 'Presentation de Genisurv' in txt:
        ch1_start = i
    if i > 160 and txt == '':
        # Look for section break that marks end of ch1
        pPr = p._p.find('w:pPr', ns)
        if pPr is not None and pPr.find('w:sectPr', ns) is not None:
            if ch2_divider is None and i > 200:
                ch2_divider = i

print(f"Ch1 content starts at P{ch1_start}")
print(f"Ch2 divider near P{ch2_divider}")

# Replace Chapter 1 heading
if ch1_start is not None:
    p = doc.paragraphs[ch1_start]
    for r in p.runs:
        r.text = r.text.replace('Genisurv', 'TrustDesk')

# Replace ALL Genisurv references throughout the document
for p in doc.paragraphs:
    for r in p.runs:
        if r.text:
            r.text = r.text.replace('Genisurv', 'TrustDesk')
            r.text = r.text.replace('GENISURV', 'TrustDesk')

# Also fix table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text:
                        r.text = r.text.replace('Genisurv', 'TrustDesk')
                        r.text = r.text.replace('GENISURV', 'TrustDesk')

# === STEP 4: Now delete everything from after Chapter 1's last section break ===
# We want to keep: Page de Garde + Dedicaces + Sommaire + Remerciments + Intro + Ch1
# And remove Ch2, Ch3, Ch4, Conclusion, Annexes, Biblio

# Find the section break that ends Chapter 1 (around P275)
print("\n=== Finding Ch1 end ===")
ch1_end = None
for i, p in enumerate(doc.paragraphs):
    if i < 250: continue
    pPr = p._p.find('w:pPr', ns)
    if pPr is not None and pPr.find('w:sectPr', ns) is not None:
        ch1_end = i
        print(f"Ch1 section break at P{i}: '{p.text[:50]}'")
        break

# Remove all body elements after ch1_end paragraph
if ch1_end is not None:
    # Get the XML element for the ch1_end paragraph
    target_elem = doc.paragraphs[ch1_end]._p
    
    # Find all elements after this one in body
    found = False
    to_remove = []
    for child in list(body):
        if found:
            # Don't remove the final sectPr (document-level)
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'sectPr':
                continue
            to_remove.append(child)
        if child is target_elem:
            found = True
    
    print(f"Removing {len(to_remove)} elements after Ch1")
    for elem in to_remove:
        body.remove(elem)

# Save
out = r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx'
doc.save(out)
print(f'\n[OK] Saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
print(f'Remaining paragraphs: {len(doc.paragraphs)}')
