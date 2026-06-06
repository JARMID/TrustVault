# -*- coding: utf-8 -*-
"""
Build TrustDesk memoire v9 - EXPANDED.
Ch1: ~140+ paragraphs (template has 130)
Ch2: ~140+ paragraphs (template has 133)
Ch3: ~60 paragraphs (start)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = Path(r"D:\TrustDesk\memoire_assets")
DIAGRAMS = ASSETS / "diagrams"
OUT = Path(r"D:\TrustDesk\Memoire_TrustDesk_FINAL_v9.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_margins(sec, top=2.4, bot=1.9):
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.2)
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bot)

def shade(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def invis_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for s in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{s}")
                b.set(qn("w:val"), "nil")
                tcB.append(b)
            tcPr.append(tcB)

def add_page_number(sec):
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    for r in [run, r2, r3]:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

def new_section(doc, page_num=True, top=2.4, bot=1.9):
    doc.add_section()
    sec = doc.sections[-1]
    set_margins(sec, top, bot)
    if page_num:
        add_page_number(sec)
    return sec

def C(doc, text, sz=12, bold=False, sa=0, sb=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold; r.font.name = "Times New Roman"
    return p

def B(doc, text, bold=False, italic=False, sa=6, sb=0):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.font.bold = bold; r.font.italic = italic
    return p

def LP(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run("\u2022  " + text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def T(doc, headers, rows, caption=""):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h; shade(c, "D9D9D9")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(val)
            p = c.paragraphs[0]
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = "Times New Roman"
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_image(doc, path, w=14, caption=""):
    p = Path(path)
    if p.exists():
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Cm(w))
    else:
        B(doc, f"[Image: {p.name}]", italic=True)
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_toc(doc):
    doc.add_heading("Sommaire", 1)
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    B(doc, "(Dans Word : Ctrl+A puis F9 pour mettre a jour le sommaire)", italic=True, sa=12)

# ── Init ─────────────────────────────────────────────────────────────────────

def init_doc():
    doc = Document()
    sec = doc.sections[0]
    set_margins(sec, top=2.4, bot=0.5)
    for sn in ["Normal", "Body Text"]:
        st = doc.styles[sn]
        st.font.name = "Times New Roman"; st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rPr = st._element.get_or_add_rPr()
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = OxmlElement("w:rFonts"); rPr.insert(0, rF)
        rF.set(qn("w:ascii"), "Times New Roman")
        rF.set(qn("w:hAnsi"), "Times New Roman")
        rF.set(qn("w:cs"), "Times New Roman")
    for level, sz in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        h = doc.styles[level]
        h.font.name = "Times New Roman"; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        if level == "Heading 1":
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

# ── Cover Page ───────────────────────────────────────────────────────────────

def page_de_garde(doc, with_logo=True):
    if with_logo:
        logo = ASSETS / "logob.png"
        if logo.exists():
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(logo), width=Cm(5))
            p.paragraph_format.space_after = Pt(8)

    C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 13, True, 2)
    C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 13, True, 2)
    C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE EN AUDIOVISUELS", 10, True, 2)
    C(doc, "", sa=4)
    C(doc, "Echahid Ahmed Mehdi \u2014 Ouled Fayet \u2014", 12, True, 12)
    C(doc, "", sa=8)
    C(doc, "Memoire De Fin De Formation Pour L'obtention Du Diplome De Technicien Superieur En Informatique", 14, True, 4)
    C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 14, True, 20)
    C(doc, "", sa=4)
    C(doc, "Theme :", 16, True, 8)
    C(doc, "", sa=4)
    C(doc, "Conception Et Realisation D'une", 18, True, 2)
    C(doc, "Plateforme Web et Mobile de Portefeuille", 18, True, 2)
    C(doc, "Electronique Securise pour Institutions Bancaires", 18, True, 2)
    C(doc, "\u00ab TrustDesk \u00bb", 18, True, 20)
    C(doc, "", sa=4)
    C(doc, "Organisme d'accueil :", 14, True, 4)

    if with_logo:
        C(doc, "Entreprise BEYN", 14, True, 20)
    else:
        C(doc, "", sa=20)

    # Authors & Supervisor table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(8); tbl.columns[1].width = Cm(8)
    invis_borders(tbl)
    c1 = tbl.rows[0].cells[0]
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("Realise Par :\n"); r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(12)
    r2 = p1.add_run("M. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(12); r2.bold = True
    c2 = tbl.rows[0].cells[1]
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p2.add_run("Suivi par :\n"); r3.bold = True; r3.font.name = "Times New Roman"; r3.font.size = Pt(12)
    r4 = p2.add_run("Mme. Himeur"); r4.font.name = "Times New Roman"; r4.font.size = Pt(12); r4.bold = True
    C(doc, "", sa=20)
    C(doc, "Promotion : 2025 / 2026", 12, True, 0)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = init_doc()

    # ── PAGE DE GARDE 1 (with BEYN logo) ─────────────────────────────────────
    page_de_garde(doc, with_logo=True)

    # ── PAGE DE GARDE 2 (without logo) ───────────────────────────────────────
    new_section(doc, page_num=False, top=2.4, bot=0.5)
    page_de_garde(doc, with_logo=False)

    # ── DEDICACES 1 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 1]", italic=True, sa=20)

    # ── DEDICACES 2 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 2]", italic=True, sa=20)

    # ── REMERCIEMENTS ────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir accorde la sante, la volonte et la perseverance necessaires pour mener a bien ce projet.")
    B(doc, "Nos remerciements les plus sinceres vont a notre encadreur Mme. Himeur, pour ses precieux conseils, sa disponibilite constante et son suivi rigoureux tout au long de ce memoire. Son expertise et ses orientations nous ont ete d'une aide inestimable.")
    B(doc, "Nous remercions egalement M. Reda Benbouzid, CEO de BEYN, et l'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur collaboration durant notre stage. Leur expertise dans les solutions bancaires numeriques a ete une source d'apprentissage precieuse.")
    B(doc, "Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants tout au long de notre formation.")
    B(doc, "Enfin, nous adressons nos remerciements a tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'evaluer ce travail.")
    B(doc, "A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
    C(doc, "", sa=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Abdelmadjid K., Wassim D.")
    r.font.name = "Times New Roman"; r.bold = True

    # ── SOMMAIRE ─────────────────────────────────────────────────────────────
    new_section(doc)
    add_toc(doc)

    # ── RESUME / ABSTRACT ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Resume", 1)
    B(doc, "Ce memoire presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre d'un stage au sein de l'entreprise BEYN. La plateforme se compose de trois composants principaux : une API RESTful centralisee developpee avec Laravel 11 et PHP 8.2, un tableau de bord d'administration web construit avec React.js 18 et TypeScript, et une application mobile cross-platform concue avec Flutter 3 et le framework de gestion d'etat Riverpod.", sb=12)
    B(doc, "TrustDesk integre plusieurs fonctionnalites avancees : un parcours de verification d'identite numerique (eKYC) incluant la detection de vivacite biometrique, un systeme de cartes virtuelles securisees par chiffrement asymetrique RSA-2048, un moteur de triage intelligent par intelligence artificielle pour la detection d'anomalies transactionnelles, et une piste d'audit complete garantissant la tracabilite de toutes les operations sensibles. L'architecture trois tiers adoptee assure la separation des responsabilites et la scalabilite horizontale de la solution.")
    B(doc, "Mots-cles : portefeuille electronique, fintech, eKYC, chiffrement RSA-2048, Flutter, Laravel, React.js, triage IA, securite bancaire.", bold=True)
    C(doc, "", sa=20)
    doc.add_heading("Abstract", 1)
    B(doc, "This thesis presents the design and implementation of TrustDesk, an integrated secure electronic wallet platform intended for Algerian banking institutions, developed during an internship at BEYN. The platform comprises three main components: a centralized RESTful API developed with Laravel 11 and PHP 8.2, a web-based administration dashboard built with React.js 18 and TypeScript, and a cross-platform mobile application designed with Flutter 3 and the Riverpod state management framework.", sb=12)
    B(doc, "TrustDesk integrates several advanced features: a digital identity verification process (eKYC) including biometric liveness detection, a virtual card system secured with RSA-2048 asymmetric encryption, an AI-powered intelligent triage engine for real-time transaction anomaly detection, and a comprehensive audit trail ensuring full traceability of all sensitive operations. The adopted three-tier architecture ensures separation of concerns and horizontal scalability.")
    B(doc, "Keywords: electronic wallet, fintech, eKYC, RSA-2048 encryption, Flutter, Laravel, React.js, AI triage, banking security.", bold=True)

    # ── LISTE DES FIGURES ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Figures", 1)
    figs = [
        "Organigramme de BEYN",
        "Cycle de vie Scrum",
        "Internet",
        "Exemple de navigateurs web",
        "Exemple d'une application web",
        "Exemple d'une application mobile",
        "Types de diagrammes UML",
        "Diagramme de cas d'utilisation global - TrustDesk",
        "Diagramme de cas d'utilisation - Client",
        "Diagramme de cas d'utilisation - Administrateur",
        "Composants d'un diagramme de sequence",
        "Diagramme de sequence - Authentification Client",
        "Architecture globale trois tiers",
        "Diagramme de classes",
        "Modele Conceptuel des Donnees (MCD)",
        "Diagramme de sequence - Inscription et eKYC",
        "Diagramme de sequence - Transfert P2P",
    ]
    for i, f in enumerate(figs, 1):
        LP(doc, f"Figure N\u00b0{i} : {f}")

    # ── LISTE DES TABLEAUX ───────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Tableaux", 1)
    tbls_list = [
        "Fiche technique de BEYN",
        "Principales solutions de BEYN",
        "Ressources humaines du projet",
        "Ressources materielles du projet",
        "Acteurs du systeme TrustDesk",
        "Besoins fonctionnels par module",
        "Besoins non fonctionnels",
        "Description textuelle UC - Inscription eKYC",
        "Description textuelle UC - Transfert P2P",
        "Description textuelle UC - Gestion des cartes virtuelles",
        "Couches de l'architecture API Laravel",
    ]
    for i, t in enumerate(tbls_list, 1):
        LP(doc, f"Tableau N\u00b0{i} : {t}")

    # ── INTRODUCTION GENERALE ────────────────────────────────────────────────
    new_section(doc)
    C(doc, "INTRODUCTION", 16, True, 20)
    B(doc, "Depuis toujours, les institutions financieres cherchent a ameliorer leurs methodes de gestion et a optimiser leurs services pour mieux repondre aux besoins de leurs clients. Avec l'essor des technologies de l'information et de la communication, les banques ont entrepris une transformation numerique profonde, passant des guichets traditionnels aux services en ligne, puis aux applications mobiles.", sb=12)
    B(doc, "Par ailleurs, la generalisation de l'Internet et le developpement des technologies mobiles ont radicalement modifie les attentes des clients bancaires. Ceux-ci exigent desormais un acces instantane a leurs comptes, des transferts rapides et securises, et des outils de gestion financiere personnalises, le tout accessible depuis leur smartphone.")
    B(doc, "Dans ce contexte, les portefeuilles electroniques (e-wallets) s'imposent comme un levier strategique. Ils permettent aux banques d'offrir des services de paiement dematerialises, des transferts de pair a pair (P2P), et la gestion de cartes virtuelles securisees.")
    B(doc, "En Algerie, cette mutation s'accelere sous l'impulsion de la Banque d'Algerie et des reformes reglementaires visant a promouvoir l'inclusion financiere. Cependant, les solutions existantes sur le marche algerien presentent plusieurs lacunes : absence d'un parcours eKYC entierement automatise, securite insuffisante pour les cartes virtuelles, et absence de detection de fraude par intelligence artificielle.")
    B(doc, "C'est dans cette perspective que s'inscrit le projet TrustDesk, realise au sein de l'entreprise BEYN, un acteur majeur de la fintech algerienne. La problematique centrale de ce memoire est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)
    B(doc, "Pour repondre a cette problematique, nous avons structure notre memoire en trois chapitres :")
    B(doc, "Le premier chapitre, intitule \u00ab Etude Prealable \u00bb, presente la societe BEYN, le contexte du marche fintech algerien, la problematique identifiee, les concepts theoriques utilises et la methodologie de developpement adoptee.")
    B(doc, "Le deuxieme chapitre, intitule \u00ab Analyse et Specification des Besoins \u00bb, est consacre a l'identification des acteurs, la formalisation des besoins fonctionnels et non fonctionnels, et la modelisation UML du systeme a l'aide des diagrammes de cas d'utilisation.")
    B(doc, "Le troisieme chapitre, intitule \u00ab Conception \u00bb, presente la conception architecturale de la solution : architecture trois tiers, modele de donnees (MCD et MLD), diagrammes de classes et de sequence, et architecture de chiffrement.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE I : ETUDE PREALABLE  (~140+ paragraphs)
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch1_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre I : Etude Prealable", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce premier chapitre constitue le socle fondateur de notre travail. Il presente l'organisme d'accueil BEYN, le contexte du marche fintech algerien dans lequel s'inscrit notre projet, la problematique identifiee, le public vise, les ressources humaines et materielles mobilisees, les concepts theoriques necessaires a la comprehension du sujet, ainsi que la methodologie de developpement adoptee.", sb=12)
    B(doc, "L'objectif principal de cette etude prealable est de poser les bases methodologiques et contextuelles indispensables a la bonne comprehension des choix techniques et fonctionnels effectues dans les chapitres suivants.")

    # ── 1.1 Presentation de BEYN ─────────────────────────────────────────────
    doc.add_heading("1.1  Presentation de l'organisme d'accueil : BEYN", 2)
    B(doc, "Fondee en 2004 et basee a Cheraga (Alger), BEYN est une entreprise specialisee dans la conception et le developpement de solutions technologiques pour les institutions bancaires et financieres algeriennes. Dirigee par M. Reda Benbouzid, elle se positionne comme un partenaire technologique strategique pour les banques partenaires, notamment la Banque de Developpement Local (BDL), la Banque Nationale d'Algerie (BNA) et Al Salam Bank Algeria.")
    B(doc, "L'entreprise compte une equipe pluridisciplinaire de developpeurs, d'ingenieurs systeme, de designers UX/UI et d'experts en securite bancaire. Son siege social est situe a Cheraga, dans la wilaya d'Alger. BEYN s'est imposee comme un acteur incontournable de l'ecosysteme fintech algerien grace a son expertise dans la digitalisation des services bancaires.")
    B(doc, "Au fil des annees, BEYN a developpe des partenariats solides avec les principales banques algeriennes, se specialisant dans trois domaines cles : le paiement mobile, la banque digitale, et les passerelles de paiement. Cette experience accumulee constitue le fondement sur lequel s'appuie le projet TrustDesk.")

    doc.add_heading("Fiche Technique de BEYN", 3)
    T(doc, ["Designation", "Information"], [
        ["Nom de l'entreprise",  "BEYN"],
        ["Date de creation",     "2004"],
        ["Siege social",         "Cheraga, Alger, Algerie"],
        ["Directeur general",    "M. Reda Benbouzid"],
        ["Secteur d'activite",   "Technologies financieres (Fintech)"],
        ["Domaine",              "Solutions bancaires numeriques"],
        ["Effectif",             "50+ collaborateurs"],
        ["Partenaires bancaires","BDL, BNA, Al Salam Bank Algeria"],
    ], "Tableau N\u00b01 : Fiche technique de BEYN")

    doc.add_heading("Organigramme de BEYN", 3)
    B(doc, "L'organigramme ci-dessous illustre la structure organisationnelle de BEYN. L'entreprise est organisee en quatre departements principaux : la Direction Generale, le Departement Technique (developpement et infrastructure), le Departement Commercial (partenariats bancaires), et le Departement Support (qualite et maintenance). Cette organisation permet une collaboration etroite entre les equipes techniques et commerciales.")
    B(doc, "[Figure N\u00b01 : Organigramme de BEYN - A inserer]", italic=True)

    doc.add_heading("Les Solutions de BEYN", 3)
    B(doc, "BEYN propose trois plateformes phares couvrant l'ensemble des besoins bancaires numeriques :")
    LP(doc, "WimPay : Solution leader de paiement mobile par QR code en Algerie, permettant les transferts P2P instantanes, le paiement de factures (telephone, internet, electricite), et le paiement chez les commercants partenaires. WimPay est deployee sur Android et iOS et integree avec plusieurs banques algeriennes. Elle utilise une architecture microservices et un chiffrement de bout en bout pour securiser les transactions.")
    LP(doc, "SELA : Plateforme de banque digitale omnicanale, orientee retail banking. SELA offre aux clients bancaires un acces complet a leurs comptes, virements, et services via une interface web et mobile unifiee. La plateforme prend en charge la gestion des beneficiaires, la consultation des releves, et les alertes personnalisees.")
    LP(doc, "KANTARA : Solution de corporate banking et de passerelles de paiement securisees pour les entreprises. KANTARA facilite les transactions inter-entreprises et l'integration avec les systemes de paiement nationaux tels que le Systeme de Paiement Electronique Interbancaire (SEPI) et la Societe d'Automatisation des Transactions Interbancaires et de Monetique (SATIM).")

    T(doc, ["Plateforme", "Cible", "Fonctionnalite Principale"], [
        ["WimPay",  "Grand public",      "Paiement QR code, transferts P2P, paiement factures"],
        ["SELA",    "Retail banking",    "Banque digitale omnicanale (web + mobile)"],
        ["KANTARA", "Corporate banking", "Passerelles de paiement, transactions inter-entreprises"],
    ], "Tableau N\u00b02 : Principales solutions de BEYN")

    # ── 1.2 Presentation du Projet ───────────────────────────────────────────
    doc.add_heading("1.2  Presentation du Projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme integree de portefeuille electronique securise, concue pour repondre aux besoins des institutions bancaires algeriennes en matiere de services financiers numeriques. Le projet a ete initie dans le cadre de notre stage chez BEYN, en reponse aux lacunes identifiees dans les solutions existantes sur le marche algerien.")
    B(doc, "Le nom TrustDesk reflete la vocation de la plateforme : offrir un espace de confiance (Trust) et un tableau de bord (Desk) complet pour la gestion des services financiers numeriques. La plateforme ambitionne de devenir un standard dans le domaine des portefeuilles electroniques bancaires en Algerie.")
    B(doc, "La plateforme se compose de trois composants principaux :")
    LP(doc, "Une API RESTful centralisee : Developpee avec Laravel 11 et PHP 8.2, elle constitue le coeur metier de la plateforme. Elle gere l'authentification Sanctum, la logique transactionnelle, le chiffrement RSA/AES, et la piste d'audit. L'API expose plus de 40 endpoints couvrant l'ensemble des fonctionnalites du systeme.")
    LP(doc, "Un tableau de bord d'administration web : Construit avec React.js 18, TypeScript et Vite, il permet aux administrateurs bancaires de superviser les clients, valider les dossiers eKYC, gerer les incidents de securite, et monitorer les KPIs en temps reel. Le dashboard offre des graphiques dynamiques et des filtres avances.")
    LP(doc, "Une application mobile cross-platform : Developpee avec Flutter 3 et Dart, utilisant Riverpod pour la gestion d'etat. Elle offre aux clients l'acces a leur portefeuille, les transferts P2P, la gestion des cartes virtuelles, et les outils d'analytique financiere. L'application supporte la biometrie (empreinte digitale et reconnaissance faciale).")
    B(doc, "Les fonctionnalites differenciantes de TrustDesk par rapport aux solutions existantes comprennent :")
    LP(doc, "Un parcours eKYC complet avec detection de vivacite biometrique permettant la verification d'identite a distance.")
    LP(doc, "Un systeme de cartes virtuelles avec chiffrement asymetrique RSA-2048 et approche Zero-Knowledge cote backend.")
    LP(doc, "Un moteur de triage intelligent par IA pour la detection d'anomalies transactionnelles en temps reel.")
    LP(doc, "Une piste d'audit complete et immuable pour la tracabilite de toutes les operations sensibles.")
    LP(doc, "Un systeme de paiement par QR code compatible avec l'ecosysteme WimPay de BEYN.")

    # ── 1.3 Problematique ────────────────────────────────────────────────────
    doc.add_heading("1.3  Problematique", 2)
    B(doc, "Le secteur bancaire algerien connait une transformation numerique acceleree, poussee par les directives de la Banque d'Algerie en matiere de paiement electronique et d'inclusion financiere. Cependant, malgre les avancees realisees, plusieurs lacunes persistent dans les solutions de portefeuille electronique actuellement disponibles sur le marche :")
    LP(doc, "Absence d'un parcours eKYC entierement automatise : La verification d'identite reste souvent manuelle ou semi-manuelle, entrainant des delais importants (parfois plusieurs jours) et une experience utilisateur mediocre. Les clients doivent frequemment se deplacer en agence pour completer leur identification.")
    LP(doc, "Securite insuffisante pour les cartes virtuelles : Les solutions existantes ne mettent pas en oeuvre de chiffrement asymetrique pour la protection des numeros de carte et des CVV. Les donnees sensibles sont souvent stockees en clair ou avec un chiffrement symetrique simple, exposant les clients a des risques de compromission.")
    LP(doc, "Absence de detection de fraude par IA : Aucune solution locale n'integre de moteur d'intelligence artificielle capable de detecter les anomalies transactionnelles en temps reel. La detection de fraude repose generalement sur des regles statiques qui ne s'adaptent pas aux nouveaux schemas de fraude.")
    LP(doc, "Manque de tracabilite : L'absence de piste d'audit detaillee rend difficile le suivi des operations sensibles et la conformite aux exigences reglementaires de la Banque d'Algerie en matiere de lutte contre le blanchiment d'argent et le financement du terrorisme.")
    LP(doc, "Experience utilisateur fragmentee : Les solutions existantes proposent des interfaces disparates entre le web et le mobile, sans coherence de design ni de parcours utilisateur. Les clients doivent souvent basculer entre plusieurs applications pour acceder a l'ensemble des services.")
    B(doc, "Face a ces constats, la problematique centrale de notre travail est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)

    # ── 1.4 Le public vise ───────────────────────────────────────────────────
    doc.add_heading("1.4  Le public vise", 2)
    B(doc, "La plateforme TrustDesk cible deux categories d'utilisateurs distinctes, chacune avec des besoins et des parcours specifiques :")
    LP(doc, "Les clients bancaires (utilisateurs finaux) : Particuliers disposant d'un compte dans l'une des banques partenaires de BEYN, souhaitant acceder a des services de portefeuille electronique via l'application mobile. Ce public comprend des utilisateurs ages de 18 a 65 ans, avec des niveaux de maitrise technologique variables. L'application mobile doit donc etre intuitive et accessible.")
    LP(doc, "Les administrateurs bancaires (back-office) : Responsables au sein des banques partenaires charges de superviser les comptes clients, valider les verifications d'identite, gerer les incidents de securite, et produire des rapports de conformite via le tableau de bord web. Ces utilisateurs sont des professionnels necessitant des outils d'analyse avances et des workflows optimises.")
    B(doc, "En plus de ces deux categories principales, le systeme interagit egalement avec des acteurs secondaires :")
    LP(doc, "Le systeme IA : Composant automatise qui effectue le triage des requetes, la detection d'anomalies, et le routage des incidents sans intervention humaine.")
    LP(doc, "Les systemes externes : APIs bancaires partenaires (BDL, BNA), passerelles de paiement SATIM, et services de notification push (Firebase Cloud Messaging).")

    # ── 1.5 Ressources ───────────────────────────────────────────────────────
    doc.add_heading("1.5  Ressources humaines et materielles", 2)

    doc.add_heading("Ressources humaines", 3)
    B(doc, "Le projet TrustDesk a ete realise par une equipe de deux stagiaires, sous l'encadrement d'une enseignante de l'INSFP et avec le soutien technique de l'equipe BEYN :")
    T(doc, ["Membre", "Role", "Responsabilites"], [
        ["M. Khalef Abdelmadjid", "Developpeur Full-Stack", "API Laravel, base de donnees, securite, chiffrement RSA/AES"],
        ["M. Dai Khaled Wassim",  "Developpeur Full-Stack", "Application Flutter, dashboard React.js, integration API"],
        ["Mme. Himeur",           "Encadreur INSFP",        "Suivi pedagogique, revisions, orientations methodologiques"],
        ["M. Reda Benbouzid",     "Superviseur BEYN",       "Cadrage fonctionnel, acces aux ressources techniques"],
    ], "Tableau N\u00b03 : Ressources humaines du projet")

    doc.add_heading("Ressources materielles", 3)
    B(doc, "Les ressources materielles et logicielles mobilisees pour la realisation du projet sont les suivantes :")
    T(doc, ["Ressource", "Specification"], [
        ["Ordinateur portable 1",  "Lenovo IdeaPad, Intel Core i7, 16 Go RAM, SSD 512 Go, Windows 11"],
        ["Ordinateur portable 2",  "HP ProBook, Intel Core i5, 8 Go RAM, SSD 256 Go, Windows 11"],
        ["Smartphone de test",     "Samsung Galaxy A54, Android 14, 6 Go RAM"],
        ["IDE Backend",            "PHPStorm 2024 / VS Code avec extensions PHP et Laravel"],
        ["IDE Frontend",           "VS Code avec extensions React, TypeScript, ESLint"],
        ["IDE Mobile",             "Android Studio / VS Code avec extensions Flutter et Dart"],
        ["Serveur de dev",         "Supabase (PostgreSQL cloud), Herd (serveur local Laravel)"],
        ["Outil UML",              "Draw.io, StarUML, PlantUML"],
        ["Versionnage",            "Git + GitHub (depot prive)"],
        ["Design",                 "Figma (maquettes UI/UX)"],
    ], "Tableau N\u00b04 : Ressources materielles du projet")

    # ── 1.6 Concepts theoriques ──────────────────────────────────────────────
    doc.add_heading("1.6  Concepts theoriques", 2)
    B(doc, "Avant d'entamer la phase d'analyse et de conception, il est essentiel de definir les concepts theoriques fondamentaux sur lesquels repose notre projet. Cette section presente les definitions generales, les notions de modelisation UML, et les technologies cles utilisees.")

    doc.add_heading("Les definitions generales", 3)

    B(doc, "Internet :", bold=True)
    B(doc, "Internet est un ensemble de reseaux mondiaux interconnectes qui permet a des ordinateurs et a des serveurs de communiquer efficacement au moyen d'un protocole de communication commun (TCP/IP). Ses principaux protocoles incluent HTTP/HTTPS pour le web, SMTP pour le courrier electronique, et FTP pour le transfert de fichiers. Internet permet l'echange de donnees sous forme de texte, d'images, de videos et d'applications a travers le monde entier.")
    B(doc, "[Figure N\u00b03 : Internet - A inserer]", italic=True)

    B(doc, "Navigateur web :", bold=True)
    B(doc, "Un navigateur web (en anglais web browser) est un logiciel client permettant d'acceder aux ressources du World Wide Web. Il interprete le code HTML, CSS et JavaScript pour afficher les pages web de maniere visuelle et interactive. Les navigateurs modernes integrent des fonctionnalites avancees telles que les outils de developpement, la gestion des cookies, et le support des Progressive Web Applications (PWA). Les principaux navigateurs sont Google Chrome, Mozilla Firefox, Microsoft Edge et Apple Safari.")
    B(doc, "[Figure N\u00b04 : Exemple de navigateurs web - A inserer]", italic=True)

    B(doc, "Serveur web :", bold=True)
    B(doc, "Un serveur web est un logiciel ou une machine physique qui heberge des sites web et des applications web, et qui repond aux requetes HTTP des clients (navigateurs). Les serveurs web les plus repandus sont Apache, Nginx et Microsoft IIS. Dans le cadre de notre projet, nous utilisons le serveur integre de Laravel (Artisan serve) pour le developpement local et Nginx pour la production.")

    B(doc, "Base de donnees :", bold=True)
    B(doc, "Une base de donnees est un systeme organise de stockage et de gestion de donnees structurees. Les bases de donnees relationnelles (SGBDR) comme MySQL, PostgreSQL et SQLite organisent les donnees sous forme de tables liees par des relations. Les operations sur les donnees sont effectuees a l'aide du langage SQL (Structured Query Language). Dans notre projet, nous utilisons PostgreSQL via Supabase pour beneficier de fonctionnalites temps reel et d'un stockage cloud securise.")

    B(doc, "Application web :", bold=True)
    B(doc, "Une application web est un logiciel qui s'execute dans un navigateur web. Contrairement aux applications de bureau, elle ne necessite pas d'installation sur l'ordinateur de l'utilisateur. Les applications web utilisent une architecture client-serveur : le navigateur (client) envoie des requetes au serveur, qui traite les donnees et renvoie les resultats sous forme de pages HTML, JSON ou XML. Les applications web modernes (SPA - Single Page Applications) utilisent des frameworks JavaScript comme React.js ou Vue.js pour offrir une experience utilisateur fluide.")
    B(doc, "[Figure N\u00b05 : Exemple d'une application web - A inserer]", italic=True)

    B(doc, "Application mobile :", bold=True)
    B(doc, "Une application mobile est un logiciel concu pour fonctionner sur des appareils mobiles tels que les smartphones et les tablettes. Elle est generalement telechargee depuis un magasin d'applications (Google Play Store, Apple App Store) et installee directement sur l'appareil. Les applications mobiles peuvent etre classees en trois categories :")
    LP(doc, "Applications natives : Developpees specifiquement pour un systeme d'exploitation (Java/Kotlin pour Android, Swift pour iOS). Elles offrent les meilleures performances et un acces complet aux fonctionnalites materielles.")
    LP(doc, "Applications hybrides : Developpees avec des technologies web (HTML, CSS, JavaScript) et encapsulees dans un conteneur natif (Apache Cordova, Ionic).")
    LP(doc, "Applications cross-platform : Developpees avec un framework unique permettant le deploiement sur plusieurs plateformes (Flutter, React Native). Flutter, notre choix pour TrustDesk, compile en code natif ARM, offrant des performances proches du natif.")
    B(doc, "[Figure N\u00b06 : Exemple d'une application mobile - A inserer]", italic=True)

    B(doc, "Framework :", bold=True)
    B(doc, "Un framework est un ensemble d'outils, de bibliotheques et de conventions qui fournit une structure de base pour le developpement d'applications. Il impose une architecture et des patterns de conception (MVC, MVVM) qui facilitent la maintenance et la collaboration. Les frameworks utilises dans notre projet sont Laravel (PHP, backend), React.js (JavaScript, frontend web), et Flutter (Dart, mobile).")

    B(doc, "API REST :", bold=True)
    B(doc, "Une API (Application Programming Interface) REST (Representational State Transfer) est une interface de programmation qui permet a des applications de communiquer entre elles via le protocole HTTP. Les APIs REST utilisent les methodes HTTP standard (GET, POST, PUT, DELETE) pour effectuer des operations CRUD (Create, Read, Update, Delete) sur des ressources identifiees par des URLs. Elles sont stateless, c'est-a-dire que chaque requete contient toutes les informations necessaires a son traitement, sans dependance a un etat stocke cote serveur.")
    B(doc, "Les avantages d'une API REST incluent : la simplicite d'integration, l'independance entre le client et le serveur, la scalabilite horizontale, et la compatibilite avec tous les langages de programmation.")

    B(doc, "Portefeuille electronique (e-wallet) :", bold=True)
    B(doc, "Un portefeuille electronique est une application ou un service qui permet de stocker, envoyer et recevoir de l'argent de maniere dematerialisee. Il peut etre lie a un compte bancaire ou fonctionner de maniere autonome. Les e-wallets offrent generalement des fonctionnalites de paiement en ligne, de transfert P2P (Peer-to-Peer), de gestion de cartes virtuelles et de suivi des depenses. Les exemples internationaux incluent Apple Pay, Google Pay et PayPal.")

    B(doc, "eKYC (Electronic Know Your Customer) :", bold=True)
    B(doc, "Le eKYC est un processus numerique de verification d'identite qui permet aux institutions financieres de verifier l'identite de leurs clients a distance, sans necessite de presence physique. Ce processus implique generalement trois etapes :")
    LP(doc, "La capture d'une piece d'identite officielle (carte d'identite, passeport) via la camera du smartphone.")
    LP(doc, "La prise d'un selfie pour la comparaison faciale avec le document d'identite.")
    LP(doc, "Une verification de vivacite (liveness detection) pour s'assurer que la personne est physiquement presente et qu'il ne s'agit pas d'une photo ou d'une video.")
    B(doc, "Le eKYC est conforme aux exigences reglementaires de la Banque d'Algerie en matiere de connaissance du client et de lutte contre le blanchiment d'argent (LCB-FT).")

    B(doc, "Chiffrement asymetrique (RSA) :", bold=True)
    B(doc, "Le chiffrement asymetrique, aussi appele chiffrement a cle publique, utilise une paire de cles mathematiquement liees : une cle publique pour le chiffrement et une cle privee pour le dechiffrement. L'algorithme RSA (Rivest-Shamir-Adleman), invente en 1977, est l'un des plus utilises en cryptographie moderne. Avec des tailles de cle de 2048 bits ou plus, il offre un niveau de securite adapte aux exigences bancaires. Le chiffrement asymetrique resout le probleme fondamental de distribution des cles inherent au chiffrement symetrique.")

    B(doc, "Chiffrement symetrique (AES) :", bold=True)
    B(doc, "L'Advanced Encryption Standard (AES) est un algorithme de chiffrement symetrique qui utilise une seule cle pour le chiffrement et le dechiffrement. AES-256 (cle de 256 bits) est considere comme incassable avec les technologies actuelles. Dans notre projet, AES est utilise en complement de RSA dans un schema de chiffrement hybride : RSA chiffre la cle AES, et AES chiffre les donnees volumineuses.")

    doc.add_heading("Definition UML", 3)
    B(doc, "Le Langage de Modelisation Unifie, de l'anglais Unified Modeling Language (UML), est un langage de modelisation graphique a base de pictogrammes concu pour fournir une methode normalisee pour visualiser la conception d'un systeme. Il est couramment utilise en developpement logiciel et en conception orientee objet. UML a ete normalise par l'Object Management Group (OMG) et sa version actuelle est la 2.5.")
    B(doc, "UML propose 14 types de diagrammes divises en deux categories :")
    B(doc, "Diagrammes structurels (statiques) :", bold=True)
    LP(doc, "Diagramme de classes : Represente les classes du systeme, leurs attributs, methodes et relations (association, heritage, composition).")
    LP(doc, "Diagramme d'objets : Represente les instances des classes a un instant donne.")
    LP(doc, "Diagramme de composants : Represente l'organisation physique des composants logiciels.")
    LP(doc, "Diagramme de deploiement : Represente l'architecture materielle et le deploiement des composants.")
    LP(doc, "Diagramme de paquetages : Represente l'organisation logique des elements en paquetages.")
    LP(doc, "Diagramme de structures composites : Represente la structure interne d'une classe complexe.")
    LP(doc, "Diagramme de profils : Represente les extensions personnalisees du metamodele UML.")
    B(doc, "Diagrammes comportementaux (dynamiques) :", bold=True)
    LP(doc, "Diagramme de cas d'utilisation : Represente les fonctionnalites du systeme vues par les acteurs externes.")
    LP(doc, "Diagramme d'activites : Represente les flux de controle et de donnees sous forme d'actions et de decisions.")
    LP(doc, "Diagramme d'etats (etats-transitions) : Represente les differents etats d'un objet et les transitions entre eux.")
    LP(doc, "Diagramme de sequence : Represente les interactions entre objets dans un ordre chronologique.")
    LP(doc, "Diagramme de communication : Represente les interactions entre objets en mettant l'accent sur les liens.")
    LP(doc, "Diagramme de temps : Represente le comportement temporel des objets.")
    LP(doc, "Diagramme de vue d'ensemble des interactions : Combine les diagrammes d'activites et de sequence.")
    B(doc, "[Figure N\u00b07 : Types de diagrammes UML - A inserer]", italic=True)
    B(doc, "Dans le cadre de ce memoire, nous utiliserons principalement les diagrammes de cas d'utilisation (chapitre II), de classes, de sequence, et le modele conceptuel des donnees (chapitre III).")

    # ── 1.7 Methodologie ─────────────────────────────────────────────────────
    doc.add_heading("1.7  Methodologie de developpement : Agile Scrum", 2)
    B(doc, "Pour la gestion de notre projet, nous avons adopte la methode Agile Scrum, une approche iterative et incrementale de gestion de projet logiciel. Scrum se distingue par sa flexibilite et sa capacite a s'adapter aux changements de besoins en cours de projet. Cette methode est particulierement adaptee aux projets de developpement logiciel ou les exigences peuvent evoluer rapidement.")

    B(doc, "Les roles definis dans Scrum :", bold=True)
    LP(doc, "Product Owner : Responsable de la definition et de la priorisation du Product Backlog. Dans notre cas, ce role a ete assure conjointement par notre encadreur et le superviseur BEYN.")
    LP(doc, "Scrum Master : Facilitateur du processus Scrum, il veille au respect des ceremonies et a l'elimination des obstacles. Ce role a ete partage entre les deux membres de l'equipe.")
    LP(doc, "Equipe de developpement : Les developpeurs qui realisent les increments du produit a chaque sprint. Notre equipe etait composee de deux developpeurs full-stack.")

    B(doc, "Les artefacts Scrum :", bold=True)
    LP(doc, "Product Backlog : Liste priorisee de toutes les fonctionnalites souhaitees pour le produit, organisee par epics et user stories.")
    LP(doc, "Sprint Backlog : Sous-ensemble du Product Backlog selectionne pour un sprint specifique, avec les taches decomposees.")
    LP(doc, "Increment : Le resultat fonctionnel livre a la fin de chaque sprint, potentiellement livrable en production.")

    B(doc, "Les ceremonies Scrum :", bold=True)
    LP(doc, "Sprint Planning : Reunion de planification au debut de chaque sprint pour selectionner les elements du backlog a realiser.")
    LP(doc, "Daily Standup : Reunion quotidienne de 15 minutes pour synchroniser l'equipe sur l'avancement, les obstacles et les priorites.")
    LP(doc, "Sprint Review : Presentation de l'increment au Product Owner et aux parties prenantes a la fin de chaque sprint.")
    LP(doc, "Sprint Retrospective : Reflexion d'equipe sur les ameliorations possibles du processus pour le sprint suivant.")

    B(doc, "Notre projet a ete decoupe en quatre sprints principaux de deux semaines chacun :")
    LP(doc, "Sprint 1 (Semaines 1-2) : Mise en place de l'infrastructure technique, authentification Sanctum, modele de donnees, et parcours eKYC complet.")
    LP(doc, "Sprint 2 (Semaines 3-4) : Module portefeuille, transferts P2P, paiements QR code, et debut du tableau de bord React.js.")
    LP(doc, "Sprint 3 (Semaines 5-6) : Cartes virtuelles avec chiffrement RSA-2048, integration du systeme de notifications push, et finalisation du dashboard.")
    LP(doc, "Sprint 4 (Semaines 7-8) : Moteur de triage IA, piste d'audit, tests unitaires et d'integration, et preparation du deploiement.")
    B(doc, "[Figure N\u00b02 : Cycle de vie Scrum - A inserer]", italic=True)

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce premier chapitre a pose le cadre complet de notre travail. Nous avons presente l'organisme d'accueil BEYN, un acteur reconnu de la fintech algerienne, ainsi que ses solutions phares. La problematique identifiee met en lumiere les lacunes des solutions existantes en matiere d'eKYC, de securite des cartes, et de detection de fraude par IA. Nous avons defini les concepts theoriques fondamentaux (Internet, applications web et mobile, API REST, UML, eKYC, chiffrement RSA/AES) et presente les ressources mobilisees. Enfin, la methodologie Agile Scrum adoptee garantit une approche iterative et collaborative. Le chapitre suivant sera consacre a l'analyse detaillee et a la specification des besoins fonctionnels et non fonctionnels de la plateforme TrustDesk.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE II : ANALYSE ET SPECIFICATION DES BESOINS (~140+ paragraphs)
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch2_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre II : Analyse et Specification des Besoins", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce deuxieme chapitre est consacre a l'analyse detaillee des besoins de la plateforme TrustDesk. Nous commencerons par les rappels theoriques sur les diagrammes UML utilises dans cette phase d'analyse, puis nous identifierons les acteurs du systeme, nous formaliserons les exigences fonctionnelles et non fonctionnelles, et nous presenterons les diagrammes de cas d'utilisation et les descriptions textuelles detaillees des principaux scenarios.", sb=12)
    B(doc, "L'objectif de ce chapitre est de fournir une specification complete et sans ambiguite des besoins, servant de contrat entre l'equipe de developpement et les parties prenantes du projet.")

    # ── 2.1 Rappels theoriques UC ────────────────────────────────────────────
    doc.add_heading("2.1  Rappels theoriques : Diagramme de cas d'utilisation", 2)
    B(doc, "Le diagramme de cas d'utilisation (Use Case Diagram) est un diagramme UML fonctionnel utilise pour representer les interactions entre les acteurs externes et le systeme. Ce diagramme est generalement utilise dans les premieres phases de l'analyse afin d'identifier les besoins fonctionnels et de definir le perimetre du systeme. Il constitue un outil de communication privilegie entre les developpeurs et les utilisateurs finaux.")
    B(doc, "Le diagramme de cas d'utilisation repond a trois questions fondamentales : Qui utilise le systeme ? Que fait le systeme ? Quelles sont les limites du systeme ?")
    B(doc, "Les elements contenus dans un diagramme de cas d'utilisation sont :")

    B(doc, "1. Acteurs (Actors)", bold=True)
    B(doc, "Les acteurs representent les entites externes qui interagissent avec le systeme. Ils peuvent etre des utilisateurs humains, des systemes externes, ou des organisations externes. L'acteur est represente graphiquement par un bonhomme (stickman). On distingue deux types d'acteurs :")
    LP(doc, "Acteur primaire : C'est l'acteur qui initie l'interaction avec le systeme pour atteindre un objectif. Il se situe generalement a gauche du diagramme.")
    LP(doc, "Acteur secondaire : C'est l'acteur sollicite par le systeme pour completer un cas d'utilisation. Il se situe generalement a droite du diagramme.")

    B(doc, "2. Cas d'utilisation (Use Cases)", bold=True)
    B(doc, "Les cas d'utilisation representent les fonctionnalites ou services fournis par le systeme aux acteurs. Ils sont representes par des ovales (ellipses) contenant le nom de la fonctionnalite. Chaque cas d'utilisation decrit une sequence d'actions que le systeme execute pour produire un resultat observable et utile pour un acteur. Exemples : S'authentifier, Effectuer un transfert, Gerer les cartes virtuelles.")

    B(doc, "3. Frontiere du systeme (System Boundary)", bold=True)
    B(doc, "La frontiere du systeme represente la limite du systeme etudie et contient tous les cas d'utilisation. Elle est representee par un rectangle englobant les ovales. Les acteurs se situent en dehors de cette frontiere, indiquant qu'ils sont externes au systeme.")

    B(doc, "4. Relations (Relationships)", bold=True)
    B(doc, "Les relations definissent les liens entre les acteurs et les cas d'utilisation, ou entre les cas d'utilisation eux-memes :")
    LP(doc, "Association : Lien direct entre un acteur et un cas d'utilisation, represente par un trait continu. Elle indique que l'acteur participe a ce cas d'utilisation.")
    LP(doc, "Include (<<include>>) : Relation d'inclusion obligatoire entre deux cas d'utilisation. Le cas d'utilisation inclus est systematiquement execute lorsque le cas d'utilisation de base est declenche. Exemple : 'Effectuer un transfert' inclut 'Verifier le solde'.")
    LP(doc, "Extend (<<extend>>) : Relation d'extension conditionnelle. Le cas d'utilisation etendu n'est execute que dans certaines conditions specifiees par un point d'extension. Exemple : 'S'authentifier' peut etre etendu par 'Reinitialiser le mot de passe'.")
    LP(doc, "Generalisation : Relation d'heritage entre deux acteurs ou entre deux cas d'utilisation. Elle permet de factoriser les comportements communs. Exemple : 'Utilisateur' generalise 'Client' et 'Administrateur'.")

    # ── 2.2 Rappels theoriques Sequence ──────────────────────────────────────
    doc.add_heading("2.2  Rappels theoriques : Diagramme de sequence", 2)
    B(doc, "Le diagramme de sequence est un diagramme UML d'interaction qui represente les echanges de messages entre les objets participants dans un ordre chronologique. Il est utilise pour modeliser le comportement dynamique du systeme pour un scenario specifique d'un cas d'utilisation.")
    B(doc, "Les composants et symboles fondamentaux d'un diagramme de sequence sont les suivants :")

    B(doc, "Acteur (Actor) :", bold=True)
    B(doc, "L'acteur represente une entite externe au systeme qui interagit avec celui-ci. Il peut s'agir d'un utilisateur humain, d'un autre systeme informatique ou d'un dispositif materiel. L'acteur declenche generalement le scenario represente dans le diagramme de sequence et echange des messages avec le systeme pour realiser une fonctionnalite specifique. Il est represente par le meme pictogramme (stickman) que dans le diagramme de cas d'utilisation.")

    B(doc, "Objet / Participant :", bold=True)
    B(doc, "Le participant (ou objet) represente une instance d'une classe prenant part a l'interaction. Il correspond a un composant du systeme (controleur, service, modele, vue) qui recoit et envoie des messages. Le participant est represente par un rectangle contenant le nom de l'objet et/ou de sa classe, suivi d'une ligne de vie verticale en pointilles.")

    B(doc, "Ligne de vie (Lifeline) :", bold=True)
    B(doc, "La ligne de vie est une ligne verticale en pointilles qui descend a partir d'un acteur ou d'un objet. Elle represente l'existence de l'objet dans le temps, du haut (debut de l'interaction) vers le bas (fin de l'interaction). L'epaisseur de la ligne peut varier pour indiquer les periodes d'activite (barres d'activation).")

    B(doc, "Barre d'activation :", bold=True)
    B(doc, "La barre d'activation est un rectangle mince place sur la ligne de vie. Elle represente la periode pendant laquelle un objet execute une operation ou attend un retour. La superposition de barres d'activation indique des appels recursifs ou imbriques.")

    B(doc, "Message :", bold=True)
    B(doc, "Le message represente une communication entre deux participants. Il indique l'appel d'une operation, l'envoi d'un signal, ou la creation/destruction d'un objet. Les messages sont representes par des fleches horizontales entre les lignes de vie.")
    B(doc, "Types de messages :", bold=True)
    LP(doc, "Message synchrone : Represente par une fleche pleine avec pointe noire. L'expediteur attend obligatoirement que le destinataire ait termine le traitement avant de continuer. C'est le type le plus courant dans les architectures request/response.")
    LP(doc, "Message asynchrone : Represente par une fleche avec pointe ouverte. L'expediteur n'attend pas la fin du traitement du destinataire et continue son execution. Utilise pour les notifications push, les evenements, et les appels non-bloquants.")
    LP(doc, "Message de retour : Represente par une fleche en pointilles. Il indique la valeur ou le resultat renvoye par le destinataire a l'expediteur apres le traitement.")
    LP(doc, "Message de creation : Represente par une fleche pointillee vers le rectangle du nouvel objet. Il indique la creation d'une nouvelle instance.")
    LP(doc, "Message de destruction : Represente par un symbole X a la fin de la ligne de vie. Il indique la destruction de l'objet.")

    B(doc, "Fragment combine :", bold=True)
    B(doc, "Les fragments combines permettent de representer des structures de controle dans le diagramme :")
    LP(doc, "alt (alternative) : Equivalent d'un if/else. Permet de representer des scenarios alternatifs.")
    LP(doc, "opt (optional) : Equivalent d'un if sans else. Le fragment n'est execute que si la condition est vraie.")
    LP(doc, "loop : Represente une boucle. Le fragment est repete tant que la condition est satisfaite.")
    LP(doc, "break : Represente une interruption du flux normal (exception, erreur).")

    B(doc, "[Figure N\u00b011 : Composants d'un diagramme de sequence - A inserer]", italic=True)

    # ── 2.3 Identification des Acteurs ───────────────────────────────────────
    doc.add_heading("2.3  Identification des Acteurs", 2)
    B(doc, "L'identification des acteurs constitue la premiere etape de l'analyse des besoins. Un acteur est toute entite externe au systeme qui interagit avec celui-ci. La plateforme TrustDesk met en interaction trois categories d'acteurs :")

    doc.add_heading("Acteur 1 : Client (Utilisateur Bancaire)", 3)
    B(doc, "Le client est l'utilisateur final du portefeuille electronique. Il est titulaire d'un compte bancaire dans l'une des banques partenaires de BEYN. Il accede a la plateforme exclusivement via l'application mobile Flutter. Le client est un acteur primaire car il initie les interactions avec le systeme.")
    B(doc, "Ses interactions principales avec le systeme incluent :")
    LP(doc, "Inscription avec verification eKYC complete (piece d'identite + selfie + liveness detection).")
    LP(doc, "Connexion securisee par token Sanctum avec support de la biometrie sur mobile (empreinte digitale ou reconnaissance faciale).")
    LP(doc, "Consultation du solde et de l'historique detaille des transactions en temps reel avec filtres par date, categorie et montant.")
    LP(doc, "Transferts P2P instantanes entre utilisateurs de la plateforme avec confirmation biometrique.")
    LP(doc, "Generation et gestion de cartes virtuelles securisees (gel/degel, definition de limites, creation de cartes a usage unique).")
    LP(doc, "Paiement par QR code chez les commercants partenaires du reseau BEYN.")
    LP(doc, "Consultation des insights financiers avec categorisation automatique des depenses et graphiques analytiques.")
    LP(doc, "Gestion du profil personnel (photo, telephone, mot de passe) et parametres de notification.")

    doc.add_heading("Acteur 2 : Administrateur", 3)
    B(doc, "L'administrateur est le responsable bancaire ou l'operateur BEYN qui supervise la plateforme via le tableau de bord web React.js. Il est egalement un acteur primaire. L'administrateur possede des privileges eleves lui permettant d'agir sur les comptes clients et la configuration du systeme.")
    B(doc, "Ses responsabilites detaillees incluent :")
    LP(doc, "Gestion complete des comptes clients : activation, desactivation temporaire, blocage permanent, et modification des informations.")
    LP(doc, "Validation ou rejet des dossiers eKYC avec possibilite de demander des documents complementaires.")
    LP(doc, "Suivi des incidents de securite : creation, assignation a un operateur, changement de priorite, et cloture avec rapport.")
    LP(doc, "Generation et exportation de rapports detailles (clients actifs, volume de transactions, incidents, audit trail).")
    LP(doc, "Consultation du dashboard avec KPIs en temps reel : nombre d'utilisateurs, volume transactionnel, taux de validation eKYC, incidents ouverts.")
    LP(doc, "Configuration des parametres systeme : limites de transfert, seuils d'alerte IA, et politiques de securite.")

    doc.add_heading("Acteur 3 : Systeme (Moteur IA + API)", 3)
    B(doc, "Le systeme represente les composants automatises de la plateforme qui operent sans intervention humaine directe. C'est un acteur secondaire car il est sollicite par les cas d'utilisation des acteurs primaires pour completer certaines operations.")
    B(doc, "Les fonctions automatisees du systeme incluent :")
    LP(doc, "Triage intelligent des requetes par IA : classification de l'intention de l'utilisateur et routage vers le module metier appropriate.")
    LP(doc, "Detection automatique d'anomalies transactionnelles basee sur des modeles statistiques (montant inhabituel, frequence anormale, localisation suspecte).")
    LP(doc, "Gel preventif automatique du portefeuille en cas de suspicion de fraude avec notification immediate a l'administrateur.")
    LP(doc, "Envoi de notifications push en temps reel aux utilisateurs concernes (confirmation de transfert, alerte de securite, validation eKYC).")
    LP(doc, "Journalisation exhaustive de toutes les operations sensibles dans la piste d'audit avec horodatage, adresse IP, et metadonnees.")

    T(doc, ["Acteur", "Type", "Interface", "Responsabilite Principale"], [
        ["Client",         "Primaire",   "Mobile (Flutter)", "Gestion du portefeuille, transactions, cartes"],
        ["Administrateur", "Primaire",   "Web (React.js)",   "Supervision, gestion clients, validation eKYC"],
        ["Systeme IA",     "Secondaire", "API (Laravel)",    "Automatisation, triage IA, audit trail"],
    ], "Tableau N\u00b05 : Acteurs du systeme TrustDesk")

    # ── 2.4 Diagramme de cas d'utilisation ───────────────────────────────────
    doc.add_heading("2.4  Diagramme de cas d'utilisation global", 2)
    B(doc, "Le diagramme de cas d'utilisation global ci-dessous represente l'ensemble des fonctionnalites offertes par la plateforme TrustDesk et les interactions entre les trois acteurs identifies et le systeme. La frontiere du systeme englobe tous les modules fonctionnels de la plateforme.")
    add_image(doc, DIAGRAMS / "use_case.png", 14, "Figure N\u00b08 : Diagramme de cas d'utilisation global - TrustDesk")

    B(doc, "Interpretation du diagramme :", bold=True)
    B(doc, "Le diagramme met en evidence la repartition des fonctionnalites entre les trois acteurs. Le Client interagit principalement avec les modules d'authentification, de portefeuille, et de cartes virtuelles. L'Administrateur gere les clients et les incidents via le tableau de bord. Le Systeme IA intervient automatiquement pour le triage des transactions et la detection d'anomalies.")

    doc.add_heading("Diagramme de cas d'utilisation pour le Client", 3)
    B(doc, "Le diagramme suivant detaille les cas d'utilisation specifiques au Client. L'acteur Client interagit avec sept cas d'utilisation principaux, dont certains incluent des sous-cas d'utilisation obligatoires (<<include>>) :")
    LP(doc, "S'inscrire / Se connecter : Inclut obligatoirement 'Verifier les credentials' et peut etre etendu par 'Reinitialiser le mot de passe'.")
    LP(doc, "Completer le eKYC : Inclut 'Capturer la piece d'identite', 'Prendre un selfie', et 'Effectuer la detection de vivacite'.")
    LP(doc, "Consulter le portefeuille : Inclut 'Afficher le solde' et 'Afficher l'historique des transactions'.")
    LP(doc, "Effectuer un transfert P2P : Inclut 'Verifier le solde' et 'Confirmer par biometrie'. Peut etre etendu par 'Annuler le transfert'.")
    LP(doc, "Paiement par QR code : Inclut 'Scanner le QR code' et 'Verifier le solde'.")
    LP(doc, "Gerer les cartes virtuelles : Inclut 'Generer une carte' et peut etre etendu par 'Geler/Degeler une carte'.")
    LP(doc, "Consulter les insights financiers : Affiche les graphiques de depenses par categorie et les tendances mensuelles.")
    B(doc, "[Figure N\u00b09 : Diagramme de cas d'utilisation - Client - A inserer]", italic=True)

    doc.add_heading("Diagramme de cas d'utilisation pour l'Administrateur", 3)
    B(doc, "Le diagramme suivant detaille les cas d'utilisation specifiques a l'Administrateur. L'acteur Administrateur interagit avec cinq cas d'utilisation principaux via le tableau de bord web :")
    LP(doc, "S'authentifier (Admin) : Connexion au dashboard avec credentials et double facteur d'authentification (2FA) optionnel.")
    LP(doc, "Gerer les clients / KYC : Inclut 'Lister les clients', 'Valider/Rejeter un dossier eKYC', et 'Modifier le statut d'un client'.")
    LP(doc, "Gerer les incidents : Inclut 'Creer un incident', 'Assigner un incident', 'Changer la priorite', et 'Cloturer un incident'.")
    LP(doc, "Consulter le dashboard : Affiche les KPIs en temps reel (utilisateurs actifs, volume transactionnel, incidents ouverts, taux de validation).")
    LP(doc, "Generer des rapports : Inclut 'Exporter en CSV/PDF' et peut etre filtre par date, type, et statut.")
    B(doc, "[Figure N\u00b010 : Diagramme de cas d'utilisation - Administrateur - A inserer]", italic=True)

    # ── 2.5 Besoins fonctionnels ─────────────────────────────────────────────
    doc.add_heading("2.5  Specification des Besoins Fonctionnels", 2)
    B(doc, "Les besoins fonctionnels decrivent les fonctionnalites que le systeme doit fournir aux utilisateurs. Ils sont organises par module fonctionnel :")

    doc.add_heading("Module Authentification et eKYC", 3)
    B(doc, "Ce module gere l'ensemble du processus d'identification et d'authentification des utilisateurs :")
    LP(doc, "Inscription par email avec verification OTP (One-Time Password) envoye par email.")
    LP(doc, "Connexion securisee par token Sanctum avec gestion des sessions multiples et deconnexion a distance.")
    LP(doc, "Support de la biometrie sur mobile (empreinte digitale via BiometricPrompt Android, Face ID iOS).")
    LP(doc, "Parcours eKYC complet : upload de piece d'identite (recto/verso), selfie avec detection de vivacite (liveness detection).")
    LP(doc, "Gestion des statuts KYC : en attente (pending), verifie (verified), rejete (rejected), expire (expired).")
    LP(doc, "Reinitialisation de mot de passe par lien securise avec token a usage unique et expiration de 60 minutes.")
    LP(doc, "Deconnexion multi-appareils avec revocation de tous les tokens Sanctum actifs.")

    doc.add_heading("Module Portefeuille et Transactions", 3)
    B(doc, "Ce module constitue le coeur fonctionnel de la plateforme :")
    LP(doc, "Consultation du solde en temps reel avec mise a jour instantanee via Supabase Realtime.")
    LP(doc, "Historique des transactions avec filtres avances par date, categorie, montant, et type (credit/debit).")
    LP(doc, "Transferts P2P instantanes entre utilisateurs de la plateforme avec confirmation biometrique obligatoire.")
    LP(doc, "Rechargement du portefeuille (top-up) via integration avec les APIs bancaires partenaires.")
    LP(doc, "Paiement de factures (telephone, internet, electricite) via les passerelles de paiement integrees.")
    LP(doc, "Paiement par QR code chez les commercants partenaires du reseau BEYN.")
    LP(doc, "Definition de limites de depenses quotidiennes et mensuelles parametrables par l'utilisateur.")
    LP(doc, "Categorisation automatique des depenses (alimentation, transport, loisirs, sante, etc.).")

    doc.add_heading("Module Cartes Virtuelles", 3)
    B(doc, "Ce module permet la generation et la gestion de cartes virtuelles securisees :")
    LP(doc, "Generation de cartes virtuelles avec chiffrement asymetrique RSA-2048 des numeros de carte et CVV.")
    LP(doc, "Support de deux types de cartes : cartes permanentes (reutilisables) et cartes jetables (disposable) a usage unique.")
    LP(doc, "Gel et degel instantane de carte depuis l'application mobile avec notification de confirmation.")
    LP(doc, "Definition de limites de depenses par carte (montant maximum par transaction, plafond mensuel).")
    LP(doc, "Affichage securise des informations de carte avec masquage partiel et revelation temporaire.")

    doc.add_heading("Module Triage IA et Securite", 3)
    B(doc, "Ce module integre l'intelligence artificielle pour la securite transactionnelle :")
    LP(doc, "Classificateur d'intention IA pour le routage automatique des requetes vers le module metier appropriate.")
    LP(doc, "Detection d'anomalies transactionnelles en temps reel basee sur l'analyse statistique (montant inhabituel, frequence anormale, localisation geographique suspecte, changement de pattern de depenses).")
    LP(doc, "Gel preventif automatique du portefeuille en cas de suspicion de fraude avec score de risque.")
    LP(doc, "Piste d'audit complete et immuable pour toutes les operations sensibles (creation, lecture, modification, suppression).")
    LP(doc, "Journalisation des tentatives d'acces echouees et detection de brute-force.")

    doc.add_heading("Module Administration", 3)
    B(doc, "Ce module fournit aux administrateurs les outils de supervision et de gestion :")
    LP(doc, "Dashboard avec KPIs en temps reel : utilisateurs actifs, volume de transactions journalier/mensuel, incidents ouverts, taux de validation eKYC.")
    LP(doc, "Gestion CRUD des clients avec filtres par statut KYC, date d'inscription, et niveau d'activite.")
    LP(doc, "Interface de validation et rejet des documents eKYC avec visualisation des pieces d'identite et des selfies.")
    LP(doc, "Gestion complete du cycle de vie des incidents : creation, assignation, suivi, escalade, et cloture avec rapport.")
    LP(doc, "Generation et exportation de rapports detailles aux formats CSV et PDF (clients, transactions, audit, tickets).")
    LP(doc, "Historique des actions administrateur avec piste d'audit dediee pour la conformite reglementaire.")

    T(doc, ["Module", "Fonctionnalites Principales"], [
        ["Authentification et eKYC",   "Inscription OTP, connexion Sanctum, biometrie, parcours eKYC, liveness"],
        ["Portefeuille et Transactions","Solde temps reel, transferts P2P, paiements QR, factures, categorisation"],
        ["Cartes Virtuelles",          "Generation RSA-2048, cartes jetables, gel/degel, limites par carte"],
        ["Triage IA et Securite",      "Classification IA, detection anomalies, gel preventif, audit trail"],
        ["Administration",             "Dashboard KPIs, gestion clients, validation eKYC, rapports, audit admin"],
    ], "Tableau N\u00b06 : Besoins fonctionnels par module")

    # ── 2.6 Besoins non fonctionnels ─────────────────────────────────────────
    doc.add_heading("2.6  Specification des Besoins Non Fonctionnels", 2)
    B(doc, "Les besoins non fonctionnels definissent les contraintes de qualite que la plateforme doit respecter. Ils sont aussi importants que les besoins fonctionnels car ils determinent la viabilite et l'acceptabilite de la solution en environnement de production bancaire.")

    B(doc, "Securite :", bold=True)
    B(doc, "La securite est la contrainte la plus critique pour une application bancaire. La plateforme doit implementer le chiffrement RSA-2048 pour les cartes virtuelles, AES-256 pour les donnees en transit, et la conformite OWASP Top 10 pour la protection contre les vulnerabilites web courantes (injection SQL, XSS, CSRF, mass-assignment).")

    B(doc, "Performance :", bold=True)
    B(doc, "Le temps de reponse de l'API doit etre inferieur a 300 millisecondes pour 95% des requetes, avec un temps de reponse P99 inferieur a 1 seconde. L'application mobile doit atteindre 60 FPS pour les animations et les transitions.")

    B(doc, "Disponibilite :", bold=True)
    B(doc, "L'objectif de disponibilite est de 99.9% (SLA bancaire), ce qui correspond a un temps d'indisponibilite maximal de 8.76 heures par an. Un mecanisme de health-check et de monitoring doit etre mis en place.")

    B(doc, "Scalabilite :", bold=True)
    B(doc, "L'architecture doit etre horizontalement extensible pour supporter l'ajout de nouvelles banques partenaires sans modification du code existant. Le systeme doit supporter au minimum 10 000 utilisateurs simultanes.")

    B(doc, "Ergonomie :", bold=True)
    B(doc, "L'interface utilisateur doit etre conforme aux standards Material Design 3 (mobile) et aux bonnes pratiques de design web (dashboard). L'application doit etre utilisable par des personnes avec un niveau de maitrise technologique basique.")

    B(doc, "Compatibilite :", bold=True)
    B(doc, "L'application mobile doit supporter Android 8.0 (API 26) et superieur, ainsi qu'iOS 12 et superieur. Le dashboard web doit fonctionner sur les navigateurs modernes : Chrome 90+, Firefox 88+, Edge 90+, et Safari 14+.")

    B(doc, "Maintenabilite :", bold=True)
    B(doc, "Le code doit suivre une architecture modulaire avec separation des responsabilites (Controllers, Services, Models). La documentation du code, les tests automatises (unitaires et d'integration), et le versionning Git sont obligatoires.")

    B(doc, "Auditabilite :", bold=True)
    B(doc, "Toutes les operations sensibles doivent etre journalisees avec horodatage, identifiant utilisateur, adresse IP, type d'action, et metadonnees contextuelles. Les journaux d'audit doivent etre immuables et conserves pendant une duree minimale de 5 ans.")

    T(doc, ["Exigence", "Description"], [
        ["Securite",      "Chiffrement RSA-2048/AES-256, conformite OWASP Top 10, protection mass-assignment"],
        ["Performance",   "Temps de reponse API < 300ms (P95), < 1s (P99), 60 FPS mobile"],
        ["Disponibilite", "SLA 99.9%, health-check, monitoring"],
        ["Scalabilite",   "Architecture horizontale, 10 000+ utilisateurs simultanes"],
        ["Ergonomie",     "Material Design 3, accessibilite, parcours utilisateur intuitif"],
        ["Compatibilite", "Android 8.0+, iOS 12+, Chrome 90+, Firefox 88+, Edge 90+, Safari 14+"],
        ["Maintenabilite","Architecture modulaire, tests automatises, documentation, Git"],
        ["Auditabilite",  "Journalisation complete, conservation 5 ans, immuabilite"],
        ["Conformite",    "Directives Banque d'Algerie, LCB-FT, protection des donnees personnelles"],
    ], "Tableau N\u00b07 : Besoins non fonctionnels")

    # ── 2.7 Descriptions textuelles ──────────────────────────────────────────
    doc.add_heading("2.7  Descriptions Textuelles des Cas d'Utilisation", 2)
    B(doc, "Les descriptions textuelles detaillent les scenarios nominaux et alternatifs des principaux cas d'utilisation. Elles servent de specification pour l'implementation et les tests.")

    B(doc, "Description textuelle du cas d'utilisation : S'inscrire et completer le eKYC", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "S'inscrire et completer le eKYC"],
        ["Acteur principal",    "Client"],
        ["Acteurs secondaires", "Systeme IA, Administrateur"],
        ["Precondition",        "Application mobile installee, connexion internet disponible, client non inscrit"],
        ["Scenario nominal",    "1. Le client ouvre l'application et choisit 'Creer un compte'\n2. Il saisit ses informations personnelles (nom, prenom, email, mot de passe)\n3. Le systeme envoie un code OTP a l'adresse email\n4. Le client saisit le code OTP pour valider son email\n5. Il uploade sa piece d'identite (recto et verso)\n6. Il prend un selfie avec activation de la detection de vivacite\n7. Le systeme verifie la concordance faciale document/selfie\n8. Le compte est cree avec statut KYC 'en attente de validation'\n9. Un administrateur examine le dossier et valide\n10. Le statut KYC passe a 'verifie', le portefeuille est active"],
        ["Scenario alternatif", "3a. Email deja utilise -> message d'erreur, proposition de connexion\n4a. Code OTP expire -> renvoi automatique d'un nouveau code\n6a. Detection de vivacite echouee -> nouvelle tentative (max 3)\n7a. Document illisible ou non conforme -> demande de re-upload\n9a. Dossier rejete -> notification au client avec motif de rejet"],
        ["Postcondition",       "Compte cree et verifie, portefeuille initialise a 0 DZD, client pret a utiliser les services"],
    ], "Tableau N\u00b08 : Description UC - Inscription eKYC")

    B(doc, "Description textuelle du cas d'utilisation : Effectuer un transfert P2P", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "Effectuer un transfert P2P"],
        ["Acteur principal",    "Client"],
        ["Acteurs secondaires", "Systeme IA"],
        ["Precondition",        "Client authentifie, KYC verifie, solde suffisant, destinataire existant"],
        ["Scenario nominal",    "1. Le client selectionne 'Transfert' dans le menu principal\n2. Il saisit l'identifiant du destinataire (email ou telephone)\n3. Le systeme verifie l'existence du destinataire\n4. Le client saisit le montant a transferer\n5. Le systeme verifie le solde disponible et les limites\n6. Le moteur IA analyse la transaction (detection d'anomalies)\n7. Le client confirme avec authentification biometrique\n8. La transaction est executee atomiquement (debit + credit)\n9. Les deux parties recoivent une notification push\n10. La transaction est journalisee dans l'audit trail"],
        ["Scenario alternatif", "3a. Destinataire inexistant -> message d'erreur\n5a. Solde insuffisant -> message d'erreur avec solde actuel\n5b. Limite depassee -> message d'erreur avec limite restante\n6a. Transaction flaggee comme suspecte -> blocage + alerte admin\n7a. Biometrie echouee (3 tentatives) -> annulation et blocage temporaire"],
        ["Postcondition",       "Soldes mis a jour, transaction enregistree, notifications envoyees, audit trail mis a jour"],
    ], "Tableau N\u00b09 : Description UC - Transfert P2P")

    B(doc, "Description textuelle du cas d'utilisation : Gerer les cartes virtuelles", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "Gerer les cartes virtuelles"],
        ["Acteur principal",    "Client"],
        ["Acteurs secondaires", "Systeme (EncryptionService)"],
        ["Precondition",        "Client authentifie, KYC verifie, au moins une carte virtuelle active ou possibilite d'en creer"],
        ["Scenario nominal",    "1. Le client accede a la section 'Cartes' de l'application\n2. Il choisit 'Generer une nouvelle carte'\n3. Il selectionne le type de carte (permanente ou jetable)\n4. Le systeme genere une paire de cles RSA-2048\n5. Le numero de carte et le CVV sont generes et chiffres avec la cle publique\n6. La cle privee est stockee dans le Secure Enclave du smartphone\n7. La carte est creee et affichee avec le numero masque\n8. Le client peut reveler temporairement les informations completes"],
        ["Scenario alternatif", "2a. Limite de cartes atteinte -> message d'erreur\n4a. Erreur de generation des cles -> nouvelle tentative\n7a. Carte jetable utilisee -> destruction automatique apres usage"],
        ["Postcondition",       "Carte virtuelle creee et securisee, prete a l'utilisation pour les paiements en ligne"],
    ], "Tableau N\u00b010 : Description UC - Gestion des cartes virtuelles")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce deuxieme chapitre a permis de formaliser de maniere exhaustive les besoins de la plateforme TrustDesk. Nous avons commence par les rappels theoriques sur les diagrammes de cas d'utilisation et de sequence, puis nous avons identifie trois acteurs cles : le Client (utilisateur mobile), l'Administrateur (operateur back-office), et le Systeme IA (composant automatise).")
    B(doc, "Les diagrammes de cas d'utilisation, tant global que par acteur, ont permis de visualiser l'ensemble des interactions fonctionnelles. Les cinq modules fonctionnels (Authentification/eKYC, Portefeuille/Transactions, Cartes Virtuelles, Triage IA/Securite, Administration) ont ete detailles avec leurs besoins specifiques. Les besoins non fonctionnels (securite, performance, disponibilite, scalabilite, ergonomie, auditabilite) ont ete specifies avec des criteres mesurables.")
    B(doc, "Enfin, les descriptions textuelles des trois principaux cas d'utilisation (Inscription eKYC, Transfert P2P, Gestion des cartes virtuelles) ont precise les scenarios nominaux et alternatifs. Le chapitre suivant presente la conception architecturale detaillee de la solution.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE III : CONCEPTION (start only)
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch3_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre III : Conception", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce chapitre presente la conception detaillee de la plateforme TrustDesk. Nous decrirons l'architecture globale du systeme, les couches de l'API Laravel, le modele de donnees (MCD et MLD), les diagrammes de classes, les diagrammes de sequence detailles, ainsi que l'architecture de chiffrement garantissant la securite des donnees bancaires.", sb=12)

    # ── 3.1 Architecture globale ─────────────────────────────────────────────
    doc.add_heading("3.1  Architecture Globale du Systeme", 2)
    B(doc, "La plateforme TrustDesk adopte une architecture trois tiers (Three-Tier Architecture) composee de trois couches distinctes communicant via des APIs RESTful securisees :")
    LP(doc, "Couche Presentation : Application mobile (Flutter 3 / Dart / Riverpod) et Dashboard web (React.js 18 / TypeScript / Vite). Ces deux clients communiquent avec la couche metier exclusivement via des appels API REST.")
    LP(doc, "Couche Metier : API RESTful Laravel 11 avec authentification Sanctum, modules de services metier, EncryptionService pour le chiffrement RSA/AES, et moteur IA de triage. Cette couche centralise toute la logique applicative.")
    LP(doc, "Couche Donnees : Base de donnees relationnelle (PostgreSQL via Supabase) pour le stockage persistant, avec Supabase Realtime pour la synchronisation en temps reel et Supabase Storage pour les fichiers KYC.")
    add_image(doc, DIAGRAMS / "architecture.png", 15, "Figure N\u00b013 : Architecture globale trois tiers - TrustDesk")

    doc.add_heading("3.1.1  Architecture Detaillee de l'API Laravel", 3)
    B(doc, "L'API Laravel suit une architecture en couches strictement separees :")
    T(doc, ["Couche", "Composant Laravel", "Role"], [
        ["Routes",          "routes/api.php",              "Definition des endpoints REST"],
        ["Middleware",      "app/Http/Middleware",          "Authentification, rate limiting, CORS"],
        ["Form Requests",   "app/Http/Requests",           "Validation stricte + rejet champs inattendus"],
        ["Controllers",     "app/Http/Controllers",        "Logique de coordination"],
        ["Services",        "app/Services",                "Logique metier (portefeuille, cartes, triage)"],
        ["Models",          "app/Models",                  "Entites Eloquent ORM + relations"],
        ["EncryptionSvc",   "app/Services/EncryptionService", "Chiffrement RSA/AES des donnees sensibles"],
        ["Events/Listeners","app/Events, app/Listeners",   "Journalisation audit, notifications push"],
    ], "Tableau N\u00b011 : Couches de l'architecture API Laravel")

    # ── 3.2 Diagramme de classes ─────────────────────────────────────────────
    doc.add_heading("3.2  Diagramme de Classes", 2)
    B(doc, "Le diagramme de classes modelise les entites principales du systeme TrustDesk et leurs relations. Les classes centrales sont User, Wallet, Transaction, VirtualCard, KycDocument, Incident et AuditLog :")
    add_image(doc, DIAGRAMS / "class_diagram.png", 15, "Figure N\u00b014 : Diagramme de classes - TrustDesk")

    doc.add_heading("3.2.1  Relations Entre les Classes", 3)
    LP(doc, "User (1) <-> (1) Wallet : Chaque utilisateur possede un portefeuille unique (relation 1:1).")
    LP(doc, "User (1) <-> (0..*) VirtualCard : Un utilisateur peut avoir plusieurs cartes virtuelles.")
    LP(doc, "User (1) <-> (0..*) KycDocument : Un utilisateur soumet un ou plusieurs documents eKYC.")
    LP(doc, "Wallet (1) <-> (0..*) Transaction : Un portefeuille est associe a plusieurs transactions.")
    LP(doc, "User (1) <-> (0..*) Incident : Un utilisateur peut etre lie a plusieurs incidents de securite.")
    LP(doc, "User (1) <-> (0..*) AuditLog : Chaque action sensible genere un enregistrement d'audit.")

    # ── 3.3 MCD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.3  Modele Conceptuel des Donnees (MCD)", 2)
    B(doc, "Le MCD represente les entites du systeme et leurs associations dans un formalisme Entite-Association. Les cardinalites refletent les regles metier du portefeuille electronique bancaire :")
    add_image(doc, DIAGRAMS / "mcd.png", 15, "Figure N\u00b015 : Modele Conceptuel des Donnees (MCD)")

    # ── 3.4 MLD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.4  Modele Logique des Donnees (MLD)", 2)
    B(doc, "La transformation du MCD en MLD produit les tables relationnelles suivantes :")
    B(doc, "users (id, name, email, phone, password, role, avatar, kyc_status, is_active, created_at)", bold=True)
    B(doc, "wallets (id, #user_id, balance, currency, status, frozen_at, daily_limit, created_at)", bold=True)
    B(doc, "transactions (id, #wallet_id, type, amount, #recipient_id, description, status, reference, category, created_at)", bold=True)
    B(doc, "virtual_cards (id, #user_id, card_number_enc, cvv_enc, expiry, status, is_disposable, public_key, created_at)", bold=True)
    B(doc, "kyc_documents (id, #user_id, document_type, document_url, selfie_url, liveness_score, status, created_at)", bold=True)
    B(doc, "incidents (id, #user_id, type, description, priority, status, #assigned_to, ai_triage_result, created_at)", bold=True)
    B(doc, "audit_logs (id, #user_id, action, entity_type, entity_id, ip_address, metadata, created_at)", bold=True)
    B(doc, "Note : Le symbole # indique une cle etrangere (FK).", italic=True)

    # ── 3.5 Diagrammes de sequence ───────────────────────────────────────────
    doc.add_heading("3.5  Diagrammes de Sequence", 2)
    B(doc, "Les diagrammes de sequence detaillent le flux d'interactions entre les composants du systeme pour les deux cas d'utilisation principaux.")

    doc.add_heading("3.5.1  Sequence : Inscription et eKYC", 3)
    B(doc, "Ce diagramme montre le flux complet d'inscription d'un nouveau client, incluant la verification OTP, l'upload des documents d'identite, la detection de vivacite, et la validation par un administrateur :")
    add_image(doc, DIAGRAMS / "sequence_ekyc.png", 14, "Figure N\u00b016 : Diagramme de sequence - Inscription et eKYC")

    doc.add_heading("3.5.2  Sequence : Transfert P2P", 3)
    B(doc, "Ce diagramme illustre le processus de transfert d'argent entre deux utilisateurs, incluant la verification du solde, l'analyse IA de la transaction, la confirmation biometrique, et la mise a jour des soldes :")
    add_image(doc, DIAGRAMS / "sequence_p2p.png", 14, "Figure N\u00b017 : Diagramme de sequence - Transfert P2P")

    # ── 3.6 Architecture de chiffrement ──────────────────────────────────────
    doc.add_heading("3.6  Architecture de Chiffrement", 2)
    B(doc, "La securite des donnees sensibles (numeros de carte, CVV) repose sur un modele de chiffrement hybride combinant chiffrement asymetrique et symetrique :")
    doc.add_heading("Chiffrement Asymetrique (RSA-2048)", 3)
    LP(doc, "Utilise pour le chiffrement des numeros de cartes virtuelles et CVV.")
    LP(doc, "Chaque client possede une paire de cles (publique/privee) generee sur son appareil.")
    LP(doc, "La cle publique est transmise au serveur pour le chiffrement.")
    LP(doc, "La cle privee reste exclusivement dans le Secure Enclave du smartphone (approche Zero-Knowledge cote backend).")
    doc.add_heading("Chiffrement Symetrique (AES-256)", 3)
    LP(doc, "Utilise pour le chiffrement des payloads volumineux en transit.")
    LP(doc, "La cle AES est elle-meme chiffree par RSA pour le transport (enveloppe hybride).")
    LP(doc, "AES-256 en mode CBC avec IV aleatoire pour chaque operation.")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce troisieme chapitre a permis de detailler la conception complete de TrustDesk : architecture trois tiers, couches de l'API Laravel, modele de donnees (MCD et MLD), diagrammes de classes et de sequence, ainsi que l'architecture de chiffrement hybride RSA/AES garantissant la securite des donnees bancaires. Les chapitres suivants presenteront la realisation concrete et les resultats des tests.")

    # ── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    size = OUT.stat().st_size // 1024
    print(f"\nMemoire sauvegarde : {OUT}  ({size} KB)")

    # Stats
    secs = len(doc.sections)
    paras = len(doc.paragraphs)
    tbls = len(doc.tables)
    imgs = len(doc.inline_shapes)
    print(f"Sections: {secs}, Paragraphs: {paras}, Tables: {tbls}, Images: {imgs}")

    # Count by chapter
    ch1_s = ch1_e = ch2_s = ch2_e = ch3_s = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip(); s = p.style.name
        if 'Chapitre I :' in t and s == 'Heading 1': ch1_s = i
        elif 'Chapitre II :' in t and s == 'Heading 1': ch2_s = i; ch1_e = i - 1
        elif 'Chapitre III :' in t and s == 'Heading 1': ch3_s = i; ch2_e = i - 1
    if ch1_s and ch1_e:
        c = sum(1 for i in range(ch1_s, ch1_e+1) if doc.paragraphs[i].text.strip())
        print(f"Ch1 non-empty: {c}")
    if ch2_s and ch2_e:
        c = sum(1 for i in range(ch2_s, ch2_e+1) if doc.paragraphs[i].text.strip())
        print(f"Ch2 non-empty: {c}")
    if ch3_s:
        c = sum(1 for i in range(ch3_s, len(doc.paragraphs)) if doc.paragraphs[i].text.strip())
        print(f"Ch3 non-empty: {c}")

if __name__ == "__main__":
    print("Building memoire v9 (EXPANDED)...")
    build()
    print("Done.")
