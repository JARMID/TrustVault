# -*- coding: utf-8 -*-
"""TrustDesk Memoire - Complete Generator with proper Algerian formatting"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page Setup
for s in doc.sections:
    s.top_margin = Cm(2.36)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2)
    s.right_margin = Cm(1.25)

# Styles
sn = doc.styles['Normal']; sn.font.name='Times New Roman'; sn.font.size=Pt(12); sn.paragraph_format.line_spacing=1.5
for lvl,sz,clr in [('Heading 1',16,True),('Heading 2',14,True),('Heading 3',12,True)]:
    h=doc.styles[lvl]; h.font.size=Pt(sz); h.font.name='Times New Roman'; h.font.color.rgb=RGBColor(0x1F,0x49,0x7D)
bt=doc.styles['Body Text']; bt.font.name='Times New Roman'; bt.font.size=Pt(12); bt.paragraph_format.line_spacing=1.5

def C(t,sz=12,b=False,sa=0,sb=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(sb)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
    return p
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
    return p
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'
def IMG(caption):
    """Placeholder for image"""
    B(f"\n[>>> INSERER ICI : {caption} <<<]\n", it=True)
def new_section(header_text=""):
    """Add section break with header"""
    doc.add_section()
    sec = doc.sections[-1]
    sec.top_margin=Cm(2.36); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2); sec.right_margin=Cm(1.25)
    if header_text:
        hdr = sec.header
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(header_text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.italic = True

# ════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
C("LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 11, True, 2)
C("MINISTERE DE LA FORMATION ET DE L\u2019ENSEIGNEMENT PROFESSIONNELS", 11, False, 6)
C("INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE", 11, False, 2)
C("EN AUDIOVISUELS", 11, False, 2)
C("Echahid Ahmed Mehdi \u2013 Ouled Fayet \u2013", 11, True, 12)
E()
C("M\u00e9moire De Fin De Formation Pour L\u2019obtention Du Dipl\u00f4me", 13, True, 2)
C("de Technicien Sup\u00e9rieur en Informatique", 13, True, 4)
C("Option : DEVELOPPEMENT WEB ET MOBILE", 12, True, 12)
E()
C("Th\u00e8me :", 14, True, 8)
E()
C("Conception Et R\u00e9alisation D\u2019une", 16, True, 2)
C("Plateforme Web et Mobile", 16, True, 2)
C("de Gestion des Incidents de S\u00e9curit\u00e9", 16, True, 2)
C("en Milieu Universitaire", 16, True, 4)
C("\u00ab TrustDesk \u00bb", 18, True, 12)
E()
C("Organisme d\u2019accueil :", 12, False, 2)
C("Entreprise BEYN", 12, True, 12)
E()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
r=p.add_run("R\u00e9alis\u00e9 Par :"); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=True
p=doc.add_paragraph()
r=p.add_run("M. Khalef Abdelmadjid"); r.font.name='Times New Roman'; r.font.size=Pt(12)
p=doc.add_paragraph()
r=p.add_run("M. Dai Khaled Wassim"); r.font.name='Times New Roman'; r.font.size=Pt(12)
E()
p=doc.add_paragraph()
r=p.add_run("Suivi par :"); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=True
p=doc.add_paragraph()
r=p.add_run("Mme. Himeur"); r.font.name='Times New Roman'; r.font.size=Pt(12)
for _ in range(3): E()
C("Promotion : 2025 / 2026", 14, True, 0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# DEDICACES - Person 1
# ════════════════════════════════════════════════════════════════
new_section()
H1("D\u00e9dicaces", WD_ALIGN_PARAGRAPH.CENTER)
E()
B("Nous tenons \u00e0 exprimer notre gratitude envers toutes les personnes qui ont contribu\u00e9, de pr\u00e8s ou de loin, \u00e0 la r\u00e9alisation de ce travail. Ce m\u00e9moire est le fruit de nombreuses heures de recherche, de d\u00e9veloppement et de r\u00e9flexion, et il n\u2019aurait pas pu voir le jour sans le soutien pr\u00e9cieux de notre entourage.")
B("Un merci particulier \u00e0 nos familles pour leur soutien ind\u00e9fectible et leurs encouragements constants. Leur confiance en nous a \u00e9t\u00e9 une source de motivation inestimable tout au long de ce parcours.")
B("Enfin, nous souhaitons remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont rendu ce travail possible. Vos conseils, vos critiques constructives et votre bienveillance nous ont permis de mener ce projet \u00e0 bien.")
B("\u00c0 toutes et \u00e0 tous, nous vous adressons nos plus sinc\u00e8res remerciements.")
for _ in range(6): E()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=p.add_run("Abdelmadjid K."); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.italic=True

doc.add_page_break()

# DEDICACES - Person 2
H1("D\u00e9dicaces", WD_ALIGN_PARAGRAPH.CENTER)
E()
B("Nous tenons \u00e0 exprimer notre gratitude envers toutes les personnes qui ont contribu\u00e9, de pr\u00e8s ou de loin, \u00e0 la r\u00e9alisation de ce travail. Ce m\u00e9moire est le fruit de nombreuses heures de recherche, de d\u00e9veloppement et de r\u00e9flexion, et il n\u2019aurait pas pu voir le jour sans le soutien pr\u00e9cieux de notre entourage.")
B("Un merci particulier \u00e0 nos familles pour leur soutien ind\u00e9fectible et leurs encouragements constants. Leur confiance en nous a \u00e9t\u00e9 une source de motivation inestimable tout au long de ce parcours.")
B("Enfin, nous souhaitons remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont rendu ce travail possible.")
B("\u00c0 toutes et \u00e0 tous, nous vous adressons nos plus sinc\u00e8res remerciements.")
for _ in range(6): E()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=p.add_run("Khaled Wassim D."); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.italic=True

doc.add_page_break()

# REMERCIEMENTS
H1("REMERCIEMENTS", WD_ALIGN_PARAGRAPH.CENTER)
E()
B("Nous tenons \u00e0 exprimer notre gratitude envers toutes les personnes qui ont contribu\u00e9, de pr\u00e8s ou de loin, \u00e0 la r\u00e9alisation de ce travail.")
B("Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir donn\u00e9 la sant\u00e9, la volont\u00e9 et la patience pour mener \u00e0 bien ce projet.")
B("Nos remerciements les plus sinc\u00e8res vont \u00e0 notre encadreur, Mme. Himeur, pour ses pr\u00e9cieux conseils, sa disponibilit\u00e9 et son suivi rigoureux tout au long de la r\u00e9alisation de ce m\u00e9moire.")
B("Nous remercions \u00e9galement l\u2019ensemble du personnel de l\u2019entreprise BEYN pour leur accueil chaleureux et leur collaboration durant notre stage.")
B("Enfin, nous adressons nos remerciements \u00e0 tous les membres du jury pour l\u2019honneur qu\u2019ils nous font en acceptant d\u2019\u00e9valuer ce travail.")
for _ in range(4): E()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=p.add_run("Abdelmadjid K., Khaled Wassim D."); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.italic=True

doc.add_page_break()

# SOMMAIRE
C("Sommaire", 16, True, 12)
E()
B("[Le sommaire sera g\u00e9n\u00e9r\u00e9 automatiquement dans Word : R\u00e9f\u00e9rences > Table des mati\u00e8res]", it=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════
new_section("Introduction")
C("INTRODUCTION", 16, True, 12)
E()
B("Depuis toujours, les organisations cherchent \u00e0 am\u00e9liorer leurs m\u00e9thodes de gestion et de suivi des activit\u00e9s critiques. Dans le domaine de la s\u00e9curit\u00e9 des campus universitaires, cette qu\u00eate d\u2019am\u00e9lioration se heurte \u00e0 des d\u00e9fis sp\u00e9cifiques li\u00e9s \u00e0 la diversit\u00e9 des acteurs, \u00e0 l\u2019\u00e9tendue g\u00e9ographique des sites, et \u00e0 la nature vari\u00e9e des incidents pouvant survenir.")
B("Avec l\u2019essor des technologies de l\u2019information et de la communication, les \u00e9tablissements d\u2019enseignement sup\u00e9rieur disposent aujourd\u2019hui d\u2019outils puissants pour transformer leurs processus de gestion de la s\u00e9curit\u00e9. Les frameworks de d\u00e9veloppement modernes, les architectures API-first, et les applications mobiles offrent des possibilit\u00e9s de centralisation et de r\u00e9activit\u00e9 sans pr\u00e9c\u00e9dent.")
B("Par ailleurs, la g\u00e9n\u00e9ralisation de l\u2019Internet et le d\u00e9veloppement des technologies mobiles ont profond\u00e9ment modifi\u00e9 les attentes des usagers en mati\u00e8re de r\u00e9activit\u00e9 et d\u2019accessibilit\u00e9 des services. Un \u00e9tudiant confront\u00e9 \u00e0 une situation d\u2019urgence doit pouvoir signaler un incident instantan\u00e9ment depuis son smartphone, tandis que les agents de s\u00e9curit\u00e9 doivent pouvoir recevoir et traiter ces alertes en temps r\u00e9el.")
B("Dans ce contexte, la gestion des incidents de s\u00e9curit\u00e9 constitue un aspect particuli\u00e8rement sensible. Les processus actuels, souvent bas\u00e9s sur des registres papier, des appels t\u00e9l\u00e9phoniques, ou des \u00e9changes informels, ne permettent pas d\u2019assurer une tra\u00e7abilit\u00e9 ad\u00e9quate ni une coordination efficace entre les diff\u00e9rents intervenants.")
B("C\u2019est dans cette perspective que s\u2019inscrit le projet TrustDesk, qui vise \u00e0 concevoir et r\u00e9aliser une plateforme int\u00e9gr\u00e9e de type Security Operations Center (SOC), compos\u00e9e d\u2019une application web de supervision, d\u2019une application mobile de terrain, et d\u2019une API RESTful centralis\u00e9e.")
B("Ainsi, la r\u00e9alisation de ce projet nous am\u00e8ne \u00e0 nous poser la probl\u00e9matique suivante :")
B("\u00ab Comment concevoir et r\u00e9aliser une plateforme web et mobile centralis\u00e9e permettant d\u2019assurer une gestion compl\u00e8te, s\u00e9curis\u00e9e et tra\u00e7able du cycle de vie des incidents de s\u00e9curit\u00e9 en milieu universitaire ? \u00bb", b=True)
B("Pour r\u00e9pondre \u00e0 cette probl\u00e9matique, nous avons structur\u00e9 notre m\u00e9moire en quatre chapitres.")
B("Le premier chapitre, intitul\u00e9 \u00ab \u00c9tude pr\u00e9alable \u00bb, pr\u00e9sente l\u2019organisme d\u2019accueil, le contexte du projet, la probl\u00e9matique identifi\u00e9e, ainsi que la m\u00e9thodologie de travail adopt\u00e9e.")
B("Le deuxi\u00e8me chapitre, intitul\u00e9 \u00ab Analyse et sp\u00e9cification des besoins \u00bb, est consacr\u00e9 \u00e0 l\u2019identification et \u00e0 la formalisation des exigences fonctionnelles et non fonctionnelles du syst\u00e8me.")
B("Le troisi\u00e8me chapitre, intitul\u00e9 \u00ab Conception de l\u2019application \u00bb, pr\u00e9sente la conception globale du syst\u00e8me \u00e0 travers les diagrammes UML et l\u2019architecture logicielle retenue.")
B("Le quatri\u00e8me chapitre, intitul\u00e9 \u00ab R\u00e9alisation et impl\u00e9mentation \u00bb, d\u00e9crit les diff\u00e9rentes \u00e9tapes du d\u00e9veloppement, les interfaces r\u00e9alis\u00e9es, et les r\u00e9sultats des tests de validation.")
B("Enfin, nous terminerons ce m\u00e9moire par une conclusion g\u00e9n\u00e9rale, dans laquelle nous pr\u00e9senterons le bilan du travail r\u00e9alis\u00e9 et les perspectives d\u2019am\u00e9lioration envisag\u00e9es.")

doc.add_page_break()

# Save part 1
out = r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v4.docx'
doc.save(out)
print(f'[OK] Part 1 saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
