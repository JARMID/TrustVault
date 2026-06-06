"""
TrustDesk Mémoire DOCX Generator
Generates Page de Garde + Chapter 1 with Algerian university formatting
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Helper functions
def add_centered(text, size=12, bold=False, color=None, space_after=0, space_before=0, arabic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if arabic:
        run.font.name = 'Traditional Arabic'
    else:
        run.font.name = 'Times New Roman'
    return p

def add_body(text, size=12, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_chapter_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_section_title(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f'{num}  {text}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_subsection_title(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{num}  {text}')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    run = p.add_run(f'• {text}')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def add_numbered(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'{num}. ')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(text)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        # Dark blue header bg
        shading = cell._tc.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(bg)
        run.font.color.rgb = RGBColor(255, 255, 255)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            if ri % 2 == 1:
                shading = cell._tc.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F4F8',
                    qn('w:val'): 'clear'
                })
                shading.append(bg)
    doc.add_paragraph()
    return table

def add_line():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

# ════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════

add_centered('الجمهورية الجزائرية الديمقراطية الشعبية', 14, True, None, 2, 0, True)
add_centered('République Algérienne Démocratique et Populaire', 13, True, (0,51,102), 8)
add_line()
add_centered('وزارة التعليم العالي و البحث العلمي', 13, True, None, 2, 6, True)
add_centered("Ministère de l'Enseignement Supérieur et de la Recherche Scientifique", 12, True, (0,51,102), 12)

add_centered('[Logo de l\'Université]', 11, False, (150,150,150), 8, 12)

add_centered('Université [Nom de votre Université]', 14, True, (0,51,102), 4, 4)
add_centered("Faculté des Sciences Exactes et de l'Informatique", 12, False, None, 2)
add_centered("Département d'Informatique", 12, False, None, 16)

add_line()
add_centered("Mémoire de Fin d'Études", 14, True, None, 4, 12)
add_centered("En vue de l'obtention du diplôme de Master en Informatique", 12, False, None, 2)
add_centered("Spécialité : Génie Logiciel", 12, True, None, 16)
add_line()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(4)

add_centered("Conception et Réalisation d'une Plateforme Web et Mobile", 17, True, (0,51,102), 2, 12)
add_centered("de Gestion des Incidents de Sécurité en Milieu Universitaire", 17, True, (0,51,102), 6)
add_centered("Cas d'étude : TrustDesk — Security Operations Center", 13, True, (100,100,100), 20)

# Author / Supervisor table
table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
table.rows[0].cells[0].text = ''
p = table.rows[0].cells[0].paragraphs[0]
r = p.add_run('Réalisé par :')
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

table.rows[0].cells[1].text = ''
p = table.rows[0].cells[1].paragraphs[0]
r = p.add_run('Encadré par :')
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

table.rows[1].cells[0].text = ''
p = table.rows[1].cells[0].paragraphs[0]
r = p.add_run('[Votre Nom et Prénom]')
r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(150,150,150)

table.rows[1].cells[1].text = ''
p = table.rows[1].cells[1].paragraphs[0]
r = p.add_run('[Dr. / Pr. Nom et Prénom]')
r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(150,150,150)

# Remove borders from author table
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge in ['top','left','bottom','right']:
            b = borders.makeelement(qn(f'w:{edge}'), {qn('w:val'): 'none', qn('w:sz'): '0'})
            borders.append(b)
        tcPr.append(borders)

doc.add_paragraph()

# Jury table
add_centered('Membres du Jury :', 12, True, None, 8, 12)
jury_table = doc.add_table(rows=4, cols=3)
jury_table.style = 'Table Grid'
jury_table.alignment = WD_TABLE_ALIGNMENT.CENTER

jury_headers = ['Nom et Prénom', 'Grade', 'Qualité']
for i, h in enumerate(jury_headers):
    cell = jury_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(255,255,255)
    shading = cell._tc.get_or_add_tcPr()
    bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): '003366', qn('w:val'): 'clear'})
    shading.append(bg)

jury_data = [
    ['[Nom Président]', '[Grade]', 'Président'],
    ['[Nom Encadreur]', '[Grade]', 'Encadreur'],
    ['[Nom Examinateur]', '[Grade]', 'Examinateur'],
]
for ri, row in enumerate(jury_data):
    for ci, val in enumerate(row):
        cell = jury_table.rows[ri+1].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.font.size = Pt(11); r.font.name = 'Times New Roman'
        if '[' in val:
            r.font.color.rgb = RGBColor(150,150,150)

doc.add_paragraph()
add_centered('Année Universitaire : 2025 / 2026', 13, True, (0,51,102), 0, 20)

# Page break
doc.add_page_break()

# ════════════════════════════════════════════════════════
# CHAPTER 1
# ════════════════════════════════════════════════════════

add_chapter_title('CHAPITRE 1')
add_chapter_title('Cadre Général du Projet')
add_line()

# 1.1 Introduction
add_section_title('1.1', 'Introduction')

add_body("La sécurité des campus universitaires constitue un enjeu majeur dans le contexte algérien actuel. Face à la diversité des incidents pouvant survenir — urgences médicales, problèmes comportementaux, atteintes aux biens, ou violations de l'intégrité académique — les établissements d'enseignement supérieur se trouvent souvent démunis, s'appuyant sur des processus manuels lents, des canaux de communication fragmentés, et une absence quasi-totale de traçabilité numérique.")

add_body("Le présent mémoire s'inscrit dans le cadre de la conception et du développement de TrustDesk, une plateforme intégrée de type Security Operations Center (SOC), dédiée à la gestion centralisée des incidents de sécurité en milieu universitaire. Cette solution se décline en trois composantes complémentaires : une application web de supervision, une application mobile de terrain, et une API RESTful assurant la communication entre les différents acteurs.")

add_body("Ce premier chapitre pose le cadre général du projet : l'organisme d'accueil, la problématique identifiée, les objectifs visés, ainsi que la méthodologie adoptée pour mener à bien cette réalisation.")

# 1.2 Organisme d'accueil
add_section_title('1.2', "Présentation de l'Organisme d'Accueil")

add_body("L'université [Nom de votre Université], fondée en [année], est un établissement public d'enseignement supérieur situé dans la wilaya de [wilaya]. Elle accueille plus de [nombre] étudiants répartis sur [nombre] facultés et [nombre] campus.", italic=True)

add_body("Le département d'Informatique, rattaché à la Faculté des Sciences Exactes et de l'Informatique, forme des ingénieurs et des chercheurs dans les domaines du génie logiciel, des systèmes d'information, de l'intelligence artificielle et des réseaux. C'est dans ce cadre académique que s'inscrit notre projet de fin d'études, visant à apporter une solution technologique innovante aux problèmes de sécurité du campus.")

# 1.3 Problématique
add_section_title('1.3', 'Problématique')

add_body("Malgré les avancées technologiques dans le domaine de la sécurité informationnelle, la gestion des incidents de sécurité dans les campus universitaires algériens souffre de nombreuses lacunes que nous avons identifiées à travers notre étude préliminaire :")

add_numbered(1, "Absence de centralisation : Les signalements d'incidents transitent par des canaux hétérogènes (appels téléphoniques, déplacements physiques, e-mails), sans système unifié de réception et de suivi.")

add_numbered(2, "Temps de réponse élevé : L'absence d'outils d'alerte instantanée (type panic button) retarde significativement l'intervention des services compétents lors de situations d'urgence.")

add_numbered(3, "Manque de traçabilité : Les incidents ne sont généralement pas documentés numériquement, rendant impossible toute analyse statistique, identification de tendances récurrentes, ou amélioration des protocoles de sécurité.")

add_numbered(4, "Cloisonnement des acteurs : Les étudiants, le personnel de sécurité, les coordinateurs de conduite et les administrateurs opèrent en silos, sans plateforme collaborative permettant le partage d'informations en temps réel.")

add_numbered(5, "Absence de mobilité : Les systèmes existants, lorsqu'ils existent, ne sont pas accessibles depuis un terminal mobile, limitant la réactivité des agents de terrain.")

add_body("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
run = p.add_run("Problématique centrale : Comment concevoir et réaliser une plateforme web et mobile intégrée, sécurisée et performante, permettant la gestion complète du cycle de vie des incidents de sécurité dans un environnement universitaire, depuis le signalement jusqu'à la résolution ?")
run.font.bold = True
run.font.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# 1.4 Objectifs
add_section_title('1.4', 'Objectifs du Projet')

add_body("Le projet TrustDesk vise à atteindre un ensemble d'objectifs fonctionnels et techniques clairement définis.")

add_subsection_title('1.4.1', 'Objectifs Fonctionnels')

add_bullet("Gestion centralisée des incidents : Permettre le signalement, le tri (triage), l'assignation, le suivi et la clôture des incidents à travers une interface web unifiée.")
add_bullet("Mécanisme de panique instantané : Offrir aux utilisateurs un bouton d'urgence déclenchable depuis l'application mobile, transmettant instantanément la géolocalisation et les informations de l'utilisateur aux services de sécurité.")
add_bullet("Système communautaire de signalement : Mettre en place un radar communautaire permettant aux membres du campus de signaler des activités suspectes (phishing, intrusion, vol) avec un système de niveaux de confiance.")
add_bullet("Tableau de bord analytique : Fournir aux administrateurs un tableau de bord interactif avec des indicateurs clés de performance (KPIs) : nombre d'incidents ouverts, temps moyen de réponse, taux de résolution, répartition par catégorie.")
add_bullet("Gestion des profils et rôles : Implémenter un système d'authentification sécurisé avec gestion des rôles (admin, étudiant, faculté, personnel, dispatch sécurité, coordinateur Title IX).")
add_bullet("Notifications en temps réel : Alerter les acteurs concernés via des notifications push et un panneau de notifications intégré.")

add_subsection_title('1.4.2', 'Objectifs Techniques')

add_bullet("Développer une API RESTful robuste avec le framework Laravel 11, assurant une séparation nette entre le backend et les interfaces utilisateur.")
add_bullet("Concevoir une application web responsive avec React.js 18, TypeScript et un système de design glassmorphique premium.")
add_bullet("Développer une application mobile cross-platform avec Flutter 3, offrant une expérience native sur Android et iOS.")
add_bullet("Garantir la sécurité des données via l'authentification par token (Laravel Sanctum), le chiffrement des communications, et la gestion des permissions par rôle.")
add_bullet("Assurer la scalabilité de l'architecture pour une adoption future par d'autres établissements universitaires.")

# 1.5 Méthodologie
add_section_title('1.5', 'Méthodologie de Travail')

add_body("Pour la réalisation de ce projet, nous avons adopté une approche de développement agile itérative, combinant les principes de la méthode Scrum avec une architecture logicielle de type client-serveur à trois tiers.")

add_subsection_title('1.5.1', 'Approche de Développement')

add_body("Le développement a été structuré en sprints itératifs, chacun se concentrant sur un module fonctionnel spécifique :")

add_table(
    ['Sprint', 'Module', 'Durée'],
    [
        ['Sprint 1', 'Architecture API & Modèle de données', '2 semaines'],
        ['Sprint 2', 'Authentification & Gestion des utilisateurs', '1 semaine'],
        ['Sprint 3', 'Module Incidents & Triage', '2 semaines'],
        ['Sprint 4', 'Module Panic & Géolocalisation', '1 semaine'],
        ['Sprint 5', 'Module Communautaire & Signaux', '2 semaines'],
        ['Sprint 6', 'Application Web — Dashboard & Pages', '3 semaines'],
        ['Sprint 7', 'Application Mobile — Toutes les fonctionnalités', '3 semaines'],
        ['Sprint 8', 'Tests, corrections & déploiement', '2 semaines'],
    ]
)

add_subsection_title('1.5.2', 'Architecture Globale du Système')

add_body("Le système TrustDesk repose sur une architecture à trois composantes distinctes mais interconnectées, communicant via des API RESTful sécurisées :")

add_body("• Couche Présentation (Frontend) : L'application web React.js pour l'administration et la supervision, et l'application mobile Flutter pour les agents de terrain et les étudiants.", space_after=3)
add_body("• Couche Métier (Backend) : L'API Laravel 11 centralisant toute la logique métier, la validation des données, l'authentification et l'autorisation.", space_after=3)
add_body("• Couche Données : La base de données MySQL stockant l'ensemble des entités du système (utilisateurs, incidents, événements de panique, signaux communautaires, journaux d'audit).")

add_centered('[Figure 1.1 — Architecture globale du système TrustDesk]', 11, False, (150,150,150), 6, 12)

add_subsection_title('1.5.3', 'Outils et Technologies Utilisés')

add_body("Le tableau suivant résume l'ensemble des technologies et outils utilisés dans le cadre de ce projet :")

add_table(
    ['Composante', 'Technologie', 'Version', 'Rôle'],
    [
        ['Backend', 'Laravel', '11.x', 'Framework PHP pour API REST'],
        ['Auth', 'Laravel Sanctum', '4.x', 'Authentification par tokens'],
        ['BDD', 'MySQL', '8.x', 'Stockage relationnel'],
        ['Web Frontend', 'React.js', '18.x', 'Interface web SPA'],
        ['Langage Web', 'TypeScript', '5.x', 'Typage statique'],
        ['Bundler', 'Vite', '5.x', 'Build & dev server'],
        ['Routing Web', 'React Router', '6.x', 'Navigation SPA'],
        ['Mobile', 'Flutter', '3.x', 'Application cross-platform'],
        ['Langage Mobile', 'Dart', '3.x', 'Langage Flutter'],
        ['État Mobile', 'Riverpod', '2.x', 'Gestion d\'état réactif'],
        ['Nav Mobile', 'GoRouter', '14.x', 'Navigation déclarative'],
        ['HTTP', 'Dio', '5.x', 'Client HTTP avancé'],
        ['Géolocalisation', 'Geolocator', '13.x', 'Services de localisation'],
        ['IDE', 'VS Code', 'Latest', 'Environnement de développement'],
        ['VCS', 'Git', 'Latest', 'Contrôle de version'],
    ]
)

add_centered('[Tableau 1.2 — Technologies et outils utilisés]', 11, False, (150,150,150), 6)

# 1.6 Structure du mémoire
add_section_title('1.6', 'Structure du Mémoire')

add_body("Le présent mémoire est organisé en quatre chapitres :")

add_bullet("Chapitre 1 — Cadre Général du Projet : (Ce chapitre) Présentation du contexte, de la problématique, des objectifs et de la méthodologie adoptée.")
add_bullet("Chapitre 2 — État de l'Art : Étude bibliographique et analyse comparative des solutions existantes de gestion d'incidents. Revue des technologies web et mobiles modernes, et justification des choix technologiques.")
add_bullet("Chapitre 3 — Analyse et Conception : Modélisation UML du système (diagrammes de cas d'utilisation, de séquence, de classes), conception de la base de données, et spécification de l'architecture logicielle.")
add_bullet("Chapitre 4 — Réalisation et Tests : Présentation de l'environnement de développement, captures d'écran des interfaces réalisées, description des fonctionnalités implémentées, et résultats des tests de validation.")

# 1.7 Conclusion
add_section_title('1.7', 'Conclusion')

add_body("Ce premier chapitre a permis de situer le contexte du projet TrustDesk, d'identifier clairement la problématique liée à la gestion des incidents de sécurité en milieu universitaire, et de définir les objectifs fonctionnels et techniques que cette plateforme vise à atteindre.")

add_body("La solution proposée se distingue par son approche multi-plateforme intégrée (web + mobile + API), son système de réponse d'urgence instantanée, et son radar communautaire collaboratif — des fonctionnalités rarement réunies dans les outils de gestion d'incidents existants dans le contexte universitaire algérien.")

add_body("Le chapitre suivant sera consacré à l'état de l'art, où nous examinerons les solutions existantes, les concepts théoriques sous-jacents, et les fondements technologiques sur lesquels repose notre réalisation.")

# References
doc.add_page_break()
add_chapter_title('Références Bibliographiques')
add_line()

refs = [
    '[1] Laravel Documentation, "Laravel 11.x — The PHP Framework for Web Artisans," 2025. [En ligne]. Disponible : https://laravel.com/docs/11.x',
    '[2] React Documentation, "React — A JavaScript library for building user interfaces," Meta, 2024. [En ligne]. Disponible : https://react.dev',
    '[3] Flutter Documentation, "Flutter — Build apps for any screen," Google, 2024. [En ligne]. Disponible : https://flutter.dev/docs',
    '[4] K. Schwaber et J. Sutherland, "The Scrum Guide," Scrum.org, 2020. [En ligne]. Disponible : https://scrumguides.org',
    '[5] OWASP Foundation, "OWASP Application Security Verification Standard (ASVS) 4.0," 2021.',
    '[6] Ministère de l\'Enseignement Supérieur et de la Recherche Scientifique, "Guide méthodologique de rédaction des mémoires," Algérie, 2023.',
]

for ref in refs:
    add_body(ref, size=11, space_after=4)

# Save
output_path = r'D:\TrustDesk\Memoire_TrustDesk_Chapitre1.docx'
doc.save(output_path)
print(f'[OK] Memoire saved to: {output_path}')
print(f'     File size: {os.path.getsize(output_path) / 1024:.1f} KB')
