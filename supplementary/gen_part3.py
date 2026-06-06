# Part 3: Add Chapter 2 and Chapter 3
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_v3.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t):
    p=doc.add_heading(t,level=1)
    for r in p.runs: r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'

doc.add_page_break()

# ═══ CHAPITRE II TITLE PAGE ═══
for _ in range(8): E()
C('CHAPITRE II:',14,True,4)
C('Analyse et Specification des Besoins',14,False,8)
for _ in range(8): E()
doc.add_page_break()

# ═══ CHAPITRE 2 CONTENT ═══
H1("Introduction :")
B("Apres avoir presente le contexte general du projet et identifie la problematique dans le chapitre precedent, nous consacrons ce deuxieme chapitre a l\u2019analyse et la specification des besoins. Cette etape est fondamentale dans le cycle de developpement logiciel car elle permet de formaliser les attentes des differents acteurs du systeme et de definir avec precision les fonctionnalites a implementer.")
B("Nous commencerons par identifier les acteurs du systeme, puis nous detaillerons les besoins fonctionnels et non fonctionnels. Enfin, nous presenterons la modelisation UML des cas d\u2019utilisation.")

H2("Identification des acteurs :")
B("Un acteur est une entite externe qui interagit avec le systeme. Dans le cadre de TrustDesk, nous avons identifie trois categories d\u2019acteurs :")
T(['Acteur','Description','Interface'],[
    ['Administrateur','Superviseur global du systeme. Gere les utilisateurs, les incidents, les statistiques et la configuration.','Application Web'],
    ['Agent de securite','Personnel de terrain. Recoit les alertes, intervient sur les incidents, utilise le bouton panique.','Application Mobile'],
    ['Utilisateur (Etudiant/Personnel)','Usager du campus. Signale les incidents et les activites suspectes via le radar communautaire.','Application Mobile'],
])
C("Tableau N\u00b06 : Identification des acteurs",12,True,8)

H2("Besoins fonctionnels :")
B("Les besoins fonctionnels decrivent les fonctionnalites que le systeme doit offrir a ses utilisateurs. Nous les avons organises par module :")

H3("Module Authentification :")
LP("Inscription avec validation de l\u2019adresse e-mail universitaire.")
LP("Connexion securisee avec token d\u2019authentification (Laravel Sanctum).")
LP("Gestion des sessions et deconnexion.")
LP("Reinitialisation du mot de passe par e-mail.")
LP("Attribution automatique des roles selon le profil.")

H3("Module Gestion des Incidents :")
LP("Creation d\u2019un incident avec titre, description, categorie, priorite et localisation.")
LP("Attribution automatique ou manuelle d\u2019un incident a un agent.")
LP("Suivi du cycle de vie : Nouveau > En cours > Resolu > Cloture.")
LP("Ajout de commentaires et de pieces jointes a un incident.")
LP("Filtrage et recherche avancee des incidents.")
LP("Historique complet des modifications d\u2019un incident.")

H3("Module Bouton Panique :")
LP("Declenchement d\u2019une alerte d\u2019urgence en un seul appui.")
LP("Capture automatique des coordonnees GPS.")
LP("Notification instantanee aux agents disponibles les plus proches.")
LP("Enregistrement horodate de l\u2019evenement dans la base de donnees.")

H3("Module Radar Communautaire :")
LP("Signalement d\u2019activites suspectes avec geolocalisation.")
LP("Categorisation des signalements (vol, vandalisme, comportement suspect, autre).")
LP("Visualisation des signalements sur une carte interactive.")
LP("Systeme de vote pour confirmer ou infirmer un signalement.")

H3("Module Dashboard (Tableau de bord) :")
LP("Affichage des KPIs : nombre d\u2019incidents, temps moyen de resolution, taux de resolution.")
LP("Graphiques dynamiques (incidents par categorie, par priorite, par periode).")
LP("Liste des incidents recents avec filtrage par statut.")
LP("Carte thermique des zones a risque.")
LP("Notifications en temps reel pour les nouvelles alertes.")

H3("Module Administration :")
LP("Gestion CRUD des utilisateurs (creation, modification, suppression).")
LP("Attribution et modification des roles et permissions.")
LP("Configuration des categories d\u2019incidents et des niveaux de priorite.")
LP("Export des donnees en format CSV/PDF.")
LP("Journalisation des actions administratives (audit log).")

H2("Besoins non fonctionnels :")
B("Les besoins non fonctionnels definissent les contraintes de qualite du systeme :")
T(['Categorie','Exigence','Critere'],[
    ['Performance','Temps de reponse API','< 500ms pour 95% des requetes'],
    ['Performance','Temps de chargement page','< 3 secondes'],
    ['Securite','Authentification','Tokens Sanctum avec expiration'],
    ['Securite','Communications','HTTPS obligatoire (TLS 1.2+)'],
    ['Securite','Donnees sensibles','Chiffrement des mots de passe (bcrypt)'],
    ['Disponibilite','Uptime cible','99.5% minimum'],
    ['Ergonomie','Design responsive','Mobile-first, compatible multi-ecrans'],
    ['Ergonomie','Accessibilite','Conformite WCAG 2.1 niveau AA'],
    ['Maintenabilite','Architecture','Separation API / Frontend / Mobile'],
    ['Compatibilite','Navigateurs','Chrome, Firefox, Edge, Safari'],
    ['Compatibilite','Mobile','Android 8.0+ / iOS 12+'],
])
C("Tableau N\u00b07 : Besoins non fonctionnels",12,True,8)

H2("Diagramme de cas d\u2019utilisation global :")
B("Le diagramme de cas d\u2019utilisation global represente les interactions entre les acteurs et les principales fonctionnalites du systeme TrustDesk.")
B("[Inserer ici le diagramme de cas d\u2019utilisation global - a generer avec StarUML, Lucidchart ou draw.io]",it=True)
C("Figure N\u00b01 : Diagramme de cas d\u2019utilisation global",12,True,8)

H2("Description textuelle des cas d\u2019utilisation :")

H3("CU01 : S\u2019authentifier")
T(['Element','Description'],[
    ['Titre',"S\u2019authentifier"],
    ['Acteur principal','Administrateur / Agent / Utilisateur'],
    ['Precondition','L\u2019utilisateur possede un compte valide'],
    ['Scenario principal','1. L\u2019utilisateur saisit son email et mot de passe\n2. Le systeme verifie les identifiants\n3. Le systeme genere un token Sanctum\n4. L\u2019utilisateur est redirige vers son espace'],
    ['Scenario alternatif','Identifiants incorrects : message d\u2019erreur affiche'],
    ['Postcondition','L\u2019utilisateur est authentifie avec son role'],
])
C("Tableau N\u00b08 : CU01 - S\u2019authentifier",12,True,8)

H3("CU02 : Signaler un incident")
T(['Element','Description'],[
    ['Titre','Signaler un incident'],
    ['Acteur principal','Agent / Utilisateur'],
    ['Precondition','L\u2019utilisateur est authentifie'],
    ['Scenario principal','1. L\u2019utilisateur clique sur \"Signaler un incident\"\n2. Il remplit le formulaire (titre, description, categorie, priorite)\n3. La geolocalisation est capturee automatiquement\n4. Le systeme enregistre l\u2019incident\n5. Les agents concernes sont notifies'],
    ['Postcondition','L\u2019incident est cree avec le statut \"Nouveau\"'],
])
C("Tableau N\u00b09 : CU02 - Signaler un incident",12,True,8)

H3("CU03 : Declencher le bouton panique")
T(['Element','Description'],[
    ['Titre','Declencher le bouton panique'],
    ['Acteur principal','Agent / Utilisateur'],
    ['Precondition','L\u2019utilisateur est authentifie et la localisation est activee'],
    ['Scenario principal','1. L\u2019utilisateur appuie sur le bouton panique\n2. Le systeme capture les coordonnees GPS\n3. Une alerte d\u2019urgence est creee automatiquement\n4. Tous les agents disponibles sont notifies immediatement\n5. La localisation est partagee en temps reel'],
    ['Postcondition','Alerte d\u2019urgence creee et agents notifies'],
])
C("Tableau N\u00b010 : CU03 - Declencher le bouton panique",12,True,8)

H3("CU04 : Gerer les incidents (Administrateur)")
T(['Element','Description'],[
    ['Titre','Gerer les incidents'],
    ['Acteur principal','Administrateur'],
    ['Precondition','L\u2019administrateur est authentifie'],
    ['Scenario principal','1. L\u2019administrateur consulte la liste des incidents\n2. Il filtre par statut, priorite ou categorie\n3. Il assigne un incident a un agent\n4. Il modifie le statut ou la priorite\n5. Il cloture l\u2019incident apres resolution'],
    ['Postcondition','L\u2019incident est mis a jour dans le systeme'],
])
C("Tableau N\u00b011 : CU04 - Gerer les incidents",12,True,8)

H2("Diagramme de sequence :")
B("Les diagrammes de sequence illustrent les interactions entre les acteurs et le systeme pour les principaux cas d\u2019utilisation.")

H3("Diagramme de sequence : Authentification")
B("[Inserer ici le diagramme de sequence de l\u2019authentification]",it=True)
C("Figure N\u00b02 : Diagramme de sequence - Authentification",12,True,8)

H3("Diagramme de sequence : Signalement d\u2019un incident")
B("[Inserer ici le diagramme de sequence du signalement d\u2019incident]",it=True)
C("Figure N\u00b03 : Diagramme de sequence - Signalement d\u2019un incident",12,True,8)

H3("Diagramme de sequence : Bouton Panique")
B("[Inserer ici le diagramme de sequence du bouton panique]",it=True)
C("Figure N\u00b04 : Diagramme de sequence - Bouton Panique",12,True,8)

H2("Conclusion :")
B("Ce chapitre nous a permis d\u2019identifier et de formaliser l\u2019ensemble des besoins fonctionnels et non fonctionnels du systeme TrustDesk. L\u2019analyse des cas d\u2019utilisation et les diagrammes de sequence nous fournissent une base solide pour la phase de conception qui sera detaillee dans le chapitre suivant.")
B("Le chapitre suivant sera consacre a la conception de l\u2019application, ou nous presenterons l\u2019architecture globale du systeme, le modele de donnees, et les diagrammes de classes.")

doc.add_page_break()

# ═══ CHAPITRE III TITLE PAGE ═══
for _ in range(8): E()
C('CHAPITRE III:',14,True,4)
C('Conception de l\u2019Application',14,False,8)
for _ in range(8): E()
doc.add_page_break()

# ═══ CHAPITRE 3 CONTENT ═══
H1("Introduction :")
B("Ce troisieme chapitre est consacre a la conception de la plateforme TrustDesk. Apres avoir identifie et formalise les besoins dans le chapitre precedent, nous allons maintenant definir l\u2019architecture globale du systeme, concevoir le modele de donnees, et elaborer les diagrammes UML necessaires a la realisation du projet.")

H2("Architecture globale du systeme :")
B("TrustDesk repose sur une architecture client-serveur a trois tiers, separant clairement les couches de presentation, de logique metier, et de donnees.")

H3("Architecture a trois tiers :")
B("Couche Presentation (Frontend) :",b=True)
LP("Application Web : React.js 18 avec TypeScript, interfacee via Vite.")
LP("Application Mobile : Flutter 3 avec Dart, utilisant Riverpod pour la gestion d\u2019etat.")
B("Couche Logique Metier (Backend) :",b=True)
LP("API RESTful : Laravel 11 (PHP 8.2) avec authentification Sanctum.")
LP("Controllers, Services, Repositories suivant le pattern MVC.")
LP("Middleware d\u2019authentification et d\u2019autorisation basee sur les roles.")
B("Couche Donnees (Base de donnees) :",b=True)
LP("MySQL 8.x : Stockage relationnel des donnees.")
LP("Migrations Laravel pour le versioning du schema.")
LP("Eloquent ORM pour les operations CRUD.")

B("[Inserer ici le schema de l\u2019architecture globale a trois tiers]",it=True)
C("Figure N\u00b05 : Architecture globale du systeme TrustDesk",12,True,8)

H2("Architecture de l\u2019API REST :")
B("L\u2019API RESTful constitue le coeur du systeme TrustDesk. Elle expose des endpoints structures selon les conventions REST standard :")
T(['Methode','Endpoint','Description','Role requis'],[
    ['POST','/api/auth/login','Authentification','Public'],
    ['POST','/api/auth/register','Inscription','Public'],
    ['GET','/api/incidents','Liste des incidents','Agent/Admin'],
    ['POST','/api/incidents','Creer un incident','Agent/User'],
    ['PUT','/api/incidents/{id}','Modifier un incident','Agent/Admin'],
    ['DELETE','/api/incidents/{id}','Supprimer un incident','Admin'],
    ['POST','/api/panic','Declencher alerte panique','Agent/User'],
    ['GET','/api/signals','Liste des signaux radar','Agent/Admin'],
    ['POST','/api/signals','Creer un signal','Agent/User'],
    ['GET','/api/dashboard/stats','Statistiques globales','Admin'],
    ['GET','/api/users','Liste des utilisateurs','Admin'],
])
C("Tableau N\u00b012 : Endpoints principaux de l\u2019API",12,True,8)

H2("Diagramme de classes :")
B("Le diagramme de classes represente la structure statique du systeme en identifiant les principales classes, leurs attributs, methodes et relations.")
B("[Inserer ici le diagramme de classes UML]",it=True)
C("Figure N\u00b06 : Diagramme de classes du systeme TrustDesk",12,True,8)

H3("Description des classes principales :")
B("Classe User :",b=True)
T(['Attribut','Type','Description'],[
    ['id','int','Identifiant unique'],['name','string','Nom complet'],
    ['email','string','Adresse e-mail'],['password','string','Mot de passe (hache)'],
    ['role','enum','admin / agent / user'],['phone','string','Numero de telephone'],
    ['is_active','boolean','Statut du compte'],
])
C("Tableau N\u00b013 : Classe User",12,True,8)

B("Classe Incident :",b=True)
T(['Attribut','Type','Description'],[
    ['id','int','Identifiant unique'],['title','string','Titre de l\u2019incident'],
    ['description','text','Description detaillee'],['category','enum','Type d\u2019incident'],
    ['priority','enum','low / medium / high / critical'],['status','enum','new / in_progress / resolved / closed'],
    ['latitude','decimal','Coordonnee GPS'],['longitude','decimal','Coordonnee GPS'],
    ['reporter_id','int','FK vers User'],['assignee_id','int','FK vers User (agent)'],
])
C("Tableau N\u00b014 : Classe Incident",12,True,8)

B("Classe PanicAlert :",b=True)
T(['Attribut','Type','Description'],[
    ['id','int','Identifiant unique'],['user_id','int','FK vers User'],
    ['latitude','decimal','Position GPS'],['longitude','decimal','Position GPS'],
    ['status','enum','active / resolved'],['resolved_by','int','FK vers User (agent)'],
    ['resolved_at','timestamp','Date de resolution'],
])
C("Tableau N\u00b015 : Classe PanicAlert",12,True,8)

H2("Modele relationnel de la base de donnees :")
B("Le schema relationnel de la base de donnees MySQL est derive du diagramme de classes. Les principales tables sont :")
LP("users : Stocke les informations des utilisateurs avec leurs roles.")
LP("incidents : Contient les incidents signales avec leur cycle de vie.")
LP("panic_alerts : Enregistre les alertes d\u2019urgence declenchees.")
LP("community_signals : Stocke les signalements du radar communautaire.")
LP("incident_comments : Commentaires associes aux incidents.")
LP("notifications : Systeme de notifications pour les alertes.")
LP("audit_logs : Journal des actions administratives.")

B("[Inserer ici le diagramme du modele relationnel (MCD/MLD)]",it=True)
C("Figure N\u00b07 : Modele relationnel de la base de donnees",12,True,8)

H2("Diagramme d\u2019activite :")
H3("Cycle de vie d\u2019un incident :")
B("Le diagramme d\u2019activite ci-dessous illustre le processus complet de gestion d\u2019un incident, depuis le signalement initial jusqu\u2019a la cloture finale.")
B("[Inserer ici le diagramme d\u2019activite - Cycle de vie d\u2019un incident]",it=True)
C("Figure N\u00b08 : Diagramme d\u2019activite - Cycle de vie d\u2019un incident",12,True,8)

H2("Maquettes des interfaces :")
B("Les maquettes suivantes representent les principales interfaces du systeme :")

H3("Interface Web - Dashboard Administrateur :")
B("[Inserer la maquette du dashboard]",it=True)
C("Figure N\u00b09 : Maquette du dashboard administrateur",12,True,8)

H3("Interface Mobile - Ecran principal :")
B("[Inserer la maquette de l\u2019ecran principal mobile]",it=True)
C("Figure N\u00b010 : Maquette de l\u2019ecran principal mobile",12,True,8)

H3("Interface Mobile - Bouton Panique :")
B("[Inserer la maquette du bouton panique]",it=True)
C("Figure N\u00b011 : Maquette du bouton panique",12,True,8)

H2("Conclusion :")
B("Ce chapitre nous a permis de definir l\u2019architecture globale du systeme TrustDesk, de concevoir le modele de donnees et d\u2019elaborer les principaux diagrammes UML. La conception detaillee des classes, des endpoints API et du schema relationnel constitue une base technique solide pour la phase de realisation.")
B("Le chapitre suivant sera consacre a la realisation et l\u2019implementation du systeme, ou nous presenterons les interfaces developpees, les extraits de code significatifs et les resultats des tests de validation.")

# SAVE
out = r'D:\TrustDesk\Memoire_TrustDesk_v4.docx'
doc.save(out)
print(f'[OK] Chapters 2+3 added. Total: {len(doc.paragraphs)} paragraphs ({os.path.getsize(out)/1024:.1f} KB)')
