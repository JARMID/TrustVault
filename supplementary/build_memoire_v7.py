# -*- coding: utf-8 -*-
"""
Build the TrustDesk memoire v7 based on EXACT feedback:
- No empty first page
- SINGLE cover page (no duplicate)
- All Times New Roman, Black text (no blue theme)
- Both student names and supervisor (Encadreur) on the cover
- Only Chapters 1, 2, and start of 3
- Compact consistent chapter banners (already generated)
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = Path(r"D:\TrustDesk\memoire_assets")
DIAGRAMS_DIR = ASSETS / "diagrams"
OUT = Path(r"D:\TrustDesk\Memoire_TrustDesk_FINAL_v7.docx")

def set_margins(sec):
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.2)
    sec.top_margin    = Cm(2.4)
    sec.bottom_margin = Cm(1.9)

def add_page_number(sec):
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_char_end)
    for r in [run, run2, run3]:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

def new_section(doc, with_page_num=True):
    doc.add_section()
    sec = doc.sections[-1]
    set_margins(sec)
    if with_page_num:
        add_page_number(sec)
    return sec

def invis_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcB.append(b)
            tcPr.append(tcB)

def C(doc, text, sz=12, bold=False, sa=0, sb=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def B(doc, text, bold=False, italic=False, sa=6, sb=0):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def LP(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run("\u2022  " + text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

def shade(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def T(doc, headers, rows, caption=""):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        shade(cell, "E0E0E0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = "Times New Roman"
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_image(doc, img_path, width_cm=14, caption=""):
    path = Path(img_path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(path), width=Cm(width_cm))
    if caption:
        C(doc, caption, 11, True, sa=10)

def chapter_cover(doc, img_path):
    sec = new_section(doc, with_page_num=False)
    # The new compact banners are wide, so we make them width=16cm
    add_image(doc, img_path, width_cm=16)

def add_toc_field(doc, heading="Sommaire"):
    doc.add_heading(heading, 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)
    r3 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3._r.append(fld_end)

def init_doc():
    doc = Document()
    # Apply margins to the very first section (no empty page!)
    sec = doc.sections[0]
    set_margins(sec)
    
    # Normal & Body Text styles
    for s_name in ["Normal", "Body Text"]:
        st = doc.styles[s_name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        rPr = st._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    
    return doc

def build():
    doc = init_doc()

    # ── PAGE DE GARDE (Exact Template Match) ─────────────────────────────────
    # Write directly to the first section
    C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 13, True, 4)
    C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 11, False, 4)
    C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE EN AUDIOVISUELS", 10, True, 2)
    C(doc, "Echahid Ahmed Mehdi \u2014 Ouled Fayet \u2014", 12, True, 20)
    
    C(doc, "Memoire De Fin De Formation Pour L'obtention Du Diplome", 14, True, 4)
    C(doc, "De Technicien Superieur en Informatique", 14, True, 8)
    C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 14, True, 30)
    
    C(doc, "Theme :", 16, True, 12)
    C(doc, "Conception Et Realisation D'une Plateforme", 18, True, 4)
    C(doc, "Web et Mobile de Portefeuille Electronique Securise", 18, True, 4)
    C(doc, "pour Institutions Bancaires", 18, True, 4)
    C(doc, "\u00ab TrustDesk \u00bb", 18, True, 30)
    
    C(doc, "Organisme d'accueil : BEYN", 14, True, 30)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(8)
    tbl.columns[1].width = Cm(8)
    invis_borders(tbl)
    
    # Left cell: Students
    c1 = tbl.rows[0].cells[0]
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("Realise Par :\n")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p1.add_run("M. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.bold = True
    
    # Right cell: Supervisor
    c2 = tbl.rows[0].cells[1]
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p2.add_run("Suivi par :\n")
    r3.bold = True
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)
    r4 = p2.add_run("Mme. Himeur")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)
    r4.bold = True
    
    C(doc, "", sa=30)
    C(doc, "Promotion : 2025 / 2026", 12, True, 0)

    # ── DEDICACES ────────────────────────────────────────────────────────────
    sec = new_section(doc, with_page_num=True)
    doc.add_heading("Dedicaces", 1)
    C(doc, "[A completer par les auteurs]", 14, True, 20)

    # ── REMERCIEMENTS ────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir accorde la sante, la volonte et la perseverance necessaires a la realisation de ce projet.")
    B(doc, "Nos remerciements les plus sinceres vont a notre encadreur Mme. Himeur, pour ses precieux conseils, sa disponibilite et son suivi rigoureux tout au long de ce memoire.")
    B(doc, "Enfin, nous adressons nos remerciements a tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'evaluer ce travail.")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Wassim D., Abdelmadjid K.")
    r.font.name = "Times New Roman"
    r.bold = True

    # ── RESUME / ABSTRACT ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Resume", 1)
    B(doc, "Ce memoire presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes. La plateforme se compose d'une API RESTful centralisee (Laravel), d'un tableau de bord d'administration web (React.js), et d'une application mobile (Flutter). TrustDesk integre un parcours de verification d'identite numerique (eKYC), un systeme de cartes virtuelles securisees par chiffrement RSA-2048, et un moteur de triage intelligent.", sb=12)
    B(doc, "Mots-cles : portefeuille electronique, eKYC, chiffrement RSA, Flutter, Laravel.", bold=True)

    C(doc, "", sa=30)
    doc.add_heading("Abstract", 1)
    B(doc, "This thesis presents the design and implementation of TrustDesk, an integrated secure electronic wallet platform for Algerian banking institutions. The platform comprises a centralized RESTful API (Laravel), a web-based administration dashboard (React.js), and a mobile application (Flutter). TrustDesk integrates a digital identity verification process (eKYC), a virtual card system secured with RSA-2048 encryption, and an intelligent triage engine.", sb=12)
    B(doc, "Keywords: electronic wallet, eKYC, RSA encryption, Flutter, Laravel.", bold=True)

    # ── SOMMAIRE ─────────────────────────────────────────────────────────────
    new_section(doc)
    add_toc_field(doc, "Sommaire")

    # ── LISTE DES FIGURES ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Figures", 1)
    figures = [
        ("Figure N\u00b01", "Diagramme de cas d'utilisation global - TrustDesk"),
        ("Figure N\u00b02", "Architecture globale trois tiers - TrustDesk"),
        ("Figure N\u00b03", "Diagramme de classes - TrustDesk"),
        ("Figure N\u00b04", "Modele Conceptuel des Donnees (MCD)"),
        ("Figure N\u00b05", "Diagramme de sequence - Inscription & eKYC"),
        ("Figure N\u00b06", "Diagramme de sequence - Transfert P2P"),
    ]
    for num, title in figures:
        LP(doc, f"{num} : {title}")

    # ── LISTE DES TABLEAUX ───────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Tableaux", 1)
    tableaux = [
        ("Tableau N\u00b01",  "Principales solutions de BEYN"),
        ("Tableau N\u00b02",  "Recapitulatif des acteurs du systeme"),
        ("Tableau N\u00b03",  "Exigences non fonctionnelles"),
        ("Tableau N\u00b04",  "Description UC - Inscription eKYC"),
        ("Tableau N\u00b05",  "Description UC - Transfert P2P"),
    ]
    for num, title in tableaux:
        LP(doc, f"{num} : {title}")

    # ── INTRODUCTION GENERALE ────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Introduction Generale", 1)
    B(doc, "L'evolution rapide des technologies financieres (fintech) transforme profondement les modes d'interaction entre les institutions bancaires et leurs clients en Algerie. Les portefeuilles electroniques (e-wallets) s'imposent comme un levier strategique, permettant aux banques d'offrir a leurs clients des services de paiement dematerialises.", sb=12)
    B(doc, "C'est dans cette perspective que s'inscrit TrustDesk, realise au sein de BEYN. La problematique centrale de ce memoire est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC) et le chiffrement des donnees ? \u00bb", bold=True)
    B(doc, "Notre travail s'articule autour des chapitres suivants :")
    LP(doc, "Chapitre I : Etude Prealable")
    LP(doc, "Chapitre II : Analyse et Specification des Besoins")
    LP(doc, "Chapitre III : Conception")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE I
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch1_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre I : Etude Prealable", 1)

    doc.add_heading("1.1 Presentation de l'organisme d'accueil : BEYN", 2)
    B(doc, "Fondee en 2004, BEYN est une entreprise specialisee dans le developpement de solutions technologiques pour les institutions bancaires. Elle se positionne comme un partenaire strategique pour les banques en Algerie.")
    T(doc, ["Plateforme", "Cible", "Fonctionnalite Principale"], [
        ["WimPay",  "Grand public",      "Paiement QR code, P2P"],
        ["SELA",    "Retail banking",    "Banque digitale omnicanale"],
    ], "Tableau N\u00b01 : Principales solutions de BEYN")

    doc.add_heading("1.2 Presentation du Projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme integree de portefeuille electronique securise composee d'une API RESTful (Laravel), d'un dashboard web (React.js), et d'une application mobile (Flutter).")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE II
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch2_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre II : Analyse et Specification des Besoins", 1)

    doc.add_heading("2.1 Identification des Acteurs", 2)
    B(doc, "La plateforme met en interaction trois acteurs principaux :")
    LP(doc, "Client (Utilisateur Bancaire) : Accede via l'application mobile Flutter.")
    LP(doc, "Administrateur : Supervise via le tableau de bord web React.js.")
    LP(doc, "Systeme : Gere le triage IA et la detection d'anomalies.")

    doc.add_heading("2.2 Diagramme de Cas d'Utilisation Global", 2)
    add_image(doc, DIAGRAMS_DIR / "use_case.png", 14, "Figure N\u00b01 : Diagramme de cas d'utilisation global")

    doc.add_heading("2.3 Specification des Besoins Fonctionnels", 2)
    doc.add_heading("Module Authentification & eKYC", 3)
    LP(doc, "Inscription avec verification OTP et parcours eKYC (piece d'identite, selfie).")

    doc.add_heading("Module Portefeuille & Transactions", 3)
    LP(doc, "Consultation du solde et transferts P2P en temps reel.")

    doc.add_heading("2.4 Descriptions Textuelles des Cas d'Utilisation", 2)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation", "Effectuer un transfert P2P"],
        ["Acteur principal",  "Client"],
        ["Scenario nominal",  "1. Saisie montant\n2. Verification solde\n3. Confirmation biometrique\n4. Execution"],
    ], "Tableau N\u00b05 : Description UC - Transfert P2P")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE III (Start only)
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch3_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre III : Conception", 1)

    doc.add_heading("3.1 Architecture Globale du Systeme", 2)
    B(doc, "La plateforme adopte une architecture trois tiers (Three-Tier Architecture).")
    add_image(doc, DIAGRAMS_DIR / "architecture.png", 15, "Figure N\u00b02 : Architecture globale")

    doc.add_heading("3.2 Diagramme de Classes", 2)
    add_image(doc, DIAGRAMS_DIR / "class_diagram.png", 15, "Figure N\u00b03 : Diagramme de classes")

    doc.add_heading("3.3 Modele Conceptuel des Donnees (MCD)", 2)
    add_image(doc, DIAGRAMS_DIR / "mcd.png", 15, "Figure N\u00b04 : MCD")

    doc.add_heading("3.4 Diagrammes de Sequence", 2)
    add_image(doc, DIAGRAMS_DIR / "sequence_ekyc.png", 14, "Figure N\u00b05 : Inscription & eKYC")
    add_image(doc, DIAGRAMS_DIR / "sequence_p2p.png", 14, "Figure N\u00b06 : Transfert P2P")

    # STOPPING AT THE START OF CHAPTER 3 AS REQUESTED

    doc.save(str(OUT))
    size = OUT.stat().st_size // 1024
    print(f"\nMemoire sauvegarde : {OUT}  ({size} KB)")

if __name__ == "__main__":
    print("Building memoire v7 (Exact feedback implementation)...")
    build()
