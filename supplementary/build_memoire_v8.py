# -*- coding: utf-8 -*-
"""
Build TrustDesk memoire v8.
Based on EXACT analysis of 'dev wm memoir finall _corrige.docx' (template):
  - TWO cover pages (1st with BEYN logo, 2nd without)
  - Both authors + Encadreur on BOTH pages
  - Dedicaces x2, Remerciements, Resume/Abstract, Sommaire, Listes
  - FULL Ch1, FULL Ch2, START of Ch3 (each ~100+ paragraphs like template)
  - All Times New Roman, ALL BLACK text, no blue
  - Template margins: L=2cm R=1.2cm T=2.4cm B=1.9cm
  - Grey table headers (not dark blue)
  - Compact landscape chapter banners
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = Path(r"D:\TrustDesk\memoire_assets")
DIAGRAMS = ASSETS / "diagrams"
OUT = Path(r"D:\TrustDesk\Memoire_TrustDesk_FINAL_v8.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_margins(sec, top=2.4, bot=1.9):
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.2)
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bot)

def shade(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def invis_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for s in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{s}")
                b.set(qn("w:val"), "nil")
                tcB.append(b)
            tcPr.append(tcB)

def add_page_number(sec):
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    for r in [run, r2, r3]:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

def new_section(doc, page_num=True, top=2.4, bot=1.9):
    doc.add_section()
    sec = doc.sections[-1]
    set_margins(sec, top, bot)
    if page_num:
        add_page_number(sec)
    return sec

def C(doc, text, sz=12, bold=False, sa=0, sb=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold; r.font.name = "Times New Roman"
    return p

def B(doc, text, bold=False, italic=False, sa=6, sb=0):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.font.bold = bold; r.font.italic = italic
    return p

def LP(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run("\u2022  " + text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def T(doc, headers, rows, caption=""):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h; shade(c, "D9D9D9")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(val)
            p = c.paragraphs[0]
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = "Times New Roman"
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_image(doc, path, w=14, caption=""):
    p = Path(path)
    if p.exists():
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Cm(w))
    else:
        B(doc, f"[Image: {p.name}]", italic=True)
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_toc(doc):
    doc.add_heading("Sommaire", 1)
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    B(doc, "(Dans Word : Ctrl+A puis F9 pour mettre a jour le sommaire)", italic=True, sa=12)

# ── Init ─────────────────────────────────────────────────────────────────────

def init_doc():
    doc = Document()
    sec = doc.sections[0]
    set_margins(sec, top=2.4, bot=0.5)
    for sn in ["Normal", "Body Text"]:
        st = doc.styles[sn]
        st.font.name = "Times New Roman"; st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rPr = st._element.get_or_add_rPr()
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = OxmlElement("w:rFonts"); rPr.insert(0, rF)
        rF.set(qn("w:ascii"), "Times New Roman")
        rF.set(qn("w:hAnsi"), "Times New Roman")
        rF.set(qn("w:cs"), "Times New Roman")
    for level, sz in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        h = doc.styles[level]
        h.font.name = "Times New Roman"; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        if level == "Heading 1":
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

# ── Cover Page ───────────────────────────────────────────────────────────────

def page_de_garde(doc, with_logo=True):
    if with_logo:
        logo = ASSETS / "logob.png"
        if logo.exists():
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(logo), width=Cm(5))
            p.paragraph_format.space_after = Pt(8)

    C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 13, True, 2)
    C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 13, True, 2)
    C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE EN AUDIOVISUELS", 10, True, 2)
    C(doc, "", sa=4)
    C(doc, "Echahid Ahmed Mehdi \u2014 Ouled Fayet \u2014", 12, True, 12)
    C(doc, "", sa=8)
    C(doc, "Memoire De Fin De Formation Pour L'obtention Du Diplome De Technicien Superieur En Informatique", 14, True, 4)
    C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 14, True, 20)
    C(doc, "", sa=4)
    C(doc, "Theme :", 16, True, 8)
    C(doc, "", sa=4)
    C(doc, "Conception Et Realisation D'une", 18, True, 2)
    C(doc, "Plateforme Web et Mobile de Portefeuille", 18, True, 2)
    C(doc, "Electronique Securise pour Institutions Bancaires", 18, True, 2)
    C(doc, "\u00ab TrustDesk \u00bb", 18, True, 20)
    C(doc, "", sa=4)
    C(doc, "Organisme d'accueil :", 14, True, 4)
    
    if with_logo:
        C(doc, "Entreprise BEYN", 14, True, 20)
    else:
        C(doc, "", sa=20)

    # Authors & Supervisor table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(8); tbl.columns[1].width = Cm(8)
    invis_borders(tbl)
    c1 = tbl.rows[0].cells[0]
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("Realise Par :\n"); r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(12)
    r2 = p1.add_run("M. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(12); r2.bold = True
    c2 = tbl.rows[0].cells[1]
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p2.add_run("Suivi par :\n"); r3.bold = True; r3.font.name = "Times New Roman"; r3.font.size = Pt(12)
    r4 = p2.add_run("Mme. Himeur"); r4.font.name = "Times New Roman"; r4.font.size = Pt(12); r4.bold = True
    C(doc, "", sa=20)
    C(doc, "Promotion : 2025 / 2026", 12, True, 0)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = init_doc()

    # ── PAGE DE GARDE 1 (with BEYN logo) ─────────────────────────────────────
    page_de_garde(doc, with_logo=True)

    # ── PAGE DE GARDE 2 (without logo) ───────────────────────────────────────
    new_section(doc, page_num=False, top=2.4, bot=0.5)
    page_de_garde(doc, with_logo=False)

    # ── DEDICACES 1 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 1]", italic=True, sa=20)

    # ── DEDICACES 2 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 2]", italic=True, sa=20)

    # ── REMERCIEMENTS ────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir accorde la sante, la volonte et la perseverance necessaires pour mener a bien ce projet.")
    B(doc, "Nos remerciements les plus sinceres vont a notre encadreur Mme. Himeur, pour ses precieux conseils, sa disponibilite constante et son suivi rigoureux tout au long de ce memoire. Son expertise et ses orientations nous ont ete d'une aide inestimable.")
    B(doc, "Nous remercions egalement M. Reda Benbouzid, CEO de BEYN, et l'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur collaboration durant notre stage. Leur expertise dans les solutions bancaires numeriques a ete une source d'apprentissage precieuse.")
    B(doc, "Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants tout au long de notre formation.")
    B(doc, "Enfin, nous adressons nos remerciements a tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'evaluer ce travail.")
    B(doc, "A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
    C(doc, "", sa=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Abdelmadjid K., Wassim D.")
    r.font.name = "Times New Roman"; r.bold = True

    # ── SOMMAIRE ─────────────────────────────────────────────────────────────
    new_section(doc)
    add_toc(doc)

    # ── RESUME / ABSTRACT ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Resume", 1)
    B(doc, "Ce memoire presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre d'un stage au sein de l'entreprise BEYN. La plateforme se compose de trois composants principaux : une API RESTful centralisee developpee avec Laravel 11 et PHP 8.2, un tableau de bord d'administration web construit avec React.js 18 et TypeScript, et une application mobile cross-platform concue avec Flutter 3 et le framework de gestion d'etat Riverpod.", sb=12)
    B(doc, "TrustDesk integre plusieurs fonctionnalites avancees : un parcours de verification d'identite numerique (eKYC) incluant la detection de vivacite biometrique, un systeme de cartes virtuelles securisees par chiffrement asymetrique RSA-2048, un moteur de triage intelligent par intelligence artificielle pour la detection d'anomalies transactionnelles, et une piste d'audit complete garantissant la tracabilite de toutes les operations sensibles. L'architecture trois tiers adoptee assure la separation des responsabilites et la scalabilite horizontale de la solution.")
    B(doc, "Mots-cles : portefeuille electronique, fintech, eKYC, chiffrement RSA-2048, Flutter, Laravel, React.js, triage IA, securite bancaire.", bold=True)
    C(doc, "", sa=20)
    doc.add_heading("Abstract", 1)
    B(doc, "This thesis presents the design and implementation of TrustDesk, an integrated secure electronic wallet platform intended for Algerian banking institutions, developed during an internship at BEYN. The platform comprises three main components: a centralized RESTful API developed with Laravel 11 and PHP 8.2, a web-based administration dashboard built with React.js 18 and TypeScript, and a cross-platform mobile application designed with Flutter 3 and the Riverpod state management framework.", sb=12)
    B(doc, "TrustDesk integrates several advanced features: a digital identity verification process (eKYC) including biometric liveness detection, a virtual card system secured with RSA-2048 asymmetric encryption, an AI-powered intelligent triage engine for real-time transaction anomaly detection, and a comprehensive audit trail ensuring full traceability of all sensitive operations. The adopted three-tier architecture ensures separation of concerns and horizontal scalability.")
    B(doc, "Keywords: electronic wallet, fintech, eKYC, RSA-2048 encryption, Flutter, Laravel, React.js, AI triage, banking security.", bold=True)

    # ── LISTE DES FIGURES ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Figures", 1)
    figs = [
        "Organigramme de BEYN",
        "Cycle de vie Scrum",
        "Internet",
        "Exemple de navigateurs web",
        "Exemple d'une application web",
        "Exemple d'une application mobile",
        "Types de diagrammes UML",
        "Diagramme de cas d'utilisation global - TrustDesk",
        "Diagramme de cas d'utilisation - Client",
        "Diagramme de cas d'utilisation - Administrateur",
        "Architecture globale trois tiers",
        "Diagramme de classes",
        "Modele Conceptuel des Donnees (MCD)",
        "Diagramme de sequence - Inscription et eKYC",
        "Diagramme de sequence - Transfert P2P",
    ]
    for i, f in enumerate(figs, 1):
        LP(doc, f"Figure N\u00b0{i} : {f}")

    # ── LISTE DES TABLEAUX ───────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Tableaux", 1)
    tbls = [
        "Fiche technique de BEYN",
        "Principales solutions de BEYN",
        "Acteurs du systeme TrustDesk",
        "Besoins fonctionnels",
        "Besoins non fonctionnels",
        "Description textuelle UC - Inscription eKYC",
        "Description textuelle UC - Transfert P2P",
        "Couches de l'architecture API Laravel",
    ]
    for i, t in enumerate(tbls, 1):
        LP(doc, f"Tableau N\u00b0{i} : {t}")

    # ── INTRODUCTION GENERALE ────────────────────────────────────────────────
    new_section(doc)
    C(doc, "INTRODUCTION", 16, True, 20)
    B(doc, "Depuis toujours, les institutions financieres cherchent a ameliorer leurs methodes de gestion et a optimiser leurs services pour mieux repondre aux besoins de leurs clients. Avec l'essor des technologies de l'information et de la communication, les banques ont entrepris une transformation numerique profonde, passant des guichets traditionnels aux services en ligne, puis aux applications mobiles.", sb=12)
    B(doc, "Par ailleurs, la generalisation de l'Internet et le developpement des technologies mobiles ont radicalement modifie les attentes des clients bancaires. Ceux-ci exigent desormais un acces instantane a leurs comptes, des transferts rapides et securises, et des outils de gestion financiere personnalises, le tout accessible depuis leur smartphone.")
    B(doc, "Dans ce contexte, les portefeuilles electroniques (e-wallets) s'imposent comme un levier strategique. Ils permettent aux banques d'offrir des services de paiement dematerialises, des transferts de pair a pair (P2P), et la gestion de cartes virtuelles securisees.")
    B(doc, "En Algerie, cette mutation s'accelere sous l'impulsion de la Banque d'Algerie et des reformes reglementaires visant a promouvoir l'inclusion financiere. Cependant, les solutions existantes sur le marche algerien presentent plusieurs lacunes : absence d'un parcours eKYC entierement automatise, securite insuffisante pour les cartes virtuelles, et absence de detection de fraude par intelligence artificielle.")
    B(doc, "C'est dans cette perspective que s'inscrit le projet TrustDesk, realise au sein de l'entreprise BEYN, un acteur majeur de la fintech algerienne. La problematique centrale de ce memoire est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)
    B(doc, "Pour repondre a cette problematique, nous avons structure notre memoire en trois chapitres :")
    B(doc, "Le premier chapitre, intitule \u00ab Etude Prealable \u00bb, presente la societe BEYN, le contexte du marche fintech algerien, la problematique identifiee, les concepts theoriques utilises et la methodologie de developpement adoptee.")
    B(doc, "Le deuxieme chapitre, intitule \u00ab Analyse et Specification des Besoins \u00bb, est consacre a l'identification des acteurs, la formalisation des besoins fonctionnels et non fonctionnels, et la modelisation UML du systeme a l'aide des diagrammes de cas d'utilisation.")
    B(doc, "Le troisieme chapitre, intitule \u00ab Conception \u00bb, presente la conception architecturale de la solution : architecture trois tiers, modele de donnees (MCD et MLD), diagrammes de classes et de sequence, et architecture de chiffrement.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE I : ETUDE PREALABLE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch1_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre I : Etude Prealable", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce premier chapitre constitue le socle fondateur de notre travail. Il presente l'organisme d'accueil BEYN, le contexte du marche fintech algerien dans lequel s'inscrit notre projet, la problematique identifiee, les concepts theoriques necessaires a la comprehension du sujet, ainsi que la methodologie de developpement adoptee.", sb=12)

    # ── 1.1 Presentation de BEYN ─────────────────────────────────────────────
    doc.add_heading("1.1  Presentation de l'organisme d'accueil : BEYN", 2)
    B(doc, "Fondee en 2004 et basee a Cheraga (Alger), BEYN est une entreprise specialisee dans la conception et le developpement de solutions technologiques pour les institutions bancaires et financieres algeriennes. Dirigee par M. Reda Benbouzid, elle se positionne comme un partenaire technologique strategique pour les banques partenaires, notamment la Banque de Developpement Local (BDL), la Banque Nationale d'Algerie (BNA) et Al Salam Bank Algeria.")
    B(doc, "L'entreprise compte une equipe pluridisciplinaire de developpeurs, d'ingenieurs systeme, de designers UX/UI et d'experts en securite bancaire. Son siege social est situe a Cheraga, dans la wilaya d'Alger.")

    doc.add_heading("Fiche Technique de BEYN", 3)
    T(doc, ["Designation", "Information"], [
        ["Nom de l'entreprise",  "BEYN"],
        ["Date de creation",     "2004"],
        ["Siege social",         "Cheraga, Alger, Algerie"],
        ["Directeur general",    "M. Reda Benbouzid"],
        ["Secteur d'activite",   "Technologies financieres (Fintech)"],
        ["Domaine",              "Solutions bancaires numeriques"],
        ["Effectif",             "50+ collaborateurs"],
        ["Partenaires bancaires","BDL, BNA, Al Salam Bank Algeria"],
    ], "Tableau N\u00b01 : Fiche technique de BEYN")

    doc.add_heading("Organigramme de BEYN", 3)
    B(doc, "L'organigramme ci-dessous illustre la structure organisationnelle de BEYN :")
    B(doc, "[Figure N\u00b01 : Organigramme de BEYN - A inserer]", italic=True)

    doc.add_heading("Les Solutions de BEYN", 3)
    B(doc, "BEYN propose trois plateformes phares couvrant l'ensemble des besoins bancaires numeriques :")
    LP(doc, "WimPay : Solution leader de paiement mobile par QR code en Algerie, permettant les transferts P2P instantanes, le paiement de factures (telephone, internet, electricite), et le paiement chez les commercants partenaires. WimPay est deployee sur Android et iOS et integree avec plusieurs banques algeriennes.")
    LP(doc, "SELA : Plateforme de banque digitale omnicanale, orientee retail banking. SELA offre aux clients bancaires un acces complet a leurs comptes, virements, et services via une interface web et mobile unifiee.")
    LP(doc, "KANTARA : Solution de corporate banking et de passerelles de paiement securisees pour les entreprises. KANTARA facilite les transactions inter-entreprises et l'integration avec les systemes de paiement nationaux.")

    T(doc, ["Plateforme", "Cible", "Fonctionnalite Principale"], [
        ["WimPay",  "Grand public",      "Paiement QR code, transferts P2P, paiement factures"],
        ["SELA",    "Retail banking",    "Banque digitale omnicanale (web + mobile)"],
        ["KANTARA", "Corporate banking", "Passerelles de paiement, transactions inter-entreprises"],
    ], "Tableau N\u00b02 : Principales solutions de BEYN")

    # ── 1.2 Presentation du Projet ───────────────────────────────────────────
    doc.add_heading("1.2  Presentation du Projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme integree de portefeuille electronique securise, concue pour repondre aux besoins des institutions bancaires algeriennes en matiere de services financiers numeriques. Le projet a ete initie dans le cadre de notre stage chez BEYN, en reponse aux lacunes identifiees dans les solutions existantes.")
    B(doc, "La plateforme se compose de trois composants principaux :")
    LP(doc, "Une API RESTful centralisee : Developpee avec Laravel 11 et PHP 8.2, elle constitue le coeur metier de la plateforme. Elle gere l'authentification Sanctum, la logique transactionnelle, le chiffrement RSA/AES, et la piste d'audit.")
    LP(doc, "Un tableau de bord d'administration web : Construit avec React.js 18, TypeScript et Vite, il permet aux administrateurs bancaires de superviser les clients, valider les dossiers eKYC, gerer les incidents de securite, et monitorer les KPIs en temps reel.")
    LP(doc, "Une application mobile cross-platform : Developpee avec Flutter 3 et Dart, utilisant Riverpod pour la gestion d'etat. Elle offre aux clients l'acces a leur portefeuille, les transferts P2P, la gestion des cartes virtuelles, et les outils d'analytique financiere.")
    B(doc, "Les fonctionnalites differenciantes de TrustDesk par rapport aux solutions existantes comprennent :")
    LP(doc, "Un parcours eKYC complet avec detection de vivacite biometrique.")
    LP(doc, "Un systeme de cartes virtuelles avec chiffrement asymetrique RSA-2048 (Zero-Knowledge cote backend).")
    LP(doc, "Un moteur de triage intelligent par IA pour la detection d'anomalies transactionnelles.")
    LP(doc, "Une piste d'audit complete et immuable pour la tracabilite des operations.")

    # ── 1.3 Problematique ────────────────────────────────────────────────────
    doc.add_heading("1.3  Problematique", 2)
    B(doc, "Malgre les avancees de la digitalisation bancaire en Algerie, plusieurs lacunes persistent dans les solutions de portefeuille electronique actuellement disponibles :")
    LP(doc, "Absence d'un parcours eKYC entierement automatise : La verification d'identite reste souvent manuelle, entrainant des delais importants et une experience utilisateur mediocre.")
    LP(doc, "Securite insuffisante pour les cartes virtuelles : Les solutions existantes ne mettent pas en oeuvre de chiffrement asymetrique pour la protection des numeros de carte et des CVV.")
    LP(doc, "Absence de detection de fraude par IA : Aucune solution locale n'integre de moteur d'intelligence artificielle capable de detecter les anomalies transactionnelles en temps reel.")
    LP(doc, "Manque de tracabilite : L'absence de piste d'audit detaillee rend difficile le suivi des operations sensibles et la conformite reglementaire.")
    B(doc, "La problematique centrale de notre travail est donc :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)

    # ── 1.4 Le public vise ───────────────────────────────────────────────────
    doc.add_heading("1.4  Le public vise", 2)
    B(doc, "La plateforme TrustDesk cible deux categories d'utilisateurs :")
    LP(doc, "Les clients bancaires : Particuliers disposant d'un compte dans l'une des banques partenaires de BEYN, souhaitant acceder a des services de portefeuille electronique via l'application mobile.")
    LP(doc, "Les administrateurs bancaires : Responsables au sein des banques partenaires charges de superviser les comptes clients, valider les verifications d'identite, et gerer les incidents de securite via le tableau de bord web.")

    # ── 1.5 Concepts theoriques ──────────────────────────────────────────────
    doc.add_heading("1.5  Concepts theoriques", 2)

    doc.add_heading("Les definitions generales", 3)
    B(doc, "Internet :", bold=True)
    B(doc, "Internet est un ensemble de reseaux mondiaux interconnectes qui permet a des ordinateurs et a des serveurs de communiquer efficacement au moyen d'un protocole de communication commun (TCP/IP). Internet permet l'echange de donnees sous forme de texte, d'images, de videos et d'applications a travers le monde entier.")
    B(doc, "Application web :", bold=True)
    B(doc, "Une application web est un logiciel qui s'execute dans un navigateur web. Contrairement aux applications de bureau, elle ne necessite pas d'installation sur l'ordinateur de l'utilisateur. Les applications web utilisent une architecture client-serveur : le navigateur (client) envoie des requetes au serveur, qui traite les donnees et renvoie les resultats sous forme de pages HTML, JSON ou XML.")
    B(doc, "Application mobile :", bold=True)
    B(doc, "Une application mobile est un logiciel concu pour fonctionner sur des appareils mobiles tels que les smartphones et les tablettes. Elle est generalement telechargee depuis un magasin d'applications (Google Play Store, Apple App Store) et installee directement sur l'appareil. Les applications mobiles peuvent etre natives (developpees pour un OS specifique), hybrides, ou cross-platform (comme Flutter).")
    B(doc, "API REST :", bold=True)
    B(doc, "Une API (Application Programming Interface) REST (Representational State Transfer) est une interface de programmation qui permet a des applications de communiquer entre elles via le protocole HTTP. Les APIs REST utilisent les methodes HTTP standard (GET, POST, PUT, DELETE) pour effectuer des operations CRUD sur des ressources identifiees par des URLs. Elles sont stateless, c'est-a-dire que chaque requete contient toutes les informations necessaires a son traitement.")
    B(doc, "Portefeuille electronique (e-wallet) :", bold=True)
    B(doc, "Un portefeuille electronique est une application ou un service qui permet de stocker, envoyer et recevoir de l'argent de maniere dematerialisee. Il peut etre lie a un compte bancaire ou fonctionne de maniere autonome. Les e-wallets offrent generalement des fonctionnalites de paiement en ligne, de transfert P2P, de gestion de cartes virtuelles et de suivi des depenses.")
    B(doc, "eKYC (Electronic Know Your Customer) :", bold=True)
    B(doc, "Le eKYC est un processus numerique de verification d'identite qui permet aux institutions financieres de verifier l'identite de leurs clients a distance, sans necessite de presence physique. Ce processus implique generalement la capture d'une piece d'identite officielle, la prise d'un selfie, et une verification de vivacite (liveness detection) pour s'assurer que la personne est bien presente physiquement.")
    B(doc, "Chiffrement asymetrique (RSA) :", bold=True)
    B(doc, "Le chiffrement asymetrique, aussi appele chiffrement a cle publique, utilise une paire de cles mathematiquement liees : une cle publique pour le chiffrement et une cle privee pour le dechiffrement. L'algorithme RSA (Rivest-Shamir-Adleman) est l'un des plus utilises, avec des tailles de cle de 2048 bits ou plus pour garantir la securite. Le chiffrement asymetrique resout le probleme de distribution des cles inherent au chiffrement symetrique.")

    doc.add_heading("Definition UML", 3)
    B(doc, "Le Langage de Modelisation Unifie, de l'anglais Unified Modeling Language (UML), est un langage de modelisation graphique a base de pictogrammes concu pour fournir une methode normalisee pour visualiser la conception d'un systeme. Il est couramment utilise en developpement logiciel et en conception orientee objet.")
    B(doc, "UML propose 14 types de diagrammes divises en deux categories :")
    LP(doc, "Diagrammes structurels (statiques) : diagramme de classes, de composants, de deploiement, d'objets, de paquetages, de structures composites, de profils.")
    LP(doc, "Diagrammes comportementaux (dynamiques) : diagramme de cas d'utilisation, d'activites, d'etats, de sequence, de communication, de temps, de vue d'ensemble des interactions.")
    B(doc, "Dans le cadre de ce memoire, nous utiliserons principalement les diagrammes de cas d'utilisation, de classes, et de sequence.")

    # ── 1.6 Methodologie ─────────────────────────────────────────────────────
    doc.add_heading("1.6  Methodologie de developpement : Agile Scrum", 2)
    B(doc, "Pour la gestion de notre projet, nous avons adopte la methode Agile Scrum, une approche iterative et incrementale de gestion de projet logiciel. Scrum se distingue par sa flexibilite et sa capacite a s'adapter aux changements de besoins en cours de projet.")
    B(doc, "Les principes fondamentaux de Scrum sont :")
    LP(doc, "Sprint : Periode de developpement fixe (generalement 2 a 4 semaines) au terme de laquelle un increment fonctionnel du produit est livre.")
    LP(doc, "Product Backlog : Liste priorisee de toutes les fonctionnalites souhaitees pour le produit.")
    LP(doc, "Sprint Backlog : Sous-ensemble du Product Backlog selectionne pour un sprint specifique.")
    LP(doc, "Daily Standup : Reunion quotidienne de 15 minutes pour synchroniser l'equipe.")
    LP(doc, "Sprint Review : Presentation de l'increment au Product Owner a la fin de chaque sprint.")
    LP(doc, "Sprint Retrospective : Reflexion d'equipe sur les ameliorations possibles du processus.")
    B(doc, "Notre projet a ete decoupe en quatre sprints principaux :")
    LP(doc, "Sprint 1 : Authentification Sanctum et parcours eKYC (2 semaines).")
    LP(doc, "Sprint 2 : Portefeuille, transferts P2P et tableau de bord (2 semaines).")
    LP(doc, "Sprint 3 : Cartes virtuelles et chiffrement RSA (2 semaines).")
    LP(doc, "Sprint 4 : Triage IA, piste d'audit et tests (2 semaines).")
    B(doc, "[Figure N\u00b02 : Cycle de vie Scrum - A inserer]", italic=True)

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce premier chapitre a pose le cadre de notre travail : un projet ambitieux, realise au sein de BEYN, un acteur reconnu de la fintech algerienne. Nous avons presente l'organisme d'accueil, la problematique identifiee, les concepts theoriques fondamentaux necessaires a la comprehension du sujet, ainsi que la methodologie Agile Scrum adoptee pour la gestion du projet. Le chapitre suivant sera consacre a l'analyse detaillee et a la specification des besoins fonctionnels et non fonctionnels de la plateforme TrustDesk.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE II : ANALYSE ET SPECIFICATION DES BESOINS
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch2_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre II : Analyse et Specification des Besoins", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce deuxieme chapitre est consacre a l'analyse detaillee des besoins de la plateforme TrustDesk. Nous identifierons les acteurs du systeme, formaliserons les exigences fonctionnelles et non fonctionnelles, et presenterons les diagrammes UML de cas d'utilisation modelisant le comportement attendu du systeme.", sb=12)

    # ── 2.1 Rappels theoriques ───────────────────────────────────────────────
    doc.add_heading("2.1  Rappels theoriques : Diagramme de cas d'utilisation", 2)
    B(doc, "Le diagramme de cas d'utilisation (Use Case Diagram) est un diagramme UML fonctionnel utilise pour representer les interactions entre les acteurs externes et le systeme. Ce diagramme est generalement utilise dans les premieres phases de l'analyse afin d'identifier les besoins fonctionnels et de definir le perimetre du systeme.")
    B(doc, "Les elements contenus dans un diagramme de cas d'utilisation sont :")
    B(doc, "1. Acteurs (Actors)", bold=True)
    B(doc, "Les acteurs representent les entites externes qui interagissent avec le systeme. Ils peuvent etre des utilisateurs humains, des systemes externes, ou des organisations externes. L'acteur est represente graphiquement par un bonhomme (stickman).")
    B(doc, "2. Cas d'utilisation (Use Cases)", bold=True)
    B(doc, "Les cas d'utilisation representent les fonctionnalites ou services fournis par le systeme aux acteurs. Ils sont representes par des ovales contenant le nom de la fonctionnalite. Exemples : S'authentifier, Effectuer un transfert, Gerer les cartes virtuelles.")
    B(doc, "3. Frontiere du systeme (System Boundary)", bold=True)
    B(doc, "Elle represente la limite du systeme etudie et contient tous les cas d'utilisation. Elle est representee par un rectangle englobant les ovales.")
    B(doc, "4. Relations (Relationships)", bold=True)
    LP(doc, "Association : Lien entre un acteur et un cas d'utilisation indiquant que l'acteur utilise cette fonctionnalite.")
    LP(doc, "Include (<<include>>) : Indique qu'un cas d'utilisation inclut obligatoirement un autre cas d'utilisation.")
    LP(doc, "Extend (<<extend>>) : Indique qu'un cas d'utilisation peut etre etendu par un autre dans certaines conditions.")
    LP(doc, "Generalisation : Permet de representer l'heritage entre acteurs ou entre cas d'utilisation.")

    # ── 2.2 Identification des Acteurs ───────────────────────────────────────
    doc.add_heading("2.2  Identification des Acteurs", 2)
    B(doc, "La plateforme TrustDesk met en interaction trois categories d'acteurs :")

    doc.add_heading("Acteur 1 : Client (Utilisateur Bancaire)", 3)
    B(doc, "Le client est l'utilisateur final du portefeuille electronique. Il accede a la plateforme via l'application mobile Flutter. Ses interactions principales incluent :")
    LP(doc, "Inscription avec verification eKYC (piece d'identite + selfie + liveness).")
    LP(doc, "Connexion securisee (token Sanctum + biometrie sur mobile).")
    LP(doc, "Consultation du solde et de l'historique des transactions en temps reel.")
    LP(doc, "Transferts P2P entre utilisateurs de la plateforme.")
    LP(doc, "Generation et gestion de cartes virtuelles (gel/degel, limites).")
    LP(doc, "Paiement par QR code chez les commercants partenaires.")
    LP(doc, "Consultation des insights financiers et categorisation des depenses.")

    doc.add_heading("Acteur 2 : Administrateur", 3)
    B(doc, "L'administrateur est le responsable bancaire ou l'operateur BEYN qui supervise la plateforme via le tableau de bord web React.js. Ses responsabilites incluent :")
    LP(doc, "Gestion des comptes clients (activation, desactivation, modification).")
    LP(doc, "Validation ou rejet des dossiers eKYC.")
    LP(doc, "Suivi des incidents de securite et gestion du workflow de triage.")
    LP(doc, "Generation de rapports (clients, transactions, audit, incidents).")
    LP(doc, "Consultation des KPIs en temps reel (dashboard).")

    doc.add_heading("Acteur 3 : Systeme (Moteur IA + API)", 3)
    B(doc, "Le systeme represente les composants automatises de la plateforme :")
    LP(doc, "Triage intelligent des requetes par IA (classification d'intention).")
    LP(doc, "Detection automatique d'anomalies transactionnelles.")
    LP(doc, "Gel preventif du portefeuille en cas de suspicion de fraude.")
    LP(doc, "Envoi de notifications push aux utilisateurs concernes.")
    LP(doc, "Journalisation de toutes les operations dans la piste d'audit.")

    T(doc, ["Acteur", "Type", "Interface", "Responsabilite Principale"], [
        ["Client",         "Primaire",   "Mobile (Flutter)", "Gestion du portefeuille et transactions"],
        ["Administrateur", "Primaire",   "Web (React.js)",   "Supervision, gestion clients et securite"],
        ["Systeme IA",     "Secondaire", "API (Laravel)",    "Automatisation, triage IA et audit"],
    ], "Tableau N\u00b03 : Acteurs du systeme TrustDesk")

    # ── 2.3 Diagramme de cas d'utilisation ───────────────────────────────────
    doc.add_heading("2.3  Diagramme de cas d'utilisation global", 2)
    B(doc, "Le diagramme de cas d'utilisation global ci-dessous represente l'ensemble des fonctionnalites offertes par la plateforme TrustDesk et les interactions entre les trois acteurs identifies et le systeme :")
    add_image(doc, DIAGRAMS / "use_case.png", 14, "Figure N\u00b08 : Diagramme de cas d'utilisation global - TrustDesk")

    doc.add_heading("Les Acteurs : les acteurs de notre projet sont", 3)
    LP(doc, "Client : Acces aux fonctionnalites de portefeuille, transferts, cartes virtuelles et eKYC.")
    LP(doc, "Administrateur : Acces a la gestion des clients, validation eKYC, gestion des incidents.")

    # ── 2.4 Besoins fonctionnels ─────────────────────────────────────────────
    doc.add_heading("2.4  Specification des Besoins Fonctionnels", 2)
    doc.add_heading("Module Authentification & eKYC", 3)
    LP(doc, "Inscription par email avec verification OTP.")
    LP(doc, "Connexion securisee (token Sanctum + biometrie sur mobile).")
    LP(doc, "Parcours eKYC : upload de piece d'identite, selfie avec detection de vivacite (liveness).")
    LP(doc, "Gestion des statuts KYC : en attente, verifie, rejete.")
    LP(doc, "Reinitialisation de mot de passe et deconnexion multi-appareils.")

    doc.add_heading("Module Portefeuille & Transactions", 3)
    LP(doc, "Consultation du solde et de l'historique des transactions en temps reel.")
    LP(doc, "Transferts P2P entre utilisateurs de la plateforme.")
    LP(doc, "Rechargement du portefeuille (top-up) via integration bancaire.")
    LP(doc, "Paiement de factures (telephone, internet, electricite).")
    LP(doc, "Paiement par QR code chez les commercants partenaires.")
    LP(doc, "Definition de limites de depenses quotidiennes et mensuelles.")

    doc.add_heading("Module Cartes Virtuelles", 3)
    LP(doc, "Generation de cartes virtuelles avec chiffrement RSA-2048.")
    LP(doc, "Support de cartes jetables (disposable) a usage unique.")
    LP(doc, "Gel et degel instantane de carte depuis l'application mobile.")

    doc.add_heading("Module Triage IA & Securite", 3)
    LP(doc, "Classificateur d'intention IA pour le routage des requetes.")
    LP(doc, "Detection d'anomalies transactionnelles (montant inhabituel, localisation suspecte).")
    LP(doc, "Gel preventif automatique du portefeuille en cas de suspicion.")
    LP(doc, "Piste d'audit complete et immuable pour toutes les operations sensibles.")

    doc.add_heading("Module Administration", 3)
    LP(doc, "Dashboard avec KPIs en temps reel (utilisateurs actifs, volume transactions, incidents).")
    LP(doc, "Gestion CRUD des clients avec filtres par statut KYC.")
    LP(doc, "Validation et rejet des documents eKYC.")
    LP(doc, "Gestion des incidents : creation, assignation, suivi, cloture.")
    LP(doc, "Generation de rapports (clients, transactions, audit, tickets).")

    T(doc, ["Module", "Fonctionnalites Principales"], [
        ["Authentification & eKYC",   "Inscription OTP, connexion Sanctum, parcours eKYC, liveness"],
        ["Portefeuille & Transactions","Solde, transferts P2P, paiements QR, factures, limites"],
        ["Cartes Virtuelles",          "Generation RSA-2048, cartes jetables, gel/degel"],
        ["Triage IA & Securite",       "Classification IA, detection anomalies, audit trail"],
        ["Administration",             "Dashboard KPIs, gestion clients, validation eKYC, rapports"],
    ], "Tableau N\u00b04 : Besoins fonctionnels")

    # ── 2.5 Besoins non fonctionnels ─────────────────────────────────────────
    doc.add_heading("2.5  Specification des Besoins Non Fonctionnels", 2)
    B(doc, "Les besoins non fonctionnels definissent les contraintes de qualite que la plateforme doit respecter :")
    T(doc, ["Exigence", "Description"], [
        ["Securite",      "Chiffrement RSA-2048/AES-256, conformite OWASP Top 10, protection contre le mass-assignment"],
        ["Performance",   "Temps de reponse API < 300ms pour 95% des requetes"],
        ["Disponibilite", "Objectif de disponibilite de 99.9% (SLA bancaire)"],
        ["Scalabilite",   "Architecture horizontalement extensible pour multi-banque"],
        ["Ergonomie",     "Interface conforme aux standards Material Design 3"],
        ["Compatibilite", "Android 8.0+, iOS 12+, navigateurs modernes (Chrome, Firefox, Edge)"],
        ["Maintenabilite","Architecture modulaire, code documente, tests automatises"],
        ["Auditabilite",  "Journalisation complete de toutes les operations sensibles"],
        ["Conformite",    "Respect des directives Banque d'Algerie (paiement electronique)"],
    ], "Tableau N\u00b05 : Besoins non fonctionnels")

    # ── 2.6 Descriptions textuelles ──────────────────────────────────────────
    doc.add_heading("2.6  Descriptions Textuelles des Cas d'Utilisation", 2)

    B(doc, "Description textuelle du cas d'utilisation : S'inscrire et completer le eKYC", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "S'inscrire et completer le eKYC"],
        ["Acteur principal",    "Client"],
        ["Precondition",        "Application mobile installee, connexion internet disponible"],
        ["Scenario nominal",    "1. Le client ouvre l'application et choisit 'Creer un compte'\n2. Il saisit ses informations (nom, email, mot de passe)\n3. Il recoit et valide le code OTP par email\n4. Il uploade sa piece d'identite (recto/verso)\n5. Il prend un selfie avec detection de vivacite\n6. Le systeme verifie les documents automatiquement\n7. Le compte est cree avec statut KYC 'en attente'\n8. Un administrateur valide le dossier\n9. Le statut passe a 'verifie'"],
        ["Scenario alternatif", "3a. Code OTP expire -> renvoi automatique\n5a. Liveness echoue -> nouvelle tentative\n6a. Document illisible -> demande de re-upload"],
        ["Postcondition",       "Compte cree, portefeuille initialise a 0 DZD"],
    ], "Tableau N\u00b06 : Description UC - Inscription eKYC")

    B(doc, "Description textuelle du cas d'utilisation : Effectuer un transfert P2P", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "Effectuer un transfert P2P"],
        ["Acteur principal",    "Client"],
        ["Precondition",        "Client authentifie, KYC verifie, solde suffisant"],
        ["Scenario nominal",    "1. Le client selectionne 'Transfert' dans le menu\n2. Il saisit l'email du destinataire et le montant\n3. Le systeme verifie le solde disponible\n4. Le moteur IA analyse la transaction (anomalies)\n5. Le client confirme avec authentification biometrique\n6. La transaction est executee atomiquement\n7. Les deux parties recoivent une notification push"],
        ["Scenario alternatif", "3a. Solde insuffisant -> message d'erreur\n4a. Transaction suspecte -> blocage + alerte admin\n5a. Biometrie echouee -> annulation"],
        ["Postcondition",       "Soldes mis a jour, transaction journalisee dans l'audit trail"],
    ], "Tableau N\u00b07 : Description UC - Transfert P2P")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce deuxieme chapitre a formalise les besoins fonctionnels et non fonctionnels de TrustDesk. Nous avons identifie trois acteurs cles \u2014 Client, Administrateur et Systeme \u2014 et decrit les cas d'utilisation couvrant l'integralite des modules : eKYC, portefeuille, cartes virtuelles, triage IA et administration. Les descriptions textuelles detaillees des principaux cas d'utilisation ont permis de preciser les scenarios nominaux et alternatifs. Le chapitre suivant presente la conception architecturale de la solution.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE III : CONCEPTION (start only)
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch3_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre III : Conception", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce chapitre presente la conception detaillee de la plateforme TrustDesk. Nous decrirons l'architecture globale du systeme, les couches de l'API Laravel, le modele de donnees (MCD et MLD), les diagrammes de classes, les diagrammes de sequence detailles, ainsi que l'architecture de chiffrement garantissant la securite des donnees bancaires.", sb=12)

    # ── 3.1 Architecture globale ─────────────────────────────────────────────
    doc.add_heading("3.1  Architecture Globale du Systeme", 2)
    B(doc, "La plateforme TrustDesk adopte une architecture trois tiers (Three-Tier Architecture) composee de trois couches distinctes communicant via des APIs RESTful securisees :")
    LP(doc, "Couche Presentation : Application mobile (Flutter 3 / Dart / Riverpod) et Dashboard web (React.js 18 / TypeScript / Vite). Ces deux clients communiquent avec la couche metier exclusivement via des appels API REST.")
    LP(doc, "Couche Metier : API RESTful Laravel 11 avec authentification Sanctum, modules de services metier, EncryptionService pour le chiffrement RSA/AES, et moteur IA de triage. Cette couche centralise toute la logique applicative.")
    LP(doc, "Couche Donnees : Base de donnees relationnelle (MySQL 8.x / Supabase) pour le stockage persistant, avec Supabase Realtime pour la synchronisation en temps reel et Supabase Storage pour les fichiers KYC.")
    add_image(doc, DIAGRAMS / "architecture.png", 15, "Figure N\u00b011 : Architecture globale trois tiers - TrustDesk")

    doc.add_heading("3.1.1  Architecture Detaillee de l'API Laravel", 3)
    B(doc, "L'API Laravel suit une architecture en couches strictement separees :")
    T(doc, ["Couche", "Composant Laravel", "Role"], [
        ["Routes",          "routes/api.php",              "Definition des endpoints REST"],
        ["Middleware",      "app/Http/Middleware",          "Authentification, rate limiting, CORS"],
        ["Form Requests",   "app/Http/Requests",           "Validation stricte + rejet champs inattendus"],
        ["Controllers",     "app/Http/Controllers",        "Logique de coordination"],
        ["Services",        "app/Services",                "Logique metier (portefeuille, cartes, triage)"],
        ["Models",          "app/Models",                  "Entites Eloquent ORM + relations"],
        ["EncryptionSvc",   "app/Services/EncryptionService", "Chiffrement RSA/AES des donnees sensibles"],
        ["Events/Listeners","app/Events, app/Listeners",   "Journalisation audit, notifications push"],
    ], "Tableau N\u00b08 : Couches de l'architecture API Laravel")

    # ── 3.2 Diagramme de classes ─────────────────────────────────────────────
    doc.add_heading("3.2  Diagramme de Classes", 2)
    B(doc, "Le diagramme de classes modelise les entites principales du systeme TrustDesk et leurs relations. Les classes centrales sont User, Wallet, Transaction, VirtualCard, KycDocument, Incident et AuditLog :")
    add_image(doc, DIAGRAMS / "class_diagram.png", 15, "Figure N\u00b012 : Diagramme de classes - TrustDesk")

    doc.add_heading("3.2.1  Relations Entre les Classes", 3)
    LP(doc, "User (1) <-> (1) Wallet : Chaque utilisateur possede un portefeuille unique (relation 1:1).")
    LP(doc, "User (1) <-> (0..*) VirtualCard : Un utilisateur peut avoir plusieurs cartes virtuelles.")
    LP(doc, "User (1) <-> (0..*) KycDocument : Un utilisateur soumet un ou plusieurs documents eKYC.")
    LP(doc, "Wallet (1) <-> (0..*) Transaction : Un portefeuille est associe a plusieurs transactions.")
    LP(doc, "User (1) <-> (0..*) Incident : Un utilisateur peut etre lie a plusieurs incidents de securite.")
    LP(doc, "User (1) <-> (0..*) AuditLog : Chaque action sensible genere un enregistrement d'audit.")

    # ── 3.3 MCD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.3  Modele Conceptuel des Donnees (MCD)", 2)
    B(doc, "Le MCD represente les entites du systeme et leurs associations dans un formalisme Entite-Association. Les cardinalites refletent les regles metier du portefeuille electronique bancaire :")
    add_image(doc, DIAGRAMS / "mcd.png", 15, "Figure N\u00b013 : Modele Conceptuel des Donnees (MCD)")

    # ── 3.4 MLD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.4  Modele Logique des Donnees (MLD)", 2)
    B(doc, "La transformation du MCD en MLD produit les tables relationnelles suivantes :")
    B(doc, "users (id, name, email, phone, password, role, avatar, kyc_status, is_active, created_at)", bold=True)
    B(doc, "wallets (id, #user_id, balance, currency, status, frozen_at, daily_limit, created_at)", bold=True)
    B(doc, "transactions (id, #wallet_id, type, amount, #recipient_id, description, status, reference, category, created_at)", bold=True)
    B(doc, "virtual_cards (id, #user_id, card_number_enc, cvv_enc, expiry, status, is_disposable, public_key, created_at)", bold=True)
    B(doc, "kyc_documents (id, #user_id, document_type, document_url, selfie_url, liveness_score, status, created_at)", bold=True)
    B(doc, "incidents (id, #user_id, type, description, priority, status, #assigned_to, ai_triage_result, created_at)", bold=True)
    B(doc, "audit_logs (id, #user_id, action, entity_type, entity_id, ip_address, metadata, created_at)", bold=True)
    B(doc, "Note : Le symbole # indique une cle etrangere (FK).", italic=True)

    # ── 3.5 Diagrammes de sequence ───────────────────────────────────────────
    doc.add_heading("3.5  Diagrammes de Sequence", 2)
    B(doc, "Les diagrammes de sequence detaillent le flux d'interactions entre les composants du systeme pour les deux cas d'utilisation principaux.")

    doc.add_heading("3.5.1  Sequence : Inscription & eKYC", 3)
    B(doc, "Ce diagramme montre le flux complet d'inscription d'un nouveau client, incluant la verification OTP, l'upload des documents d'identite, la detection de vivacite, et la validation par un administrateur :")
    add_image(doc, DIAGRAMS / "sequence_ekyc.png", 14, "Figure N\u00b014 : Diagramme de sequence - Inscription & eKYC")

    doc.add_heading("3.5.2  Sequence : Transfert P2P", 3)
    B(doc, "Ce diagramme illustre le processus de transfert d'argent entre deux utilisateurs, incluant la verification du solde, l'analyse IA de la transaction, la confirmation biometrique, et la mise a jour des soldes :")
    add_image(doc, DIAGRAMS / "sequence_p2p.png", 14, "Figure N\u00b015 : Diagramme de sequence - Transfert P2P")

    # ── 3.6 Architecture de chiffrement ──────────────────────────────────────
    doc.add_heading("3.6  Architecture de Chiffrement", 2)
    B(doc, "La securite des donnees sensibles (numeros de carte, CVV) repose sur un modele de chiffrement hybride combinant chiffrement asymetrique et symetrique :")
    doc.add_heading("Chiffrement Asymetrique (RSA-2048)", 3)
    LP(doc, "Utilise pour le chiffrement des numeros de cartes virtuelles et CVV.")
    LP(doc, "Chaque client possede une paire de cles (publique/privee) generee sur son appareil.")
    LP(doc, "La cle publique est transmise au serveur pour le chiffrement.")
    LP(doc, "La cle privee reste exclusivement dans le Secure Enclave du smartphone (approche Zero-Knowledge cote backend).")
    doc.add_heading("Chiffrement Symetrique (AES-256)", 3)
    LP(doc, "Utilise pour le chiffrement des payloads volumineux en transit.")
    LP(doc, "La cle AES est elle-meme chiffree par RSA pour le transport (enveloppe hybride).")
    LP(doc, "AES-256 en mode CBC avec IV aleatoire pour chaque operation.")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce troisieme chapitre a permis de detailler la conception complete de TrustDesk : architecture trois tiers, couches de l'API Laravel, modele de donnees (MCD et MLD), diagrammes de classes et de sequence, ainsi que l'architecture de chiffrement hybride RSA/AES garantissant la securite des donnees bancaires. Les chapitres suivants presenteront la realisation concrete et les resultats des tests.")

    # ── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    size = OUT.stat().st_size // 1024
    print(f"\nMemoire sauvegarde : {OUT}  ({size} KB)")

    # Stats
    secs = len(doc.sections)
    paras = len(doc.paragraphs)
    tbls = len(doc.tables)
    imgs = len(doc.inline_shapes)
    print(f"Sections: {secs}, Paragraphs: {paras}, Tables: {tbls}, Images: {imgs}")

if __name__ == "__main__":
    print("Building memoire v8...")
    build()
    print("Done.")
