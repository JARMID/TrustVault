"""Pass 3: Update tables with TrustDesk project data"""
from docx import Document
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx')

# === Table 0: Fiche technique ===
t0 = doc.tables[0]
table_data = {
    1: 'Projet Acad\u00e9mique - Ann\u00e9e 2024/2025',
    2: 'Projet de Fin de Formation',
    3: 'INSFP Audiovisuel \u2013 Ouled Fayet, Alger, Alg\u00e9rie',
    4: 'D\u00e9veloppement Web et Mobile \u2013 Gestion de la s\u00e9curit\u00e9',
    5: 'Environnement universitaire / Campus',
    6: '2 Stagiaires',
    7: 'Tel: N/A',
    8: 'contact@trustdesk.dz',
}
for ri, val in table_data.items():
    cell = t0.rows[ri].cells[1]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = val
            break

# === Table 1: Ressources humaines ===
t1 = doc.tables[1]
rh_data = [
    ('Service', 'Postes', 'Nombre'),
    ('\u00c9quipe Projet', 'D\u00e9veloppeur Full-Stack (Backend)', '1'),
    ('\u00c9quipe Projet', 'D\u00e9veloppeur Full-Stack (Frontend + Mobile)', '1'),
    ('Encadrement', 'Enseignante encadreur', '1'),
    ('Encadrement', 'Tuteur en entreprise', '1'),
    ('\u00c9tablissement', 'INSFP Audiovisuel', '1'),
    ('\u00c9tablissement', 'Section D\u00e9veloppement Web et Mobile', '1'),
]
for ri, (s, p_text, n) in enumerate(rh_data):
    if ri < len(t1.rows):
        row = t1.rows[ri]
        cells_data = [s, p_text, n]
        for ci, val in enumerate(cells_data):
            if ci < len(row.cells):
                cell = row.cells[ci]
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = val
                        break

# === Table 2: Ressources materielles ===
t2 = doc.tables[2]
mat_data = [
    ('D\u00e9signation', 'Caract\u00e9ristiques', 'Nombre', 'Usage'),
    ('Laptop D\u00e9veloppement', 'Intel Core i7 / 16 Go RAM / SSD 512 Go', '2', 'D\u00e9veloppement'),
    ('Serveur Local (XAMPP)', 'Apache + MySQL 8.0 + PHP 8.2', '1', 'API Backend'),
    ('Smartphone Android', 'Android 12+ / 6 Go RAM', '1', 'Tests Mobile'),
    ('Connexion Internet', 'ADSL / 4G Haut D\u00e9bit', '1', 'D\u00e9ploiement'),
]
for ri, vals in enumerate(mat_data):
    if ri < len(t2.rows):
        row = t2.rows[ri]
        for ci, val in enumerate(vals):
            if ci < len(row.cells):
                cell = row.cells[ci]
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = val
                        break

# Final catch: remaining Genisurv references
for p in doc.paragraphs:
    for r in p.runs:
        if r.text:
            r.text = r.text.replace('Genisurv', 'TrustDesk').replace('GENISURV', 'TrustDesk')
            r.text = r.text.replace('genisurv', 'trustdesk')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text:
                        r.text = r.text.replace('Genisurv', 'TrustDesk').replace('GENISURV', 'TrustDesk')

out = r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'D:\TrustDesk\Memoire_TrustDesk_CH1_v3.docx'
    doc.save(out)
print(f'[OK] Saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
