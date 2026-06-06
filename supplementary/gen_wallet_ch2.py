# -*- coding: utf-8 -*-
"""Chapter 2 - Analyse et Specification des Besoins (Wallet)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t);r.font.size=Pt(sz);r.font.bold=b;r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text');p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa);p.paragraph_format.line_spacing=1.5
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.bold=b;r.font.italic=it
def E():doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph');p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr();numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}));pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h);r.font.bold=True;r.font.size=Pt(12);r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val));r.font.size=Pt(12);r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1);p.alignment=a
    for r in p.runs:r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs:r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs:r.font.name='Times New Roman'

# ═══════ CHAPTER COVER ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)
for _ in range(8):E()
C('CHAPITRE II :',18,True,6)
C('Analyse et Sp\u00e9cification des Besoins',16,False,8)
for _ in range(8):E()
doc.add_page_break()

# ═══════ CHAPTER CONTENT ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)

H1("Introduction :")
B("Ce chapitre est consacr\u00e9 \u00e0 l\u2019analyse d\u00e9taill\u00e9e des besoins de la plateforme TrustDesk. Nous identifierons les acteurs du syst\u00e8me, formaliserons les exigences fonctionnelles et non fonctionnelles, et pr\u00e9senterons les diagrammes UML mod\u00e9lisant le comportement attendu de la plateforme.")

# ─── ACTEURS ───
H1("Identification des acteurs :")
B("La plateforme TrustDesk met en interaction trois cat\u00e9gories d\u2019acteurs :")

H2("Acteur 1 : Client (Utilisateur bancaire)")
B("Le client est l\u2019utilisateur final du portefeuille \u00e9lectronique. Il acc\u00e8de \u00e0 la plateforme via l\u2019application mobile Flutter. Ses principales interactions comprennent : l\u2019inscription avec v\u00e9rification eKYC, la gestion de son portefeuille (solde, historique), les transferts P2P, les paiements par QR code, la gestion de ses cartes virtuelles (cr\u00e9ation, gel/d\u00e9gel), la consultation de ses insights financiers, et la personnalisation de son profil (avatar, param\u00e8tres de s\u00e9curit\u00e9).")

H2("Acteur 2 : Administrateur")
B("L\u2019administrateur est le responsable bancaire ou l\u2019op\u00e9rateur BEYN qui supervise la plateforme via le tableau de bord web React.js. Ses responsabilit\u00e9s incluent : la gestion des comptes clients et de leurs statuts KYC, le suivi et la validation des transactions, la gestion des incidents de s\u00e9curit\u00e9 (triage, escalade), la g\u00e9n\u00e9ration de rapports d\u2019activit\u00e9 et d\u2019audit, la gestion des tokens d\u2019acc\u00e8s et de l\u2019organigramme des op\u00e9rateurs, et le monitoring en temps r\u00e9el des KPIs.")

H2("Acteur 3 : Syst\u00e8me (Moteur IA + API)")
B("Le syst\u00e8me repr\u00e9sente les composants automatis\u00e9s de la plateforme. Il intervient dans : le triage intelligent des requ\u00eates par IA (classification d\u2019intention), la d\u00e9tection automatique d\u2019anomalies transactionnelles, l\u2019envoi de notifications push et alertes de s\u00e9curit\u00e9, la tokenisation des donn\u00e9es de carte, et la journalisation dans la piste d\u2019audit (audit logs).")

T(['Acteur','Type','Interface','Responsabilit\u00e9 principale'],[
    ['Client','Primaire','Mobile (Flutter)','Gestion du portefeuille et des transactions'],
    ['Administrateur','Primaire','Web (React.js)','Supervision, gestion clients et s\u00e9curit\u00e9'],
    ['Syst\u00e8me','Secondaire','API (Laravel)','Automatisation, triage IA et audit'],
])
C("Tableau N\u00b06 : R\u00e9capitulatif des acteurs du syst\u00e8me",12,True,8)
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de contexte des acteurs <<<]",it=True)

# ─── BESOINS FONCTIONNELS ───
H1("Sp\u00e9cification des besoins fonctionnels :")

H2("Module Authentification & eKYC :")
LP("Inscription par email/t\u00e9l\u00e9phone avec v\u00e9rification OTP.")
LP("Connexion s\u00e9curis\u00e9e (Sanctum token + biom\u00e9trie sur mobile).")
LP("Parcours eKYC : upload de pi\u00e8ce d\u2019identit\u00e9, selfie, d\u00e9tection de vivacit\u00e9.")
LP("Gestion des statuts KYC : en attente, v\u00e9rifi\u00e9, rejet\u00e9.")
LP("R\u00e9initialisation de mot de passe et d\u00e9connexion multi-appareils.")

H2("Module Portefeuille & Transactions :")
LP("Consultation du solde et de l\u2019historique des transactions en temps r\u00e9el.")
LP("Transferts P2P entre utilisateurs de la plateforme.")
LP("Rechargement du portefeuille (top-up) via int\u00e9gration bancaire.")
LP("Paiement de factures (t\u00e9l\u00e9phone, internet, \u00e9lectricit\u00e9).")
LP("Paiement par QR code chez les commer\u00e7ants partenaires.")
LP("D\u00e9finition de limites de d\u00e9penses quotidiennes/mensuelles.")

H2("Module Cartes Virtuelles :")
LP("G\u00e9n\u00e9ration de cartes virtuelles avec chiffrement RSA/AES.")
LP("Support de cartes jetables (disposable) \u00e0 usage unique.")
LP("Gel/d\u00e9gel instantan\u00e9 de carte depuis l\u2019application mobile.")
LP("Affichage s\u00e9curis\u00e9 des d\u00e9tails de carte (num\u00e9ro, CVV, expiration).")
LP("Tokenisation compatible Apple Wallet / Google Pay.")

H2("Module Triage IA & S\u00e9curit\u00e9 :")
LP("Classificateur d\u2019intention IA pour le routage automatique des requ\u00eates.")
LP("D\u00e9tection d\u2019anomalies transactionnelles (montant inhabituel, localisation suspecte).")
LP("Alertes automatiques et escalade vers l\u2019administrateur.")
LP("Gel pr\u00e9ventif automatique du portefeuille en cas de suspicion.")
LP("Piste d\u2019audit compl\u00e8te et immuable pour toutes les op\u00e9rations sensibles.")

H2("Module Administration :")
LP("Dashboard avec KPIs en temps r\u00e9el (utilisateurs actifs, volume de transactions, incidents).")
LP("Gestion CRUD des clients avec filtres par statut KYC.")
LP("Validation/rejet des documents eKYC.")
LP("Gestion des incidents : cr\u00e9ation, assignation, suivi, cl\u00f4ture.")
LP("G\u00e9n\u00e9ration de rapports (clients, transactions, audit, tickets).")
LP("Gestion des tokens d\u2019acc\u00e8s et de l\u2019organigramme des op\u00e9rateurs.")

H2("Module Analytique & Insights :")
LP("Cat\u00e9gorisation automatique des d\u00e9penses (alimentation, transport, divertissement, etc.).")
LP("Graphiques d\u2019\u00e9volution des d\u00e9penses (journalier, hebdomadaire, mensuel).")
LP("D\u00e9finition et suivi de budgets personnalis\u00e9s par cat\u00e9gorie.")
LP("Alertes de d\u00e9passement de budget.")
LP("R\u00e9sum\u00e9 financier mensuel avec insights personnalis\u00e9s.")

# ─── BESOINS NON FONCTIONNELS ───
H1("Sp\u00e9cification des besoins non fonctionnels :")
T(['Exigence','Description'],[
    ['S\u00e9curit\u00e9','Chiffrement RSA-2048/AES-256, conformit\u00e9 OWASP Top 10, PCI-DSS'],
    ['Performance','Temps de r\u00e9ponse API < 300ms pour 95% des requ\u00eates'],
    ['Disponibilit\u00e9','Objectif de disponibilit\u00e9 de 99.9% (SLA bancaire)'],
    ['Scalabilit\u00e9','Architecture horizontalement extensible pour multi-banque'],
    ['Ergonomie','Interface conforme aux standards Material Design 3'],
    ['Compatibilit\u00e9','Android 8.0+, iOS 12+, navigateurs modernes'],
    ['Maintenabilit\u00e9','Architecture modulaire, code document\u00e9, tests automatis\u00e9s'],
    ['Auditabilit\u00e9','Journalisation compl\u00e8te de toutes les op\u00e9rations sensibles'],
    ['Conformit\u00e9','Respect des directives Banque d\u2019Alg\u00e9rie (paiement \u00e9lectronique)'],
])
C("Tableau N\u00b07 : Exigences non fonctionnelles",12,True,8)

# ─── USE CASES ───
H1("Mod\u00e9lisation UML :")
B("Nous utilisons le langage UML (Unified Modeling Language) pour formaliser les besoins identifi\u00e9s. Nous pr\u00e9senterons le diagramme de cas d\u2019utilisation global, puis les descriptions textuelles des cas d\u2019utilisation principaux.")

H2("Diagramme de cas d\u2019utilisation global :")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de cas d\u2019utilisation global avec les 3 acteurs (Client, Administrateur, Syst\u00e8me) <<<]",it=True)
C("Figure N\u00b02 : Diagramme de cas d\u2019utilisation global",12,True,8)

# UC1
H2("Description textuelle : S\u2019inscrire et compl\u00e9ter le eKYC")
T(['Champ','Description'],[
    ['Cas d\u2019utilisation','S\u2019inscrire et compl\u00e9ter le eKYC'],
    ['Acteur principal','Client'],
    ['Acteur secondaire','Syst\u00e8me (v\u00e9rification IA)'],
    ['Pr\u00e9condition','L\u2019application mobile est install\u00e9e'],
    ['Sc\u00e9nario nominal','1. Le client lance l\u2019application\n2. Il saisit ses informations personnelles\n3. Il uploade sa pi\u00e8ce d\u2019identit\u00e9\n4. Il prend un selfie pour la v\u00e9rification de vivacit\u00e9\n5. Le syst\u00e8me v\u00e9rifie les documents\n6. Le compte est cr\u00e9\u00e9 avec statut KYC \u00ab en attente \u00bb\n7. L\u2019admin valide \u2192 statut \u00ab v\u00e9rifi\u00e9 \u00bb'],
    ['Sc\u00e9nario alternatif','Document illisible \u2192 demande de re-upload'],
    ['Postcondition','Compte cr\u00e9\u00e9, portefeuille initialis\u00e9 \u00e0 0 DZD'],
])
C("Tableau N\u00b08 : Description UC \u2014 Inscription eKYC",12,True,8)

# UC2
H2("Description textuelle : Effectuer un transfert P2P")
T(['Champ','Description'],[
    ['Cas d\u2019utilisation','Effectuer un transfert P2P'],
    ['Acteur principal','Client'],
    ['Acteur secondaire','Syst\u00e8me (validation + triage IA)'],
    ['Pr\u00e9condition','Client authentifi\u00e9, KYC v\u00e9rifi\u00e9, solde suffisant'],
    ['Sc\u00e9nario nominal','1. Le client s\u00e9lectionne \u00ab Transfert \u00bb\n2. Il saisit le destinataire et le montant\n3. Le syst\u00e8me v\u00e9rifie le solde et analyse la transaction (IA)\n4. Le client confirme avec biom\u00e9trie\n5. La transaction est ex\u00e9cut\u00e9e\n6. Notification aux deux parties'],
    ['Sc\u00e9nario alternatif','Solde insuffisant \u2192 message d\u2019erreur\nTransaction suspecte \u2192 blocage + alerte admin'],
    ['Postcondition','Soldes mis \u00e0 jour, transaction journalis\u00e9e dans l\u2019audit'],
])
C("Tableau N\u00b09 : Description UC \u2014 Transfert P2P",12,True,8)

# UC3
H2("Description textuelle : G\u00e9rer une carte virtuelle")
T(['Champ','Description'],[
    ['Cas d\u2019utilisation','G\u00e9rer une carte virtuelle'],
    ['Acteur principal','Client'],
    ['Pr\u00e9condition','Client authentifi\u00e9, KYC v\u00e9rifi\u00e9'],
    ['Sc\u00e9nario nominal','1. Le client acc\u00e8de \u00e0 la section \u00ab Cartes \u00bb\n2. Il cr\u00e9e une carte virtuelle (standard ou jetable)\n3. Le syst\u00e8me g\u00e9n\u00e8re la carte avec chiffrement PK/SK\n4. Le client peut consulter, geler ou supprimer la carte'],
    ['Postcondition','Carte cr\u00e9\u00e9e et chiffr\u00e9e, visible dans l\u2019app'],
])
C("Tableau N\u00b010 : Description UC \u2014 Gestion carte virtuelle",12,True,8)

# UC4
H2("Description textuelle : G\u00e9rer les clients (Admin)")
T(['Champ','Description'],[
    ['Cas d\u2019utilisation','G\u00e9rer les clients'],
    ['Acteur principal','Administrateur'],
    ['Pr\u00e9condition','Admin authentifi\u00e9 sur le dashboard web'],
    ['Sc\u00e9nario nominal','1. L\u2019admin consulte la liste des clients\n2. Il filtre par statut KYC, activit\u00e9, etc.\n3. Il peut valider/rejeter un dossier eKYC\n4. Il peut geler/d\u00e9geler un compte\n5. Il consulte l\u2019historique des transactions d\u2019un client'],
    ['Postcondition','Statut client mis \u00e0 jour, action journalis\u00e9e'],
])
C("Tableau N\u00b011 : Description UC \u2014 Gestion clients (Admin)",12,True,8)

# UC5
H2("Description textuelle : Triage IA des incidents")
T(['Champ','Description'],[
    ['Cas d\u2019utilisation','Triage IA des incidents'],
    ['Acteur principal','Syst\u00e8me (Moteur IA)'],
    ['Acteur secondaire','Administrateur'],
    ['Pr\u00e9condition','Transaction ou requ\u00eate d\u00e9clench\u00e9e'],
    ['Sc\u00e9nario nominal','1. Le syst\u00e8me re\u00e7oit une transaction ou requ\u00eate\n2. L\u2019IA analyse l\u2019intention et le niveau de risque\n3. Classification : normal / suspect / critique\n4. Si suspect : alerte admin + gel pr\u00e9ventif\n5. L\u2019admin examine et prend une d\u00e9cision'],
    ['Postcondition','Incident class\u00e9, action prise, audit log cr\u00e9\u00e9'],
])
C("Tableau N\u00b012 : Description UC \u2014 Triage IA",12,True,8)

# ─── SEQUENCE DIAGRAMS ───
H2("Diagrammes de s\u00e9quence :")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de s\u00e9quence \u2014 Inscription eKYC <<<]",it=True)
C("Figure N\u00b03 : Diagramme de s\u00e9quence \u2014 Inscription eKYC",12,True,8)
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de s\u00e9quence \u2014 Transfert P2P <<<]",it=True)
C("Figure N\u00b04 : Diagramme de s\u00e9quence \u2014 Transfert P2P",12,True,8)
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de s\u00e9quence \u2014 Triage IA <<<]",it=True)
C("Figure N\u00b05 : Diagramme de s\u00e9quence \u2014 Triage IA",12,True,8)

# ─── CONCLUSION ───
H2("Conclusion :")
B("Ce deuxi\u00e8me chapitre nous a permis de formaliser les besoins fonctionnels et non fonctionnels de la plateforme TrustDesk dans sa nouvelle orientation de portefeuille \u00e9lectronique s\u00e9curis\u00e9. Nous avons identifi\u00e9 trois acteurs cl\u00e9s \u2014 le Client, l\u2019Administrateur et le Syst\u00e8me \u2014 et d\u00e9crit en d\u00e9tail les cas d\u2019utilisation couvrant l\u2019int\u00e9gralit\u00e9 des modules fonctionnels : authentification eKYC, portefeuille et transactions, cartes virtuelles, triage IA, administration et analytique financi\u00e8re.")
B("Le chapitre suivant sera consacr\u00e9 \u00e0 la conception de l\u2019application, o\u00f9 nous pr\u00e9senterons l\u2019architecture globale du syst\u00e8me, le mod\u00e8le de donn\u00e9es relationnel, et les diagrammes de classes d\u00e9taill\u00e9s.")

doc.add_page_break()
out=r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx'
doc.save(out)
print(f'[OK] Ch1+Ch2 saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
