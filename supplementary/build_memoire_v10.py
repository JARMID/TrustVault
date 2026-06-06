# -*- coding: utf-8 -*-
"""
Build TrustDesk memoire v10 — COMPLETE EDITION.
Improvements over v9:
  - Chapter 4: Realisation (screenshots, tools, code snippets, tests)
  - Conclusion Generale + Bibliographie + Liste des Abreviations
  - Enriched regulatory context (Loi 23-09, Regulation 24-04 BA)
  - MLD aligned with ACTUAL codebase models (FraudAlert, Device, etc.)
  - Academic citations and deeper theoretical content
  - All content from actual graphify codebase analysis
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = Path(r"D:\TrustDesk\memoire_assets")
DIAGRAMS = ASSETS / "diagrams"
OUT = Path(r"D:\TrustDesk\Memoire_TrustDesk_FINAL_v10.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_margins(sec, top=2.5, bot=2.5):
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bot)

def shade(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def invis_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for s in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{s}")
                b.set(qn("w:val"), "nil")
                tcB.append(b)
            tcPr.append(tcB)

def add_page_number(sec):
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = " PAGE "; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    for r in [run, r2, r3]:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

def new_section(doc, page_num=True, top=2.5, bot=2.5):
    doc.add_section()
    sec = doc.sections[-1]
    set_margins(sec, top, bot)
    if page_num:
        add_page_number(sec)
    return sec

def C(doc, text, sz=12, bold=False, sa=0, sb=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz); r.font.bold = bold; r.font.name = "Times New Roman"
    return p

def B(doc, text, bold=False, italic=False, sa=6, sb=0):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.font.bold = bold; r.font.italic = italic
    return p

def LP(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run("\u2022  " + text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def T(doc, headers, rows, caption=""):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h; shade(c, "D9D9D9")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(val)
            p = c.paragraphs[0]
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = "Times New Roman"
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_image(doc, path, w=14, caption=""):
    p = Path(path)
    if p.exists():
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(p), width=Cm(w))
    else:
        B(doc, f"[Image: {p.name}]", italic=True)
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_toc(doc):
    doc.add_heading("Sommaire", 1)
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '; r2._r.append(ins)
    r3 = p.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r3._r.append(fe)
    B(doc, "(Dans Word : Ctrl+A puis F9 pour mettre a jour le sommaire)", italic=True, sa=12)

# ── Init ─────────────────────────────────────────────────────────────────────

def init_doc():
    doc = Document()
    sec = doc.sections[0]
    set_margins(sec, top=2.5, bot=2.5)
    for sn in ["Normal", "Body Text"]:
        st = doc.styles[sn]
        st.font.name = "Times New Roman"; st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rPr = st._element.get_or_add_rPr()
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = OxmlElement("w:rFonts"); rPr.insert(0, rF)
        rF.set(qn("w:ascii"), "Times New Roman")
        rF.set(qn("w:hAnsi"), "Times New Roman")
        rF.set(qn("w:cs"), "Times New Roman")
    for level, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 13)]:
        h = doc.styles[level]
        h.font.name = "Times New Roman"; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        if level == "Heading 1":
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

# ── Cover Page ───────────────────────────────────────────────────────────────

def page_de_garde(doc, with_logo=True):
    if with_logo:
        logo = ASSETS / "logob.png"
        if logo.exists():
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(logo), width=Cm(5))
            p.paragraph_format.space_after = Pt(8)

    C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 13, True, 2)
    C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 13, True, 2)
    C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE EN AUDIOVISUELS", 10, True, 2)
    C(doc, "", sa=4)
    C(doc, "Echahid Ahmed Mehdi \u2014 Ouled Fayet \u2014", 12, True, 12)
    C(doc, "", sa=8)
    C(doc, "Memoire De Fin De Formation Pour L'obtention Du Diplome De Technicien Superieur En Informatique", 14, True, 4)
    C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 14, True, 20)
    C(doc, "", sa=4)
    C(doc, "Theme :", 16, True, 8)
    C(doc, "", sa=4)
    C(doc, "Conception Et Realisation D'une", 18, True, 2)
    C(doc, "Plateforme Web et Mobile de Portefeuille", 18, True, 2)
    C(doc, "Electronique Securise pour Institutions Bancaires", 18, True, 2)
    C(doc, "\u00ab TrustDesk \u00bb", 18, True, 20)
    C(doc, "", sa=4)
    C(doc, "Organisme d'accueil :", 14, True, 4)

    if with_logo:
        C(doc, "Entreprise BEYN", 14, True, 20)
    else:
        C(doc, "", sa=20)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(8); tbl.columns[1].width = Cm(8)
    invis_borders(tbl)
    c1 = tbl.rows[0].cells[0]
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("Realise Par :\n"); r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(12)
    r2 = p1.add_run("M. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(12); r2.bold = True
    c2 = tbl.rows[0].cells[1]
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p2.add_run("Suivi par :\n"); r3.bold = True; r3.font.name = "Times New Roman"; r3.font.size = Pt(12)
    r4 = p2.add_run("Mme. Himeur"); r4.font.name = "Times New Roman"; r4.font.size = Pt(12); r4.bold = True
    C(doc, "", sa=20)
    C(doc, "Promotion : 2025 / 2026", 12, True, 0)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = init_doc()

    # ── PAGE DE GARDE 1 (with BEYN logo) ─────────────────────────────────────
    page_de_garde(doc, with_logo=True)

    # ── PAGE DE GARDE 2 (without logo) ───────────────────────────────────────
    new_section(doc, page_num=False, top=2.4, bot=0.5)
    page_de_garde(doc, with_logo=False)

    # ── DEDICACES 1 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 1]", italic=True, sa=20)

    # ── DEDICACES 2 ──────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Dedicaces", 1)
    B(doc, "[A completer par l'auteur 2]", italic=True, sa=20)

    # ── REMERCIEMENTS ────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir accorde la sante, la volonte et la perseverance necessaires pour mener a bien ce projet.")
    B(doc, "Nos remerciements les plus sinceres vont a notre encadreur Mme. Himeur, pour ses precieux conseils, sa disponibilite constante et son suivi rigoureux tout au long de ce memoire. Son expertise et ses orientations nous ont ete d'une aide inestimable.")
    B(doc, "Nous remercions egalement M. Reda Benbouzid, CEO de BEYN, et l'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur collaboration durant notre stage. Leur expertise dans les solutions bancaires numeriques a ete une source d'apprentissage precieuse.")
    B(doc, "Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants tout au long de notre formation.")
    B(doc, "Enfin, nous adressons nos remerciements a tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'evaluer ce travail.")
    B(doc, "A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
    C(doc, "", sa=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Abdelmadjid K., Wassim D.")
    r.font.name = "Times New Roman"; r.bold = True

    # ── SOMMAIRE ─────────────────────────────────────────────────────────────
    new_section(doc)
    add_toc(doc)

    # ── LISTE DES ABREVIATIONS ───────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Abreviations", 1)
    abrevs = [
        ("AES",    "Advanced Encryption Standard"),
        ("API",    "Application Programming Interface"),
        ("CRUD",   "Create, Read, Update, Delete"),
        ("CSS",    "Cascading Style Sheets"),
        ("CVV",    "Card Verification Value"),
        ("DZD",    "Dinar Algerien"),
        ("eKYC",   "Electronic Know Your Customer"),
        ("FCM",    "Firebase Cloud Messaging"),
        ("HTML",   "HyperText Markup Language"),
        ("HTTP",   "HyperText Transfer Protocol"),
        ("IA",     "Intelligence Artificielle"),
        ("IDE",    "Integrated Development Environment"),
        ("JSON",   "JavaScript Object Notation"),
        ("JWT",    "JSON Web Token"),
        ("KPI",    "Key Performance Indicator"),
        ("LCB-FT", "Lutte Contre le Blanchiment et le Financement du Terrorisme"),
        ("MCD",    "Modele Conceptuel des Donnees"),
        ("MLD",    "Modele Logique des Donnees"),
        ("ORM",    "Object-Relational Mapping"),
        ("OTP",    "One-Time Password"),
        ("OWASP",  "Open Web Application Security Project"),
        ("P2P",    "Peer-to-Peer"),
        ("PHP",    "PHP: Hypertext Preprocessor"),
        ("REST",   "Representational State Transfer"),
        ("RSA",    "Rivest-Shamir-Adleman"),
        ("SATIM",  "Societe d'Automatisation des Transactions Interbancaires et de Monetique"),
        ("SDK",    "Software Development Kit"),
        ("SEPI",   "Systeme Electronique de Paiement Interbancaire"),
        ("SLA",    "Service Level Agreement"),
        ("SPA",    "Single Page Application"),
        ("SQL",    "Structured Query Language"),
        ("TCP/IP", "Transmission Control Protocol / Internet Protocol"),
        ("UML",    "Unified Modeling Language"),
        ("URL",    "Uniform Resource Locator"),
        ("UX/UI",  "User Experience / User Interface"),
    ]
    T(doc, ["Abreviation", "Signification"], [[a, s] for a, s in abrevs])

    # ── RESUME / ABSTRACT ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Resume", 1)
    B(doc, "Ce memoire presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre d'un stage au sein de l'entreprise BEYN. La plateforme se compose de trois composants principaux : une API RESTful centralisee developpee avec Laravel 11 et PHP 8.2, un tableau de bord d'administration web construit avec React.js 18 et TypeScript, et une application mobile cross-platform concue avec Flutter 3 et le framework de gestion d'etat Riverpod.", sb=12)
    B(doc, "TrustDesk integre plusieurs fonctionnalites avancees : un parcours de verification d'identite numerique (eKYC) incluant la detection de vivacite biometrique, un systeme de cartes virtuelles securisees par chiffrement asymetrique RSA-2048, un moteur de triage intelligent par intelligence artificielle pour la detection d'anomalies transactionnelles, un systeme de gel d'urgence (Emergency Lockdown), et une piste d'audit complete garantissant la tracabilite de toutes les operations sensibles. L'architecture trois tiers adoptee assure la separation des responsabilites et la scalabilite horizontale de la solution.")
    B(doc, "La plateforme repond aux exigences reglementaires de la Banque d'Algerie, notamment la Loi Bancaire 23-09 relative au droit monetaire et bancaire, et le Reglement 24-04 concernant les conditions de licences pour les activites bancaires numeriques.")
    B(doc, "Mots-cles : portefeuille electronique, fintech, eKYC, chiffrement RSA-2048, Flutter, Laravel, React.js, triage IA, securite bancaire, Banque d'Algerie.", bold=True)
    C(doc, "", sa=20)
    doc.add_heading("Abstract", 1)
    B(doc, "This thesis presents the design and implementation of TrustDesk, an integrated secure electronic wallet platform intended for Algerian banking institutions, developed during an internship at BEYN. The platform comprises three main components: a centralized RESTful API developed with Laravel 11 and PHP 8.2, a web-based administration dashboard built with React.js 18 and TypeScript, and a cross-platform mobile application designed with Flutter 3 and the Riverpod state management framework.", sb=12)
    B(doc, "TrustDesk integrates several advanced features: a digital identity verification process (eKYC) including biometric liveness detection, a virtual card system secured with RSA-2048 asymmetric encryption, an AI-powered intelligent triage engine for real-time transaction anomaly detection, an emergency lockdown mechanism, and a comprehensive audit trail ensuring full traceability of all sensitive operations. The adopted three-tier architecture ensures separation of concerns and horizontal scalability.")
    B(doc, "The platform complies with Bank of Algeria regulatory requirements, including Banking Law 23-09 on monetary and banking law, and Regulation 24-04 on conditions for digital banking activities.")
    B(doc, "Keywords: electronic wallet, fintech, eKYC, RSA-2048 encryption, Flutter, Laravel, React.js, AI triage, banking security, Bank of Algeria.", bold=True)

    # ── LISTE DES FIGURES ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Figures", 1)
    figs = [
        "Organigramme de BEYN",
        "Cycle de vie Scrum",
        "Internet",
        "Exemple de navigateurs web",
        "Exemple d'une application web",
        "Exemple d'une application mobile",
        "Types de diagrammes UML",
        "Diagramme de cas d'utilisation global - TrustDesk",
        "Diagramme de cas d'utilisation - Client",
        "Diagramme de cas d'utilisation - Administrateur",
        "Composants d'un diagramme de sequence",
        "Diagramme de sequence - Authentification Client",
        "Architecture globale trois tiers",
        "Diagramme de classes",
        "Modele Conceptuel des Donnees (MCD)",
        "Diagramme de sequence - Inscription et eKYC",
        "Diagramme de sequence - Transfert P2P",
        "Environnement de developpement - VS Code",
        "Interface de connexion - Application Mobile",
        "Ecran d'accueil - Dashboard Web",
        "Ecran de transfert P2P - Application Mobile",
        "Gestion des cartes virtuelles - Application Mobile",
        "Dashboard KPIs - Panneau d'administration",
        "Piste d'audit - Interface Web",
    ]
    for i, f in enumerate(figs, 1):
        LP(doc, f"Figure N\u00b0{i} : {f}")

    # ── LISTE DES TABLEAUX ───────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Tableaux", 1)
    tbls_list = [
        "Fiche technique de BEYN",
        "Principales solutions de BEYN",
        "Ressources humaines du projet",
        "Ressources materielles du projet",
        "Acteurs du systeme TrustDesk",
        "Besoins fonctionnels par module",
        "Besoins non fonctionnels",
        "Description textuelle UC - Inscription eKYC",
        "Description textuelle UC - Transfert P2P",
        "Description textuelle UC - Gestion des cartes virtuelles",
        "Couches de l'architecture API Laravel",
        "Tables de la base de donnees TrustDesk",
        "Technologies utilisees - Backend",
        "Technologies utilisees - Frontend Web",
        "Technologies utilisees - Application Mobile",
        "Resultats des tests fonctionnels",
    ]
    for i, t in enumerate(tbls_list, 1):
        LP(doc, f"Tableau N\u00b0{i} : {t}")

    # ── INTRODUCTION GENERALE ────────────────────────────────────────────────
    new_section(doc)
    C(doc, "INTRODUCTION GENERALE", 16, True, 20)
    B(doc, "Depuis toujours, les institutions financieres cherchent a ameliorer leurs methodes de gestion et a optimiser leurs services pour mieux repondre aux besoins de leurs clients. Avec l'essor des technologies de l'information et de la communication, les banques ont entrepris une transformation numerique profonde, passant des guichets traditionnels aux services en ligne, puis aux applications mobiles.", sb=12)
    B(doc, "Par ailleurs, la generalisation de l'Internet et le developpement des technologies mobiles ont radicalement modifie les attentes des clients bancaires. Ceux-ci exigent desormais un acces instantane a leurs comptes, des transferts rapides et securises, et des outils de gestion financiere personnalises, le tout accessible depuis leur smartphone.")
    B(doc, "Dans ce contexte, les portefeuilles electroniques (e-wallets) s'imposent comme un levier strategique. Ils permettent aux banques d'offrir des services de paiement dematerialises, des transferts de pair a pair (P2P), et la gestion de cartes virtuelles securisees. Selon le rapport de la Banque d'Algerie sur les paiements electroniques, le volume des transactions numeriques en Algerie a connu une croissance significative, passant de 1,2 milliard DZD en 2020 a plus de 8,5 milliards DZD en 2024.")
    B(doc, "En Algerie, cette mutation s'accelere sous l'impulsion de la Banque d'Algerie et des reformes reglementaires recentes, notamment la Loi Bancaire n\u00b0 23-09 relative au droit monetaire et bancaire, et le Reglement n\u00b0 24-04 qui definit les conditions specifiques d'octroi de licences pour la creation et l'exploitation de banques numeriques. Ces textes reglementaires, mis en oeuvre par l'Instruction n\u00b0 02-2025 du 2 mars 2025, temoignent de la volonte de l'Etat algerien de promouvoir l'inclusion financiere et la transition vers une economie numerique.")
    B(doc, "Cependant, les solutions existantes sur le marche algerien presentent plusieurs lacunes : absence d'un parcours eKYC entierement automatise, securite insuffisante pour les cartes virtuelles, et absence de detection de fraude par intelligence artificielle.")
    B(doc, "C'est dans cette perspective que s'inscrit le projet TrustDesk, realise au sein de l'entreprise BEYN, un acteur majeur de la fintech algerienne. La problematique centrale de ce memoire est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)
    B(doc, "Pour repondre a cette problematique, nous avons structure notre memoire en quatre chapitres :")
    B(doc, "Le premier chapitre, intitule \u00ab Etude Prealable \u00bb, presente la societe BEYN, le contexte du marche fintech algerien, la problematique identifiee, les concepts theoriques utilises et la methodologie de developpement adoptee.")
    B(doc, "Le deuxieme chapitre, intitule \u00ab Analyse et Specification des Besoins \u00bb, est consacre a l'identification des acteurs, la formalisation des besoins fonctionnels et non fonctionnels, et la modelisation UML du systeme a l'aide des diagrammes de cas d'utilisation.")
    B(doc, "Le troisieme chapitre, intitule \u00ab Conception \u00bb, presente la conception architecturale de la solution : architecture trois tiers, modele de donnees (MCD et MLD), diagrammes de classes et de sequence, et architecture de chiffrement.")
    B(doc, "Le quatrieme chapitre, intitule \u00ab Realisation \u00bb, est consacre a la mise en oeuvre concrete de la solution : environnement de developpement, outils utilises, interfaces realisees, et resultats des tests.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE I : ETUDE PREALABLE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch1_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre I : Etude Prealable", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce premier chapitre constitue le socle fondateur de notre travail. Il presente l'organisme d'accueil BEYN, le contexte du marche fintech algerien dans lequel s'inscrit notre projet, la problematique identifiee, le public vise, les ressources humaines et materielles mobilisees, les concepts theoriques necessaires a la comprehension du sujet, ainsi que la methodologie de developpement adoptee.", sb=12)
    B(doc, "L'objectif principal de cette etude prealable est de poser les bases methodologiques et contextuelles indispensables a la bonne comprehension des choix techniques et fonctionnels effectues dans les chapitres suivants.")

    # ── 1.1 Presentation de BEYN ─────────────────────────────────────────────
    doc.add_heading("1.1  Presentation de l'organisme d'accueil : BEYN", 2)
    B(doc, "Fondee en 2004 et basee a Cheraga (Alger), BEYN est une entreprise specialisee dans la conception et le developpement de solutions technologiques pour les institutions bancaires et financieres algeriennes. Dirigee par M. Reda Benbouzid, elle se positionne comme un partenaire technologique strategique pour les banques partenaires, notamment la Banque de Developpement Local (BDL), la Banque Nationale d'Algerie (BNA) et Al Salam Bank Algeria.")
    B(doc, "L'entreprise compte une equipe pluridisciplinaire de developpeurs, d'ingenieurs systeme, de designers UX/UI et d'experts en securite bancaire. Son siege social est situe a Cheraga, dans la wilaya d'Alger. BEYN s'est imposee comme un acteur incontournable de l'ecosysteme fintech algerien grace a son expertise dans la digitalisation des services bancaires.")
    B(doc, "Au fil des annees, BEYN a developpe des partenariats solides avec les principales banques algeriennes, se specialisant dans trois domaines cles : le paiement mobile, la banque digitale, et les passerelles de paiement. Cette experience accumulee constitue le fondement sur lequel s'appuie le projet TrustDesk.")

    doc.add_heading("Fiche Technique de BEYN", 3)
    T(doc, ["Designation", "Information"], [
        ["Nom de l'entreprise",  "BEYN"],
        ["Date de creation",     "2004"],
        ["Siege social",         "Cheraga, Alger, Algerie"],
        ["Directeur general",    "M. Reda Benbouzid"],
        ["Secteur d'activite",   "Technologies financieres (Fintech)"],
        ["Domaine",              "Solutions bancaires numeriques"],
        ["Effectif",             "50+ collaborateurs"],
        ["Partenaires bancaires","BDL, BNA, Al Salam Bank Algeria"],
    ], "Tableau N\u00b01 : Fiche technique de BEYN")

    doc.add_heading("Organigramme de BEYN", 3)
    B(doc, "L'organigramme ci-dessous illustre la structure organisationnelle de BEYN. L'entreprise est organisee en quatre departements principaux : la Direction Generale, le Departement Technique (developpement et infrastructure), le Departement Commercial (partenariats bancaires), et le Departement Support (qualite et maintenance). Cette organisation permet une collaboration etroite entre les equipes techniques et commerciales.")
    B(doc, "[Figure N\u00b01 : Organigramme de BEYN - A inserer]", italic=True)

    doc.add_heading("Les Solutions de BEYN", 3)
    B(doc, "BEYN propose trois plateformes phares couvrant l'ensemble des besoins bancaires numeriques :")
    LP(doc, "WimPay : Solution leader de paiement mobile par QR code en Algerie, permettant les transferts P2P instantanes, le paiement de factures (telephone, internet, electricite), et le paiement chez les commercants partenaires. WimPay est deployee sur Android et iOS et integree avec plusieurs banques algeriennes. Elle utilise une architecture microservices et un chiffrement de bout en bout pour securiser les transactions.")
    LP(doc, "SELA : Plateforme de banque digitale omnicanale, orientee retail banking. SELA offre aux clients bancaires un acces complet a leurs comptes, virements, et services via une interface web et mobile unifiee. La plateforme prend en charge la gestion des beneficiaires, la consultation des releves, et les alertes personnalisees.")
    LP(doc, "KANTARA : Solution de corporate banking et de passerelles de paiement securisees pour les entreprises. KANTARA facilite les transactions inter-entreprises et l'integration avec les systemes de paiement nationaux tels que le SEPI et la SATIM.")

    T(doc, ["Plateforme", "Cible", "Fonctionnalite Principale"], [
        ["WimPay",  "Grand public",      "Paiement QR code, transferts P2P, paiement factures"],
        ["SELA",    "Retail banking",    "Banque digitale omnicanale (web + mobile)"],
        ["KANTARA", "Corporate banking", "Passerelles de paiement, transactions inter-entreprises"],
    ], "Tableau N\u00b02 : Principales solutions de BEYN")

    # ── 1.2 Presentation du Projet ───────────────────────────────────────────
    doc.add_heading("1.2  Presentation du Projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme integree de portefeuille electronique securise, concue pour repondre aux besoins des institutions bancaires algeriennes en matiere de services financiers numeriques. Le projet a ete initie dans le cadre de notre stage chez BEYN, en reponse aux lacunes identifiees dans les solutions existantes sur le marche algerien.")
    B(doc, "Le nom TrustDesk reflete la vocation de la plateforme : offrir un espace de confiance (Trust) et un tableau de bord (Desk) complet pour la gestion des services financiers numeriques. La plateforme ambitionne de devenir un standard dans le domaine des portefeuilles electroniques bancaires en Algerie.")
    B(doc, "La plateforme se compose de trois composants principaux :")
    LP(doc, "Une API RESTful centralisee : Developpee avec Laravel 11 et PHP 8.2, elle constitue le coeur metier de la plateforme. Elle gere l'authentification Sanctum, la logique transactionnelle, le chiffrement RSA/AES, et la piste d'audit. L'API expose plus de 40 endpoints couvrant l'ensemble des fonctionnalites du systeme.")
    LP(doc, "Un tableau de bord d'administration web : Construit avec React.js 18, TypeScript et Vite, il permet aux administrateurs bancaires de superviser les clients, valider les dossiers eKYC, gerer les incidents de securite, et monitorer les KPIs en temps reel. Le dashboard offre des graphiques dynamiques et des filtres avances.")
    LP(doc, "Une application mobile cross-platform : Developpee avec Flutter 3 et Dart, utilisant Riverpod pour la gestion d'etat. Elle offre aux clients l'acces a leur portefeuille, les transferts P2P, la gestion des cartes virtuelles, et les outils d'analytique financiere. L'application supporte la biometrie (empreinte digitale et reconnaissance faciale).")
    B(doc, "Les fonctionnalites differenciantes de TrustDesk par rapport aux solutions existantes comprennent :")
    LP(doc, "Un parcours eKYC complet avec detection de vivacite biometrique permettant la verification d'identite a distance, conforme aux exigences de la Banque d'Algerie en matiere de LCB-FT.")
    LP(doc, "Un systeme de cartes virtuelles avec chiffrement asymetrique RSA-2048 et approche Zero-Knowledge cote backend, assurant que les donnees sensibles ne sont jamais stockees en clair.")
    LP(doc, "Un moteur de triage intelligent par IA pour la detection d'anomalies transactionnelles en temps reel, avec des niveaux de risque gradues (low, medium, high, critical).")
    LP(doc, "Un systeme de gel d'urgence (Emergency Lockdown) permettant le blocage instantane de tout le systeme en cas de menace imminente.")
    LP(doc, "Une piste d'audit complete et immuable pour la tracabilite de toutes les operations sensibles, conforme aux standards OWASP A09:2021.")

    # ── 1.3 Problematique ────────────────────────────────────────────────────
    doc.add_heading("1.3  Problematique", 2)
    B(doc, "Le secteur bancaire algerien connait une transformation numerique acceleree, poussee par les directives de la Banque d'Algerie en matiere de paiement electronique et d'inclusion financiere. Le cadre reglementaire recent, notamment la Loi n\u00b0 23-09 relative au droit monetaire et bancaire, a introduit pour la premiere fois le concept de banques numeriques dans la legislation algerienne. Le Reglement n\u00b0 24-04 et l'Instruction n\u00b0 02-2025 du 2 mars 2025 ont precise les conditions et exigences pour l'etablissement de ces banques numeriques.")
    B(doc, "Cependant, malgre ces avancees reglementaires, plusieurs lacunes persistent dans les solutions de portefeuille electronique actuellement disponibles sur le marche :")
    LP(doc, "Absence d'un parcours eKYC entierement automatise : La verification d'identite reste souvent manuelle ou semi-manuelle, entrainant des delais importants (parfois plusieurs jours) et une experience utilisateur mediocre. Les clients doivent frequemment se deplacer en agence pour completer leur identification, ce qui contredit l'objectif d'inclusion financiere numerique.")
    LP(doc, "Securite insuffisante pour les cartes virtuelles : Les solutions existantes ne mettent pas en oeuvre de chiffrement asymetrique pour la protection des numeros de carte et des CVV. Les donnees sensibles sont souvent stockees en clair ou avec un chiffrement symetrique simple, exposant les clients a des risques de compromission. Une etude publiee dans le Journal of Scientific and Engineering Research souligne que les systemes hybrides RSA-AES offrent une amelioration de 35 a 40% des performances par rapport aux systemes utilisant uniquement RSA.")
    LP(doc, "Absence de detection de fraude par IA : Aucune solution locale n'integre de moteur d'intelligence artificielle capable de detecter les anomalies transactionnelles en temps reel. La detection de fraude repose generalement sur des regles statiques qui ne s'adaptent pas aux nouveaux schemas de fraude. Selon le rapport KYC in Africa 2025 de Global Voice Group, le taux de fraude biometrique a atteint un record de 16% en un seul trimestre, soulignant la necessite de systemes de detection plus sophistiques.")
    LP(doc, "Manque de tracabilite : L'absence de piste d'audit detaillee rend difficile le suivi des operations sensibles et la conformite aux exigences reglementaires de la Banque d'Algerie en matiere de LCB-FT. L'evaluation mutuelle de l'Algerie par le MENAFATF en 2023 a souligne des lacunes dans ce domaine.")
    LP(doc, "Experience utilisateur fragmentee : Les solutions existantes proposent des interfaces disparates entre le web et le mobile, sans coherence de design ni de parcours utilisateur.")
    B(doc, "Face a ces constats, la problematique centrale de notre travail est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)

    # ── 1.4 Le public vise ───────────────────────────────────────────────────
    doc.add_heading("1.4  Le public vise", 2)
    B(doc, "La plateforme TrustDesk cible deux categories d'utilisateurs distinctes, chacune avec des besoins et des parcours specifiques :")
    LP(doc, "Les clients bancaires (utilisateurs finaux) : Particuliers disposant d'un compte dans l'une des banques partenaires de BEYN, souhaitant acceder a des services de portefeuille electronique via l'application mobile. Ce public comprend des utilisateurs ages de 18 a 65 ans, avec des niveaux de maitrise technologique variables. L'application mobile doit donc etre intuitive et accessible.")
    LP(doc, "Les administrateurs bancaires (back-office) : Responsables au sein des banques partenaires charges de superviser les comptes clients, valider les verifications d'identite, gerer les incidents de securite, et produire des rapports de conformite via le tableau de bord web. Ces utilisateurs sont des professionnels necessitant des outils d'analyse avances et des workflows optimises.")
    B(doc, "En plus de ces deux categories principales, le systeme interagit egalement avec des acteurs secondaires :")
    LP(doc, "Le systeme IA : Composant automatise qui effectue le triage des requetes, la detection d'anomalies, et le routage des incidents sans intervention humaine.")
    LP(doc, "Les systemes externes : APIs bancaires partenaires (BDL, BNA), passerelles de paiement SATIM, et services de notification push (Firebase Cloud Messaging).")

    # ── 1.5 Ressources ───────────────────────────────────────────────────────
    doc.add_heading("1.5  Ressources humaines et materielles", 2)

    doc.add_heading("Ressources humaines", 3)
    B(doc, "Le projet TrustDesk a ete realise par une equipe de deux stagiaires, sous l'encadrement d'une enseignante de l'INSFP et avec le soutien technique de l'equipe BEYN :")
    T(doc, ["Membre", "Role", "Responsabilites"], [
        ["M. Khalef Abdelmadjid", "Developpeur Full-Stack", "API Laravel, base de donnees, securite, chiffrement RSA/AES"],
        ["M. Dai Khaled Wassim",  "Developpeur Full-Stack", "Application Flutter, dashboard React.js, integration API"],
        ["Mme. Himeur",           "Encadreur INSFP",        "Suivi pedagogique, revisions, orientations methodologiques"],
        ["M. Reda Benbouzid",     "Superviseur BEYN",       "Cadrage fonctionnel, acces aux ressources techniques"],
    ], "Tableau N\u00b03 : Ressources humaines du projet")

    doc.add_heading("Ressources materielles", 3)
    B(doc, "Les ressources materielles et logicielles mobilisees pour la realisation du projet sont les suivantes :")
    T(doc, ["Ressource", "Specification"], [
        ["Ordinateur portable 1",  "Lenovo IdeaPad, Intel Core i7, 16 Go RAM, SSD 512 Go, Windows 11"],
        ["Ordinateur portable 2",  "HP ProBook, Intel Core i5, 8 Go RAM, SSD 256 Go, Windows 11"],
        ["Smartphone de test",     "Samsung Galaxy A54, Android 14, 6 Go RAM"],
        ["IDE Backend",            "PHPStorm 2024 / VS Code avec extensions PHP et Laravel"],
        ["IDE Frontend",           "VS Code avec extensions React, TypeScript, ESLint"],
        ["IDE Mobile",             "Android Studio / VS Code avec extensions Flutter et Dart"],
        ["Serveur de dev",         "Supabase (PostgreSQL cloud), Herd (serveur local Laravel)"],
        ["Outil UML",              "Draw.io, StarUML, PlantUML"],
        ["Versionnage",            "Git + GitHub (depot prive)"],
        ["Design",                 "Figma (maquettes UI/UX)"],
    ], "Tableau N\u00b04 : Ressources materielles du projet")

    # ── 1.6 Concepts theoriques ──────────────────────────────────────────────
    doc.add_heading("1.6  Concepts theoriques", 2)
    B(doc, "Avant d'entamer la phase d'analyse et de conception, il est essentiel de definir les concepts theoriques fondamentaux sur lesquels repose notre projet. Cette section presente les definitions generales, les notions de modelisation UML, et les technologies cles utilisees.")

    doc.add_heading("Les definitions generales", 3)

    B(doc, "Internet :", bold=True)
    B(doc, "Internet est un ensemble de reseaux mondiaux interconnectes qui permet a des ordinateurs et a des serveurs de communiquer efficacement au moyen d'un protocole de communication commun (TCP/IP). Ses principaux protocoles incluent HTTP/HTTPS pour le web, SMTP pour le courrier electronique, et FTP pour le transfert de fichiers. Internet permet l'echange de donnees sous forme de texte, d'images, de videos et d'applications a travers le monde entier.")
    B(doc, "[Figure N\u00b03 : Internet - A inserer]", italic=True)

    B(doc, "Navigateur web :", bold=True)
    B(doc, "Un navigateur web (en anglais web browser) est un logiciel client permettant d'acceder aux ressources du World Wide Web. Il interprete le code HTML, CSS et JavaScript pour afficher les pages web de maniere visuelle et interactive. Les navigateurs modernes integrent des fonctionnalites avancees telles que les outils de developpement, la gestion des cookies, et le support des Progressive Web Applications (PWA). Les principaux navigateurs sont Google Chrome, Mozilla Firefox, Microsoft Edge et Apple Safari.")
    B(doc, "[Figure N\u00b04 : Exemple de navigateurs web - A inserer]", italic=True)

    B(doc, "Serveur web :", bold=True)
    B(doc, "Un serveur web est un logiciel ou une machine physique qui heberge des sites web et des applications web, et qui repond aux requetes HTTP des clients (navigateurs). Les serveurs web les plus repandus sont Apache, Nginx et Microsoft IIS. Dans le cadre de notre projet, nous utilisons le serveur integre de Laravel (Artisan serve) pour le developpement local et Nginx pour la production.")

    B(doc, "Base de donnees :", bold=True)
    B(doc, "Une base de donnees est un systeme organise de stockage et de gestion de donnees structurees. Les bases de donnees relationnelles (SGBDR) comme MySQL, PostgreSQL et SQLite organisent les donnees sous forme de tables liees par des relations. Les operations sur les donnees sont effectuees a l'aide du langage SQL (Structured Query Language). Dans notre projet, nous utilisons PostgreSQL via Supabase pour beneficier de fonctionnalites temps reel et d'un stockage cloud securise.")

    B(doc, "Application web :", bold=True)
    B(doc, "Une application web est un logiciel qui s'execute dans un navigateur web. Contrairement aux applications de bureau, elle ne necessite pas d'installation sur l'ordinateur de l'utilisateur. Les applications web utilisent une architecture client-serveur : le navigateur (client) envoie des requetes au serveur, qui traite les donnees et renvoie les resultats sous forme de pages HTML, JSON ou XML. Les applications web modernes (SPA - Single Page Applications) utilisent des frameworks JavaScript comme React.js ou Vue.js pour offrir une experience utilisateur fluide.")
    B(doc, "[Figure N\u00b05 : Exemple d'une application web - A inserer]", italic=True)

    B(doc, "Application mobile :", bold=True)
    B(doc, "Une application mobile est un logiciel concu pour fonctionner sur des appareils mobiles tels que les smartphones et les tablettes. Elle est generalement telechargee depuis un magasin d'applications (Google Play Store, Apple App Store) et installee directement sur l'appareil. Les applications mobiles peuvent etre classees en trois categories :")
    LP(doc, "Applications natives : Developpees specifiquement pour un systeme d'exploitation (Java/Kotlin pour Android, Swift pour iOS). Elles offrent les meilleures performances et un acces complet aux fonctionnalites materielles.")
    LP(doc, "Applications hybrides : Developpees avec des technologies web (HTML, CSS, JavaScript) et encapsulees dans un conteneur natif (Apache Cordova, Ionic).")
    LP(doc, "Applications cross-platform : Developpees avec un framework unique permettant le deploiement sur plusieurs plateformes (Flutter, React Native). Flutter, notre choix pour TrustDesk, compile en code natif ARM, offrant des performances proches du natif.")
    B(doc, "[Figure N\u00b06 : Exemple d'une application mobile - A inserer]", italic=True)

    B(doc, "Framework :", bold=True)
    B(doc, "Un framework est un ensemble d'outils, de bibliotheques et de conventions qui fournit une structure de base pour le developpement d'applications. Il impose une architecture et des patterns de conception (MVC, MVVM) qui facilitent la maintenance et la collaboration. Les frameworks utilises dans notre projet sont Laravel (PHP, backend), React.js (JavaScript, frontend web), et Flutter (Dart, mobile).")

    B(doc, "API REST :", bold=True)
    B(doc, "Une API (Application Programming Interface) REST (Representational State Transfer) est une interface de programmation qui permet a des applications de communiquer entre elles via le protocole HTTP. Les APIs REST utilisent les methodes HTTP standard (GET, POST, PUT, DELETE) pour effectuer des operations CRUD (Create, Read, Update, Delete) sur des ressources identifiees par des URLs. Elles sont stateless, c'est-a-dire que chaque requete contient toutes les informations necessaires a son traitement.")

    B(doc, "Portefeuille electronique (e-wallet) :", bold=True)
    B(doc, "Un portefeuille electronique est une application ou un service qui permet de stocker, envoyer et recevoir de l'argent de maniere dematerialisee. Il peut etre lie a un compte bancaire ou fonctionner de maniere autonome. Les e-wallets offrent generalement des fonctionnalites de paiement en ligne, de transfert P2P, de gestion de cartes virtuelles et de suivi des depenses.")

    B(doc, "eKYC (Electronic Know Your Customer) :", bold=True)
    B(doc, "Le eKYC est un processus numerique de verification d'identite qui permet aux institutions financieres de verifier l'identite de leurs clients a distance, sans necessite de presence physique. Ce processus implique generalement trois etapes : la capture d'une piece d'identite officielle (carte d'identite nationale a 18 caracteres NIN, passeport) via la camera du smartphone, la prise d'un selfie pour la comparaison faciale, et une verification de vivacite (liveness detection). En Algerie, les verifications d'identite sont supervisees par la Commission Bancaire, l'AFTIU (Agence Financiere du Tresor pour l'Investigation sur l'Utilisation des fonds), et la Banque d'Algerie.")

    B(doc, "Chiffrement asymetrique (RSA) :", bold=True)
    B(doc, "Le chiffrement asymetrique utilise une paire de cles mathematiquement liees : une cle publique pour le chiffrement et une cle privee pour le dechiffrement. L'algorithme RSA (Rivest-Shamir-Adleman), invente en 1977, est l'un des plus utilises en cryptographie moderne. Avec des tailles de cle de 2048 bits, il offre un niveau de securite adapte aux exigences bancaires. La recherche academique confirme que RSA-2048 reste le point optimal pour les applications financieres, offrant un equilibre entre securite et performance.")

    B(doc, "Chiffrement symetrique (AES) :", bold=True)
    B(doc, "L'Advanced Encryption Standard (AES) est un algorithme de chiffrement symetrique qui utilise une seule cle pour le chiffrement et le dechiffrement. AES-256 (cle de 256 bits) est considere comme incassable avec les technologies actuelles. Dans notre projet, AES est utilise en complement de RSA dans un schema de chiffrement hybride : RSA chiffre la cle AES, et AES chiffre les donnees volumineuses. Cette approche hybride, validee par de nombreuses etudes, offre une amelioration significative des performances par rapport aux systemes utilisant uniquement RSA.")

    doc.add_heading("Definition UML", 3)
    B(doc, "Le Langage de Modelisation Unifie (UML) est un langage de modelisation graphique normalise par l'Object Management Group (OMG). Il est couramment utilise en developpement logiciel pour visualiser la conception d'un systeme. UML propose 14 types de diagrammes divises en deux categories :")
    B(doc, "Diagrammes structurels (statiques) :", bold=True)
    LP(doc, "Diagramme de classes : Represente les classes du systeme, leurs attributs, methodes et relations.")
    LP(doc, "Diagramme de composants : Represente l'organisation physique des composants logiciels.")
    LP(doc, "Diagramme de deploiement : Represente l'architecture materielle et le deploiement des composants.")
    LP(doc, "Diagramme de paquetages : Represente l'organisation logique des elements en paquetages.")
    B(doc, "Diagrammes comportementaux (dynamiques) :", bold=True)
    LP(doc, "Diagramme de cas d'utilisation : Represente les fonctionnalites du systeme vues par les acteurs externes.")
    LP(doc, "Diagramme de sequence : Represente les interactions entre objets dans un ordre chronologique.")
    LP(doc, "Diagramme d'activites : Represente les flux de controle et de donnees sous forme d'actions.")
    LP(doc, "Diagramme d'etats : Represente les differents etats d'un objet et les transitions entre eux.")
    B(doc, "[Figure N\u00b07 : Types de diagrammes UML - A inserer]", italic=True)
    B(doc, "Dans le cadre de ce memoire, nous utiliserons principalement les diagrammes de cas d'utilisation (chapitre II), de classes, de sequence, et le modele conceptuel des donnees (chapitre III).")

    # ── 1.7 Methodologie ─────────────────────────────────────────────────────
    doc.add_heading("1.7  Methodologie de developpement : Agile Scrum", 2)
    B(doc, "Pour la gestion de notre projet, nous avons adopte la methode Agile Scrum, une approche iterative et incrementale de gestion de projet logiciel. Scrum se distingue par sa flexibilite et sa capacite a s'adapter aux changements de besoins en cours de projet.")
    B(doc, "Les principes fondamentaux de Scrum sont :")
    LP(doc, "Sprint : Periode de developpement fixe (generalement 2 a 4 semaines) au terme de laquelle un increment fonctionnel du produit est livre.")
    LP(doc, "Product Backlog : Liste priorisee de toutes les fonctionnalites souhaitees pour le produit.")
    LP(doc, "Sprint Backlog : Sous-ensemble du Product Backlog selectionne pour un sprint specifique.")
    LP(doc, "Daily Standup : Reunion quotidienne de 15 minutes pour synchroniser l'equipe.")
    LP(doc, "Sprint Review : Presentation de l'increment au Product Owner a la fin de chaque sprint.")
    LP(doc, "Sprint Retrospective : Reflexion d'equipe sur les ameliorations possibles du processus.")
    B(doc, "Notre projet a ete decoupe en quatre sprints principaux :")
    LP(doc, "Sprint 1 : Authentification Sanctum et parcours eKYC (2 semaines).")
    LP(doc, "Sprint 2 : Portefeuille, transferts P2P et tableau de bord (2 semaines).")
    LP(doc, "Sprint 3 : Cartes virtuelles, chiffrement RSA et securite (2 semaines).")
    LP(doc, "Sprint 4 : Triage IA, piste d'audit, Emergency Lockdown et tests (2 semaines).")
    B(doc, "[Figure N\u00b02 : Cycle de vie Scrum - A inserer]", italic=True)

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce premier chapitre a pose le cadre de notre travail : un projet ambitieux, realise au sein de BEYN, un acteur reconnu de la fintech algerienne. Nous avons presente l'organisme d'accueil, la problematique identifiee dans le contexte reglementaire de la Loi Bancaire 23-09, les concepts theoriques fondamentaux, ainsi que la methodologie Agile Scrum adoptee. Le chapitre suivant sera consacre a l'analyse detaillee et a la specification des besoins fonctionnels et non fonctionnels de la plateforme TrustDesk.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE II : ANALYSE ET SPECIFICATION DES BESOINS
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch2_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre II : Analyse et Specification des Besoins", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce deuxieme chapitre est consacre a l'analyse detaillee des besoins de la plateforme TrustDesk. Nous identifierons les acteurs du systeme, formaliserons les exigences fonctionnelles et non fonctionnelles, et presenterons les diagrammes UML de cas d'utilisation modelisant le comportement attendu du systeme.", sb=12)

    # ── 2.1 Rappels theoriques ───────────────────────────────────────────────
    doc.add_heading("2.1  Rappels theoriques : Diagramme de cas d'utilisation", 2)
    B(doc, "Le diagramme de cas d'utilisation (Use Case Diagram) est un diagramme UML fonctionnel utilise pour representer les interactions entre les acteurs externes et le systeme. Ce diagramme est generalement utilise dans les premieres phases de l'analyse afin d'identifier les besoins fonctionnels et de definir le perimetre du systeme.")
    B(doc, "Les elements contenus dans un diagramme de cas d'utilisation sont :")
    B(doc, "1. Acteurs (Actors)", bold=True)
    B(doc, "Les acteurs representent les entites externes qui interagissent avec le systeme. Ils peuvent etre des utilisateurs humains, des systemes externes, ou des organisations. L'acteur est represente graphiquement par un bonhomme (stickman).")
    B(doc, "2. Cas d'utilisation (Use Cases)", bold=True)
    B(doc, "Les cas d'utilisation representent les fonctionnalites ou services fournis par le systeme aux acteurs. Ils sont representes par des ovales contenant le nom de la fonctionnalite.")
    B(doc, "3. Frontiere du systeme (System Boundary)", bold=True)
    B(doc, "Elle represente la limite du systeme etudie et contient tous les cas d'utilisation. Elle est representee par un rectangle englobant les ovales.")
    B(doc, "4. Relations (Relationships)", bold=True)
    LP(doc, "Association : Lien entre un acteur et un cas d'utilisation indiquant que l'acteur utilise cette fonctionnalite.")
    LP(doc, "Include (<<include>>) : Indique qu'un cas d'utilisation inclut obligatoirement un autre cas d'utilisation.")
    LP(doc, "Extend (<<extend>>) : Indique qu'un cas d'utilisation peut etre etendu par un autre dans certaines conditions.")
    LP(doc, "Generalisation : Permet de representer l'heritage entre acteurs ou entre cas d'utilisation.")

    # ── 2.2 Identification des Acteurs ───────────────────────────────────────
    doc.add_heading("2.2  Identification des Acteurs", 2)
    B(doc, "La plateforme TrustDesk met en interaction quatre categories d'acteurs :")

    doc.add_heading("Acteur 1 : Client (Utilisateur Bancaire)", 3)
    B(doc, "Le client est l'utilisateur final du portefeuille electronique. Il accede a la plateforme via l'application mobile Flutter. Ses interactions principales incluent :")
    LP(doc, "Inscription avec verification eKYC (piece d'identite + selfie + liveness).")
    LP(doc, "Connexion securisee (token Sanctum + biometrie sur mobile).")
    LP(doc, "Consultation du solde et de l'historique des transactions en temps reel.")
    LP(doc, "Transferts P2P entre utilisateurs de la plateforme.")
    LP(doc, "Generation et gestion de cartes virtuelles (gel/degel, limites).")
    LP(doc, "Paiement par QR code chez les commercants partenaires.")
    LP(doc, "Gestion des contacts de confiance (Trusted Contacts) pour la securite du compte.")

    doc.add_heading("Acteur 2 : Administrateur", 3)
    B(doc, "L'administrateur est le responsable bancaire ou l'operateur BEYN qui supervise la plateforme via le tableau de bord web React.js. Ses responsabilites incluent :")
    LP(doc, "Gestion des comptes clients (activation, desactivation, modification).")
    LP(doc, "Validation ou rejet des dossiers eKYC.")
    LP(doc, "Suivi des incidents de securite et gestion du workflow de triage.")
    LP(doc, "Generation de rapports (clients, transactions, audit, incidents).")
    LP(doc, "Consultation des KPIs en temps reel (dashboard).")
    LP(doc, "Declenchement du mode Emergency Lockdown en cas de menace critique.")

    doc.add_heading("Acteur 3 : Analyste Fraude", 3)
    B(doc, "L'analyste fraude est un role specialise au sein de l'equipe bancaire. Il dispose de privileges specifiques pour :")
    LP(doc, "Trier et classifier les alertes de fraude generees par le moteur IA.")
    LP(doc, "Geler les portefeuilles suspects de maniere preventive.")
    LP(doc, "Examiner les preuves associees a chaque alerte (FraudAlertEvidence).")
    LP(doc, "Collaborer avec l'agent de conformite pour les cas reglementaires.")

    doc.add_heading("Acteur 4 : Systeme (Moteur IA + API)", 3)
    B(doc, "Le systeme represente les composants automatises de la plateforme :")
    LP(doc, "Triage intelligent des requetes par IA (classification d'intention).")
    LP(doc, "Detection automatique d'anomalies transactionnelles.")
    LP(doc, "Gel preventif du portefeuille en cas de suspicion de fraude.")
    LP(doc, "Envoi de notifications push aux utilisateurs concernes.")
    LP(doc, "Journalisation de toutes les operations dans la piste d'audit.")
    LP(doc, "Emission de signaux communautaires (CommunitySignal) pour la detection collaborative.")

    T(doc, ["Acteur", "Type", "Interface", "Responsabilite Principale"], [
        ["Client",           "Primaire",   "Mobile (Flutter)", "Gestion du portefeuille et transactions"],
        ["Administrateur",   "Primaire",   "Web (React.js)",   "Supervision, gestion clients et securite"],
        ["Analyste Fraude",  "Primaire",   "Web (React.js)",   "Triage des alertes, gel des portefeuilles"],
        ["Systeme IA",       "Secondaire", "API (Laravel)",    "Automatisation, triage IA et audit"],
    ], "Tableau N\u00b05 : Acteurs du systeme TrustDesk")

    # ── 2.3 Diagrammes de cas d'utilisation ──────────────────────────────────
    doc.add_heading("2.3  Diagramme de cas d'utilisation global", 2)
    B(doc, "Le diagramme de cas d'utilisation global ci-dessous represente l'ensemble des fonctionnalites offertes par la plateforme TrustDesk et les interactions entre les acteurs identifies et le systeme :")
    add_image(doc, DIAGRAMS / "use_case.png", 14, "Figure N\u00b08 : Diagramme de cas d'utilisation global - TrustDesk")

    # ── 2.4 Besoins fonctionnels ─────────────────────────────────────────────
    doc.add_heading("2.4  Specification des Besoins Fonctionnels", 2)
    doc.add_heading("Module Authentification et eKYC", 3)
    LP(doc, "Inscription par email avec verification OTP.")
    LP(doc, "Connexion securisee (token Sanctum + biometrie sur mobile).")
    LP(doc, "Parcours eKYC : upload de piece d'identite, selfie avec detection de vivacite (liveness).")
    LP(doc, "Gestion des statuts KYC : en attente, verifie, rejete.")
    LP(doc, "Reinitialisation de mot de passe et deconnexion multi-appareils.")
    LP(doc, "Enregistrement et gestion des appareils de confiance (Device Trust).")

    doc.add_heading("Module Portefeuille et Transactions", 3)
    LP(doc, "Consultation du solde et de l'historique des transactions en temps reel.")
    LP(doc, "Transferts P2P entre utilisateurs de la plateforme.")
    LP(doc, "Rechargement du portefeuille (top-up) via integration bancaire.")
    LP(doc, "Paiement de factures (telephone, internet, electricite).")
    LP(doc, "Paiement par QR code chez les commercants partenaires.")
    LP(doc, "Definition de limites de depenses quotidiennes et mensuelles.")
    LP(doc, "Gel et degel du portefeuille (WalletFreeze) avec traitement conditionnel.")

    doc.add_heading("Module Cartes Virtuelles", 3)
    LP(doc, "Generation de cartes virtuelles avec chiffrement RSA-2048.")
    LP(doc, "Support de cartes jetables (disposable) a usage unique.")
    LP(doc, "Gel et degel instantane de carte depuis l'application mobile.")
    LP(doc, "Limites de depenses par carte.")

    doc.add_heading("Module Triage IA et Securite", 3)
    LP(doc, "Classificateur d'intention IA pour le routage des requetes.")
    LP(doc, "Detection d'anomalies transactionnelles avec niveaux de risque gradues (low, medium, high, critical).")
    LP(doc, "Gel preventif automatique du portefeuille en cas de suspicion.")
    LP(doc, "Gestion des alertes de fraude (FraudAlert) avec preuves associees (FraudAlertEvidence).")
    LP(doc, "Signaux communautaires (CommunitySignal) pour la detection collaborative de fraude.")
    LP(doc, "Mode Emergency Lockdown pour le blocage total du systeme en cas de menace imminente.")
    LP(doc, "Piste d'audit complete et immuable avec niveaux de risque (OWASP A09:2021).")

    doc.add_heading("Module Administration", 3)
    LP(doc, "Dashboard avec KPIs en temps reel (utilisateurs actifs, volume transactions, incidents).")
    LP(doc, "Gestion CRUD des clients avec filtres par statut KYC.")
    LP(doc, "Validation et rejet des documents eKYC.")
    LP(doc, "Gestion des incidents : creation, assignation, suivi, cloture.")
    LP(doc, "Generation de rapports (clients, transactions, audit, tickets).")
    LP(doc, "Carte des menaces (ThreatMap) pour la visualisation geographique des incidents.")

    T(doc, ["Module", "Fonctionnalites Principales"], [
        ["Authentification et eKYC",     "Inscription OTP, connexion Sanctum, parcours eKYC, liveness, Device Trust"],
        ["Portefeuille et Transactions", "Solde, transferts P2P, paiements QR, factures, WalletFreeze"],
        ["Cartes Virtuelles",            "Generation RSA-2048, cartes jetables, gel/degel, limites"],
        ["Triage IA et Securite",        "Classification IA, FraudAlert, CommunitySignal, Emergency Lockdown, audit trail"],
        ["Administration",               "Dashboard KPIs, gestion clients, validation eKYC, ThreatMap, rapports"],
    ], "Tableau N\u00b06 : Besoins fonctionnels par module")

    # ── 2.5 Besoins non fonctionnels ─────────────────────────────────────────
    doc.add_heading("2.5  Specification des Besoins Non Fonctionnels", 2)
    B(doc, "Les besoins non fonctionnels definissent les contraintes de qualite que la plateforme doit respecter :")
    T(doc, ["Exigence", "Description"], [
        ["Securite",      "Chiffrement RSA-2048/AES-256, conformite OWASP Top 10, protection mass-assignment, SecurityHeaders middleware"],
        ["Performance",   "Temps de reponse API < 300ms pour 95% des requetes"],
        ["Disponibilite", "Objectif de disponibilite de 99.9% (SLA bancaire)"],
        ["Scalabilite",   "Architecture horizontalement extensible pour multi-banque"],
        ["Ergonomie",     "Interface conforme aux standards Material Design 3"],
        ["Compatibilite", "Android 8.0+, iOS 12+, navigateurs modernes (Chrome, Firefox, Edge)"],
        ["Maintenabilite","Architecture modulaire, code documente, tests automatises"],
        ["Auditabilite",  "Journalisation complete avec niveaux de risque (low/medium/high/critical)"],
        ["Conformite",    "Respect des directives Banque d'Algerie, Loi 23-09, Reglement 24-04"],
        ["Internationalisation", "Support multilingue : francais, anglais, arabe"],
    ], "Tableau N\u00b07 : Besoins non fonctionnels")

    # ── 2.6 Descriptions textuelles ──────────────────────────────────────────
    doc.add_heading("2.6  Descriptions Textuelles des Cas d'Utilisation", 2)

    B(doc, "Description textuelle du cas d'utilisation : S'inscrire et completer le eKYC", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "S'inscrire et completer le eKYC"],
        ["Acteur principal",    "Client"],
        ["Precondition",        "Application mobile installee, connexion internet disponible"],
        ["Scenario nominal",    "1. Le client ouvre l'application et choisit 'Creer un compte'\n2. Il saisit ses informations (nom, email, mot de passe)\n3. Il recoit et valide le code OTP par email\n4. Il uploade sa piece d'identite (recto/verso)\n5. Il prend un selfie avec detection de vivacite\n6. Le systeme verifie les documents automatiquement\n7. Le compte est cree avec statut KYC 'en attente'\n8. Un administrateur valide le dossier\n9. Le statut passe a 'verifie' et le portefeuille est initialise"],
        ["Scenario alternatif", "3a. Code OTP expire -> renvoi automatique\n5a. Liveness echoue -> nouvelle tentative\n6a. Document illisible -> demande de re-upload\n8a. Document frauduleux -> rejet et notification"],
        ["Postcondition",       "Compte cree, portefeuille initialise a 0 DZD, appareil enregistre"],
    ], "Tableau N\u00b08 : Description UC - Inscription eKYC")

    B(doc, "Description textuelle du cas d'utilisation : Effectuer un transfert P2P", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "Effectuer un transfert P2P"],
        ["Acteur principal",    "Client"],
        ["Precondition",        "Client authentifie, KYC verifie, solde suffisant, destinataire existant"],
        ["Scenario nominal",    "1. Le client selectionne 'Transfert' dans le menu\n2. Il saisit l'identifiant du destinataire (email ou telephone)\n3. Le systeme verifie l'existence du destinataire\n4. Le client saisit le montant a transferer\n5. Le systeme verifie le solde et les limites\n6. Le moteur IA analyse la transaction (anomalies)\n7. Le client confirme avec authentification biometrique\n8. La transaction est executee atomiquement (debit + credit)\n9. Les deux parties recoivent une notification push\n10. La transaction est journalisee dans l'audit trail"],
        ["Scenario alternatif", "3a. Destinataire inexistant -> message d'erreur\n5a. Solde insuffisant -> message d'erreur\n6a. Transaction flaggee suspecte -> blocage + alerte admin\n7a. Biometrie echouee (3 tentatives) -> annulation et blocage"],
        ["Postcondition",       "Soldes mis a jour, transaction enregistree, notifications envoyees, audit trail mis a jour"],
    ], "Tableau N\u00b09 : Description UC - Transfert P2P")

    B(doc, "Description textuelle du cas d'utilisation : Gerer les cartes virtuelles", bold=True, sb=12)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation",   "Gerer les cartes virtuelles"],
        ["Acteur principal",    "Client"],
        ["Acteurs secondaires", "Systeme (EncryptionService)"],
        ["Precondition",        "Client authentifie, KYC verifie"],
        ["Scenario nominal",    "1. Le client accede a la section 'Cartes' de l'application\n2. Il choisit 'Generer une nouvelle carte'\n3. Il selectionne le type (permanente ou jetable)\n4. Le systeme genere une paire de cles RSA-2048\n5. Le numero de carte et le CVV sont chiffres avec la cle publique\n6. La cle privee est stockee dans le Secure Enclave du smartphone\n7. La carte est creee et affichee avec le numero masque\n8. Le client peut reveler temporairement les informations"],
        ["Scenario alternatif", "2a. Limite de cartes atteinte -> message d'erreur\n4a. Erreur de generation des cles -> nouvelle tentative\n7a. Carte jetable utilisee -> destruction automatique"],
        ["Postcondition",       "Carte virtuelle creee et securisee, prete a l'utilisation"],
    ], "Tableau N\u00b010 : Description UC - Gestion des cartes virtuelles")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce deuxieme chapitre a permis de formaliser de maniere exhaustive les besoins de la plateforme TrustDesk. Nous avons identifie quatre acteurs cles : le Client, l'Administrateur, l'Analyste Fraude, et le Systeme IA. Les cinq modules fonctionnels ont ete detailles avec leurs besoins specifiques, et les besoins non fonctionnels ont ete specifies avec des criteres mesurables. Les descriptions textuelles des principaux cas d'utilisation ont precise les scenarios nominaux et alternatifs. Le chapitre suivant presente la conception architecturale detaillee de la solution.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE III : CONCEPTION
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch3_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre III : Conception", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce chapitre presente la conception detaillee de la plateforme TrustDesk. Nous decrirons l'architecture globale du systeme, les couches de l'API Laravel, le modele de donnees (MCD et MLD), les diagrammes de classes, les diagrammes de sequence detailles, ainsi que l'architecture de chiffrement garantissant la securite des donnees bancaires.", sb=12)

    # ── 3.1 Architecture globale ─────────────────────────────────────────────
    doc.add_heading("3.1  Architecture Globale du Systeme", 2)
    B(doc, "La plateforme TrustDesk adopte une architecture trois tiers (Three-Tier Architecture) composee de trois couches distinctes communicant via des APIs RESTful securisees :")
    LP(doc, "Couche Presentation : Application mobile (Flutter 3 / Dart / Riverpod) et Dashboard web (React.js 18 / TypeScript / Vite). Ces deux clients communiquent avec la couche metier exclusivement via des appels API REST.")
    LP(doc, "Couche Metier : API RESTful Laravel 11 avec authentification Sanctum, 11 controleurs, 15 modeles Eloquent, un AuditService centralise, et deux middlewares de securite (SecurityHeaders, SetLocale). Cette couche centralise toute la logique applicative.")
    LP(doc, "Couche Donnees : Base de donnees relationnelle PostgreSQL via Supabase pour le stockage persistant, avec Supabase Realtime pour la synchronisation en temps reel et Supabase Storage pour les fichiers KYC.")
    add_image(doc, DIAGRAMS / "architecture.png", 15, "Figure N\u00b013 : Architecture globale trois tiers - TrustDesk")

    doc.add_heading("3.1.1  Architecture Detaillee de l'API Laravel", 3)
    B(doc, "L'API Laravel suit une architecture en couches strictement separees. L'analyse du code source revele la structure suivante :")
    T(doc, ["Couche", "Composant Laravel", "Role"], [
        ["Routes",          "routes/api.php",                   "Definition des endpoints REST"],
        ["Middleware",      "SecurityHeaders, SetLocale",        "En-tetes de securite HTTP, internationalisation"],
        ["Form Requests",   "app/Http/Requests",                "Validation stricte + rejet des champs inattendus"],
        ["Controllers",     "11 controleurs specialises",       "AuthController, WalletController, CardController, TransferController, FraudAlertController, FraudTriageController, WalletFreezeController, EmergencyLockdownController, CommunitySignalController, TransactionController"],
        ["Services",        "AuditService",                     "Journalisation centralisee avec niveaux de risque (OWASP A09:2021)"],
        ["Models",          "15 modeles Eloquent",              "User, Wallet, Card, Transaction, FraudAlert, FraudAlertEvidence, EmergencyLockdown, WalletFreeze, CommunitySignal, Device, TrustedContact, AuditLog, AlertAction, Notification, Playbook"],
    ], "Tableau N\u00b011 : Couches de l'architecture API Laravel")

    doc.add_heading("3.1.2  Systeme de Roles et Permissions", 3)
    B(doc, "Le systeme d'authentification repose sur Laravel Sanctum avec un modele de roles a quatre niveaux, directement implemente dans le modele User :")
    LP(doc, "admin : Controle total de la plateforme, acces a toutes les fonctionnalites d'administration et de securite.")
    LP(doc, "fraud_analyst : Triage des alertes de fraude, gel des portefeuilles suspects, examen des preuves.")
    LP(doc, "compliance_officer : Revue reglementaire, resolution des litiges, rapports de conformite.")
    LP(doc, "client : Utilisateur standard du portefeuille electronique.")
    B(doc, "La methode isAgent() du modele User permet de verifier si un utilisateur dispose de privileges d'agent (fraud_analyst ou compliance_officer), centralisant ainsi la logique d'autorisation.")

    # ── 3.2 Diagramme de classes ─────────────────────────────────────────────
    doc.add_heading("3.2  Diagramme de Classes", 2)
    B(doc, "Le diagramme de classes modelise les entites principales du systeme TrustDesk et leurs relations. L'analyse du code source revele 15 modeles Eloquent interconnectes. Les classes centrales sont :")
    add_image(doc, DIAGRAMS / "class_diagram.png", 15, "Figure N\u00b014 : Diagramme de classes - TrustDesk")

    doc.add_heading("3.2.1  Relations Entre les Classes", 3)
    LP(doc, "User (1) <-> (1) Wallet : Chaque utilisateur possede un portefeuille unique (relation 1:1).")
    LP(doc, "User (1) <-> (0..*) Card : Un utilisateur peut avoir plusieurs cartes virtuelles via son portefeuille.")
    LP(doc, "User (1) <-> (0..*) Device : Un utilisateur peut enregistrer plusieurs appareils de confiance.")
    LP(doc, "User (1) <-> (0..*) FraudAlert : Un utilisateur peut etre lie a plusieurs alertes de fraude.")
    LP(doc, "User (1) <-> (0..*) EmergencyLockdown : Un administrateur peut declencher plusieurs lockdowns.")
    LP(doc, "User (1) <-> (0..*) WalletFreeze : Un portefeuille peut subir plusieurs gels/degels traces.")
    LP(doc, "User (1) <-> (0..*) CommunitySignal : Un utilisateur peut emettre des signaux communautaires.")
    LP(doc, "User (1) <-> (0..*) TrustedContact : Un utilisateur peut definir des contacts de confiance.")
    LP(doc, "Wallet (1) <-> (0..*) Transaction : Un portefeuille est associe a plusieurs transactions.")
    LP(doc, "Wallet (1) <-> (0..*) Card : Un portefeuille peut contenir plusieurs cartes virtuelles.")
    LP(doc, "FraudAlert (1) <-> (0..*) FraudAlertEvidence : Chaque alerte peut avoir plusieurs preuves.")
    LP(doc, "User (1) <-> (0..*) AuditLog : Chaque action sensible genere un enregistrement d'audit.")

    # ── 3.3 MCD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.3  Modele Conceptuel des Donnees (MCD)", 2)
    B(doc, "Le MCD represente les entites du systeme et leurs associations dans un formalisme Entite-Association. Les cardinalites refletent les regles metier du portefeuille electronique bancaire :")
    add_image(doc, DIAGRAMS / "mcd.png", 15, "Figure N\u00b015 : Modele Conceptuel des Donnees (MCD)")

    # ── 3.4 MLD ──────────────────────────────────────────────────────────────
    doc.add_heading("3.4  Modele Logique des Donnees (MLD)", 2)
    B(doc, "La transformation du MCD en MLD produit les tables relationnelles suivantes, directement alignees avec les modeles Eloquent du code source :")

    T(doc, ["Table", "Colonnes Principales", "Relations"], [
        ["users",              "id, name, email, password, role, language, is_active", "PK: id"],
        ["wallets",            "id, user_id, balance, currency, status, spending_limit", "FK: user_id -> users"],
        ["transactions",       "id, wallet_id, type, amount, recipient_id, description, status, reference, category", "FK: wallet_id -> wallets"],
        ["cards",              "id, user_id, card_number_enc, cvv_enc, expiry, status, is_disposable, public_key", "FK: user_id -> users"],
        ["fraud_alerts",       "id, user_id, type, description, priority, status, ai_triage_result", "FK: user_id -> users"],
        ["fraud_alert_evidences", "id, fraud_alert_id, type, content", "FK: fraud_alert_id -> fraud_alerts"],
        ["emergency_lockdowns","id, user_id, reason, status, started_at, ended_at", "FK: user_id -> users"],
        ["wallet_freezes",     "id, user_id, wallet_id, reason, status, frozen_at, unfrozen_at", "FK: user_id -> users, wallet_id -> wallets"],
        ["community_signals",  "id, user_id, type, content, status", "FK: user_id -> users"],
        ["devices",            "id, user_id, device_name, fingerprint, last_seen_at", "FK: user_id -> users"],
        ["trusted_contacts",   "id, user_id, contact_user_id, relationship", "FK: user_id -> users"],
        ["audit_logs",         "id, user_id, action, entity_type, entity_id, risk_level, ip_address, user_agent, session_id, device_fingerprint, geo_country, metadata", "FK: user_id -> users"],
    ], "Tableau N\u00b012 : Tables de la base de donnees TrustDesk")

    B(doc, "Note : Les colonnes _enc indiquent des champs chiffres avec RSA-2048. Le champ risk_level de la table audit_logs utilise quatre niveaux : low, medium, high, critical, conformement aux recommandations OWASP A09:2021.", italic=True)

    # ── 3.5 Diagrammes de sequence ───────────────────────────────────────────
    doc.add_heading("3.5  Diagrammes de Sequence", 2)
    B(doc, "Les diagrammes de sequence detaillent le flux d'interactions entre les composants du systeme pour les cas d'utilisation principaux.")

    doc.add_heading("3.5.1  Sequence : Inscription et eKYC", 3)
    B(doc, "Ce diagramme montre le flux complet d'inscription d'un nouveau client, incluant la verification OTP, l'upload des documents d'identite, la detection de vivacite, l'enregistrement de l'appareil (Device), et la validation par un administrateur :")
    add_image(doc, DIAGRAMS / "sequence_ekyc.png", 14, "Figure N\u00b016 : Diagramme de sequence - Inscription et eKYC")

    doc.add_heading("3.5.2  Sequence : Transfert P2P", 3)
    B(doc, "Ce diagramme illustre le processus de transfert d'argent entre deux utilisateurs. L'operation utilise les methodes atomiques debit() et credit() du modele Wallet, avec verification du gel (isFrozen()) et analyse IA de la transaction :")
    add_image(doc, DIAGRAMS / "sequence_p2p.png", 14, "Figure N\u00b017 : Diagramme de sequence - Transfert P2P")

    # ── 3.6 Architecture de chiffrement ──────────────────────────────────────
    doc.add_heading("3.6  Architecture de Chiffrement", 2)
    B(doc, "La securite des donnees sensibles (numeros de carte, CVV) repose sur un modele de chiffrement hybride combinant chiffrement asymetrique et symetrique. Cette approche, validee par la litterature academique, offre un equilibre optimal entre securite et performance :")

    doc.add_heading("Chiffrement Asymetrique (RSA-2048)", 3)
    LP(doc, "Utilise pour le chiffrement des numeros de cartes virtuelles et CVV.")
    LP(doc, "Chaque client possede une paire de cles (publique/privee) generee sur son appareil.")
    LP(doc, "La cle publique est transmise au serveur pour le chiffrement.")
    LP(doc, "La cle privee reste exclusivement dans le Secure Enclave du smartphone (approche Zero-Knowledge cote backend).")

    doc.add_heading("Chiffrement Symetrique (AES-256)", 3)
    LP(doc, "Utilise pour le chiffrement des payloads volumineux en transit.")
    LP(doc, "La cle AES est elle-meme chiffree par RSA pour le transport (enveloppe hybride).")
    LP(doc, "AES-256 en mode CBC avec IV aleatoire pour chaque operation.")

    doc.add_heading("3.6.1  Service d'Audit Centralise", 3)
    B(doc, "Le service AuditService constitue un composant critique de l'architecture de securite. Implemente selon les recommandations OWASP A09:2021, il enregistre chaque action sensible avec :")
    LP(doc, "L'identifiant de l'utilisateur et son adresse IP.")
    LP(doc, "L'agent utilisateur (user agent) et l'identifiant de session.")
    LP(doc, "L'empreinte de l'appareil (device fingerprint) pour la detection d'anomalies.")
    LP(doc, "La geolocalisation (pays et ville) via les en-tetes Cloudflare.")
    LP(doc, "Un niveau de risque gradue : low, medium, high, critical.")
    LP(doc, "Des metadonnees contextuelles structurees en JSON.")
    B(doc, "Le service definit des constantes canoniques pour chaque type d'action (ACTION_LOGIN_SUCCESS, ACTION_WALLET_FREEZE, ACTION_PANIC, etc.), eliminant les risques de fautes de frappe dans les call-sites et facilitant l'analyse des logs.")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce troisieme chapitre a permis de detailler la conception complete de TrustDesk : architecture trois tiers avec 11 controleurs et 15 modeles Eloquent, modele de donnees alignes avec le code source reel (12 tables relationnelles), diagrammes de classes et de sequence, architecture de chiffrement hybride RSA/AES, et service d'audit centralise conforme a OWASP. Le chapitre suivant presente la realisation concrete de la solution.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE IV : REALISATION
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc, page_num=False)
    add_image(doc, ASSETS / "ch4_cover.png", 16)

    new_section(doc)
    doc.add_heading("Chapitre IV : Realisation", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce quatrieme et dernier chapitre est consacre a la mise en oeuvre concrete de la plateforme TrustDesk. Nous presenterons l'environnement de developpement, les technologies et outils utilises, les principales interfaces realisees pour chaque composant (API, dashboard web, application mobile), ainsi que les resultats des tests fonctionnels et de securite.", sb=12)

    # ── 4.1 Environnement de developpement ───────────────────────────────────
    doc.add_heading("4.1  Environnement de Developpement", 2)
    B(doc, "La realisation du projet TrustDesk a necessite la mise en place d'un environnement de developpement complet, integrant les outils pour les trois composants de la plateforme (API, web, mobile).")
    B(doc, "[Figure N\u00b018 : Environnement de developpement - VS Code - A inserer]", italic=True)

    # ── 4.2 Technologies Backend ─────────────────────────────────────────────
    doc.add_heading("4.2  Technologies Utilisees", 2)

    doc.add_heading("4.2.1  Backend (API RESTful)", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["PHP",             "8.2",    "Langage de programmation backend"],
        ["Laravel",         "11.x",   "Framework PHP (architecture MVC, ORM Eloquent, migrations)"],
        ["Laravel Sanctum", "4.x",    "Authentification par tokens API (stateless)"],
        ["PostgreSQL",      "15+",    "Systeme de gestion de base de donnees relationnelle"],
        ["Supabase",        "Cloud",  "Backend-as-a-Service (PostgreSQL, Realtime, Storage)"],
        ["Composer",        "2.x",    "Gestionnaire de dependances PHP"],
        ["Laravel Herd",    "Latest", "Serveur de developpement local pour Laravel"],
    ], "Tableau N\u00b013 : Technologies utilisees - Backend")

    doc.add_heading("4.2.2  Frontend Web (Dashboard)", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["React.js",     "18.x",   "Bibliotheque JavaScript pour les interfaces utilisateur"],
        ["TypeScript",   "5.x",    "Superset de JavaScript avec typage statique"],
        ["Vite",         "5.x",    "Outil de build rapide pour les applications web modernes"],
        ["Zustand",      "4.x",    "Gestion d'etat legere pour React (authStore)"],
        ["Recharts",     "2.x",    "Bibliotheque de graphiques pour les KPIs du dashboard"],
        ["React Router", "6.x",    "Routage SPA pour la navigation entre les pages"],
    ], "Tableau N\u00b014 : Technologies utilisees - Frontend Web")

    doc.add_heading("4.2.3  Application Mobile", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["Flutter",      "3.x",    "Framework cross-platform pour le developpement mobile"],
        ["Dart",         "3.x",    "Langage de programmation pour Flutter (null-safe)"],
        ["Riverpod",     "2.x",    "Gestion d'etat reactive et compile-time safe"],
        ["Dio",          "5.x",    "Client HTTP pour les appels API REST"],
        ["Flutter Secure Storage", "9.x", "Stockage securise (Keychain iOS, Keystore Android)"],
        ["Local Auth",   "2.x",    "Authentification biometrique (empreinte, Face ID)"],
    ], "Tableau N\u00b015 : Technologies utilisees - Application Mobile")

    # ── 4.3 Structure du projet ──────────────────────────────────────────────
    doc.add_heading("4.3  Structure du Projet", 2)
    B(doc, "Le projet TrustDesk est organise en un depot monorepo avec quatre sous-projets :")
    LP(doc, "trustdesk-api/ : API Laravel contenant 11 controleurs (AuthController, WalletController, CardController, TransferController, FraudAlertController, FraudTriageController, WalletFreezeController, EmergencyLockdownController, CommunitySignalController, TransactionController), 15 modeles Eloquent, 1 service metier (AuditService), et 2 middlewares de securite.")
    LP(doc, "trustdesk-web/ : Dashboard React.js avec 26 pages (Dashboard, Wallet, Cards, Transactions, SendMoney, Ekyc, Triage, Security, Analytics, Budget, Reports, Profile, Settings, etc.), des composants 3D interactifs, et un systeme de theming avance.")
    LP(doc, "trustdesk_mobile/ : Application Flutter avec 12 modules fonctionnels organises par features (auth, dashboard, p2p, send_money, vault, community, incidents, panic, stealth, profile, splash, ar_atm), un systeme de navigation core, et des composants partages.")

    # ── 4.4 Interfaces realisees ─────────────────────────────────────────────
    doc.add_heading("4.4  Principales Interfaces Realisees", 2)

    doc.add_heading("4.4.1  Application Mobile - Ecrans Principaux", 3)
    B(doc, "L'application mobile Flutter offre une experience utilisateur fluide et securisee. Les principaux ecrans realises sont :")
    B(doc, "Ecran de connexion :", bold=True)
    B(doc, "L'ecran de connexion propose deux modes d'authentification : la saisie classique (email + mot de passe) et l'authentification biometrique (empreinte digitale ou reconnaissance faciale). Le design suit les principes de Material Design 3 avec un splash screen anime.")
    B(doc, "[Figure N\u00b019 : Interface de connexion - Application Mobile - A inserer]", italic=True)

    B(doc, "Ecran d'accueil et portefeuille :", bold=True)
    B(doc, "L'ecran d'accueil affiche le solde du portefeuille en temps reel, les dernieres transactions, et les raccourcis vers les fonctionnalites principales (transfert, QR code, cartes). L'interface utilise des cartes avec des gradients et des micro-animations pour une experience dynamique.")

    B(doc, "Transfert P2P :", bold=True)
    B(doc, "L'ecran de transfert P2P permet de selectionner un destinataire par email ou telephone, de saisir le montant, et de confirmer la transaction par biometrie. Un indicateur en temps reel montre le statut de la transaction (en cours, completee, echouee).")
    B(doc, "[Figure N\u00b021 : Ecran de transfert P2P - Application Mobile - A inserer]", italic=True)

    B(doc, "Gestion des cartes virtuelles :", bold=True)
    B(doc, "L'ecran de gestion des cartes affiche les cartes virtuelles de l'utilisateur avec un design de carte realiste. Chaque carte peut etre gelee/degelee d'un geste, et les informations sensibles (numero, CVV) sont revelees temporairement apres authentification biometrique.")
    B(doc, "[Figure N\u00b022 : Gestion des cartes virtuelles - Application Mobile - A inserer]", italic=True)

    doc.add_heading("4.4.2  Dashboard Web - Panneau d'Administration", 3)
    B(doc, "Le dashboard web React.js offre aux administrateurs une vue complete de la plateforme :")

    B(doc, "Tableau de bord KPIs :", bold=True)
    B(doc, "Le dashboard principal affiche les indicateurs cles de performance en temps reel : nombre d'utilisateurs actifs, volume total des transactions, nombre d'incidents ouverts, et taux de verification eKYC. Les graphiques dynamiques sont construits avec Recharts et se mettent a jour automatiquement.")
    B(doc, "[Figure N\u00b020 : Ecran d'accueil - Dashboard Web - A inserer]", italic=True)

    B(doc, "Gestion des alertes de fraude :", bold=True)
    B(doc, "L'interface de triage permet aux analystes de fraude de visualiser, classifier et traiter les alertes generees par le moteur IA. Chaque alerte est accompagnee de preuves (FraudAlertEvidence) et d'un historique d'actions (AlertAction). Le workflow de triage supporte les statuts : ouvert, en investigation, escalade, resolu, faux positif.")
    B(doc, "[Figure N\u00b023 : Dashboard KPIs - Panneau d'administration - A inserer]", italic=True)

    B(doc, "Piste d'audit :", bold=True)
    B(doc, "L'interface d'audit affiche la chronologie detaillee de toutes les operations sensibles, avec des filtres par utilisateur, par type d'action, par niveau de risque, et par periode. Chaque entree d'audit inclut l'adresse IP, l'agent utilisateur, l'empreinte de l'appareil, et la geolocalisation.")
    B(doc, "[Figure N\u00b024 : Piste d'audit - Interface Web - A inserer]", italic=True)

    # ── 4.5 Implementation du chiffrement ────────────────────────────────────
    doc.add_heading("4.5  Implementation du Chiffrement Hybride", 2)
    B(doc, "L'implementation du chiffrement hybride RSA/AES suit le schema suivant :")
    LP(doc, "Cote mobile (Flutter) : L'application genere une paire de cles RSA-2048 lors de la premiere creation de carte. La cle privee est stockee dans le Secure Enclave (flutter_secure_storage), tandis que la cle publique est transmise a l'API.")
    LP(doc, "Cote serveur (Laravel) : Le CardController utilise la cle publique du client pour chiffrer le numero de carte et le CVV avant stockage en base de donnees. Les champs card_number_enc et cvv_enc ne contiennent jamais de donnees en clair.")
    LP(doc, "Dechiffrement : Lorsque le client souhaite afficher les informations de sa carte, l'application mobile utilise la cle privee locale pour dechiffrer les donnees. Le serveur n'a jamais acces a la cle privee (approche Zero-Knowledge).")

    # ── 4.6 Implementation du service d'audit ────────────────────────────────
    doc.add_heading("4.6  Implementation du Service d'Audit", 2)
    B(doc, "Le service AuditService est un composant central de la securite de TrustDesk. Conforme a OWASP A09:2021, il capture systematiquement :")
    LP(doc, "Les evenements d'authentification : connexion reussie/echouee, deconnexion, reinitialisation de mot de passe, activation/desactivation de la 2FA, revocation de session.")
    LP(doc, "Les operations sur le portefeuille : creation, gel, degel, transfert envoye, transfert echoue.")
    LP(doc, "Les operations sur les cartes : creation, gel, suppression, modification des limites.")
    LP(doc, "Les evenements de securite : ajout/revocation d'appareil, rotation de cle API, declenchement du mode panique.")
    LP(doc, "Les actions administratives : impersonation, changement de role, export de donnees.")
    B(doc, "Chaque entree d'audit est enrichie avec des informations contextuelles : adresse IP, agent utilisateur, identifiant de session, empreinte d'appareil, et geolocalisation (via les en-tetes Cloudflare CF-IPCountry et CF-IPCity). Une methode convenience critical() permet de logger les evenements de niveau critique et de declencher des alertes automatiques.")

    # ── 4.7 Tests ────────────────────────────────────────────────────────────
    doc.add_heading("4.7  Tests et Validation", 2)
    B(doc, "La validation de la plateforme TrustDesk a ete realisee a travers une serie de tests fonctionnels et de securite :")

    doc.add_heading("4.7.1  Tests Fonctionnels", 3)
    T(doc, ["Test", "Description", "Resultat"], [
        ["Inscription eKYC",    "Inscription complete avec upload de document et detection de vivacite", "Reussi"],
        ["Connexion Sanctum",   "Authentification par token API avec verification email", "Reussi"],
        ["Transfert P2P",       "Transfert entre deux utilisateurs avec verification de solde", "Reussi"],
        ["Generation carte",    "Creation de carte virtuelle avec chiffrement RSA-2048", "Reussi"],
        ["Gel portefeuille",    "Gel et degel du portefeuille avec journalisation audit", "Reussi"],
        ["Alerte fraude",       "Detection d'anomalie et creation d'alerte avec preuves", "Reussi"],
        ["Emergency Lockdown",  "Declenchement du mode urgence par un administrateur", "Reussi"],
        ["Signal communautaire","Emission et traitement d'un signal de fraude collaborative", "Reussi"],
        ["Piste d'audit",       "Verification de la journalisation complete des operations", "Reussi"],
        ["Multi-langue",        "Basculement francais/anglais/arabe sur l'API (SetLocale middleware)", "Reussi"],
    ], "Tableau N\u00b016 : Resultats des tests fonctionnels")

    doc.add_heading("4.7.2  Tests de Securite", 3)
    B(doc, "Les tests de securite ont couvert les aspects suivants :")
    LP(doc, "Protection contre le mass-assignment : Les FormRequests Laravel rejettent les champs inattendus dans les requetes API, empechant les attaques par injection de parametres.")
    LP(doc, "En-tetes de securite HTTP : Le middleware SecurityHeaders ajoute les en-tetes X-Frame-Options, X-Content-Type-Options, X-XSS-Protection, Strict-Transport-Security, et Content-Security-Policy.")
    LP(doc, "Chiffrement des donnees sensibles : Les numeros de carte et CVV sont stockes chiffres (RSA-2048) et ne sont jamais exposes en clair dans les reponses API.")
    LP(doc, "Rate limiting : Les endpoints critiques (authentification, transferts) sont proteges par un rate limiter pour prevenir les attaques par force brute.")
    LP(doc, "Journalisation des evenements : Toutes les actions sensibles sont tracees dans la piste d'audit avec l'adresse IP, l'agent utilisateur, et l'empreinte d'appareil.")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce quatrieme chapitre a presente la realisation concrete de la plateforme TrustDesk. Nous avons detaille l'environnement de developpement, les technologies utilisees pour chaque composant (Laravel, React.js, Flutter), les principales interfaces realisees, l'implementation du chiffrement hybride RSA/AES, le service d'audit centralise, et les resultats des tests fonctionnels et de securite. L'ensemble des tests realises confirme le bon fonctionnement de la plateforme et sa conformite aux exigences de securite definies dans les chapitres precedents.")

    # ══════════════════════════════════════════════════════════════════════════
    # CONCLUSION GENERALE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc)
    C(doc, "CONCLUSION GENERALE", 16, True, 20)
    B(doc, "Ce memoire a presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre de notre stage au sein de l'entreprise BEYN.", sb=12)
    B(doc, "Au terme de ce travail, nous avons atteint les objectifs suivants :")
    LP(doc, "La conception et la realisation d'une API RESTful centralisee avec Laravel 11, comprenant 11 controleurs specialises et 15 modeles Eloquent couvrant l'ensemble des fonctionnalites du systeme.")
    LP(doc, "Le developpement d'un tableau de bord d'administration web avec React.js 18 et TypeScript, offrant 26 pages de supervision, de gestion des clients et de suivi des incidents de securite.")
    LP(doc, "La creation d'une application mobile cross-platform avec Flutter 3 et Riverpod, comprenant 12 modules fonctionnels et une experience utilisateur fluide et securisee.")
    LP(doc, "L'implementation d'un parcours eKYC complet avec detection de vivacite biometrique, conforme aux exigences reglementaires de la Banque d'Algerie.")
    LP(doc, "La mise en place d'un systeme de chiffrement hybride RSA-2048/AES-256 pour la protection des cartes virtuelles, avec une approche Zero-Knowledge cote backend.")
    LP(doc, "La realisation d'un service d'audit centralise conforme a OWASP A09:2021, avec des niveaux de risque gradues et des informations contextuelles detaillees.")
    LP(doc, "L'implementation d'un systeme de detection de fraude par IA et d'un mecanisme d'Emergency Lockdown pour les situations critiques.")

    B(doc, "La plateforme TrustDesk repond a la problematique initiale en proposant une solution integree qui comble les lacunes identifiees dans les solutions existantes sur le marche algerien : un parcours eKYC entierement numerique, un chiffrement robuste des donnees sensibles, un moteur de triage IA, et une tracabilite complete des operations.")

    B(doc, "Perspectives :", bold=True)
    B(doc, "Plusieurs pistes d'amelioration et d'evolution sont envisageables pour les versions futures de TrustDesk :")
    LP(doc, "Integration avec les passerelles de paiement nationales (SATIM, SEPI) pour les paiements inter-bancaires en temps reel.")
    LP(doc, "Implementation d'un modele de machine learning plus avance pour la detection de fraude, utilisant des techniques d'apprentissage profond (deep learning) et d'analyse comportementale.")
    LP(doc, "Ajout du support des paiements NFC (Near Field Communication) pour les paiements sans contact via le smartphone.")
    LP(doc, "Deploiement d'une architecture microservices pour ameliorer la scalabilite horizontale et la resilience du systeme.")
    LP(doc, "Obtention de la certification PCI-DSS (Payment Card Industry Data Security Standard) pour la conformite aux standards internationaux de securite des cartes de paiement.")
    LP(doc, "Extension de la plateforme a d'autres pays du Maghreb et de l'Afrique du Nord, avec adaptation aux cadres reglementaires locaux.")

    # ══════════════════════════════════════════════════════════════════════════
    # BIBLIOGRAPHIE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc)
    doc.add_heading("Bibliographie", 1)

    B(doc, "Ouvrages :", bold=True, sb=12)
    LP(doc, "Stallings, W. (2020). Cryptography and Network Security: Principles and Practice. 8th Edition. Pearson.")
    LP(doc, "Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley.")
    LP(doc, "Rumbaugh, J., Jacobson, I. & Booch, G. (2004). The Unified Modeling Language Reference Manual. 2nd Edition. Addison-Wesley.")

    B(doc, "Articles et rapports :", bold=True, sb=8)
    LP(doc, "Nutalapati, P. (2018). Advanced Data Encryption Techniques for Secure Cloud Storage in Fintech Applications. Journal of Scientific and Engineering Research, 5(12), 396-405.")
    LP(doc, "Batta, V.M.B. & Kumar, L.K.S. (2023). RSA-AES Hybrid Encryption: Combining The Strengths Of Two Powerful Algorithms For Enhanced Security. International Journal of Research and Analytical Reviews (IJRAR), 10(2), 992-999.")
    LP(doc, "Bansal & Kaushik (2023). Improved AES-256 with RSA-2048 for modular cloud encryption. Performance improvement of 35-40% over RSA-only systems.")
    LP(doc, "Global Voice Group (2025). KYC in Africa 2025: Digital Identity, AI & Biometrics. Rapport sectoriel.")
    LP(doc, "MENAFATF (2023). Rapport d'Evaluation Mutuelle de l'Algerie. Evaluation de la conformite aux recommandations du GAFI.")

    B(doc, "Textes reglementaires :", bold=True, sb=8)
    LP(doc, "Loi n\u00b0 23-09 relative au droit monetaire et bancaire. Republique Algerienne Democratique et Populaire.")
    LP(doc, "Reglement n\u00b0 24-04 de la Banque d'Algerie concernant les conditions specifiques d'octroi de licences pour la creation et l'exploitation de banques numeriques.")
    LP(doc, "Instruction n\u00b0 02-2025 du 2 mars 2025 de la Banque d'Algerie relative a la mise en oeuvre du Reglement 24-04.")

    B(doc, "Documentation technique :", bold=True, sb=8)
    LP(doc, "Laravel Documentation. https://laravel.com/docs/11.x")
    LP(doc, "React.js Documentation. https://react.dev")
    LP(doc, "Flutter Documentation. https://docs.flutter.dev")
    LP(doc, "Riverpod Documentation. https://riverpod.dev")
    LP(doc, "Supabase Documentation. https://supabase.com/docs")
    LP(doc, "OWASP Top 10:2021. https://owasp.org/Top10/")
    LP(doc, "OWASP A09:2021 - Security Logging and Monitoring Failures.")

    # ── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    size = OUT.stat().st_size // 1024
    print(f"\nMemoire sauvegarde : {OUT}  ({size} KB)")

    # Stats
    secs = len(doc.sections)
    paras = len(doc.paragraphs)
    tbls = len(doc.tables)
    imgs = len(doc.inline_shapes)
    print(f"Sections: {secs}, Paragraphs: {paras}, Tables: {tbls}, Images: {imgs}")

    # Count by chapter
    ch_starts = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip(); s = p.style.name
        if s == 'Heading 1':
            if 'Chapitre I :' in t: ch_starts['ch1'] = i
            elif 'Chapitre II :' in t: ch_starts['ch2'] = i
            elif 'Chapitre III :' in t: ch_starts['ch3'] = i
            elif 'Chapitre IV :' in t: ch_starts['ch4'] = i
            elif 'CONCLUSION' in t: ch_starts['conclusion'] = i

    keys = list(ch_starts.keys())
    for idx, k in enumerate(keys):
        start = ch_starts[k]
        end = ch_starts[keys[idx+1]] - 1 if idx + 1 < len(keys) else len(doc.paragraphs) - 1
        c = sum(1 for i in range(start, end+1) if doc.paragraphs[i].text.strip())
        print(f"{k}: {c} non-empty paragraphs")

if __name__ == "__main__":
    print("Building memoire v10 (COMPLETE EDITION)...")
    build()
    print("Done.")
