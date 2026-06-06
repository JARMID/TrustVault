import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_shading_to_cell(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=50, bottom=50, start=50, end=50):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', start), ('w:right', end)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_academic_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

def init_doc():
    doc = Document()
    for s in doc.sections:
        apply_academic_margins(s)
        
    # Default Normal style
    sn = doc.styles['Normal']
    sn.font.name = 'Times New Roman'
    sn.font.size = Pt(12)
    sn.paragraph_format.line_spacing = 1.5
    sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Body Text style
    bt = doc.styles['Body Text']
    bt.font.name = 'Times New Roman'
    bt.font.size = Pt(12)
    bt.paragraph_format.line_spacing = 1.5
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    return doc

def C(doc, t, sz=12, b=False, sa=0, sb=0, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(t)
    r.font.size = Pt(sz)
    r.font.bold = b
    r.font.name = 'Times New Roman'
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def B(doc, t, b=False, it=False, sa=6, sb=0):
    p = doc.add_paragraph(style='Body Text')
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(t)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = b
    r.font.italic = it
    return p

def LP(doc, t):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    numPr.append(pPr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'}))
    numPr.append(pPr.makeelement(qn('w:numId'), {qn('w:val'): '1'}))
    pPr.append(numPr)
    
def T(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        add_shading_to_cell(cell, '003366')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1:
                add_shading_to_cell(cell, 'F0F4F8')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'

def add_page_de_garde(doc, with_logo=True):
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    
    if with_logo:
        logo_path = r'D:\TrustDesk\memoire_assets\logob.png'
        if os.path.exists(logo_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(logo_path, width=Cm(4))
            p.paragraph_format.space_after = Pt(12)
    
    C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 12, True, 6)
    C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 12, False, 6)
    C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE", 12, False, 2)
    C(doc, "EN AUDIOVISUELS", 12, False, 2)
    C(doc, "Echahid Ahmed Mehdi - Ouled Fayet -", 12, True, 24)
    
    C(doc, "Mémoire De Fin De Formation Pour L'obtention Du Diplôme", 14, True, 6)
    C(doc, "de Technicien Supérieur en Informatique", 14, True, 12)
    C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 14, True, 36)
    
    C(doc, "Thème :", 16, True, 12, color=(0x00, 0x33, 0x66))
    
    C(doc, "Conception Et Réalisation D'une", 18, True, 6, color=(0x00, 0x33, 0x66))
    C(doc, "Plateforme Web et Mobile", 18, True, 6, color=(0x00, 0x33, 0x66))
    C(doc, "de Portefeuille Électronique Sécurisé", 18, True, 6, color=(0x00, 0x33, 0x66))
    C(doc, "pour Institutions Bancaires", 18, True, 12, color=(0x00, 0x33, 0x66))
    C(doc, "« TrustDesk »", 22, True, 36, color=(0x00, 0x33, 0x66))
    
    C(doc, "Organisme d'accueil :", 12, False, 6)
    C(doc, "Entreprise BEYN", 14, True, 36, color=(0x00, 0x33, 0x66))
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    
    # Remove borders (invisible table)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    c1 = table.rows[0].cells[0]
    p1 = c1.paragraphs[0]
    p1.add_run("Réalisé Par :").bold = True
    p1.add_run("\nM. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    c2 = table.rows[0].cells[1]
    p2 = c2.paragraphs[0]
    p2.add_run("Encadré par :").bold = True
    p2.add_run("\nMme. Himeur")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    C(doc, "", sa=36)
    C(doc, "Promotion : 2025 / 2026", 14, True, 0)
    
def add_chapter_cover(doc, img_path):
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        # Full width image for chapter cover
        r.add_picture(img_path, width=Cm(16))
    else:
        C(doc, "CHAPITRE", 24, True, 24)

def build_memoire():
    doc = init_doc()
    
    # Page de garde 1 (with logo)
    add_page_de_garde(doc, with_logo=True)
    
    # Page de garde 2 (without logo)
    add_page_de_garde(doc, with_logo=False)
    
    # Remerciements
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons à exprimer notre gratitude envers toutes les personnes qui ont contribué, de près ou de loin, à la réalisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir donné la santé, la volonté et la patience pour mener à bien ce projet.")
    B(doc, "Nos remerciements les plus sincères vont à notre encadreur, Mme. Himeur, pour ses précieux conseils, sa disponibilité et son suivi rigoureux tout au long de la réalisation de ce mémoire.")
    B(doc, "Nous remercions également M. Reda Benbouzid, CEO de BEYN, et l'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur collaboration durant notre projet de fin d'études. Leur expertise dans le domaine des technologies bancaires (notamment sur les plateformes WimPay, SELA et KANTARA) a été une source d'apprentissage inestimable pour le développement de notre portefeuille électronique TrustDesk.")
    B(doc, "Enfin, nous adressons nos remerciements à tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'évaluer ce travail.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Abdelmadjid K., Khaled Wassim D.").italic = True
    
    # Introduction Generale
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("Introduction Générale", 1)
    B(doc, "L'évolution rapide des technologies financières (fintech) a profondément transformé les modes d'interaction entre les institutions bancaires et leurs clients. En Algérie, cette transformation s'accélère sous l'impulsion de la Banque d'Algérie et des réformes réglementaires visant à moderniser le système de paiement national et à promouvoir l'inclusion financière.", sb=12)
    B(doc, "Dans ce contexte, les portefeuilles électroniques (e-wallets) s'imposent comme un levier stratégique majeur. Ils permettent aux banques d'offrir à leurs clients des services de paiement dématérialisés, des transferts instantanés de pair à pair (P2P), la gestion de cartes virtuelles, ainsi que des outils d'analyse financière personnalisée — le tout accessible depuis un smartphone.")
    B(doc, "C'est dans cette perspective que s'inscrit le projet TrustDesk, réalisé au sein de l'entreprise BEYN — un acteur majeur de la fintech algérienne, partenaire technologique de plus de dix banques nationales. TrustDesk vise à concevoir et réaliser une plateforme intégrée de portefeuille électronique sécurisé, combinant une application mobile (Flutter), un tableau de bord web d'administration (React.js), et une API RESTful centralisée (Laravel).")
    B(doc, "Ainsi, la problématique centrale de ce mémoire est la suivante :")
    B(doc, "« Comment concevoir et réaliser une plateforme web et mobile de portefeuille électronique sécurisé, intégrant la vérification d'identité numérique (eKYC), le chiffrement asymétrique des données sensibles, et un système de triage intelligent par IA, destinée aux institutions bancaires algériennes ? »", b=True)
    B(doc, "Pour répondre à cette problématique, nous avons structuré notre mémoire en quatre chapitres :")
    LP(doc, "Le premier chapitre, intitulé « Étude préalable », présente l'organisme d'accueil BEYN, le contexte du marché fintech algérien, la problématique identifiée, ainsi que la méthodologie de travail adoptée.")
    LP(doc, "Le deuxième chapitre, « Analyse et spécification des besoins », est consacré à l'identification des acteurs, la formalisation des exigences, et la modélisation UML.")
    LP(doc, "Le troisième chapitre, « Conception de l'application », présente l'architecture globale, le modèle de données, et les diagrammes détaillés.")
    LP(doc, "Le quatrième chapitre, « Réalisation et implémentation », décrit les technologies utilisées, les interfaces réalisées et les résultats des tests.")

    # Chapitre 1
    add_chapter_cover(doc, r'D:\TrustDesk\memoire_assets\ch1_cover.png')
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("1. Présentation de l'organisme d'accueil", 2)
    B(doc, "Dans le cadre de notre projet de fin d'études, nous avons été accueillis au sein de l'entreprise BEYN, un acteur majeur de l'écosystème fintech algérien. Fondée en 2004 et basée à Chéraga (Alger), BEYN est spécialisée dans la conception et le développement de solutions technologiques destinées aux institutions bancaires et financières.", sb=12)
    B(doc, "Dirigée par son CEO, M. Reda Benbouzid, l'entreprise se positionne comme un partenaire technologique stratégique pour les banques, en leur fournissant des solutions digitales robustes qui s'intègrent avec leurs systèmes core banking existants. BEYN collabore avec de nombreuses banques partenaires en Algérie, notamment la BDL, la BNA, et Al Salam Bank Algeria.")
    B(doc, "L'expertise de BEYN couvre plusieurs plateformes phares :")
    LP(doc, "WimPay : Une solution leader de paiement mobile par QR code, adoptée par de nombreuses banques pour les transferts P2P et les paiements de factures.")
    LP(doc, "SELA : Plateforme de banque digitale omnicanale orientée retail banking.")
    LP(doc, "KANTARA : Solution orientée corporate banking et passerelles de paiement sécurisées.")
    
    doc.add_heading("2. Présentation du projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme intégrée de portefeuille électronique sécurisé (Secure Digital Wallet), conçue pour les institutions bancaires clientes de BEYN. La plateforme combine une API backend Laravel, un tableau de bord React.js, et une application mobile Flutter.")
    
    doc.add_heading("3. Problématique", 2)
    B(doc, "Malgré les progrès réalisés dans la digitalisation du secteur bancaire algérien, les solutions existantes présentent des lacunes significatives : absence de vérification eKYC automatisée, sécurité asymétrique insuffisante pour les cartes virtuelles, et manque d'intelligence artificielle pour la détection de fraudes en temps réel.")
    
    doc.add_heading("4. Méthodologie", 2)
    B(doc, "Nous avons adopté la méthode Agile Scrum, décomposant le développement en sprints successifs pour livrer incrémentalement l'API, le backend web et l'application mobile.")

    # Chapitre 2
    add_chapter_cover(doc, r'D:\TrustDesk\memoire_assets\ch2_cover.png')
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("1. Identification des acteurs", 2)
    B(doc, "La plateforme interagit avec trois acteurs principaux :", sb=12)
    LP(doc, "Client : Utilisateur final de l'application mobile (gestion de portefeuille, transferts, cartes).")
    LP(doc, "Administrateur : Employé bancaire ou opérateur BEYN gérant les comptes et la sécurité via le web.")
    LP(doc, "Système : Entité automatisée gérant l'IA de triage et l'audit logging.")

    doc.add_heading("2. Spécification des besoins fonctionnels", 2)
    LP(doc, "Authentification et eKYC : L'utilisateur doit pouvoir scanner ses documents et valider son identité via détection de vivacité.")
    LP(doc, "Gestion de portefeuille : Suivi du solde et transferts P2P instantanés.")
    LP(doc, "Cartes virtuelles : Création et gestion de cartes chiffrées (y compris jetables).")
    LP(doc, "Dashboard Administrateur : Suivi KPI, validation KYC et triage d'incidents.")
    
    doc.add_heading("3. Spécification des besoins non fonctionnels", 2)
    T(doc, ['Exigence', 'Description'], [
        ['Sécurité', 'Chiffrement hybride RSA/AES, conformité OWASP.'],
        ['Performance', 'Temps de réponse de l\'API inférieur à 300ms.'],
        ['Compatibilité', 'Mobile (Flutter) iOS/Android, Web (React) multisupport.']
    ])

    doc.add_heading("4. Diagramme de cas d'utilisation", 2)
    B(doc, "Le diagramme suivant illustre les interactions principales des acteurs avec TrustDesk.")
    B(doc, "[ INSERER LE DIAGRAMME DE CAS D'UTILISATION ICI ]", it=True)

    # Chapitre 3
    add_chapter_cover(doc, r'D:\TrustDesk\memoire_assets\ch3_cover.png')
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("1. Architecture globale", 2)
    B(doc, "L'architecture s'appuie sur un modèle en trois tiers (Three-Tier) : un frontend client lourd (Flutter/React), une API centralisée (Laravel 11), et une base de données relationnelle (MySQL 8).", sb=12)
    
    doc.add_heading("2. Modèle Conceptuel des Données (MCD)", 2)
    B(doc, "Les entités principales incluent User, Wallet, Transaction, VirtualCard, Incident, et AuditLog.")
    B(doc, "[ INSERER LE DIAGRAMME MCD ICI ]", it=True)
    
    doc.add_heading("3. Diagramme de classes", 2)
    B(doc, "Le diagramme de classes illustre les relations (1..N) entre un utilisateur et ses multiples cartes, transactions et documents KYC.")
    B(doc, "[ INSERER LE DIAGRAMME DE CLASSES ICI ]", it=True)
    
    doc.add_heading("4. Architecture de chiffrement", 2)
    B(doc, "Les numéros de cartes et CVV sont chiffrés asymétriquement via RSA-2048. La clé publique est stockée sur le serveur, tandis que la clé privée réside dans le Secure Enclave du smartphone du client, garantissant une protection Zero-Knowledge de la part du backend.")

    # Chapitre 4
    add_chapter_cover(doc, r'D:\TrustDesk\memoire_assets\ch4_cover.png')
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("1. Environnement et Outils", 2)
    B(doc, "Le développement a nécessité une synergie d'outils modernes :", sb=12)
    T(doc, ['Composant', 'Technologie Principale'], [
        ['Backend API', 'PHP 8.2 / Laravel 11 / Sanctum'],
        ['Base de Données', 'MySQL 8.x'],
        ['Frontend Web', 'React.js 18 / TypeScript / Tailwind / Vite'],
        ['Application Mobile', 'Flutter 3.x / Dart / Riverpod 3.x']
    ])
    
    doc.add_heading("2. Réalisation des interfaces", 2)
    B(doc, "L'application mobile a été conçue avec une esthétique « Dark Mode » aux teintes vert/émeraude, reprenant la charte graphique de BEYN. Les animations fluides (glassmorphism, micro-interactions) optimisent l'expérience utilisateur.")
    B(doc, "[ INSERER LES CAPTURES D'ECRAN DE L'APPLICATION MOBILE ICI ]", it=True)
    
    doc.add_heading("3. Tests et Validation", 2)
    B(doc, "Des tests unitaires stricts ont été implémentés dans Laravel. L'injection de dépendances et l'utilisation de FormRequests (avec strict validation) garantissent l'immunité aux failles classiques (Mass Assignment). Côté mobile, la migration vers le pattern Notifier de Riverpod a assuré une gestion d'état robuste et réactive.")

    # Conclusion Generale
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("Conclusion Générale", 1)
    B(doc, "Ce projet de fin d'études au sein de BEYN nous a permis de développer TrustDesk, une solution fintech de pointe répondant aux enjeux sécuritaires et technologiques du marché algérien.", sb=12)
    B(doc, "La combinaison de Laravel pour l'API, React.js pour l'administration et Flutter pour le mobile forme un écosystème hautement évolutif. L'intégration de la cryptographie asymétrique, couplée au triage par intelligence artificielle, confère à cette solution un niveau de sécurité digne des standards bancaires mondiaux.")
    B(doc, "Les perspectives futures incluent l'intégration directe de TrustDesk avec la passerelle WimPay et les systèmes KANTARA, ainsi que le déploiement de modèles d'apprentissage profond pour une détection de fraude en temps réel encore plus affinée.")

    # Bibliographie
    doc.add_section()
    apply_academic_margins(doc.sections[-1])
    doc.add_heading("Bibliographie", 1)
    B(doc, "1. BEYN Official. (2025). Solutions bancaires (WimPay, SELA, KANTARA). Alger, Algérie.", sb=12)
    B(doc, "2. Laravel Documentation (v11.x). Taylor Otwell et al. https://laravel.com/docs")
    B(doc, "3. Flutter Documentation (v3.22). Google. https://docs.flutter.dev/")
    B(doc, "4. React Documentation. Meta. https://react.dev/")
    B(doc, "5. OWASP Top 10 (2021). The Open Worldwide Application Security Project.")
    
    doc.save(r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v5.docx')

if __name__ == '__main__':
    build_memoire()
    print("Memoire generated successfully at D:\\TrustDesk\\Memoire_TrustDesk_FINAL_v5.docx")
