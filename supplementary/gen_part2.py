# Part 2: Add Introduction content + Chapter 1
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\_tmp_memoire.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=2); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'

# ═══ INTRODUCTION CONTENT ═══
B("Depuis toujours, les organisations cherchent a ameliorer leurs methodes de gestion et de suivi des activites critiques. Dans le domaine de la securite des campus universitaires, cette quete d\u2019amelioration se heurte a des defis specifiques lies a la diversite des acteurs, a l\u2019etendue geographique des sites, et a la nature variee des incidents pouvant survenir.")
B("Avec l\u2019essor des technologies de l\u2019information et de la communication, les etablissements d\u2019enseignement superieur disposent aujourd\u2019hui d\u2019outils puissants pour transformer leurs processus de gestion de la securite. Les frameworks de developpement modernes, les architectures API-first, et les applications mobiles offrent des possibilites de centralisation et de reactivite sans precedent.")
B("Par ailleurs, la generalisation de l\u2019Internet et le developpement des technologies mobiles ont profondement modifie les attentes des usagers en matiere de reactivite et d\u2019accessibilite des services. Un etudiant confronte a une situation d\u2019urgence doit pouvoir signaler un incident instantanement depuis son smartphone, tandis que les agents de securite doivent pouvoir recevoir et traiter ces alertes en temps reel.")
B("Dans ce contexte, la gestion des incidents de securite constitue un aspect particulierement sensible. Les processus actuels, souvent bases sur des registres papier, des appels telephoniques, ou des echanges informels, ne permettent pas d\u2019assurer une tracabilite adequate ni une coordination efficace entre les differents intervenants.")
B("C\u2019est dans cette perspective que s\u2019inscrit le projet TrustDesk, qui vise a concevoir et realiser une plateforme integree de type Security Operations Center (SOC), composee d\u2019une application web de supervision, d\u2019une application mobile de terrain, et d\u2019une API RESTful centralisee.")
B("Ainsi, la realisation de ce projet nous amene a nous poser la problematique suivante :")
B("\u00AB Comment concevoir et realiser une plateforme web et mobile centralisee permettant d\u2019assurer une gestion complete, securisee et tracable du cycle de vie des incidents de securite en milieu universitaire ? \u00BB",b=True)
B("Pour repondre a cette problematique, nous avons structure notre memoire en quatre chapitres.")
B("Le premier chapitre, intitule \u00AB Etude prealable \u00BB, presente le contexte du projet, la problematique identifiee, ainsi que la methodologie de travail adoptee.")
B("Le deuxieme chapitre, intitule \u00AB Analyse et specification des besoins \u00BB, est consacre a l\u2019identification et a la formalisation des exigences fonctionnelles et non fonctionnelles du systeme.")
B("Le troisieme chapitre, intitule \u00AB Conception de l\u2019application \u00BB, presente la conception globale du systeme a travers les diagrammes UML et l\u2019architecture logicielle retenue.")
B("Le quatrieme chapitre, intitule \u00AB Realisation et implementation \u00BB, decrit les differentes etapes du developpement, les interfaces realisees, et les resultats des tests de validation.")
B("Enfin, nous terminerons ce memoire par une conclusion generale, dans laquelle nous presenterons le bilan du travail realise et les perspectives d\u2019amelioration envisagees.")

doc.add_page_break()

# ═══ CHAPITRE I TITLE PAGE ═══
for _ in range(8): E()
C('CHAPITRE I:',14,True,4)
C('Etude Prealable',14,False,8)
for _ in range(8): E()
doc.add_page_break()

# ═══ CHAPITRE 1 CONTENT ═══
H1("Presentation du projet :")
B("TrustDesk est une plateforme integree de type Security Operations Center (SOC), specialement concue pour la gestion des incidents de securite en milieu universitaire. Elle se compose de trois modules complementaires : une API RESTful backend, une application web de supervision, et une application mobile de terrain.",b=True)

H2("Fiche Technique du projet :")
T(['Designation','Description'],[
    ['Nom du projet','TrustDesk'],
    ['Nature','Plateforme web et mobile de gestion des incidents'],
    ['Domaine',"Securite des campus universitaires"],
    ['Architecture','API REST + SPA Web + Application Mobile'],
    ['Backend','Laravel 11 (PHP 8.2)'],
    ['Frontend Web','React.js 18 + TypeScript'],
    ['Application Mobile','Flutter 3 (Dart)'],
    ['Base de donnees','MySQL 8.x'],
])
C("Tableau N\u00b01 : Fiche technique du projet TrustDesk",12,True,8)

H2("Problematique :")
B("Malgre les avancees technologiques dans le domaine de la securite informationnelle, la gestion des incidents de securite dans les campus universitaires algeriens souffre de nombreuses lacunes que nous avons identifiees a travers notre etude preliminaire.")
B("Les principales insuffisances constatees sont les suivantes :")
LP("Absence de centralisation : Les signalements d\u2019incidents transitent par des canaux heterogenes (appels telephoniques, deplacements physiques, e-mails), sans systeme unifie de reception et de suivi.")
LP("Temps de reponse eleve : L\u2019absence d\u2019outils d\u2019alerte instantanee retarde significativement l\u2019intervention des services competents lors de situations d\u2019urgence.")
LP("Manque de tracabilite : Les incidents ne sont pas documentes numeriquement, rendant impossible toute analyse statistique.")
LP("Cloisonnement des acteurs : Les etudiants, le personnel de securite et les administrateurs operent en silos.")
LP("Absence de mobilite : Les systemes existants ne sont pas accessibles depuis un terminal mobile.")
B("Ces constats nous amenent a formuler la question centrale suivante :")
B("\u00AB Comment concevoir et realiser une plateforme web et mobile integree, securisee et performante, permettant la gestion complete du cycle de vie des incidents de securite dans un environnement universitaire ? \u00BB",b=True)

H2("Objectifs du projet :")
B("Objectifs generaux :",b=True)
B("Nous devons concevoir et developper une plateforme (web et mobile) pour la gestion centralisee des incidents de securite en milieu universitaire.")
LP("Mettre en oeuvre un systeme de gestion centralisee des incidents (signalement, triage, assignation, suivi, cloture).")
LP("Implementer un mecanisme de panique instantane avec geolocalisation depuis l\u2019application mobile.")
LP("Mettre en place un radar communautaire permettant aux usagers de signaler des activites suspectes.")
LP("Fournir un tableau de bord analytique avec indicateurs cles de performance (KPIs).")
LP("Implementer un systeme d\u2019authentification securise base sur les roles.")
LP("Assurer les notifications en temps reel pour alerter les acteurs concernes.")

B("Objectifs Espace Administrateur :",b=True)
LP("Superviser l\u2019ensemble des incidents declares sur la plateforme.")
LP("Gerer les utilisateurs, les roles et les permissions.")
LP("Consulter les statistiques globales et les indicateurs de performance.")
LP("Configurer les parametres du systeme et les workflows de triage.")

B("Objectifs Espace Agent de terrain (Mobile) :",b=True)
LP("Recevoir les alertes d\u2019incidents en temps reel avec geolocalisation.")
LP("Declencher un bouton de panique en situation d\u2019urgence.")
LP("Consulter et mettre a jour le statut des incidents assignes.")
LP("Contribuer au radar communautaire via des signalements.")

B("Les Contraintes :",b=True)
LP("L\u2019application doit etre entierement responsive.")
LP("Mise en place d\u2019un systeme d\u2019authentification securise (Laravel Sanctum).")
LP("Compatibilite avec les navigateurs modernes (Chrome, Firefox, Edge, Safari).")
LP("Fonctionnement sur Android 8.0+ et iOS 12+.")
LP("Temps de reponse API inferieur a 500ms.")

H3("Le public vise :")
B("Les administrateurs universitaires, les agents de securite, le personnel enseignant et administratif, ainsi que les etudiants de l\u2019etablissement.")

H2("Ressources humaines et materiels :")
H3("Ressources humaines :")
T(['Nom','Role'],[
    ['Khalef Abdelmadjid','Developpeur Full-Stack'],
    ['Dai Khaled Wassim','Developpeur Full-Stack'],
    ['Mme.Himeur','Encadreur / Suivi'],
])
C("Tableau N\u00b02 : Ressources humaines",12,True,8)

H3("Ressources materielles :")
T(['Materiel','Specifications'],[
    ['PC Portable 1','[Specifications]'],
    ['PC Portable 2','[Specifications]'],
    ['Smartphone Android','Tests application mobile'],
    ['Connexion Internet','Developpement et deploiement'],
])
C("Tableau N\u00b03 : Les ressources Materielles",12,True,8)

H2("Methodologie de travail :")
B("Pour la realisation de ce projet, nous avons adopte une approche de developpement agile iterative, combinant les principes de la methode Scrum avec une architecture logicielle de type client-serveur a trois tiers.")
H3("Planning des sprints :")
T(['Sprint','Module','Duree'],[
    ['Sprint 1','Architecture API & Modele de donnees','2 semaines'],
    ['Sprint 2','Authentification & Gestion des utilisateurs','1 semaine'],
    ['Sprint 3','Module Incidents & Triage','2 semaines'],
    ['Sprint 4','Module Panic & Geolocalisation','1 semaine'],
    ['Sprint 5','Module Communautaire & Signaux','2 semaines'],
    ['Sprint 6','Application Web - Dashboard & Pages','3 semaines'],
    ['Sprint 7','Application Mobile','3 semaines'],
    ['Sprint 8','Tests, corrections & deploiement','2 semaines'],
])
C("Tableau N\u00b04 : Planning des sprints",12,True,8)

H2("Concepts theoriques :")
H3("Les definitions generales :")
B("Application web :",b=True)
B("Une application web est un logiciel qui s\u2019execute dans un navigateur web. Contrairement aux logiciels de bureau traditionnels, les applications web ne necessitent pas d\u2019installation locale et sont accessibles depuis n\u2019importe quel appareil disposant d\u2019une connexion Internet.")
B("Application Mobile :",b=True)
B("Une application mobile est un logiciel concu pour fonctionner sur des appareils mobiles. Elle peut etre native ou cross-platform (fonctionnant sur plusieurs systemes a partir d\u2019un code source unique, comme Flutter).")
B("API REST :",b=True)
B("Une API RESTful est une interface de programmation qui respecte les contraintes architecturales du style REST. Elle permet a differentes applications de communiquer via HTTP en utilisant des methodes standard (GET, POST, PUT, DELETE).")
B("SOC (Security Operations Center) :",b=True)
B("Un Centre d\u2019Operations de Securite est une structure centralisee dediee a la surveillance, la detection, l\u2019analyse et la reponse aux incidents de securite.")

H2("Outils et technologies utilises :")
T(['Composante','Technologie','Version','Role'],[
    ['Backend','Laravel (PHP)','11.x','Framework API REST'],
    ['Auth','Laravel Sanctum','4.x','Tokens API securises'],
    ['BDD','MySQL','8.x','Stockage relationnel'],
    ['Frontend Web','React.js','18.x','Interface SPA'],
    ['Typage','TypeScript','5.x','Typage statique'],
    ['Bundler','Vite','5.x','Build & dev server'],
    ['Routing Web','React Router','6.x','Navigation SPA'],
    ['Mobile','Flutter','3.x','App cross-platform'],
    ['Langage Mobile','Dart','3.x','Langage Flutter'],
    ['Etat Mobile','Riverpod','2.x','Gestion d\u2019etat'],
    ['Navigation','GoRouter','14.x','Nav declarative'],
    ['HTTP Client','Dio','5.x','Requetes HTTP'],
    ['GPS','Geolocator','13.x','Geolocalisation'],
    ['IDE','VS Code','Latest','Editeur de code'],
    ['VCS','Git','Latest','Controle de version'],
])
C("Tableau N\u00b05 : Technologies et outils utilises",12,True,8)

H2("Conclusion :")
B("Ce premier chapitre nous a permis de situer le contexte general du projet TrustDesk, d\u2019identifier clairement la problematique liee a la gestion des incidents de securite en milieu universitaire, et de definir les objectifs fonctionnels et techniques que cette plateforme vise a atteindre.")
B("La solution proposee se distingue par son approche multi-plateforme integree (web + mobile + API), son systeme de reponse d\u2019urgence instantanee, et son radar communautaire collaboratif.")
B("Le chapitre suivant sera consacre a l\u2019analyse et la specification des besoins, ou nous formaliserons les exigences fonctionnelles et non fonctionnelles du systeme a travers les outils de modelisation UML.")

# SAVE FINAL
out = r'D:\TrustDesk\Memoire_TrustDesk_v3.docx'
doc.save(out)
print(f'[OK] Final saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
# Cleanup
os.remove(r'D:\TrustDesk\_tmp_memoire.docx')
