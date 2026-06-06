"""
Final cleanup: 
1. Remove organigramme image (P162) + its caption/heading
2. Remove 'Structure d'accueil' image (P168) 
3. Update company to BEYN as organisme d'accueil
4. Rewrite Ch1 section 1 as 'Presentation de l'organisme d'accueil' about BEYN
5. Clean all template-leftover content
"""
from docx import Document
from docx.oxml.ns import qn
import os, copy

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx')
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# =====================================================================
# 1. Remove organigramme image from P162
# =====================================================================
p162 = doc.paragraphs[162]
drawings = p162._p.findall('.//w:drawing', ns)
for drw in drawings:
    drw.getparent().remove(drw)
print(f"Removed {len(drawings)} drawing(s) from P162 (organigramme)")

# 2. Remove structure d'accueil image from P168
p168 = doc.paragraphs[168]
drawings2 = p168._p.findall('.//w:drawing', ns)
for drw in drawings2:
    drw.getparent().remove(drw)
print(f"Removed {len(drawings2)} drawing(s) from P168 (structure)")

# Also check P151 for any leftover image
p151 = doc.paragraphs[151]
drawings3 = p151._p.findall('.//w:drawing', ns)
for drw in drawings3:
    drw.getparent().remove(drw)
print(f"Removed {len(drawings3)} drawing(s) from P151")

# =====================================================================
# 3. Rewrite Chapter 1 Section 1: Presentation de l'organisme d'accueil
# =====================================================================
def set_para_text(p, text):
    """Set paragraph text preserving formatting of first run"""
    for r in p.runs:
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    elif text:
        run = p.add_run(text)

# P154: Change heading from "Presentation du projet" to "Presentation de l'organisme d'accueil"
set_para_text(doc.paragraphs[154], "Pr\u00e9sentation de l\u2019organisme d\u2019accueil :")

# P155: Rewrite about BEYN
set_para_text(doc.paragraphs[155],
    "BEYN est une entreprise alg\u00e9rienne sp\u00e9cialis\u00e9e dans le domaine de la communication "
    "et du d\u00e9veloppement num\u00e9rique. Fond\u00e9e avec la vocation d\u2019accompagner les entreprises "
    "et les institutions dans leur transformation digitale, BEYN offre un \u00e9ventail de services "
    "couvrant la communication institutionnelle, le d\u00e9veloppement web et mobile, le design graphique, "
    "le marketing digital et le conseil strat\u00e9gique. L\u2019entreprise s\u2019appuie sur une \u00e9quipe de "
    "d\u00e9veloppeurs, de designers et de consultants pour concevoir des solutions num\u00e9riques modernes "
    "et adapt\u00e9es aux besoins sp\u00e9cifiques de ses clients."
)

# P156: Fiche technique heading -> about BEYN
set_para_text(doc.paragraphs[156], "Fiche Technique de BEYN :")

# P159: Table caption
set_para_text(doc.paragraphs[159], "Tableau N\u00b01 : Fiche Technique de BEYN")

# P160: Organigramme heading -> placeholder
set_para_text(doc.paragraphs[160], "Organigramme de BEYN :")

# P162: was image, now add placeholder text
set_para_text(doc.paragraphs[162], "[Ins\u00e9rer ici l\u2019organigramme de BEYN]")

# P164: Figure caption
set_para_text(doc.paragraphs[164], "Figure N\u00b01 : Organigramme de BEYN")

# P167: Structure d'accueil
set_para_text(doc.paragraphs[167], "Structure d\u2019accueil :")

# P168: was image, now placeholder
set_para_text(doc.paragraphs[168], "[Ins\u00e9rer ici la structure d\u2019accueil de BEYN]")

# =====================================================================
# 4. Update Table 0 (Fiche Technique) for BEYN
# =====================================================================
t0 = doc.tables[0]

fiche_data = [
    # Row 0: Nom
    (0, 1, "BEYN"),
    # Row 1: Date de creation
    (1, 1, "[Ann\u00e9e de cr\u00e9ation de BEYN]"),
    # Row 2: Statut juridique
    (2, 1, "[Statut juridique de BEYN]"),
    # Row 3: Adresse
    (3, 1, "[Adresse de BEYN]"),
    # Row 4: Secteur d'activite
    (4, 1, "Communication, D\u00e9veloppement Web & Mobile, Design Graphique, Marketing Digital"),
    # Row 5: Implantation
    (5, 1, "[Implantation de BEYN]"),
    # Row 6: Effectifs
    (6, 1, "[Nombre d\u2019employ\u00e9s]"),
    # Row 7: Telephone
    (7, 1, "[T\u00e9l\u00e9phone de BEYN]"),
    # Row 8: Email
    (8, 1, "[Email de BEYN]"),
]

for ri, ci, val in fiche_data:
    cell = t0.rows[ri].cells[ci]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = val
            break

# Also update the left column labels to be correct
fiche_labels = [
    (0, 0, "Nom de l\u2019entreprise :"),
]
# Keep labels as-is, they're already correct from template

# =====================================================================
# 5. Update Table 1 (Ressources humaines) - for the project team at BEYN
# =====================================================================
t1 = doc.tables[1]
rh = [
    ('Service', 'Postes', 'Nombre'),
    ('\u00c9quipe Projet', 'D\u00e9veloppeur Full-Stack (Backend API)', '1'),
    ('\u00c9quipe Projet', 'D\u00e9veloppeur Full-Stack (Frontend & Mobile)', '1'),
    ('Encadrement BEYN', 'Tuteur de stage', '1'),
    ('Encadrement INSFP', 'Enseignante encadreur (Mme. Himeur)', '1'),
]
for ri, (s, poste, n) in enumerate(rh):
    if ri < len(t1.rows):
        row = t1.rows[ri]
        for ci, val in enumerate([s, poste, n]):
            if ci < len(row.cells):
                cell = row.cells[ci]
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ''
                    if p.runs:
                        p.runs[0].text = val
                        break
# Clear extra rows (5 and 6) if they exist
for ri in range(5, len(t1.rows)):
    row = t1.rows[ri]
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''

# =====================================================================
# 6. Fix page de garde - update "Organisme d'accueil" to BEYN
# =====================================================================
# Check paragraphs 0-35 for organisme d'accueil references
for i in range(0, 35):
    p = doc.paragraphs[i]
    if p.text.strip():
        # Update title to remove "entreprise" -> just show BEYN
        pass

# P16 currently says "TrustDesk" - this is the project name, keep it
# P15 says "Au Profit De L'entreprise" - change to "Au Profit De L'entreprise BEYN"
for i in range(10, 25):
    p = doc.paragraphs[i]
    for r in p.runs:
        if r.text and "Au Profit" in r.text:
            r.text = "Au Profit De BEYN"
        if r.text and "Organisme d" in r.text:
            pass  # Keep but we'll update related text

# =====================================================================
# 7. Rewrite the "Presentation du sujet" section (P241+) for correctness
# =====================================================================
# This is about the TrustDesk project itself, done at BEYN
set_para_text(doc.paragraphs[241], "Pr\u00e9sentation du sujet :")

set_para_text(doc.paragraphs[242],
    "Dans le cadre de notre stage de fin de formation au sein de l\u2019entreprise BEYN, "
    "nous avons \u00e9t\u00e9 charg\u00e9s de concevoir et r\u00e9aliser la plateforme TrustDesk : un syst\u00e8me "
    "int\u00e9gr\u00e9 de gestion des incidents de s\u00e9curit\u00e9 en milieu universitaire. Cette plateforme "
    "se compose de trois modules compl\u00e9mentaires : une API RESTful backend d\u00e9velopp\u00e9e avec "
    "Laravel 11, une application web de supervision construite avec React.js 18 et TypeScript, "
    "et une application mobile de terrain r\u00e9alis\u00e9e avec Flutter 3."
)

set_para_text(doc.paragraphs[243],
    "L\u2019application web constitue le c\u0153ur du syst\u00e8me de supervision et permet aux "
    "administrateurs de g\u00e9rer l\u2019ensemble des incidents signal\u00e9s, de visualiser les statistiques "
    "en temps r\u00e9el via un tableau de bord interactif, de g\u00e9rer les utilisateurs et leurs r\u00f4les, "
    "et de configurer les param\u00e8tres du syst\u00e8me. Le dashboard int\u00e8gre des indicateurs cl\u00e9s de "
    "performance (KPIs) tels que le nombre total d\u2019incidents, le temps moyen de r\u00e9solution, "
    "le taux de r\u00e9solution et la r\u00e9partition par cat\u00e9gorie et par priorit\u00e9."
)

set_para_text(doc.paragraphs[244],
    "En compl\u00e9ment, l\u2019application mobile permet aux agents de s\u00e9curit\u00e9 et aux usagers du campus "
    "d\u2019acc\u00e9der aux fonctionnalit\u00e9s essentielles directement depuis leur terminal mobile. Elle offre "
    "notamment un bouton panique permettant de d\u00e9clencher une alerte d\u2019urgence en un seul geste avec "
    "capture automatique des coordonn\u00e9es GPS, un radar communautaire pour signaler les activit\u00e9s "
    "suspectes avec g\u00e9olocalisation, et un syst\u00e8me de notifications push en temps r\u00e9el."
)

set_para_text(doc.paragraphs[245],
    "Gr\u00e2ce \u00e0 cette solution centralis\u00e9e et multi-plateforme, les \u00e9tablissements universitaires "
    "disposent d\u2019un outil performant et s\u00e9curis\u00e9 pour la gestion compl\u00e8te du cycle de vie des "
    "incidents de s\u00e9curit\u00e9, depuis le signalement initial jusqu\u2019\u00e0 la cl\u00f4ture, en passant par "
    "le triage, l\u2019assignation, le suivi et la r\u00e9solution."
)

# P248: Problematique - mention BEYN context
set_para_text(doc.paragraphs[248],
    "Durant notre stage au sein de l\u2019entreprise BEYN, nous avons identifi\u00e9 plusieurs insuffisances "
    "dans les syst\u00e8mes de gestion de la s\u00e9curit\u00e9 utilis\u00e9s dans les campus universitaires alg\u00e9riens : "
    "l\u2019absence de centralisation des signalements qui transitent par des canaux h\u00e9t\u00e9rog\u00e8nes "
    "(appels t\u00e9l\u00e9phoniques, d\u00e9placements physiques, e-mails), le temps de r\u00e9ponse \u00e9lev\u00e9 en "
    "l\u2019absence d\u2019outils d\u2019alerte instantan\u00e9e, le manque de tra\u00e7abilit\u00e9 des incidents non "
    "document\u00e9s num\u00e9riquement, et le cloisonnement entre les diff\u00e9rents acteurs (\u00e9tudiants, "
    "personnel de s\u00e9curit\u00e9, administration)."
)

# =====================================================================
# 8. Fix Introduction to mention BEYN
# =====================================================================
set_para_text(doc.paragraphs[122],
    "C\u2019est dans cette perspective que s\u2019inscrit le projet TrustDesk, r\u00e9alis\u00e9 au sein de "
    "l\u2019entreprise BEYN, qui vise \u00e0 concevoir et r\u00e9aliser une plateforme int\u00e9gr\u00e9e de type "
    "Security Operations Center (SOC), compos\u00e9e d\u2019une application web de supervision, "
    "d\u2019une application mobile de terrain et d\u2019une API RESTful centralis\u00e9e, d\u00e9di\u00e9e \u00e0 la "
    "gestion compl\u00e8te du cycle de vie des incidents de s\u00e9curit\u00e9 en milieu universitaire."
)

# =====================================================================
# 9. Final sweep: remove any remaining old references
# =====================================================================
old_refs = {
    'TrustDesk, fond\u00e9e en': 'BEYN, fond\u00e9e en',
    'la soci\u00e9t\u00e9 TrustDesk': 'l\u2019entreprise BEYN',
    'la SARL TrustDesk': 'l\u2019entreprise BEYN',
    'SARL TrustDesk': 'BEYN',
}

for p in doc.paragraphs:
    for r in p.runs:
        if r.text:
            for old, new in old_refs.items():
                if old in r.text:
                    r.text = r.text.replace(old, new)

# Save
out = r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'D:\TrustDesk\Memoire_TrustDesk_FINAL.docx'
    doc.save(out)

print(f'\n[OK] Saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
