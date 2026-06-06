# -*- coding: utf-8 -*-
"""TrustDesk Memoire Generator v2 - Matches template formatting exactly"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page Setup matching template: left=2cm, right=1.25cm, top=2.36cm, bottom=0.49cm
for section in doc.sections:
    section.top_margin = Cm(2.36)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.25)

# ── Normal style: Times New Roman 12pt
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ── Heading 1: 16pt, color 1F497D (matching template)
h1 = doc.styles['Heading 1']
h1.font.size = Pt(16)
h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
h1.font.name = 'Times New Roman'

# ── Heading 2: 14pt
h2 = doc.styles['Heading 2']
h2.font.size = Pt(14)
h2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
h2.font.name = 'Times New Roman'

# ── Heading 3: 12pt bold
h3 = doc.styles['Heading 3']
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.name = 'Times New Roman'
h3.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── Body Text: 12pt
bt = doc.styles['Body Text']
bt.font.name = 'Times New Roman'
bt.font.size = Pt(12)
bt.paragraph_format.line_spacing = 1.5

# ── Helpers ──
def centered(text, size=12, bold=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = 'Times New Roman'
    return p

def body(text, bold=False, italic=False, sa=6):
    p = doc.add_paragraph(style='Body Text')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = bold
    r.font.italic = italic
    return p

def heading1(text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_heading(text, level=1)
    p.alignment = align
    for r in p.runs:
        r.font.name = 'Times New Roman'
    return p

def heading2(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_heading(text, level=2)
    p.alignment = align
    for r in p.runs:
        r.font.name = 'Times New Roman'
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name = 'Times New Roman'
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    # Add bullet
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
    numId = numPr.makeelement(qn('w:numId'), {qn('w:val'): '1'})
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p

def empty():
    doc.add_paragraph(style='Body Text')

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'
    return t

def fig_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.name = 'Times New Roman'

def tbl_caption(text):
    fig_caption(text)

# ═══════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════
centered('LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE\nMINISTERE DE L\u2019ENSEIGNEMENT SUPERIEUR ET DE LA RECHERCHE SCIENTIFIQUE', 15, True, 6)
empty()
centered('[Nom de votre Universite / Institut]', 10, True, 4)
empty()
empty()
centered("Memoire De Fin D\u2019Etudes Pour L\u2019obtention Du Diplome\nde Master en Informatique", 14, True, 4)
centered('Option : DEVELOPPEMENT WEB ET MOBILE', 14, True, 8)
empty()
centered('Theme :', 16, True, 4)
empty()

centered("Conception Et Realisation D\u2019une", 18, True, 2)
centered("Plateforme Web et Mobile", 18, True, 2)
centered("de Gestion des Incidents de Securite", 18, True, 2)
centered("en Milieu Universitaire", 18, True, 2)
centered("\u00AB TrustDesk \u00BB", 18, True, 8)

empty()
centered("Organisme d\u2019accueil :", 14, True, 4)
empty()
centered("[Nom de l\u2019organisme d\u2019accueil]", 11, False, 8)
empty()
empty()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run("Realise Par :")
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run("[Votre Nom et Prenom]")
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

empty()

p = doc.add_paragraph()
r = p.add_run("Suivi par :")
r.font.size = Pt(12); r.font.name = 'Times New Roman'

p = doc.add_paragraph()
r = p.add_run("[Dr./Pr. Nom de l\u2019encadreur]")
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

empty()
empty()
empty()
centered('Promotion : 2025 / 2026', 12, True, 0)

doc.add_page_break()

# ═══════════════════════════════════════════
# DEDICACES
# ═══════════════════════════════════════════
heading1('Dedicaces')
empty()
body("Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce projet de fin d\u2019etudes. Nous remercions sincerement notre encadreur [Nom de l\u2019encadreur] pour ses conseils avises, son suivi regulier et sa disponibilite tout au long de ce travail.")
body("Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants tout au long de notre parcours academique.")
body("Enfin, nous souhaitons remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont rendu possible l\u2019aboutissement de ce projet.")
body("A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
empty()
empty()
p = doc.add_paragraph()
r = p.add_run("[Votre Nom]")
r.font.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_page_break()

# ═══════════════════════════════════════════
# SOMMAIRE (placeholder)
# ═══════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("Sommaire")
r.font.size = Pt(16)
r.font.bold = True
r.font.name = 'Times New Roman'
empty()
body("[Le sommaire sera genere automatiquement dans Word : References > Table des matieres]", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════
# INTRODUCTION GENERALE
# ═══════════════════════════════════════════
centered('INTRODUCTION', 14, True, 8)
empty()

body("Depuis toujours, les organisations cherchent a ameliorer leurs methodes de gestion et de suivi des activites critiques. Dans le domaine de la securite des campus universitaires, cette quete d\u2019amelioration se heurte a des defis specifiques lies a la diversite des acteurs, a l\u2019etendue geographique des sites, et a la nature variee des incidents pouvant survenir.")

body("Avec l\u2019essor des technologies de l\u2019information et de la communication, les etablissements d\u2019enseignement superieur disposent aujourd\u2019hui d\u2019outils puissants pour transformer leurs processus de gestion de la securite. Les frameworks de developpement modernes, les architectures API-first, et les applications mobiles offrent des possibilites de centralisation et de reactivite sans precedent.")

body("Par ailleurs, la generalisation de l\u2019Internet et le developpement des technologies mobiles ont profondement modifie les attentes des usagers en matiere de reactivite et d\u2019accessibilite des services. Un etudiant confronte a une situation d\u2019urgence doit pouvoir signaler un incident instantanement depuis son smartphone, tandis que les agents de securite doivent pouvoir recevoir et traiter ces alertes en temps reel.")

body("Dans ce contexte, la gestion des incidents de securite constitue un aspect particulierement sensible. Les processus actuels, souvent bases sur des registres papier, des appels telephoniques, ou des echanges informels, ne permettent pas d\u2019assurer une tracabilite adequate ni une coordination efficace entre les differents intervenants.")

body("C\u2019est dans cette perspective que s\u2019inscrit le projet TrustDesk, qui vise a concevoir et realiser une plateforme integree de type Security Operations Center (SOC), composee d\u2019une application web de supervision, d\u2019une application mobile de terrain, et d\u2019une API RESTful centralisee.")

body("Ainsi, la realisation de ce projet nous amene a nous poser la problematique suivante :")

body("\u00AB Comment concevoir et realiser une plateforme web et mobile centralisee permettant d\u2019assurer une gestion complete, securisee et tracable du cycle de vie des incidents de securite en milieu universitaire ? \u00BB", bold=True)

body("Pour repondre a cette problematique, nous avons structure notre memoire en quatre chapitres.")
body("Le premier chapitre, intitule \u00AB Etude prealable \u00BB, presente le contexte du projet, l\u2019organisme d\u2019accueil, la problematique identifiee, ainsi que la methodologie de travail adoptee.")
body("Le deuxieme chapitre, intitule \u00AB Analyse et specification des besoins \u00BB, est consacre a l\u2019identification et a la formalisation des exigences fonctionnelles et non fonctionnelles du systeme.")
body("Le troisieme chapitre, intitule \u00AB Conception de l\u2019application \u00BB, presente la conception globale du systeme a travers les diagrammes UML et l\u2019architecture logicielle retenue.")
body("Le quatrieme chapitre, intitule \u00AB Realisation et implementation \u00BB, decrit les differentes etapes du developpement, les interfaces realisees, et les resultats des tests de validation.")
body("Enfin, nous terminerons ce memoire par une conclusion generale.")

doc.add_page_break()

# ═══════════════════════════════════════════
# CHAPITRE 1 - ETUDE PREALABLE
# ═══════════════════════════════════════════
heading1('Chapitre 1 : Etude Prealable')
empty()

# 1. Introduction
heading2("Introduction :")
body("Ce premier chapitre constitue le point de depart de notre etude. Il a pour objectif de presenter le cadre general dans lequel s\u2019inscrit notre projet, de decrire l\u2019organisme d\u2019accueil, d\u2019exposer la problematique identifiee, et de definir les objectifs a atteindre. Nous y presenterons egalement la methodologie de travail adoptee ainsi que les outils et technologies retenus pour la realisation de la plateforme TrustDesk.")

# 2. Presentation de l'organisme
heading2("Presentation de l\u2019organisme d\u2019accueil :")
body("[Nom de l\u2019organisme], fonde en [annee], est un etablissement [nature] situe dans la wilaya de [wilaya]. Il accueille plus de [nombre] etudiants repartis sur [nombre] facultes et [nombre] campus.", italic=True)
body("Le departement d\u2019Informatique, rattache a la Faculte des Sciences Exactes et de l\u2019Informatique, forme des ingenieurs et des chercheurs dans les domaines du genie logiciel, des systemes d\u2019information, de l\u2019intelligence artificielle et des reseaux.")

heading3("Fiche Technique :")
table(
    ['Designation', 'Description'],
    [
        ['Raison sociale', '[Nom de l\u2019organisme]'],
        ['Date de creation', '[Annee]'],
        ['Secteur d\u2019activite', 'Enseignement superieur et recherche'],
        ['Localisation', '[Wilaya, Algerie]'],
        ['Effectif', '[Nombre] etudiants, [Nombre] enseignants'],
        ['Site web', '[URL]'],
    ]
)
tbl_caption("Tableau N\u00b01 : Fiche technique de l\u2019organisme")

heading3("Organigramme :")
body("[Inserer ici l\u2019organigramme de l\u2019organisme d\u2019accueil]", italic=True)
fig_caption("Figure N\u00b01 : Organigramme de l\u2019organisme")

heading3("Structure d\u2019accueil :")
body("Notre stage s\u2019est deroule au sein du departement d\u2019informatique, dont la mission principale est d\u2019assurer la formation des etudiants dans les specialites liees aux technologies de l\u2019information. C\u2019est dans ce cadre que nous avons identifie les besoins en matiere de gestion de la securite du campus.")

# 3. Problematique
heading2("Problematique :")
body("Malgre les avancees technologiques dans le domaine de la securite informationnelle, la gestion des incidents de securite dans les campus universitaires algeriens souffre de nombreuses lacunes que nous avons identifiees a travers notre etude preliminaire.")

body("Les principales insuffisances constatees sont les suivantes :")

bullet("Absence de centralisation : Les signalements d\u2019incidents transitent par des canaux heterogenes (appels telephoniques, deplacements physiques, e-mails), sans systeme unifie de reception et de suivi.")
bullet("Temps de reponse eleve : L\u2019absence d\u2019outils d\u2019alerte instantanee (type panic button) retarde significativement l\u2019intervention des services competents lors de situations d\u2019urgence.")
bullet("Manque de tracabilite : Les incidents ne sont generalement pas documentes numeriquement, rendant impossible toute analyse statistique ou identification de tendances recurrentes.")
bullet("Cloisonnement des acteurs : Les etudiants, le personnel de securite, les coordinateurs et les administrateurs operent en silos, sans plateforme collaborative permettant le partage d\u2019informations en temps reel.")
bullet("Absence de mobilite : Les systemes existants, lorsqu\u2019ils existent, ne sont pas accessibles depuis un terminal mobile, limitant la reactivite des agents de terrain.")

body("Ces constats nous amenent a formuler la question centrale suivante :")
body("\u00AB Comment concevoir et realiser une plateforme web et mobile integree, securisee et performante, permettant la gestion complete du cycle de vie des incidents de securite dans un environnement universitaire, depuis le signalement jusqu\u2019a la resolution ? \u00BB", bold=True)

# 4. Objectifs
heading2("Objectifs du projet :")
body("Le projet TrustDesk vise a atteindre un ensemble d\u2019objectifs fonctionnels et techniques.")

body("Objectifs generaux :", bold=True)
body("Nous devons concevoir et developper une plateforme (web et mobile) pour la gestion centralisee des incidents de securite en milieu universitaire.")

bullet("Mettre en oeuvre un systeme de gestion centralisee des incidents de securite (signalement, triage, assignation, suivi, cloture).")
bullet("Implementer un mecanisme de panique instantane avec geolocalisation depuis l\u2019application mobile.")
bullet("Mettre en place un radar communautaire permettant aux usagers de signaler des activites suspectes.")
bullet("Fournir un tableau de bord analytique avec indicateurs cles de performance (KPIs).")
bullet("Implementer un systeme d\u2019authentification securise base sur les roles (admin, etudiant, personnel, securite).")
bullet("Assurer les notifications en temps reel pour alerter les acteurs concernes.")

body("Objectifs Espace Administrateur :", bold=True)
bullet("Superviser l\u2019ensemble des incidents declares sur la plateforme.")
bullet("Gerer les utilisateurs, les roles et les permissions.")
bullet("Consulter les statistiques globales et les indicateurs de performance.")
bullet("Configurer les parametres du systeme et les workflows de triage.")

body("Objectifs Espace Agent de terrain (Mobile) :", bold=True)
bullet("Recevoir les alertes d\u2019incidents en temps reel avec geolocalisation.")
bullet("Declencher un bouton de panique en situation d\u2019urgence.")
bullet("Consulter et mettre a jour le statut des incidents assignes.")
bullet("Contribuer au radar communautaire via des signalements.")

body("Les Contraintes :", bold=True)
bullet("L\u2019application doit etre entierement responsive, s\u2019adaptant aux differents types d\u2019appareils.")
bullet("Mise en place d\u2019un systeme d\u2019authentification securise et de chiffrement des donnees (Laravel Sanctum).")
bullet("L\u2019application web doit etre compatible avec les navigateurs modernes (Chrome, Firefox, Edge, Safari).")
bullet("L\u2019application mobile doit fonctionner sur Android 8.0+ et iOS 12+.")
bullet("Le temps de reponse de l\u2019API ne doit pas depasser 500ms pour les operations courantes.")

heading3("Le public vise :")
body("Les administrateurs universitaires, les agents de securite, le personnel enseignant et administratif, ainsi que les etudiants de l\u2019etablissement.")

# 5. Methodologie
heading2("Methodologie de travail :")
body("Pour la realisation de ce projet, nous avons adopte une approche de developpement agile iterative, combinant les principes de la methode Scrum avec une architecture logicielle de type client-serveur a trois tiers.")

heading3("Planning des sprints :")

table(
    ['Sprint', 'Module', 'Duree'],
    [
        ['Sprint 1', 'Architecture API & Modele de donnees', '2 semaines'],
        ['Sprint 2', 'Authentification & Gestion des utilisateurs', '1 semaine'],
        ['Sprint 3', 'Module Incidents & Triage', '2 semaines'],
        ['Sprint 4', 'Module Panic & Geolocalisation', '1 semaine'],
        ['Sprint 5', 'Module Communautaire & Signaux', '2 semaines'],
        ['Sprint 6', 'Application Web - Dashboard & Pages', '3 semaines'],
        ['Sprint 7', 'Application Mobile - Toutes les fonctionnalites', '3 semaines'],
        ['Sprint 8', 'Tests, corrections & deploiement', '2 semaines'],
    ]
)
tbl_caption("Tableau N\u00b02 : Planning des sprints")

# 6. Concepts theoriques
heading2("Concepts theoriques :")
heading3("Les definitions generales :")

body("Application web :", bold=True)
body("Une application web est un logiciel qui s\u2019execute dans un navigateur web. Contrairement aux logiciels de bureau traditionnels, les applications web ne necessitent pas d\u2019installation locale et sont accessibles depuis n\u2019importe quel appareil disposant d\u2019une connexion Internet et d\u2019un navigateur compatible.")

body("Application Mobile :", bold=True)
body("Une application mobile est un logiciel concu pour fonctionner sur des appareils mobiles tels que les smartphones et les tablettes. Elle peut etre native (developpee pour un systeme d\u2019exploitation specifique) ou cross-platform (fonctionnant sur plusieurs systemes a partir d\u2019un code source unique).")

body("API REST :", bold=True)
body("Une API (Application Programming Interface) RESTful est une interface de programmation qui respecte les contraintes architecturales du style REST (Representational State Transfer). Elle permet a differentes applications de communiquer entre elles via le protocole HTTP en utilisant des methodes standard (GET, POST, PUT, DELETE).")

body("SOC (Security Operations Center) :", bold=True)
body("Un Centre d\u2019Operations de Securite est une structure centralisee dediee a la surveillance, la detection, l\u2019analyse et la reponse aux incidents de securite. Dans le contexte de TrustDesk, nous adaptons ce concept au milieu universitaire pour creer une plateforme de gestion proactive des incidents.")

heading3("Definition UML :")
body("Le Langage de Modelisation Unifie, de l\u2019anglais Unified Modeling Language (UML), est un langage de modelisation graphique a base de pictogrammes concu pour fournir une methode normalisee pour visualiser la conception d\u2019un systeme. Il est couramment utilise en developpement logiciel et en conception orientee objet.")

# 7. Technologies
heading2("Outils et technologies utilises :")

table(
    ['Composante', 'Technologie', 'Version', 'Role'],
    [
        ['Backend', 'Laravel (PHP)', '11.x', 'Framework API REST'],
        ['Authentification', 'Laravel Sanctum', '4.x', 'Tokens API securises'],
        ['Base de donnees', 'MySQL', '8.x', 'Stockage relationnel'],
        ['Frontend Web', 'React.js', '18.x', 'Interface SPA'],
        ['Typage', 'TypeScript', '5.x', 'Typage statique JS'],
        ['Bundler', 'Vite', '5.x', 'Build & dev server'],
        ['Routing Web', 'React Router', '6.x', 'Navigation SPA'],
        ['Mobile', 'Flutter', '3.x', 'App cross-platform'],
        ['Langage Mobile', 'Dart', '3.x', 'Langage Flutter'],
        ['Etat Mobile', 'Riverpod', '2.x', 'Gestion d\u2019etat reactif'],
        ['Navigation', 'GoRouter', '14.x', 'Nav declarative'],
        ['HTTP Client', 'Dio', '5.x', 'Requetes HTTP'],
        ['Geolocalisation', 'Geolocator', '13.x', 'Services GPS'],
        ['IDE', 'VS Code', 'Latest', 'Editeur de code'],
        ['Versioning', 'Git', 'Latest', 'Controle de version'],
    ]
)
tbl_caption("Tableau N\u00b03 : Technologies et outils utilises")

# 8. Conclusion
heading2("Conclusion :")
body("Ce premier chapitre nous a permis de situer le contexte general du projet TrustDesk, de presenter l\u2019organisme d\u2019accueil, d\u2019identifier clairement la problematique liee a la gestion des incidents de securite en milieu universitaire, et de definir les objectifs fonctionnels et techniques que cette plateforme vise a atteindre.")
body("La solution proposee se distingue par son approche multi-plateforme integree (web + mobile + API), son systeme de reponse d\u2019urgence instantanee, et son radar communautaire collaboratif.")
body("Le chapitre suivant sera consacre a l\u2019analyse et la specification des besoins, ou nous formaliserons les exigences fonctionnelles et non fonctionnelles du systeme a travers les outils de modelisation UML.")

# SAVE
output = r'D:\TrustDesk\Memoire_TrustDesk_Chapitre1.docx'
doc.save(output)
sz = os.path.getsize(output)
print(f'[OK] Saved: {output} ({sz/1024:.1f} KB)')
