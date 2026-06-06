# Part 1: Create base document with page de garde
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pickle

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.36); s.bottom_margin = Cm(2); s.left_margin = Cm(2); s.right_margin = Cm(1.25)

ns = doc.styles['Normal']; ns.font.name='Times New Roman'; ns.font.size=Pt(12); ns.paragraph_format.line_spacing=1.5
h1 = doc.styles['Heading 1']; h1.font.size=Pt(16); h1.font.color.rgb=RGBColor(0x1F,0x49,0x7D); h1.font.name='Times New Roman'
h2 = doc.styles['Heading 2']; h2.font.size=Pt(14); h2.font.color.rgb=RGBColor(0x1F,0x49,0x7D); h2.font.name='Times New Roman'
h3 = doc.styles['Heading 3']; h3.font.size=Pt(12); h3.font.bold=True; h3.font.name='Times New Roman'; h3.font.color.rgb=RGBColor(0x1F,0x49,0x7D)
bt = doc.styles['Body Text']; bt.font.name='Times New Roman'; bt.font.size=Pt(12); bt.paragraph_format.line_spacing=1.5

def C(t,sz=12,b=False,sa=0,sb=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(sb)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'; return p

def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it; return p

def E(): doc.add_paragraph(style='Body Text')

def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)

def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'

def H1(t,a=WD_ALIGN_PARAGRAPH.CENTER):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=2); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'
def FC(t): C(t,12,True)

# ═══ PAGE DE GARDE ═══
C('LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE\nMINISTERE DE L\u2019ENSEIGNEMENT SUPERIEUR ET DE LA RECHERCHE SCIENTIFIQUE',15,True,6)
E()
C('[Nom de votre Universite / Institut]',10,True,4)
E(); E()
C("Memoire De Fin D\u2019Etudes Pour L\u2019obtention Du Diplome\nde Master en Informatique",14,True,4)
C('Option : DEVELOPPEMENT WEB ET MOBILE',14,True,8)
E()
C('Theme :',16,True,4)
E(); E()
C("Conception Et Realisation D\u2019une",18,True,2)
C("Plateforme Web et Mobile",18,True,2)
C("de Gestion des Incidents de Securite",18,True,2)
C("en Milieu Universitaire",18,True,2)
C("\u00AB TrustDesk \u00BB",18,True,8)
E()
C("Organisme d\u2019accueil :",14,True,4)
E()
C("[Logo de l\u2019organisme]",11,False,8)
E(); E()

# Realise Par with left indent matching template (179070 EMU)
for txt in ["Realise Par :","M.Khalef Abdelmadjid","M.Dai Khaled Wassim"]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent=Emu(179070)
    r=p.add_run(txt); r.font.size=Pt(12); r.font.bold=True; r.font.name='Times New Roman'

# Suivi par
for txt in ["Suivi par :","Mme.Himeur"]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Emu(179070)
    r=p.add_run(txt); r.font.size=Pt(12); r.font.bold=True; r.font.name='Times New Roman'

for _ in range(4): E()
C('Promotion : 2025 / 2026',12,True)

doc.add_page_break()

# ═══ DEDICACES ═══
H1('Dedicaces')
E()
B("Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce projet de fin d\u2019etudes. Nous remercions sincerement notre encadreur Mme.Himeur pour ses conseils avises, son suivi regulier et sa disponibilite tout au long de ce travail.")
B("Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants tout au long de notre parcours academique.")
B("Enfin, nous souhaitons remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont rendu possible l\u2019aboutissement de ce projet.")
B("A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
for _ in range(5): E()
p=doc.add_paragraph(); r=p.add_run("Khalef A. , Dai K.W."); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'

doc.add_page_break()

# ═══ SOMMAIRE ═══
p=doc.add_paragraph(); r=p.add_run("Sommaire"); r.font.size=Pt(16); r.font.bold=True; r.font.name='Times New Roman'
E()
B("[Le sommaire sera genere automatiquement dans Word : References > Table des matieres]",it=True)

doc.add_page_break()

# ═══ REMERCIMENTS ═══
H1('REMERCIMENTS')
E()
B("Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce projet. Leur soutien et leurs encouragements ont ete precieux tout au long de cette aventure.")
B("Un merci particulier a nos familles pour leur soutien indefectible et leurs encouragements constants. Leur confiance en nos capacites a ete une source de motivation inestimable.")
B("Enfin, on souhaite remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont permis l\u2019aboutissement de ce travail.")
B("A toutes et a tous, nous vous adressons nos plus sinceres remerciements.")
for _ in range(5): E()
p=doc.add_paragraph(); r=p.add_run("Khalef A. , Dai K.W."); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'

doc.add_page_break()

# ═══ INTRODUCTION (title page) ═══
for _ in range(8): E()
C('INTRODUCTION',14,True,8)
for _ in range(8): E()

doc.add_page_break()

# Save intermediate
doc.save(r'D:\TrustDesk\_tmp_memoire.docx')
print('[OK] Part 1 done')
