# -*- coding: utf-8 -*-
"""Chapter 1 - Etude Prealable (Wallet Pivot)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'
def new_section(header_text=""):
    doc.add_section()
    sec = doc.sections[-1]
    sec.top_margin=Cm(2.36); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2); sec.right_margin=Cm(1.25)
    if header_text:
        hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(header_text)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.font.italic=True

# ════════════════════════════════════════════════════════════════
# CHAPITRE I - PAGE DE GARDE
# ════════════════════════════════════════════════════════════════
new_section()
for _ in range(8): E()
C('CHAPITRE I :', 18, True, 6)
C('\u00c9tude Pr\u00e9alable', 16, False, 8)
for _ in range(8): E()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CHAPITRE I - CONTENU
# ════════════════════════════════════════════════════════════════
new_section("\u00c9tude Pr\u00e9alable")

H1("Pr\u00e9sentation de l\u2019organisme d\u2019accueil :")

B("Dans le cadre de notre stage de fin de formation, nous avons \u00e9t\u00e9 accueillis au sein de l\u2019entreprise BEYN, un acteur majeur de l\u2019\u00e9cosyst\u00e8me fintech alg\u00e9rien. Fond\u00e9e en 2004 et bas\u00e9e \u00e0 Cheraga (Alger), BEYN est sp\u00e9cialis\u00e9e dans la conception et le d\u00e9veloppement de solutions technologiques destin\u00e9es aux institutions bancaires et financi\u00e8res.")
B("L\u2019entreprise se positionne comme un partenaire technologique strat\u00e9gique pour les banques, en leur fournissant des solutions frontend (web et mobile) qui s\u2019int\u00e8grent avec leurs syst\u00e8mes core banking existants. \u00c0 ce jour, BEYN collabore avec plus de dix banques partenaires en Alg\u00e9rie, parmi lesquelles la BDL (Banque de D\u00e9veloppement Local) et Al Salam Bank Algeria.")
B("L\u2019expertise de BEYN couvre plusieurs domaines cl\u00e9s :")
LP("D\u00e9veloppement de plateformes de banque digitale omnicanale (SELA pour le retail banking, KANTARA pour le corporate banking).")
LP("Solutions de paiement mobile et \u00e9lectronique (WimPay \u2014 paiement par QR code, transferts P2P, paiement de factures).")
LP("Optimisation des flux financiers inter-entreprises (TEQA \u2014 fournisseurs et distributeurs).")
LP("Accompagnement dans la transformation digitale et l\u2019inclusion financi\u00e8re.")
B("BEYN dispose \u00e9galement d\u2019une pr\u00e9sence internationale, avec une entit\u00e9 incorpor\u00e9e au Luxembourg pour soutenir ses op\u00e9rations en Europe et en Afrique.")

H2("Fiche technique de l\u2019organisme :")
T(['D\u00e9signation','Description'],[
    ['Raison sociale','BEYN SPA'],
    ['Date de cr\u00e9ation','2004'],
    ['Si\u00e8ge social','Cheraga, Alger, Alg\u00e9rie'],
    ['Secteur d\u2019activit\u00e9','Fintech \u2014 Solutions bancaires digitales'],
    ['Produits phares','SELA, WimPay, KANTARA, TEQA'],
    ['Clients','Banques et institutions financi\u00e8res (+10 partenaires)'],
    ['Pr\u00e9sence internationale','Luxembourg (Europe/Afrique)'],
])
C("Tableau N\u00b01 : Fiche technique de l\u2019entreprise BEYN", 12, True, 8)
B("[>>> INS\u00c9RER ICI : Logo de l\u2019entreprise BEYN <<<]", it=True)

H2("Organigramme de l\u2019organisme :")
B("[>>> INS\u00c9RER ICI : Figure N\u00b01 \u2013 Organigramme de l\u2019entreprise BEYN <<<]", it=True)
C("Figure N\u00b01 : Organigramme de l\u2019entreprise BEYN", 12, True, 8)

# ────────────────────────────────────────────────────────────
H1("Pr\u00e9sentation du projet :")

B("TrustDesk est une plateforme int\u00e9gr\u00e9e de portefeuille \u00e9lectronique s\u00e9curis\u00e9 (Secure Digital Wallet), con\u00e7ue sp\u00e9cifiquement pour les institutions bancaires clientes de BEYN. La plateforme combine trois composants majeurs :")
LP("Une API RESTful backend d\u00e9velopp\u00e9e avec Laravel 11 (PHP 8.2), servant de couche m\u00e9tier centralis\u00e9e pour la gestion des comptes, des transactions, de l\u2019authentification et du triage de s\u00e9curit\u00e9.")
LP("Une application web de supervision et d\u2019administration construite avec React.js 18 et TypeScript, offrant aux administrateurs bancaires un tableau de bord complet pour la gestion des clients, le suivi des transactions, et le monitoring de la s\u00e9curit\u00e9.")
LP("Une application mobile cross-platform d\u00e9velopp\u00e9e avec Flutter 3 (Dart), destin\u00e9e aux clients bancaires pour la gestion de leur portefeuille, les paiements par QR code, les transferts P2P, et l\u2019acc\u00e8s \u00e0 leurs cartes virtuelles.")

B("Le nom \u00ab TrustDesk \u00bb refl\u00e8te la double vocation de la plateforme : \u00ab Trust \u00bb pour la confiance et la s\u00e9curit\u00e9 des transactions bancaires, et \u00ab Desk \u00bb pour le centre d\u2019op\u00e9rations qui surveille et prot\u00e8ge l\u2019ensemble du syst\u00e8me.")

B("La plateforme se distingue des solutions existantes par l\u2019int\u00e9gration de fonctionnalit\u00e9s avanc\u00e9es rarement pr\u00e9sentes dans les portefeuilles \u00e9lectroniques alg\u00e9riens :")
LP("V\u00e9rification d\u2019identit\u00e9 num\u00e9rique (eKYC) avec d\u00e9tection de vivacit\u00e9 biom\u00e9trique pour l\u2019ouverture de compte \u00e0 distance.")
LP("Cartes virtuelles s\u00e9curis\u00e9es avec chiffrement asym\u00e9trique (cl\u00e9s publique/priv\u00e9e) et cartes jetables (disposable cards) pour les achats en ligne.")
LP("Triage intelligent par intelligence artificielle pour la d\u00e9tection automatique des transactions suspectes et la classification des incidents de s\u00e9curit\u00e9.")
LP("Syst\u00e8me de gel/d\u00e9gel instantan\u00e9 des comptes et cartes avec piste d\u2019audit compl\u00e8te.")
LP("Analytique financi\u00e8re personnalis\u00e9e avec insights de d\u00e9penses, cat\u00e9gorisation automatique et budg\u00e9tisation.")

H2("Fiche technique du projet :")
T(['D\u00e9signation','Description'],[
    ['Nom du projet','TrustDesk'],
    ['Nature','Plateforme de portefeuille \u00e9lectronique s\u00e9curis\u00e9 (web + mobile)'],
    ['Domaine','Fintech / Services bancaires digitaux'],
    ['Architecture','API REST + SPA Web + Application Mobile'],
    ['Backend','Laravel 11 (PHP 8.2)'],
    ['Frontend Web','React.js 18 + TypeScript + Vite'],
    ['Application Mobile','Flutter 3 (Dart)'],
    ['Base de donn\u00e9es','MySQL 8.x'],
    ['Authentification','Laravel Sanctum + eKYC biom\u00e9trique'],
    ['S\u00e9curit\u00e9','Chiffrement RSA/AES, tokenisation, audit SOC'],
])
C("Tableau N\u00b02 : Fiche technique du projet TrustDesk", 12, True, 8)

# ────────────────────────────────────────────────────────────
H2("Probl\u00e9matique :")

B("Malgr\u00e9 les progr\u00e8s r\u00e9alis\u00e9s dans la digitalisation du secteur bancaire alg\u00e9rien, les solutions de portefeuille \u00e9lectronique existantes pr\u00e9sentent des lacunes significatives que nous avons identifi\u00e9es lors de notre \u00e9tude pr\u00e9liminaire au sein de BEYN :")
LP("Absence de v\u00e9rification d\u2019identit\u00e9 num\u00e9rique (eKYC) : L\u2019ouverture de compte n\u00e9cessite encore un d\u00e9placement physique en agence, freinant l\u2019adoption et l\u2019inclusion financi\u00e8re.")
LP("S\u00e9curit\u00e9 insuffisante des donn\u00e9es sensibles : Les num\u00e9ros de cartes et donn\u00e9es de transaction ne b\u00e9n\u00e9ficient pas syst\u00e9matiquement d\u2019un chiffrement de bout en bout par cl\u00e9s asym\u00e9triques.")
LP("Absence de cartes virtuelles : Les clients ne disposent pas de cartes jetables (disposable cards) pour s\u00e9curiser leurs achats en ligne, contrairement aux solutions internationales comme Revolut.")
LP("D\u00e9tection de fraude rudimentaire : Les syst\u00e8mes existants reposent sur des r\u00e8gles statiques, sans intelligence artificielle capable de d\u00e9tecter les anomalies comportementales en temps r\u00e9el.")
LP("Manque d\u2019analytique financi\u00e8re : Les clients n\u2019ont pas acc\u00e8s \u00e0 des outils d\u2019analyse de leurs d\u00e9penses, de budg\u00e9tisation ou d\u2019insights personnalis\u00e9s.")
LP("Interface utilisateur non adapt\u00e9e aux standards actuels : Les interfaces existantes ne r\u00e9pondent pas aux exigences de design et d\u2019exp\u00e9rience utilisateur des n\u00e9o-banques modernes.")

B("Ces constats nous am\u00e8nent \u00e0 formuler la question centrale suivante :")
B("\u00ab Comment concevoir et r\u00e9aliser une plateforme web et mobile de portefeuille \u00e9lectronique s\u00e9curis\u00e9, int\u00e9grant l\u2019eKYC, le chiffrement asym\u00e9trique, les cartes virtuelles et un triage intelligent par IA, destin\u00e9e aux institutions bancaires alg\u00e9riennes ? \u00bb", b=True)

# ────────────────────────────────────────────────────────────
H2("Les objectifs :")

B("Objectifs g\u00e9n\u00e9raux :", b=True)
LP("Concevoir et d\u00e9velopper une plateforme de portefeuille \u00e9lectronique s\u00e9curis\u00e9 multi-plateforme (web + mobile + API) destin\u00e9e aux clients bancaires.")
LP("Impl\u00e9menter un module eKYC (Know Your Customer \u00e9lectronique) avec v\u00e9rification de documents d\u2019identit\u00e9 et d\u00e9tection de vivacit\u00e9 biom\u00e9trique.")
LP("Mettre en place un syst\u00e8me de cartes virtuelles s\u00e9curis\u00e9es avec chiffrement RSA/AES et support des cartes jetables.")
LP("D\u00e9velopper un moteur de triage intelligent par IA pour la d\u00e9tection d\u2019anomalies transactionnelles et la classification automatique des incidents.")
LP("Fournir un tableau de bord d\u2019administration complet pour la gestion des clients, le suivi des transactions et le monitoring de s\u00e9curit\u00e9.")
LP("Int\u00e9grer des outils d\u2019analytique financi\u00e8re personnalis\u00e9e (cat\u00e9gorisation des d\u00e9penses, budgets, insights).")

B("Objectifs Espace Administrateur (Web) :", b=True)
LP("Superviser l\u2019ensemble des comptes clients et leurs statuts KYC.")
LP("G\u00e9rer les transactions, les gels de comptes, et les incidents de s\u00e9curit\u00e9.")
LP("Consulter les rapports d\u2019activit\u00e9, les statistiques globales et les KPIs en temps r\u00e9el.")
LP("G\u00e9rer l\u2019organigramme, les tokens d\u2019acc\u00e8s et les permissions des op\u00e9rateurs.")
LP("G\u00e9n\u00e9rer des rapports clients, tickets de support et audits de conformit\u00e9.")

B("Objectifs Espace Client (Mobile) :", b=True)
LP("G\u00e9rer son portefeuille \u00e9lectronique : solde, historique, transferts P2P.")
LP("Effectuer des paiements par QR code et r\u00e8glement de factures.")
LP("G\u00e9rer ses cartes virtuelles : cr\u00e9ation, gel/d\u00e9gel, carte jetable pour achats en ligne.")
LP("Consulter ses insights financiers : cat\u00e9gorisation des d\u00e9penses, budgets, alertes de d\u00e9passement.")
LP("Compl\u00e9ter son processus eKYC depuis l\u2019application : photo de pi\u00e8ce d\u2019identit\u00e9, selfie, v\u00e9rification de vivacit\u00e9.")
LP("Personnaliser son profil : avatar, param\u00e8tres de s\u00e9curit\u00e9, pr\u00e9f\u00e9rences de notification.")

H3("Le public vis\u00e9 :")
B("La plateforme TrustDesk cible les publics suivants :")
LP("Les clients bancaires (particuliers) : utilisateurs finaux du portefeuille \u00e9lectronique mobile.")
LP("Les administrateurs bancaires : responsables de la supervision des comptes, des transactions et de la conformit\u00e9 r\u00e9glementaire.")
LP("Les op\u00e9rateurs BEYN : \u00e9quipes techniques assurant le support, la maintenance et le monitoring de la plateforme.")

# ────────────────────────────────────────────────────────────
H2("Les contraintes :")
LP("S\u00e9curit\u00e9 bancaire : Chiffrement RSA-2048/AES-256 pour les donn\u00e9es sensibles, conformit\u00e9 PCI-DSS, protection OWASP Top 10.")
LP("Performance : Temps de r\u00e9ponse API inf\u00e9rieur \u00e0 300ms pour 95% des transactions financi\u00e8res.")
LP("Compatibilit\u00e9 mobile : Android 8.0+ et iOS 12+ pour l\u2019application Flutter.")
LP("Responsivit\u00e9 web : Interface d\u2019administration compatible desktop, tablette et mobile.")
LP("Scalabilit\u00e9 : Architecture capable de supporter la mont\u00e9e en charge progressive li\u00e9e \u00e0 l\u2019ajout de nouvelles banques partenaires.")
LP("R\u00e9glementation : Conformit\u00e9 aux directives de la Banque d\u2019Alg\u00e9rie en mati\u00e8re de paiement \u00e9lectronique et d\u2019identification des clients.")

# ────────────────────────────────────────────────────────────
H2("Ressources humaines et mat\u00e9rielles :")

H3("Ressources humaines :")
T(['Nom','R\u00f4le'],[
    ['Khalef Abdelmadjid','D\u00e9veloppeur Full-Stack (Backend API + Mobile Flutter)'],
    ['Dai Khaled Wassim','D\u00e9veloppeur Full-Stack (Frontend React + Int\u00e9gration API)'],
    ['Mme. Himeur','Encadreur / Suivi p\u00e9dagogique'],
])
C("Tableau N\u00b03 : Ressources humaines", 12, True, 8)

H3("Ressources mat\u00e9rielles :")
T(['Mat\u00e9riel','Sp\u00e9cifications'],[
    ['PC Portable 1','[Sp\u00e9cifications \u00e0 compl\u00e9ter]'],
    ['PC Portable 2','[Sp\u00e9cifications \u00e0 compl\u00e9ter]'],
    ['Smartphone Android','Tests application mobile Flutter'],
    ['Connexion Internet','D\u00e9veloppement, tests API et d\u00e9ploiement'],
])
C("Tableau N\u00b04 : Ressources mat\u00e9rielles", 12, True, 8)

# ────────────────────────────────────────────────────────────
H2("M\u00e9thodologie de travail :")

B("Pour la r\u00e9alisation de ce projet, nous avons adopt\u00e9 une approche de d\u00e9veloppement agile it\u00e9rative, bas\u00e9e sur les principes de la m\u00e9thode Scrum. Cette m\u00e9thodologie nous a permis de d\u00e9composer le projet en sprints de dur\u00e9e variable, chacun d\u00e9di\u00e9 \u00e0 un module fonctionnel sp\u00e9cifique, avec des livraisons incr\u00e9mentales et des revues r\u00e9guli\u00e8res avec notre encadreur.")
B("L\u2019architecture retenue est une architecture trois tiers (three-tier) s\u00e9parant la couche de pr\u00e9sentation (React/Flutter), la couche m\u00e9tier (Laravel API), et la couche de donn\u00e9es (MySQL). Cette s\u00e9paration garantit la modularit\u00e9, la maintenabilit\u00e9 et l\u2019\u00e9volutivit\u00e9 du syst\u00e8me.")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Cycle de d\u00e9veloppement Scrum utilis\u00e9 <<<]", it=True)

H3("Planning des sprints :")
T(['Sprint','Module','Dur\u00e9e'],[
    ['Sprint 1','Architecture API & Mod\u00e8le de donn\u00e9es','2 semaines'],
    ['Sprint 2','Authentification & Module eKYC','2 semaines'],
    ['Sprint 3','Module Portefeuille & Transactions','2 semaines'],
    ['Sprint 4','Module Cartes Virtuelles & Chiffrement','2 semaines'],
    ['Sprint 5','Module Triage IA & D\u00e9tection de fraude','2 semaines'],
    ['Sprint 6','Gestion Clients & Administration','2 semaines'],
    ['Sprint 7','Application Web \u2014 Dashboard & Rapports','3 semaines'],
    ['Sprint 8','Application Mobile Flutter','3 semaines'],
    ['Sprint 9','Tests, corrections & d\u00e9ploiement','2 semaines'],
])
C("Tableau N\u00b05 : Planning des sprints", 12, True, 8)

# ────────────────────────────────────────────────────────────
H2("Conclusion :")

B("Ce premier chapitre nous a permis de situer le contexte g\u00e9n\u00e9ral du projet TrustDesk, de pr\u00e9senter l\u2019organisme d\u2019accueil BEYN et son positionnement strat\u00e9gique dans l\u2019\u00e9cosyst\u00e8me fintech alg\u00e9rien, d\u2019identifier la probl\u00e9matique li\u00e9e aux insuffisances des solutions de portefeuille \u00e9lectronique actuelles, et de d\u00e9finir les objectifs fonctionnels et techniques de la plateforme.")
B("La solution propos\u00e9e se distingue par son approche s\u00e9curis\u00e9e de bout en bout (eKYC, chiffrement asym\u00e9trique, cartes virtuelles jetables), son intelligence artificielle int\u00e9gr\u00e9e pour le triage des incidents, et son mod\u00e8le B2B adapt\u00e9 aux besoins des banques partenaires de BEYN.")
B("Le chapitre suivant sera consacr\u00e9 \u00e0 l\u2019analyse et la sp\u00e9cification des besoins, o\u00f9 nous formaliserons les exigences fonctionnelles et non fonctionnelles du syst\u00e8me \u00e0 travers les outils de mod\u00e9lisation UML.")

doc.add_page_break()

# SAVE
out = r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx'
doc.save(out)
print(f'[OK] Chapter 1 added: {out} ({os.path.getsize(out)/1024:.1f} KB)')
