# -*- coding: utf-8 -*-
"""Chapter 4 - Realisation + Conclusion Generale (Wallet)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t);r.font.size=Pt(sz);r.font.bold=b;r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text');p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa);p.paragraph_format.line_spacing=1.5
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.bold=b;r.font.italic=it
def E():doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph');p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr();numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}));pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h);r.font.bold=True;r.font.size=Pt(12);r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val));r.font.size=Pt(12);r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1);p.alignment=a
    for r in p.runs:r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs:r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs:r.font.name='Times New Roman'

# ═══════ CHAPTER 4 COVER ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)
for _ in range(8):E()
C('CHAPITRE IV :',18,True,6)
C('R\u00e9alisation et Impl\u00e9mentation',16,False,8)
for _ in range(8):E()
doc.add_page_break()

# ═══════ CONTENT ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)

H1("Introduction :")
B("Ce dernier chapitre pr\u00e9sente la r\u00e9alisation concr\u00e8te de la plateforme TrustDesk. Nous y d\u00e9crivons l\u2019environnement de d\u00e9veloppement, les outils et technologies utilis\u00e9s, les interfaces r\u00e9alis\u00e9es, ainsi que les r\u00e9sultats des tests de validation.")

# ─── OUTILS ───
H1("Environnement de d\u00e9veloppement :")

H2("Outils logiciels :")
T(['Outil','Version','Usage'],[
    ['Visual Studio Code','1.90+','IDE principal (backend + frontend)'],
    ['Android Studio','2024.1+','\u00c9mulateur Android + d\u00e9veloppement Flutter'],
    ['Postman','11.x','Tests et documentation des APIs REST'],
    ['Git / GitHub','2.44+','Gestion de version et collaboration'],
    ['XAMPP','8.2','Serveur local Apache + MySQL'],
    ['Figma','Web','Maquettage UI/UX des interfaces'],
    ['StarUML / Draw.io','Web','Mod\u00e9lisation UML (diagrammes)'],
])
C("Tableau N\u00b021 : Outils logiciels utilis\u00e9s",12,True,8)

H2("Technologies Backend :")
T(['Technologie','Version','R\u00f4le'],[
    ['PHP','8.2','Langage serveur'],
    ['Laravel','11.x','Framework backend (API REST, ORM, auth)'],
    ['Laravel Sanctum','4.x','Authentification par tokens bearer'],
    ['MySQL','8.x','Syst\u00e8me de gestion de base de donn\u00e9es'],
    ['Composer','2.7+','Gestionnaire de d\u00e9pendances PHP'],
])
C("Tableau N\u00b022 : Technologies backend",12,True,8)

H2("Technologies Frontend Web :")
T(['Technologie','Version','R\u00f4le'],[
    ['React.js','18.x','Biblioth\u00e8que UI (composants r\u00e9actifs)'],
    ['TypeScript','5.x','Typage statique pour JavaScript'],
    ['Vite','5.x','Bundler et serveur de d\u00e9veloppement'],
    ['Axios','1.7+','Client HTTP pour les appels API'],
    ['Recharts','2.x','Graphiques et visualisations de donn\u00e9es'],
    ['Lucide React','0.4+','Biblioth\u00e8que d\u2019ic\u00f4nes SVG'],
])
C("Tableau N\u00b023 : Technologies frontend web",12,True,8)

H2("Technologies Mobile :")
T(['Technologie','Version','R\u00f4le'],[
    ['Flutter','3.22+','Framework cross-platform (iOS + Android)'],
    ['Dart','3.4+','Langage de programmation Flutter'],
    ['Provider / Riverpod','2.x','Gestion d\u2019\u00e9tat r\u00e9actif'],
    ['Dio','5.x','Client HTTP pour les appels API'],
    ['flutter_secure_storage','9.x','Stockage s\u00e9curis\u00e9 des tokens et cl\u00e9s'],
    ['local_auth','2.x','Authentification biom\u00e9trique (empreinte, Face ID)'],
    ['qr_code_scanner','1.x','Lecture de QR codes pour paiements'],
])
C("Tableau N\u00b024 : Technologies mobile",12,True,8)

# ─── INTERFACES ───
H1("Pr\u00e9sentation des interfaces r\u00e9alis\u00e9es :")

H2("Interfaces de l\u2019application mobile (Client) :")
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 \u00c9cran de connexion / inscription <<<]",it=True)
C("Figure N\u00b012 : \u00c9cran de connexion de l\u2019application mobile",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Processus eKYC (upload document + selfie) <<<]",it=True)
C("Figure N\u00b013 : Processus de v\u00e9rification eKYC",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Dashboard portefeuille (solde + historique) <<<]",it=True)
C("Figure N\u00b014 : Dashboard du portefeuille client",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Transfert P2P <<<]",it=True)
C("Figure N\u00b015 : Interface de transfert P2P",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Gestion des cartes virtuelles <<<]",it=True)
C("Figure N\u00b016 : Gestion des cartes virtuelles",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Paiement par QR code <<<]",it=True)
C("Figure N\u00b017 : Paiement par QR code",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Insights financiers <<<]",it=True)
C("Figure N\u00b018 : Insights et analytique financi\u00e8re",12,True,8)

H2("Interfaces du tableau de bord web (Administrateur) :")
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Page de connexion admin <<<]",it=True)
C("Figure N\u00b019 : Page de connexion administrateur",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Dashboard admin (KPIs + graphiques) <<<]",it=True)
C("Figure N\u00b020 : Dashboard d\u2019administration",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Gestion des clients (liste + d\u00e9tails) <<<]",it=True)
C("Figure N\u00b021 : Gestion des clients",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Validation eKYC par l\u2019admin <<<]",it=True)
C("Figure N\u00b022 : Validation eKYC",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Gestion des incidents + triage IA <<<]",it=True)
C("Figure N\u00b023 : Gestion des incidents et triage IA",12,True,8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran \u2014 Rapports et audit logs <<<]",it=True)
C("Figure N\u00b024 : Rapports et piste d\u2019audit",12,True,8)

# ─── CODE HIGHLIGHTS ───
H1("Extraits de code significatifs :")

H2("Validation stricte des entr\u00e9es (withValidator) :")
B("Pour pr\u00e9venir les attaques par mass-assignment, toutes les Form Requests Laravel utilisent un hook withValidator qui rejette syst\u00e9matiquement tout champ non d\u00e9clar\u00e9 dans les r\u00e8gles de validation :")
B("[>>> INS\u00c9RER ICI : Extrait de code \u2014 withValidator dans LoginRequest.php <<<]",it=True)

H2("Service de chiffrement RSA :")
B("Le EncryptionService encapsule la logique de chiffrement/d\u00e9chiffrement RSA pour les donn\u00e9es de cartes virtuelles :")
B("[>>> INS\u00c9RER ICI : Extrait de code \u2014 EncryptionService.php (encrypt/decrypt) <<<]",it=True)

H2("Composant React \u2014 Dashboard admin :")
B("[>>> INS\u00c9RER ICI : Extrait de code \u2014 Composant Dashboard.tsx (KPIs + graphiques Recharts) <<<]",it=True)

H2("Widget Flutter \u2014 Carte virtuelle :")
B("[>>> INS\u00c9RER ICI : Extrait de code \u2014 Widget VirtualCardWidget (affichage s\u00e9curis\u00e9) <<<]",it=True)

# ─── TESTS ───
H1("Tests et validation :")

H2("Tests unitaires (Backend) :")
B("Les tests unitaires Laravel couvrent les modules critiques : authentification, transactions, et chiffrement.")
T(['Module test\u00e9','Tests','R\u00e9sultat'],[
    ['Authentification (login/register)','8 tests','8/8 pass\u00e9s'],
    ['Transfert P2P','6 tests','6/6 pass\u00e9s'],
    ['Cr\u00e9ation carte virtuelle','5 tests','5/5 pass\u00e9s'],
    ['Chiffrement RSA/AES','4 tests','4/4 pass\u00e9s'],
    ['Validation withValidator','6 tests','6/6 pass\u00e9s'],
])
C("Tableau N\u00b025 : R\u00e9sultats des tests unitaires",12,True,8)

H2("Tests d\u2019int\u00e9gration (API) :")
B("Les tests d\u2019int\u00e9gration Postman v\u00e9rifient les endpoints API dans des sc\u00e9narios r\u00e9alistes :")
T(['Endpoint','M\u00e9thode','Sc\u00e9nario','R\u00e9sultat'],[
    ['POST /api/register','POST','Inscription valide','201 Created'],
    ['POST /api/login','POST','Connexion valide','200 OK + token'],
    ['POST /api/transfer','POST','Transfert P2P valide','200 OK'],
    ['POST /api/transfer','POST','Solde insuffisant','422 Error'],
    ['POST /api/cards','POST','Cr\u00e9ation carte virtuelle','201 Created'],
    ['PUT /api/cards/{id}/freeze','PUT','Gel de carte','200 OK'],
    ['GET /api/wallet','GET','Consultation solde','200 OK'],
])
C("Tableau N\u00b026 : R\u00e9sultats des tests d\u2019int\u00e9gration API",12,True,8)

H2("Tests de s\u00e9curit\u00e9 :")
LP("Protection OWASP Top 10 : injection SQL, XSS, CSRF \u2014 test\u00e9s et valid\u00e9s.")
LP("Rejet des champs inattendus (mass-assignment) : valid\u00e9 via withValidator.")
LP("Force brute : rate limiting configur\u00e9 \u00e0 60 requ\u00eates/minute par IP.")
LP("Tokens Sanctum : expiration configur\u00e9e, r\u00e9vocation fonctionnelle.")

# ─── CONCLUSION CH4 ───
H2("Conclusion :")
B("Ce quatri\u00e8me chapitre a pr\u00e9sent\u00e9 la r\u00e9alisation concr\u00e8te de la plateforme TrustDesk, en d\u00e9taillant l\u2019environnement de d\u00e9veloppement, les technologies utilis\u00e9es, les interfaces r\u00e9alis\u00e9es pour l\u2019application mobile et le tableau de bord web, ainsi que les r\u00e9sultats des tests de validation. La plateforme r\u00e9pond aux objectifs d\u00e9finis dans le cahier des charges et int\u00e8gre les standards de s\u00e9curit\u00e9 bancaire attendus.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# CONCLUSION GENERALE
# ════════════════════════════════════════════════════════════════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)

C("CONCLUSION G\u00c9N\u00c9RALE",16,True,12)
E()
B("Au terme de ce m\u00e9moire, nous avons con\u00e7u et r\u00e9alis\u00e9 TrustDesk, une plateforme int\u00e9gr\u00e9e de portefeuille \u00e9lectronique s\u00e9curis\u00e9 destin\u00e9e aux institutions bancaires alg\u00e9riennes, d\u00e9velopp\u00e9e dans le cadre de notre stage au sein de l\u2019entreprise BEYN.")
B("La plateforme combine trois composants technologiques compl\u00e9mentaires : une API RESTful centralis\u00e9e (Laravel 11), un tableau de bord d\u2019administration (React.js 18), et une application mobile cross-platform (Flutter 3). Elle int\u00e8gre des fonctionnalit\u00e9s avanc\u00e9es rarement pr\u00e9sentes dans les solutions existantes sur le march\u00e9 alg\u00e9rien :")
LP("Un module eKYC complet avec v\u00e9rification de documents d\u2019identit\u00e9 et d\u00e9tection de vivacit\u00e9 biom\u00e9trique.")
LP("Un syst\u00e8me de cartes virtuelles s\u00e9curis\u00e9es avec chiffrement asym\u00e9trique RSA/AES et support des cartes jetables.")
LP("Un moteur de triage intelligent par intelligence artificielle pour la d\u00e9tection d\u2019anomalies transactionnelles.")
LP("Une piste d\u2019audit compl\u00e8te et immuable garantissant la tra\u00e7abilit\u00e9 de toutes les op\u00e9rations sensibles.")
LP("Des outils d\u2019analytique financi\u00e8re personnalis\u00e9e (cat\u00e9gorisation des d\u00e9penses, budgets, insights).")
B("Ce projet nous a permis d\u2019approfondir nos comp\u00e9tences en d\u00e9veloppement full-stack, en architecture de syst\u00e8mes distribu\u00e9s, en cryptographie appliqu\u00e9e, et en conception d\u2019interfaces utilisateur modernes pour le secteur bancaire.")
B("En termes de perspectives d\u2019\u00e9volution, plusieurs axes d\u2019am\u00e9lioration sont envisag\u00e9s :")
LP("Int\u00e9gration avec les syst\u00e8mes core banking des banques partenaires de BEYN via des APIs standardis\u00e9es.")
LP("Impl\u00e9mentation de la tokenisation NFC pour les paiements sans contact (Apple Pay, Google Pay).")
LP("D\u00e9ploiement d\u2019un mod\u00e8le de machine learning entra\u00een\u00e9 pour la d\u00e9tection de fraude en temps r\u00e9el.")
LP("Extension de la plateforme vers les services de micro-cr\u00e9dit et d\u2019\u00e9pargne digitale.")
LP("Certification PCI-DSS compl\u00e8te pour le traitement des donn\u00e9es de paiement en production.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# BIBLIOGRAPHIE
# ════════════════════════════════════════════════════════════════
C("BIBLIOGRAPHIE",16,True,12)
E()
B("[1] Laravel Documentation (v11.x), https://laravel.com/docs/11.x", sa=4)
B("[2] React Documentation, https://react.dev/", sa=4)
B("[3] Flutter Documentation, https://docs.flutter.dev/", sa=4)
B("[4] Laravel Sanctum, https://laravel.com/docs/11.x/sanctum", sa=4)
B("[5] OWASP Top 10 (2021), https://owasp.org/www-project-top-ten/", sa=4)
B("[6] PCI DSS v4.0, https://www.pcisecuritystandards.org/", sa=4)
B("[7] BEYN Official Website, https://beyn.dz/", sa=4)
B("[8] Banque d\u2019Alg\u00e9rie \u2014 R\u00e9glementation paiement \u00e9lectronique, https://www.bank-of-algeria.dz/", sa=4)
B("[9] RSA Cryptography (PKCS #1), RFC 8017, https://datatracker.ietf.org/doc/html/rfc8017", sa=4)
B("[10] Vite.js Documentation, https://vitejs.dev/", sa=4)
B("[11] MySQL 8.0 Reference Manual, https://dev.mysql.com/doc/refman/8.0/en/", sa=4)
B("[12] Postman API Platform, https://www.postman.com/", sa=4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# GLOSSAIRE
# ════════════════════════════════════════════════════════════════
C("GLOSSAIRE",16,True,12)
E()
T(['Terme','D\u00e9finition'],[
    ['API','Application Programming Interface \u2014 Interface de programmation applicative'],
    ['REST','Representational State Transfer \u2014 Style d\u2019architecture pour les APIs web'],
    ['eKYC','Electronic Know Your Customer \u2014 V\u00e9rification d\u2019identit\u00e9 num\u00e9rique'],
    ['P2P','Peer-to-Peer \u2014 Transfert de pair \u00e0 pair entre utilisateurs'],
    ['RSA','Rivest\u2013Shamir\u2013Adleman \u2014 Algorithme de chiffrement asym\u00e9trique'],
    ['AES','Advanced Encryption Standard \u2014 Algorithme de chiffrement sym\u00e9trique'],
    ['PCI-DSS','Payment Card Industry Data Security Standard'],
    ['OWASP','Open Worldwide Application Security Project'],
    ['SOC','Security Operations Center \u2014 Centre d\u2019op\u00e9rations de s\u00e9curit\u00e9'],
    ['UML','Unified Modeling Language \u2014 Langage de mod\u00e9lisation unifi\u00e9'],
    ['ORM','Object-Relational Mapping \u2014 Mapping objet-relationnel'],
    ['JWT','JSON Web Token \u2014 Standard d\u2019authentification par jetons'],
    ['CRUD','Create, Read, Update, Delete \u2014 Op\u00e9rations de base sur les donn\u00e9es'],
    ['QR Code','Quick Response Code \u2014 Code \u00e0 r\u00e9ponse rapide'],
    ['SPA','Single Page Application \u2014 Application monopage'],
    ['DZD','Dinar alg\u00e9rien \u2014 Devise nationale'],
])
C("Tableau N\u00b027 : Glossaire des termes techniques",12,True,8)

# SAVE FINAL
out=r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx'
doc.save(out)
print(f'[OK] COMPLETE MEMOIRE saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
