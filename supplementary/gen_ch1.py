# -*- coding: utf-8 -*-
"""Append Chapter 1 to the memoire"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v4.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'
def new_section(header_text=""):
    doc.add_section()
    sec = doc.sections[-1]
    sec.top_margin=Cm(2.36); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2); sec.right_margin=Cm(1.25)
    if header_text:
        hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(header_text)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.font.italic=True

# ════════════════════════════════════════════════════════════════
# CHAPITRE I - GUARD PAGE
# ════════════════════════════════════════════════════════════════
new_section()
for _ in range(8): E()
C('CHAPITRE I :', 18, True, 6)
C('\u00c9tude Pr\u00e9alable', 16, False, 8)
for _ in range(8): E()
doc.add_page_break()

# CHAPITRE I CONTENT
new_section("\u00c9tude Pr\u00e9alable")

H1("Pr\u00e9sentation de l\u2019organisme d\u2019accueil :")
B("Dans le cadre de notre stage de fin de formation au sein de l\u2019entreprise BEYN, nous avons \u00e9t\u00e9 charg\u00e9s de concevoir et r\u00e9aliser la plateforme TrustDesk. BEYN est une entreprise alg\u00e9rienne sp\u00e9cialis\u00e9e dans le d\u00e9veloppement de solutions informatiques innovantes, offrant des services de conseil, de d\u00e9veloppement logiciel et d\u2019int\u00e9gration de syst\u00e8mes d\u2019information pour divers secteurs d\u2019activit\u00e9.")
B("L\u2019entreprise met \u00e0 disposition de ses clients des solutions technologiques adapt\u00e9es \u00e0 leurs besoins sp\u00e9cifiques, couvrant le d\u00e9veloppement d\u2019applications web et mobiles, la mise en place d\u2019infrastructures r\u00e9seau, ainsi que l\u2019accompagnement dans la transformation digitale.")

H2("Fiche Technique de l\u2019organisme :")
T(['D\u00e9signation','Description'],[
    ['Raison sociale','BEYN'],
    ['Secteur d\u2019activit\u00e9','D\u00e9veloppement informatique & Solutions IT'],
    ['Si\u00e8ge social','Alger, Alg\u00e9rie'],
    ['Domaine d\u2019expertise','Applications web, mobile, int\u00e9gration SI'],
    ['Effectif','[Nombre d\u2019employ\u00e9s]'],
])
C("Tableau N\u00b01 : Fiche technique de l\u2019organisme d\u2019accueil", 12, True, 8)
B("[>>> INS\u00c9RER ICI : Capture d\u2019\u00e9cran ou logo de l\u2019entreprise BEYN <<<]", it=True)

H2("Organigramme de l\u2019organisme :")
B("[>>> INS\u00c9RER ICI : Figure N\u00b01 \u2013 Organigramme de l\u2019entreprise BEYN <<<]", it=True)
C("Figure N\u00b01 : Organigramme de l\u2019entreprise BEYN", 12, True, 8)

H1("Pr\u00e9sentation du projet :")
B("TrustDesk est une plateforme int\u00e9gr\u00e9e de type Security Operations Center (SOC), sp\u00e9cialement con\u00e7ue pour la gestion des incidents de s\u00e9curit\u00e9 en milieu universitaire. Elle se compose de trois modules compl\u00e9mentaires : une API RESTful backend d\u00e9velopp\u00e9e avec Laravel 11, une application web de supervision construite avec React.js 18 et TypeScript, et une application mobile de terrain r\u00e9alis\u00e9e avec Flutter 3.")
B("L\u2019application web constitue le c\u0153ur du syst\u00e8me de supervision et permet aux administrateurs de g\u00e9rer l\u2019ensemble des incidents signal\u00e9s, de visualiser les statistiques en temps r\u00e9el via un tableau de bord interactif, de g\u00e9rer les utilisateurs et leurs r\u00f4les, et de configurer les param\u00e8tres du syst\u00e8me. Le dashboard int\u00e8gre des indicateurs cl\u00e9s de performance (KPIs) tels que le nombre total d\u2019incidents, le temps moyen de r\u00e9solution, le taux de r\u00e9solution et la r\u00e9partition par cat\u00e9gorie et par priorit\u00e9.")
B("En compl\u00e9ment, l\u2019application mobile permet aux agents de s\u00e9curit\u00e9 et aux usagers du campus d\u2019acc\u00e9der aux fonctionnalit\u00e9s essentielles directement depuis leur terminal mobile. Elle offre notamment un bouton panique permettant de d\u00e9clencher une alerte d\u2019urgence en un seul geste avec capture automatique des coordonn\u00e9es GPS, un radar communautaire pour signaler les activit\u00e9s suspectes avec g\u00e9olocalisation, et un syst\u00e8me de notifications push en temps r\u00e9el.")
B("Par ailleurs, la plateforme int\u00e8gre un module de portefeuille \u00e9lectronique (wallet) permettant la gestion des transactions financi\u00e8res li\u00e9es aux services de s\u00e9curit\u00e9 du campus, avec des fonctionnalit\u00e9s de gel de fonds, d\u2019historique des transactions et de suivi des soldes en temps r\u00e9el.")

H2("Fiche Technique du projet :")
T(['D\u00e9signation','Description'],[
    ['Nom du projet','TrustDesk'],
    ['Nature','Plateforme web et mobile de gestion des incidents de s\u00e9curit\u00e9'],
    ['Domaine','S\u00e9curit\u00e9 des campus universitaires'],
    ['Architecture','API REST + SPA Web + Application Mobile'],
    ['Backend','Laravel 11 (PHP 8.2)'],
    ['Frontend Web','React.js 18 + TypeScript'],
    ['Application Mobile','Flutter 3 (Dart)'],
    ['Base de donn\u00e9es','MySQL 8.x'],
    ['Authentification','Laravel Sanctum (tokens API)'],
])
C("Tableau N\u00b02 : Fiche technique du projet TrustDesk", 12, True, 8)

H2("Probl\u00e9matique :")
B("Malgr\u00e9 les avanc\u00e9es technologiques dans le domaine de la s\u00e9curit\u00e9 informationnelle, la gestion des incidents de s\u00e9curit\u00e9 dans les campus universitaires alg\u00e9riens souffre de nombreuses lacunes que nous avons identifi\u00e9es \u00e0 travers notre \u00e9tude pr\u00e9liminaire r\u00e9alis\u00e9e au sein de l\u2019entreprise BEYN.")
B("Les principales insuffisances constat\u00e9es sont les suivantes :")
LP("Absence de centralisation : Les signalements d\u2019incidents transitent par des canaux h\u00e9t\u00e9rog\u00e8nes (appels t\u00e9l\u00e9phoniques, d\u00e9placements physiques, e-mails), sans syst\u00e8me unifi\u00e9 de r\u00e9ception et de suivi.")
LP("Temps de r\u00e9ponse \u00e9lev\u00e9 : L\u2019absence d\u2019outils d\u2019alerte instantan\u00e9e retarde significativement l\u2019intervention des services comp\u00e9tents lors de situations d\u2019urgence.")
LP("Manque de tra\u00e7abilit\u00e9 : Les incidents ne sont pas document\u00e9s num\u00e9riquement, rendant impossible toute analyse statistique ou r\u00e9trospective.")
LP("Cloisonnement des acteurs : Les \u00e9tudiants, le personnel de s\u00e9curit\u00e9 et les administrateurs op\u00e8rent en silos sans canal de communication commun.")
LP("Absence de mobilit\u00e9 : Les syst\u00e8mes existants, quand ils existent, ne sont pas accessibles depuis un terminal mobile, limitant la r\u00e9activit\u00e9 sur le terrain.")
B("Ces constats nous am\u00e8nent \u00e0 formuler la question centrale suivante :")
B("\u00ab Comment concevoir et r\u00e9aliser une plateforme web et mobile int\u00e9gr\u00e9e, s\u00e9curis\u00e9e et performante, permettant la gestion compl\u00e8te du cycle de vie des incidents de s\u00e9curit\u00e9 dans un environnement universitaire, tout en r\u00e9pondant aux objectifs assign\u00e9s et en respectant les contraintes impos\u00e9es ? \u00bb", b=True)

H2("Les objectifs :")
B("Objectifs g\u00e9n\u00e9raux :", b=True)
B("Nous devons concevoir et d\u00e9velopper une plateforme (web et mobile) pour la gestion centralis\u00e9e des incidents de s\u00e9curit\u00e9 en milieu universitaire, int\u00e9grant les fonctionnalit\u00e9s suivantes :")
LP("Mettre en \u0153uvre un syst\u00e8me de gestion centralis\u00e9e des incidents couvrant l\u2019int\u00e9gralit\u00e9 du cycle de vie : signalement, triage automatique par score de priorit\u00e9, assignation aux agents comp\u00e9tents, suivi en temps r\u00e9el, r\u00e9solution et cl\u00f4ture document\u00e9e.")
LP("Impl\u00e9menter un m\u00e9canisme de bouton panique instantan\u00e9 avec capture automatique des coordonn\u00e9es GPS depuis l\u2019application mobile, permettant une intervention d\u2019urgence rapide.")
LP("Mettre en place un radar communautaire permettant aux usagers du campus de signaler les activit\u00e9s suspectes avec g\u00e9olocalisation pr\u00e9cise.")
LP("Fournir un tableau de bord analytique avec indicateurs cl\u00e9s de performance (KPIs) : nombre total d\u2019incidents, temps moyen de r\u00e9solution, taux de r\u00e9solution, r\u00e9partition par cat\u00e9gorie.")
LP("Impl\u00e9menter un syst\u00e8me d\u2019authentification s\u00e9curis\u00e9 bas\u00e9 sur les r\u00f4les (RBAC) avec Laravel Sanctum.")
LP("Assurer les notifications en temps r\u00e9el pour alerter imm\u00e9diatement les acteurs concern\u00e9s.")
LP("Int\u00e9grer un module de portefeuille \u00e9lectronique pour la gestion des transactions li\u00e9es aux services de s\u00e9curit\u00e9.")

B("Objectifs Espace Administrateur (Web) :", b=True)
LP("Superviser l\u2019ensemble des incidents d\u00e9clar\u00e9s sur la plateforme via un tableau de bord interactif et dynamique.")
LP("G\u00e9rer les utilisateurs, les r\u00f4les et les permissions d\u2019acc\u00e8s au syst\u00e8me.")
LP("Consulter les statistiques globales et les indicateurs cl\u00e9s de performance (KPIs) en temps r\u00e9el.")
LP("Configurer les cat\u00e9gories d\u2019incidents, les niveaux de priorit\u00e9 et les workflows de triage automatique.")
LP("G\u00e9rer le module de portefeuille \u00e9lectronique et superviser les transactions.")

B("Objectifs Espace Agent de terrain (Mobile) :", b=True)
LP("Recevoir les alertes d\u2019incidents en temps r\u00e9el avec g\u00e9olocalisation pr\u00e9cise.")
LP("D\u00e9clencher un bouton de panique en situation d\u2019urgence avec capture GPS automatique.")
LP("Consulter et mettre \u00e0 jour le statut des incidents assign\u00e9s directement depuis le terrain.")
LP("Contribuer au radar communautaire via des signalements g\u00e9olocalis\u00e9s.")

H3("Le public vis\u00e9 :")
B("La plateforme TrustDesk cible les publics suivants :")
LP("Les administrateurs universitaires : responsables de la supervision g\u00e9n\u00e9rale de la s\u00e9curit\u00e9 du campus.")
LP("Les agents de s\u00e9curit\u00e9 : personnel de terrain charg\u00e9 de l\u2019intervention et de la r\u00e9solution des incidents.")
LP("Le personnel enseignant et administratif : utilisateurs pouvant signaler des incidents.")
LP("Les \u00e9tudiants : usagers principaux du campus, b\u00e9n\u00e9ficiant du bouton panique et du radar communautaire.")

H2("Les contraintes :")
LP("L\u2019application web doit \u00eatre enti\u00e8rement responsive, s\u2019adaptant aux diff\u00e9rents types d\u2019\u00e9crans (desktop, tablette, mobile).")
LP("Mise en place d\u2019un syst\u00e8me d\u2019authentification s\u00e9curis\u00e9 (Laravel Sanctum) avec chiffrement des donn\u00e9es sensibles.")
LP("Compatibilit\u00e9 avec les navigateurs modernes (Chrome, Firefox, Edge, Safari).")
LP("Fonctionnement sur Android 8.0+ et iOS 12+ pour l\u2019application mobile.")
LP("Temps de r\u00e9ponse API inf\u00e9rieur \u00e0 500ms pour 95% des requ\u00eates.")
LP("Conformit\u00e9 aux bonnes pratiques OWASP Top 10 en mati\u00e8re de s\u00e9curit\u00e9 applicative.")

H2("Ressources humaines et mat\u00e9rielles :")
H3("Ressources humaines :")
T(['Nom','R\u00f4le'],[
    ['Khalef Abdelmadjid','D\u00e9veloppeur Full-Stack (Backend + Mobile)'],
    ['Dai Khaled Wassim','D\u00e9veloppeur Full-Stack (Frontend + API)'],
    ['Mme. Himeur','Encadreur / Suivi p\u00e9dagogique'],
])
C("Tableau N\u00b03 : Ressources humaines", 12, True, 8)

H3("Ressources mat\u00e9rielles :")
T(['Mat\u00e9riel','Sp\u00e9cifications'],[
    ['PC Portable 1','[Sp\u00e9cifications \u00e0 compl\u00e9ter]'],
    ['PC Portable 2','[Sp\u00e9cifications \u00e0 compl\u00e9ter]'],
    ['Smartphone Android','Tests application mobile Flutter'],
    ['Connexion Internet','D\u00e9veloppement, tests API et d\u00e9ploiement'],
])
C("Tableau N\u00b04 : Ressources mat\u00e9rielles", 12, True, 8)

H2("M\u00e9thodologie de travail :")
B("Pour la r\u00e9alisation de ce projet, nous avons adopt\u00e9 une approche de d\u00e9veloppement agile it\u00e9rative, combinant les principes de la m\u00e9thode Scrum avec une architecture logicielle de type client-serveur \u00e0 trois tiers (three-tier architecture).")
B("Cette m\u00e9thodologie nous a permis de d\u00e9composer le projet en sprints de dur\u00e9e variable, chacun d\u00e9di\u00e9 \u00e0 un module fonctionnel sp\u00e9cifique, avec des livraisons incr\u00e9mentales et des revues r\u00e9guli\u00e8res avec notre encadreur.")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Cycle de d\u00e9veloppement Scrum utilis\u00e9 <<<]", it=True)

H3("Planning des sprints :")
T(['Sprint','Module','Dur\u00e9e'],[
    ['Sprint 1','Architecture API & Mod\u00e8le de donn\u00e9es','2 semaines'],
    ['Sprint 2','Authentification & Gestion des utilisateurs','1 semaine'],
    ['Sprint 3','Module Incidents & Triage automatique','2 semaines'],
    ['Sprint 4','Module Panic & G\u00e9olocalisation','1 semaine'],
    ['Sprint 5','Module Communautaire & Signaux','2 semaines'],
    ['Sprint 6','Module Wallet & Transactions','2 semaines'],
    ['Sprint 7','Application Web - Dashboard & Pages','3 semaines'],
    ['Sprint 8','Application Mobile Flutter','3 semaines'],
    ['Sprint 9','Tests, corrections & d\u00e9ploiement','2 semaines'],
])
C("Tableau N\u00b05 : Planning des sprints", 12, True, 8)

H2("Conclusion :")
B("Ce premier chapitre nous a permis de situer le contexte g\u00e9n\u00e9ral du projet TrustDesk, de pr\u00e9senter l\u2019organisme d\u2019accueil BEYN, d\u2019identifier clairement la probl\u00e9matique li\u00e9e \u00e0 la gestion des incidents de s\u00e9curit\u00e9 en milieu universitaire, et de d\u00e9finir les objectifs fonctionnels et techniques que cette plateforme vise \u00e0 atteindre.")
B("La solution propos\u00e9e se distingue par son approche multi-plateforme int\u00e9gr\u00e9e (web + mobile + API), son syst\u00e8me de r\u00e9ponse d\u2019urgence instantan\u00e9e via le bouton panique, son radar communautaire collaboratif, et son module de portefeuille \u00e9lectronique.")
B("Le chapitre suivant sera consacr\u00e9 \u00e0 l\u2019analyse et la sp\u00e9cification des besoins, o\u00f9 nous formaliserons les exigences fonctionnelles et non fonctionnelles du syst\u00e8me \u00e0 travers les outils de mod\u00e9lisation UML.")

doc.add_page_break()

# SAVE
out = r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v4.docx'
doc.save(out)
print(f'[OK] Chapter 1 added: {out} ({os.path.getsize(out)/1024:.1f} KB)')
