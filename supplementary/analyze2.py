from docx import Document

doc = Document(r'D:\TrustDesk\template.docx')

# Find chapter title pages and INTRODUCTION page
print('=== LOOKING FOR CHAPTER/INTRO PAGES ===')
for i, p in enumerate(doc.paragraphs[110:160]):
    idx = i + 110
    txt = p.text.strip()
    if txt:
        sz = round(p.runs[0].font.size / 12700, 1) if p.runs and p.runs[0].font.size else 'inh'
        bold = p.runs[0].font.bold if p.runs else '-'
        print(f'{idx:3d} | {p.style.name:20s} | sz={sz} bold={bold} align={p.alignment} | {txt[:90]}')

# The Suivi par positioning
print('\n=== REALISE PAR / SUIVI PAR ===')
for i in range(25, 30):
    p = doc.paragraphs[i]
    txt = p.text
    has_tab = chr(9) in txt
    indent = p.paragraph_format.left_indent
    print(f'P{i}: align={p.alignment} indent={indent} has_tab={has_tab} text="{txt}"')

# Check paragraph 28 runs (Suivi par uses tab)
p28 = doc.paragraphs[28]
for ri, r in enumerate(p28.runs):
    print(f'  Run {ri}: "{r.text}" sz={r.font.size} bold={r.font.bold}')

# Promotion
print('\n=== BOTTOM OF PAGE DE GARDE ===')
for i in range(30, 40):
    p = doc.paragraphs[i]
    txt = p.text.strip()
    if txt:
        print(f'P{i}: "{txt}"')

# Find section breaks
print('\n=== SECTIONS ===')
for i, s in enumerate(doc.sections[:10]):
    from docx.shared import Emu
    print(f'Section {i}: left={Emu(s.left_margin).cm:.2f}cm right={Emu(s.right_margin).cm:.2f}cm top={Emu(s.top_margin).cm:.2f}cm bottom={Emu(s.bottom_margin).cm:.2f}cm')
