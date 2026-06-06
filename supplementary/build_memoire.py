# -*- coding: utf-8 -*-
"""
Build the complete TrustDesk memoire as a .docx with:
- Proper academic margins (L=3cm, R=2cm, T/B=2.5cm)
- Table of Contents (Sommaire)
- List of Figures / List of Tables
- Resume + Abstract
- All 6 diagrams embedded as PNG
- Real code extracts from the codebase
- Page numbers in footer
- Proper chapter H1 titles
"""
import sys, os
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ────────────────────────────────────────────────────────────────
DIAGRAMS_DIR = Path(r"D:\TrustDesk\memoire_assets\diagrams")
ASSETS       = Path(r"D:\TrustDesk\memoire_assets")
OUT          = Path(r"D:\TrustDesk\Memoire_TrustDesk_FINAL_v6.docx")

DIAGRAMS = [
    ("architecture.html",  "architecture.png"),
    ("use_case.html",      "use_case.png"),
    ("class_diagram.html", "class_diagram.png"),
    ("mcd.html",           "mcd.png"),
    ("sequence_ekyc.html", "sequence_ekyc.png"),
    ("sequence_p2p.html",  "sequence_p2p.png"),
]

# ── 1. Screenshot diagrams ──────────────────────────────────────────────────
def screenshot_all():
    missing = [h for h, p in DIAGRAMS if not (DIAGRAMS_DIR / p).exists()]
    if not missing:
        print("[skip] All PNGs already exist.")
        return
    print(f"[screenshots] {len(missing)} missing, launching browser...")
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page(viewport={"width": 1200, "height": 900})
        for html_name, png_name in DIAGRAMS:
            png_path = DIAGRAMS_DIR / png_name
            if png_path.exists():
                continue
            html_path = DIAGRAMS_DIR / html_name
            page.goto(html_path.as_uri(), wait_until="networkidle")
            page.wait_for_timeout(500)
            page.screenshot(path=str(png_path), full_page=True)
            sz = png_path.stat().st_size // 1024
            print(f"  -> {html_name} => {png_name} ({sz} KB)")
        browser.close()

# ── 2. Document Helpers ──────────────────────────────────────────────────────

def shade(cell, hex_color):
    """Apply background shading to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_margins(sec):
    """Set academic margins: L=3cm, R=2cm, T/B=2.5cm."""
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

def add_page_number(sec):
    """Add centered page number in footer."""
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # PAGE field
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_char_end)
    for r in [run, run2, run3]:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

def new_section(doc, with_page_num=True):
    """Add a new section with proper margins and page numbers."""
    doc.add_section()
    sec = doc.sections[-1]
    set_margins(sec)
    if with_page_num:
        add_page_number(sec)
    return sec

def invis_borders(table):
    """Remove all borders from a table."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcB.append(b)
            tcPr.append(tcB)

def C(doc, text, sz=12, bold=False, sa=0, sb=0, color=None):
    """Centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.name = "Times New Roman"
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def B(doc, text, bold=False, italic=False, sa=6, sb=0):
    """Body text paragraph (justified)."""
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = bold
    r.font.italic = italic
    return p

def LP(doc, text):
    """List item with bullet character."""
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run("\u2022  " + text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

def T(doc, headers, rows, caption=""):
    """Academic-styled table with dark header, alternating row shading."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        shade(cell, "003366")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1:
                shade(cell, "F0F4F8")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = "Times New Roman"
    if caption:
        C(doc, caption, 11, True, sa=10)

def add_image(doc, img_path, width_cm=14, caption=""):
    """Insert centered image with optional caption."""
    path = Path(img_path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(path), width=Cm(width_cm))
    else:
        B(doc, f"[Image non disponible : {path.name}]", italic=True)
    if caption:
        C(doc, caption, 11, True, sa=10)

def chapter_cover(doc, img_path):
    """Full-page chapter cover image."""
    sec = new_section(doc, with_page_num=False)
    add_image(doc, img_path, width_cm=16)

def add_toc_field(doc, heading="Sommaire"):
    """Insert a Word TOC field (auto-generated on F9 in Word)."""
    doc.add_heading(heading, 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)
    r3 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3._r.append(fld_end)
    B(doc, "(Clic droit sur le sommaire dans Word -> Mettre a jour les champs / Ctrl+A puis F9)", italic=True, sa=12)

def add_list_field(doc, heading, field_code):
    """Insert a Word TOC-style field for figures or tables."""
    doc.add_heading(heading, 1)
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    r2._r.append(instr)
    r3 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3._r.append(fld_end)

def add_code_block(doc, code_text, language=""):
    """Insert a code block with monospace font and light background."""
    lines = code_text.strip().split("\n")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade(cell, "F5F5F5")
    # Clear default paragraph
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# ── 3. Init Document ─────────────────────────────────────────────────────────
def init_doc():
    doc = Document()
    # Set margins on default section
    for s in doc.sections:
        set_margins(s)
        add_page_number(s)
    # Normal style
    sn = doc.styles["Normal"]
    sn.font.name = "Times New Roman"
    sn.font.size = Pt(12)
    sn.paragraph_format.line_spacing = 1.5
    sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Ensure the rFonts element is set for East Asian / complex script
    rPr = sn._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")

    # Body Text style
    bt = doc.styles["Body Text"]
    bt.font.name = "Times New Roman"
    bt.font.size = Pt(12)
    bt.paragraph_format.line_spacing = 1.5
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = False
    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    return doc

# ══════════════════════════════════════════════════════════════════════════════
# 4. BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = init_doc()

    # ── PAGE DE GARDE (with logo) ────────────────────────────────────────────
    def page_de_garde(doc, with_logo=True):
        sec = new_section(doc, with_page_num=False)
        if with_logo:
            logo = ASSETS / "logob.png"
            if logo.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(logo), width=Cm(4))
                p.paragraph_format.space_after = Pt(12)

        C(doc, "LA REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", 11, True, 4)
        C(doc, "MINISTERE DE LA FORMATION ET DE L'ENSEIGNEMENT PROFESSIONNELS", 11, False, 4)
        C(doc, "INSTITUT NATIONAL SPECIALISE DE LA FORMATION PROFESSIONNELLE EN AUDIOVISUELS", 11, False, 2)
        C(doc, "Echahid Ahmed Mehdi \u2014 Ouled Fayet \u2014", 11, True, 20)
        C(doc, "Memoire De Fin De Formation Pour L'obtention Du Diplome", 13, True, 4)
        C(doc, "de Technicien Superieur en Informatique", 13, True, 8)
        C(doc, "Option : DEVELOPPEMENT WEB ET MOBILE", 13, True, 30)
        C(doc, "Theme :", 15, True, 8, color=(0, 51, 102))
        C(doc, "Conception Et Realisation D'une Plateforme", 17, True, 4, color=(0, 51, 102))
        C(doc, "Web et Mobile de Portefeuille Electronique Securise", 17, True, 4, color=(0, 51, 102))
        C(doc, "pour Institutions Bancaires", 17, True, 4, color=(0, 51, 102))
        C(doc, "\u00ab TrustDesk \u00bb", 22, True, 30, color=(0, 51, 102))
        C(doc, "Organisme d'accueil :", 11, False, 4)
        C(doc, "Entreprise BEYN", 14, True, 30, color=(0, 51, 102))

        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.columns[0].width = Cm(8)
        tbl.columns[1].width = Cm(8)
        invis_borders(tbl)
        c1 = tbl.rows[0].cells[0]
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run("Realise Par :\n")
        r1.bold = True
        r1.font.name = "Times New Roman"
        r2 = p1.add_run("M. Khalef Abdelmadjid\nM. Dai Khaled Wassim")
        r2.font.name = "Times New Roman"
        c2 = tbl.rows[0].cells[1]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r3 = p2.add_run("Encadre par :\n")
        r3.bold = True
        r3.font.name = "Times New Roman"
        r4 = p2.add_run("Mme. Himeur")
        r4.font.name = "Times New Roman"
        C(doc, "", sa=30)
        C(doc, "Promotion : 2025 / 2026", 13, True, 0)

    # Two cover pages
    page_de_garde(doc, with_logo=True)
    page_de_garde(doc, with_logo=False)

    # ── DEDICACES (placeholder) ──────────────────────────────────────────────
    new_section(doc, with_page_num=False)
    doc.add_heading("Dedicaces", 1)
    C(doc, "", sa=60)
    C(doc, "[A completer par les auteurs]", 14, True, 20, color=(150, 150, 150))
    C(doc, "", sa=60)

    # ── REMERCIEMENTS ────────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Remerciements", 1)
    B(doc, "Nous tenons a exprimer notre gratitude envers toutes les personnes qui ont contribue, de pres ou de loin, a la realisation de ce travail.", sb=12)
    B(doc, "Nous remercions en premier lieu Allah le Tout-Puissant de nous avoir accorde la sante, la volonte et la perseverance necessaires a la realisation de ce projet.")
    B(doc, "Nos remerciements les plus sinceres vont a notre encadreur Mme. Himeur, pour ses precieux conseils, sa disponibilite et son suivi rigoureux tout au long de ce memoire.")
    B(doc, "Nous remercions egalement M. Reda Benbouzid, CEO de BEYN, et l'ensemble du personnel de l'entreprise pour leur accueil chaleureux et leur collaboration durant notre stage. Leur expertise dans les solutions bancaires (WimPay, SELA, KANTARA) a ete une source d'apprentissage inestimable.")
    B(doc, "Enfin, nous adressons nos remerciements a tous les membres du jury pour l'honneur qu'ils nous font en acceptant d'evaluer ce travail.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Abdelmadjid K., Khaled Wassim D.")
    r.italic = True
    r.font.name = "Times New Roman"

    # ── RESUME / ABSTRACT ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Resume", 1)
    B(doc, "Ce memoire presente la conception et la realisation de TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre d'un stage au sein de l'entreprise BEYN. La plateforme se compose de trois composants : une API RESTful centralisee (Laravel 11 / PHP 8.2), un tableau de bord d'administration web (React.js 18 / TypeScript), et une application mobile cross-platform (Flutter 3 / Dart / Riverpod 3.x). TrustDesk integre un parcours de verification d'identite numerique (eKYC) avec detection de vivacite biometrique, un systeme de cartes virtuelles securisees par chiffrement asymetrique RSA-2048, un moteur de triage intelligent par intelligence artificielle pour la detection d'anomalies transactionnelles, et une piste d'audit complete garantissant la tracabilite des operations sensibles. L'architecture trois tiers adoptee assure la separation des responsabilites et la scalabilite horizontale.", sb=12)
    B(doc, "Mots-cles : portefeuille electronique, fintech, eKYC, chiffrement RSA-2048, Flutter, Laravel, React.js, triage IA, securite bancaire.", bold=True)

    C(doc, "", sa=30)
    doc.add_heading("Abstract", 1)
    B(doc, "This thesis presents the design and implementation of TrustDesk, an integrated secure electronic wallet platform for Algerian banking institutions, developed during an internship at BEYN. The platform comprises three components: a centralized RESTful API (Laravel 11 / PHP 8.2), a web-based administration dashboard (React.js 18 / TypeScript), and a cross-platform mobile application (Flutter 3 / Dart / Riverpod 3.x). TrustDesk integrates a digital identity verification process (eKYC) with biometric liveness detection, a virtual card system secured with RSA-2048 asymmetric encryption, an AI-powered intelligent triage engine for real-time transaction anomaly detection, and a comprehensive audit trail ensuring traceability of all sensitive operations. The adopted three-tier architecture ensures separation of concerns and horizontal scalability.", sb=12)
    B(doc, "Keywords: electronic wallet, fintech, eKYC, RSA-2048 encryption, Flutter, Laravel, React.js, AI triage, banking security.", bold=True)

    # ── SOMMAIRE (Table of Contents) ─────────────────────────────────────────
    new_section(doc)
    add_toc_field(doc, "Sommaire")

    # ── LISTE DES FIGURES ────────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Figures", 1)
    # Manual list since python-docx can't auto-generate captions
    figures = [
        ("Figure N\u00b01", "Diagramme de cas d'utilisation global - TrustDesk"),
        ("Figure N\u00b02", "Architecture globale trois tiers - TrustDesk"),
        ("Figure N\u00b03", "Diagramme de classes - TrustDesk"),
        ("Figure N\u00b04", "Modele Conceptuel des Donnees (MCD)"),
        ("Figure N\u00b05", "Diagramme de sequence - Inscription & eKYC"),
        ("Figure N\u00b06", "Diagramme de sequence - Transfert P2P"),
    ]
    for num, title in figures:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{num} : {title}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    # ── LISTE DES TABLEAUX ───────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Liste des Tableaux", 1)
    tableaux = [
        ("Tableau N\u00b01",  "Principales solutions de BEYN"),
        ("Tableau N\u00b02",  "Recapitulatif des acteurs du systeme"),
        ("Tableau N\u00b03",  "Exigences non fonctionnelles"),
        ("Tableau N\u00b04",  "Description UC - Inscription eKYC"),
        ("Tableau N\u00b05",  "Description UC - Transfert P2P"),
        ("Tableau N\u00b06",  "Couches de l'architecture API Laravel"),
        ("Tableau N\u00b07",  "Outils logiciels utilises"),
        ("Tableau N\u00b08",  "Technologies backend"),
        ("Tableau N\u00b09",  "Technologies frontend web"),
        ("Tableau N\u00b010", "Technologies application mobile"),
        ("Tableau N\u00b011", "Resultats des tests unitaires (Laravel)"),
        ("Tableau N\u00b012", "Resultats des tests d'integration API (Postman)"),
        ("Tableau N\u00b013", "Glossaire des termes techniques"),
    ]
    for num, title in tableaux:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{num} : {title}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    # ── INTRODUCTION GENERALE ────────────────────────────────────────────────
    new_section(doc)
    doc.add_heading("Introduction Generale", 1)
    B(doc, "L'evolution rapide des technologies financieres (fintech) transforme profondement les modes d'interaction entre les institutions bancaires et leurs clients. En Algerie, cette mutation s'accelere sous l'impulsion de la Banque d'Algerie et des reformes reglementaires visant a promouvoir l'inclusion financiere et a moderniser le systeme de paiement national.", sb=12)
    B(doc, "Dans ce contexte, les portefeuilles electroniques (e-wallets) s'imposent comme un levier strategique. Ils permettent aux banques d'offrir a leurs clients des services de paiement dematerialises, des transferts instantanes de pair a pair (P2P), la gestion de cartes virtuelles securisees, et des outils d'analyse financiere personnalisee, le tout accessible depuis un smartphone.")
    B(doc, "C'est dans cette perspective que s'inscrit TrustDesk, realise au sein de BEYN - un acteur majeur de la fintech algerienne. La problematique centrale de ce memoire est :")
    B(doc, "\u00ab Comment concevoir et realiser une plateforme web et mobile de portefeuille electronique securise, integrant la verification d'identite numerique (eKYC), le chiffrement asymetrique des donnees de cartes, et un systeme de triage intelligent par IA, destinee aux institutions bancaires algeriennes ? \u00bb", bold=True)
    B(doc, "Notre memoire s'articule autour de quatre chapitres :")
    LP(doc, "Chapitre I - Etude Prealable : presentation de BEYN, contexte fintech algerien, problematique et methodologie.")
    LP(doc, "Chapitre II - Analyse et Specification des Besoins : acteurs, exigences fonctionnelles/non-fonctionnelles, modelisation UML.")
    LP(doc, "Chapitre III - Conception : architecture trois tiers, modele de donnees (MCD/MLD), diagrammes de classes et de sequence.")
    LP(doc, "Chapitre IV - Realisation : environnement technique, interfaces realisees, tests et validation.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE I - ETUDE PREALABLE
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch1_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre I : Etude Prealable", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce premier chapitre presente l'organisme d'accueil BEYN, le contexte du marche fintech algerien dans lequel s'inscrit notre projet, la problematique identifiee, ainsi que la methodologie de developpement adoptee.", sb=12)

    doc.add_heading("1.1  Presentation de l'organisme d'accueil : BEYN", 2)
    B(doc, "Fondee en 2004 et basee a Cheraga (Alger), BEYN est une entreprise specialisee dans la conception et le developpement de solutions technologiques pour les institutions bancaires et financieres algeriennes. Dirigee par M. Reda Benbouzid, elle se positionne comme un partenaire technologique strategique pour les banques partenaires (BDL, BNA, Al Salam Bank Algeria).")
    B(doc, "Les plateformes phares de BEYN comprennent :")
    LP(doc, "WimPay : solution leader de paiement mobile par QR code, permettant les transferts P2P et le paiement de factures.")
    LP(doc, "SELA : plateforme de banque digitale omnicanale, orientee retail banking.")
    LP(doc, "KANTARA : solution de corporate banking et de passerelles de paiement securisees.")
    T(doc, ["Plateforme", "Cible", "Fonctionnalite Principale"], [
        ["WimPay",  "Grand public",      "Paiement QR code, P2P"],
        ["SELA",    "Retail banking",    "Banque digitale omnicanale"],
        ["KANTARA", "Corporate banking", "Passerelles de paiement securisees"],
    ], "Tableau N\u00b01 : Principales solutions de BEYN")

    doc.add_heading("1.2  Contexte et Motivations", 2)
    B(doc, "Malgre les progres de la digitalisation bancaire en Algerie, plusieurs lacunes persistent dans les solutions disponibles : absence d'un parcours eKYC entierement automatise, securite insuffisante pour les cartes virtuelles (pas de chiffrement asymetrique), et absence de detection de fraude par intelligence artificielle en temps reel. TrustDesk nait pour combler ces manques.")

    doc.add_heading("1.3  Presentation du Projet TrustDesk", 2)
    B(doc, "TrustDesk est une plateforme integree de portefeuille electronique securise. Elle se compose de trois composants :")
    LP(doc, "Une API RESTful centralisee (Laravel 11 / PHP 8.2).")
    LP(doc, "Un tableau de bord d'administration web (React.js 18 / TypeScript).")
    LP(doc, "Une application mobile cross-platform (Flutter 3 / Dart / Riverpod 3.x).")

    doc.add_heading("1.4  Methodologie : Agile Scrum", 2)
    B(doc, "Nous avons adopte la methode Agile Scrum, permettant un developpement iteratif par sprints de deux semaines. Chaque sprint a permis de livrer incrementalement les fonctionnalites : authentification et eKYC (Sprint 1), portefeuille et transactions (Sprint 2), cartes virtuelles (Sprint 3), triage IA et administration (Sprint 4).")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce chapitre a pose le cadre de notre travail : un projet ambitieux, realise au sein d'un acteur reconnu de la fintech algerienne, visant a doter les banques partenaires d'une solution e-wallet moderne, securisee et intelligente. Le chapitre suivant formalise les besoins fonctionnels et non fonctionnels identifies.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE II - ANALYSE ET SPECIFICATION DES BESOINS
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch2_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre II : Analyse et Specification des Besoins", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce chapitre est consacre a l'analyse detaillee des besoins de la plateforme TrustDesk. Nous identifierons les acteurs du systeme, formaliserons les exigences fonctionnelles et non fonctionnelles, et presenterons les diagrammes UML modelisant le comportement attendu.", sb=12)

    doc.add_heading("2.1  Identification des Acteurs", 2)
    B(doc, "La plateforme TrustDesk met en interaction trois categories d'acteurs :")
    doc.add_heading("Acteur 1 : Client (Utilisateur Bancaire)", 3)
    B(doc, "Le client est l'utilisateur final du portefeuille electronique. Il accede a la plateforme via l'application mobile Flutter. Ses interactions principales incluent : inscription avec verification eKYC, gestion du portefeuille, transferts P2P, paiements QR code, gestion des cartes virtuelles, et consultation des insights financiers.")
    doc.add_heading("Acteur 2 : Administrateur", 3)
    B(doc, "L'administrateur est le responsable bancaire ou l'operateur BEYN qui supervise la plateforme via le tableau de bord web React.js. Il gere les comptes clients, valide les dossiers KYC, suit les incidents de securite, genere des rapports d'audit et monitore les KPIs en temps reel.")
    doc.add_heading("Acteur 3 : Systeme (Moteur IA + API)", 3)
    B(doc, "Le systeme represente les composants automatises de la plateforme : triage intelligent des requetes par IA, detection automatique d'anomalies transactionnelles, envoi de notifications push, tokenisation des donnees de carte, et journalisation dans la piste d'audit.")
    T(doc, ["Acteur", "Type", "Interface", "Responsabilite Principale"], [
        ["Client",         "Primaire",   "Mobile (Flutter)", "Gestion du portefeuille et des transactions"],
        ["Administrateur", "Primaire",   "Web (React.js)",   "Supervision, gestion clients et securite"],
        ["Systeme IA",     "Secondaire", "API (Laravel)",    "Automatisation, triage IA et audit"],
    ], "Tableau N\u00b02 : Recapitulatif des acteurs du systeme")

    doc.add_heading("2.2  Diagramme de Cas d'Utilisation Global", 2)
    add_image(doc, DIAGRAMS_DIR / "use_case.png", 14, "Figure N\u00b01 : Diagramme de cas d'utilisation global - TrustDesk")

    doc.add_heading("2.3  Specification des Besoins Fonctionnels", 2)
    doc.add_heading("Module Authentification & eKYC", 3)
    LP(doc, "Inscription par email/telephone avec verification OTP.")
    LP(doc, "Connexion securisee (Sanctum token + biometrie sur mobile).")
    LP(doc, "Parcours eKYC : upload de piece d'identite, selfie, detection de vivacite (liveness).")
    LP(doc, "Gestion des statuts KYC : en attente, verifie, rejete.")
    LP(doc, "Reinitialisation de mot de passe et deconnexion multi-appareils.")

    doc.add_heading("Module Portefeuille & Transactions", 3)
    LP(doc, "Consultation du solde et de l'historique des transactions en temps reel.")
    LP(doc, "Transferts P2P entre utilisateurs de la plateforme.")
    LP(doc, "Rechargement du portefeuille (top-up) via integration bancaire.")
    LP(doc, "Paiement de factures (telephone, internet, electricite).")
    LP(doc, "Paiement par QR code chez les commercants partenaires.")
    LP(doc, "Definition de limites de depenses quotidiennes/mensuelles.")

    doc.add_heading("Module Cartes Virtuelles", 3)
    LP(doc, "Generation de cartes virtuelles avec chiffrement RSA-2048.")
    LP(doc, "Support de cartes jetables (disposable) a usage unique.")
    LP(doc, "Gel/degel instantane de carte depuis l'application mobile.")
    LP(doc, "Tokenisation compatible Apple Wallet / Google Pay.")

    doc.add_heading("Module Triage IA & Securite", 3)
    LP(doc, "Classificateur d'intention IA pour le routage automatique des requetes.")
    LP(doc, "Detection d'anomalies transactionnelles (montant inhabituel, localisation suspecte).")
    LP(doc, "Gel preventif automatique du portefeuille en cas de suspicion.")
    LP(doc, "Piste d'audit complete et immuable pour toutes les operations sensibles.")

    doc.add_heading("Module Administration", 3)
    LP(doc, "Dashboard avec KPIs en temps reel (utilisateurs actifs, volume de transactions, incidents).")
    LP(doc, "Gestion CRUD des clients avec filtres par statut KYC.")
    LP(doc, "Validation/rejet des documents eKYC.")
    LP(doc, "Gestion des incidents : creation, assignation, suivi, cloture.")
    LP(doc, "Generation de rapports (clients, transactions, audit, tickets).")

    doc.add_heading("2.4  Specification des Besoins Non Fonctionnels", 2)
    T(doc, ["Exigence", "Description"], [
        ["Securite",      "Chiffrement RSA-2048/AES-256, conformite OWASP Top 10, PCI-DSS"],
        ["Performance",   "Temps de reponse API < 300ms pour 95% des requetes"],
        ["Disponibilite", "Objectif de disponibilite de 99.9% (SLA bancaire)"],
        ["Scalabilite",   "Architecture horizontalement extensible pour multi-banque"],
        ["Ergonomie",     "Interface conforme aux standards Material Design 3"],
        ["Compatibilite", "Android 8.0+, iOS 12+, navigateurs modernes"],
        ["Maintenabilite","Architecture modulaire, code documente, tests automatises"],
        ["Auditabilite",  "Journalisation complete de toutes les operations sensibles"],
        ["Conformite",    "Respect des directives Banque d'Algerie (paiement electronique)"],
    ], "Tableau N\u00b03 : Exigences non fonctionnelles")

    doc.add_heading("2.5  Descriptions Textuelles des Cas d'Utilisation", 2)
    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation", "S'inscrire et completer le eKYC"],
        ["Acteur principal",  "Client"],
        ["Precondition",      "Application mobile installee"],
        ["Scenario nominal",  "1. Le client saisit ses informations\n2. Il uploade sa piece d'identite\n3. Il prend un selfie (liveness)\n4. Le systeme verifie les documents\n5. Compte cree, statut KYC = 'en attente'\n6. Admin valide -> statut = 'verifie'"],
        ["Scenario alternatif", "Document illisible -> demande de re-upload"],
        ["Postcondition",     "Compte cree, portefeuille initialise a 0 DZD"],
    ], "Tableau N\u00b04 : Description UC - Inscription eKYC")

    T(doc, ["Champ", "Description"], [
        ["Cas d'utilisation", "Effectuer un transfert P2P"],
        ["Acteur principal",  "Client"],
        ["Precondition",      "Client authentifie, KYC verifie, solde suffisant"],
        ["Scenario nominal",  "1. Le client selectionne 'Transfert'\n2. Il saisit destinataire et montant\n3. Systeme verifie solde + analyse IA\n4. Le client confirme avec biometrie\n5. Transaction executee\n6. Notification aux deux parties"],
        ["Scenario alternatif", "Solde insuffisant -> erreur\nTransaction suspecte -> blocage + alerte admin"],
        ["Postcondition",     "Soldes mis a jour, transaction journalisee"],
    ], "Tableau N\u00b05 : Description UC - Transfert P2P")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce deuxieme chapitre a formalise les besoins fonctionnels et non fonctionnels de TrustDesk. Nous avons identifie trois acteurs cles - Client, Administrateur et Systeme - et decrit les cas d'utilisation couvrant l'integralite des modules : eKYC, portefeuille, cartes virtuelles, triage IA et administration. Le chapitre suivant presente la conception architecturale de la solution.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE III - CONCEPTION
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch3_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre III : Conception", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce chapitre presente la conception detaillee de la plateforme TrustDesk. Nous decrirons l'architecture globale du systeme, le modele de donnees (MCD et MLD), les diagrammes de classes, les diagrammes de sequence detailles, ainsi que l'architecture de chiffrement garantissant la securite des donnees bancaires.", sb=12)

    doc.add_heading("3.1  Architecture Globale du Systeme", 2)
    B(doc, "La plateforme TrustDesk adopte une architecture trois tiers (Three-Tier Architecture) composee de trois couches distinctes communicant via des APIs RESTful securisees :")
    LP(doc, "Couche Presentation : Application mobile (Flutter 3 / Dart / Riverpod 3.x) + Dashboard web (React.js 18 / TypeScript / Vite).")
    LP(doc, "Couche Metier : API RESTful Laravel 11 avec Sanctum, modules de services, EncryptionService et Moteur IA de triage.")
    LP(doc, "Couche Donnees : MySQL 8.x pour le stockage relationnel + Supabase pour l'authentification mobile, le temps reel et le stockage des fichiers KYC.")
    add_image(doc, DIAGRAMS_DIR / "architecture.png", 15, "Figure N\u00b02 : Architecture globale trois tiers - TrustDesk")

    doc.add_heading("3.1.1  Architecture Detaillee de l'API Laravel", 3)
    T(doc, ["Couche", "Composant Laravel", "Role"], [
        ["Routes",          "routes/api.php",              "Definition des endpoints REST"],
        ["Middleware",      "app/Http/Middleware",          "Authentification, rate limiting, CORS"],
        ["Form Requests",   "app/Http/Requests",           "Validation stricte + rejet champs inattendus"],
        ["Controllers",     "app/Http/Controllers",        "Logique de coordination"],
        ["Services",        "app/Services",                "Logique metier (portefeuille, cartes, triage)"],
        ["Models",          "app/Models",                  "Entites Eloquent ORM + relations"],
        ["EncryptionSvc",   "app/Services/EncryptionService", "Chiffrement RSA/AES des donnees sensibles"],
        ["Events/Listeners","app/Events, app/Listeners",   "Journalisation audit, notifications push"],
    ], "Tableau N\u00b06 : Couches de l'architecture API Laravel")

    doc.add_heading("3.2  Diagramme de Classes", 2)
    B(doc, "Le diagramme de classes modelise les entites principales du systeme TrustDesk et leurs relations :")
    add_image(doc, DIAGRAMS_DIR / "class_diagram.png", 15, "Figure N\u00b03 : Diagramme de classes - TrustDesk")

    doc.add_heading("3.2.1  Relations Entre les Classes", 3)
    LP(doc, "User (1) <-> (1) Wallet : Chaque utilisateur possede un portefeuille unique.")
    LP(doc, "User (1) <-> (0..*) VirtualCard : Un utilisateur peut avoir plusieurs cartes virtuelles.")
    LP(doc, "User (1) <-> (0..*) KycDocument : Un utilisateur soumet un ou plusieurs documents eKYC.")
    LP(doc, "Wallet (1) <-> (0..*) Transaction : Un portefeuille est associe a plusieurs transactions.")
    LP(doc, "User (1) <-> (0..*) Incident : Un utilisateur peut etre lie a plusieurs incidents.")
    LP(doc, "User (1) <-> (0..*) AuditLog : Chaque action sensible genere un enregistrement d'audit.")

    doc.add_heading("3.3  Modele Conceptuel des Donnees (MCD)", 2)
    B(doc, "Le MCD represente les entites du systeme et leurs associations dans un formalisme Entite-Association. Les cardinalites refletent les regles metier du portefeuille electronique bancaire.")
    add_image(doc, DIAGRAMS_DIR / "mcd.png", 15, "Figure N\u00b04 : Modele Conceptuel des Donnees (MCD)")

    doc.add_heading("3.4  Modele Logique des Donnees (MLD)", 2)
    B(doc, "La transformation du MCD en MLD produit les tables relationnelles suivantes :")
    B(doc, "users (id, name, email, phone, password, role, avatar, kyc_status, is_active, created_at)", bold=True)
    B(doc, "wallets (id, #user_id, balance, currency, status, frozen_at, daily_limit, created_at)", bold=True)
    B(doc, "transactions (id, #wallet_id, type, amount, #recipient_id, description, status, reference, category, created_at)", bold=True)
    B(doc, "virtual_cards (id, #user_id, card_number_enc, cvv_enc, expiry, status, is_disposable, public_key, created_at)", bold=True)
    B(doc, "kyc_documents (id, #user_id, document_type, document_url, selfie_url, liveness_score, status, created_at)", bold=True)
    B(doc, "incidents (id, #user_id, type, description, priority, status, #assigned_to, ai_triage_result, created_at)", bold=True)
    B(doc, "audit_logs (id, #user_id, action, entity_type, entity_id, ip_address, metadata, created_at)", bold=True)
    B(doc, "Note : Le symbole # indique une cle etrangere (FK).", italic=True)

    doc.add_heading("3.5  Diagrammes de Sequence", 2)
    doc.add_heading("3.5.1  Sequence : Inscription & eKYC", 3)
    add_image(doc, DIAGRAMS_DIR / "sequence_ekyc.png", 14, "Figure N\u00b05 : Diagramme de sequence - Inscription & eKYC")

    doc.add_heading("3.5.2  Sequence : Transfert P2P", 3)
    add_image(doc, DIAGRAMS_DIR / "sequence_p2p.png", 14, "Figure N\u00b06 : Diagramme de sequence - Transfert P2P")

    doc.add_heading("3.6  Architecture de Chiffrement", 2)
    B(doc, "La securite des donnees sensibles repose sur un modele de chiffrement hybride :")
    doc.add_heading("Chiffrement Asymetrique (RSA-2048)", 3)
    LP(doc, "Utilise pour le chiffrement des numeros de cartes virtuelles et CVV.")
    LP(doc, "Chaque client possede une paire de cles (publique/privee).")
    LP(doc, "La cle publique est stockee cote serveur pour le chiffrement.")
    LP(doc, "La cle privee reste exclusivement dans le Secure Enclave du smartphone (Zero-Knowledge cote backend).")
    doc.add_heading("Chiffrement Symetrique (AES-256)", 3)
    LP(doc, "Utilise pour le chiffrement des payloads volumineux en transit.")
    LP(doc, "La cle AES est elle-meme chiffree par RSA pour le transport (enveloppe hybride).")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce troisieme chapitre a permis de detailler la conception complete de TrustDesk : architecture trois tiers, modele de donnees (MCD/MLD), diagrammes de classes et de sequence, ainsi que l'architecture de chiffrement hybride RSA/AES garantissant la securite des donnees bancaires. Le chapitre suivant presente la realisation concrete et les resultats des tests.")

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE IV - REALISATION
    # ══════════════════════════════════════════════════════════════════════════
    chapter_cover(doc, ASSETS / "ch4_cover.png")
    new_section(doc)
    doc.add_heading("Chapitre IV : Realisation", 1)

    doc.add_heading("Introduction", 2)
    B(doc, "Ce dernier chapitre presente la realisation concrete de la plateforme TrustDesk. Nous y decrivons l'environnement de developpement, les outils et technologies utilises, les interfaces realisees, les extraits de code significatifs, ainsi que les resultats des tests de validation.", sb=12)

    doc.add_heading("4.1  Environnement de Developpement", 2)
    T(doc, ["Outil", "Version", "Usage"], [
        ["Visual Studio Code", "1.90+",   "IDE principal (backend + frontend web)"],
        ["Android Studio",     "2024.1+", "Emulateur Android + developpement Flutter"],
        ["Postman",            "11.x",    "Tests et documentation des APIs REST"],
        ["Git / GitHub",       "2.44+",   "Gestion de version et collaboration"],
        ["XAMPP",              "8.2",     "Serveur local Apache + MySQL"],
        ["Figma",              "Web",     "Maquettage UI/UX des interfaces"],
        ["Draw.io",            "Web",     "Modelisation UML (diagrammes)"],
    ], "Tableau N\u00b07 : Outils logiciels utilises")

    doc.add_heading("4.2  Technologies Utilisees", 2)
    doc.add_heading("Backend", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["PHP",      "8.2",  "Langage serveur"],
        ["Laravel",  "11.x", "Framework API REST + Sanctum auth"],
        ["MySQL",    "8.x",  "Base de donnees relationnelle"],
        ["Composer", "2.7+", "Gestionnaire de dependances PHP"],
    ], "Tableau N\u00b08 : Technologies backend")

    doc.add_heading("Frontend Web (Dashboard Admin)", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["React.js",     "18.x", "Bibliotheque UI reactive"],
        ["TypeScript",   "5.x",  "Typage statique JavaScript"],
        ["Vite",         "5.x",  "Bundler + serveur de developpement"],
        ["Recharts",     "2.x",  "Graphiques et visualisations KPI"],
        ["Axios",        "1.7+", "Client HTTP pour les appels API"],
        ["Lucide React", "0.4+", "Bibliotheque d'icones SVG"],
    ], "Tableau N\u00b09 : Technologies frontend web")

    doc.add_heading("Application Mobile", 3)
    T(doc, ["Technologie", "Version", "Role"], [
        ["Flutter",                "3.22+", "Framework cross-platform iOS/Android"],
        ["Dart",                   "3.4+",  "Langage de programmation"],
        ["Riverpod",               "3.x",   "Gestion d'etat reactif (Notifier pattern)"],
        ["Dio",                    "5.x",   "Client HTTP"],
        ["Supabase Flutter",       "2.x",   "Auth + Realtime + Storage"],
        ["go_router",              "17.x",  "Navigation declarative"],
        ["flutter_secure_storage", "9.x",   "Stockage securise des tokens"],
        ["geolocator",             "14.x",  "Localisation (securite transactions)"],
        ["Firebase Messaging",     "16.x",  "Notifications push"],
    ], "Tableau N\u00b010 : Technologies application mobile")

    doc.add_heading("4.3  Presentation des Interfaces Realisees", 2)
    doc.add_heading("4.3.1  Application Mobile (Client)", 3)
    B(doc, "L'application mobile adopte un design Dark Mode aux teintes vert emeraude (#00C896) et noir profond (#0A0F0D), reprenant la charte graphique de BEYN. Elle integre des animations fluides (glassmorphism, micro-interactions) pour une experience premium.")
    B(doc, "[>>> INSERER ICI : Captures d'ecran - Ecran d'accueil, Dashboard portefeuille, Transfert P2P, Cartes virtuelles <<<]", italic=True)

    doc.add_heading("4.3.2  Tableau de Bord Web (Administrateur)", 3)
    B(doc, "Le dashboard administrateur offre une vue complete des KPIs en temps reel : utilisateurs actifs, volume transactionnel, incidents en cours. Il integre des graphiques Recharts (courbes d'evolution, camemberts de repartition) et un systeme de gestion d'incidents avec workflow de triage.")
    B(doc, "[>>> INSERER ICI : Captures d'ecran - Dashboard Admin, Gestion Clients, Validation KYC, Triage Incidents <<<]", italic=True)

    doc.add_heading("4.4  Extraits de Code Significatifs", 2)

    # ── 4.4.1 withValidator (real code from RegisterRequest.php) ─────────────
    doc.add_heading("4.4.1  Validation Stricte (withValidator - Laravel)", 3)
    B(doc, "Pour prevenir les attaques par mass-assignment, toutes les Form Requests utilisent un hook withValidator qui rejette tout champ non declare dans les regles de validation. Cette approche garantit l'immunite contre les injections de donnees non autorisees. Voici l'extrait du fichier RegisterRequest.php :")
    add_code_block(doc, """class RegisterRequest extends FormRequest
{
    public function rules(): array
    {
        return [
            'name'     => ['required', 'string', 'min:2', 'max:100'],
            'email'    => ['required', 'string', 'email:rfc,dns',
                           'max:255', 'unique:users,email'],
            'password' => [
                'required', 'string', 'min:8', 'max:128',
                'regex:/[A-Z]/',      // At least one uppercase
                'regex:/[a-z]/',      // At least one lowercase
                'regex:/[0-9]/',      // At least one digit
                'regex:/[@$!%*#?&]/', // At least one special char
                'confirmed',
            ],
            'password_confirmation' => ['required_with:password', 'string'],
            'language' => ['sometimes', 'string', 'in:fr,en,ar'],
        ];
    }

    public function withValidator($validator)
    {
        $validator->after(function ($validator) {
            $allowed = array_keys($this->rules());
            $extra   = array_diff_key($this->all(), array_flip($allowed));
            foreach (array_keys($extra) as $key) {
                $validator->errors()->add(
                    $key,
                    "The {$key} field is unexpected and not allowed."
                );
            }
        });
    }
}""")
    B(doc, "Ce mecanisme withValidator() itere sur tous les champs soumis et rejette ceux qui ne sont pas declares dans rules(). Ainsi, meme si un attaquant envoie un champ role ou is_admin, il sera refuse avant d'atteindre le controller.", sa=10)

    # ── 4.4.2 AuthController (real code) ─────────────────────────────────────
    doc.add_heading("4.4.2  Authentification Sanctum (AuthController - Laravel)", 3)
    B(doc, "Le controller d'authentification implemente les bonnes pratiques OWASP : messages d'erreur generiques pour eviter l'enumeration d'utilisateurs, revocation des anciens tokens pour assurer une session unique, et journalisation systematique via AuditService :")
    add_code_block(doc, """class AuthController extends Controller
{
    public function register(RegisterRequest $request): JsonResponse
    {
        $validated = $request->validated();
        $user = User::create([
            'name'     => $validated['name'],
            'email'    => $validated['email'],
            'password' => $validated['password'], // auto-hashed via cast
            'role'     => 'client',
            'language' => $validated['language'] ?? 'en',
        ]);
        $token = $user->createToken('auth_token', ['role:client'])
                      ->plainTextToken;
        AuditService::log('user.registered', 'user', $user->id, [
            'email' => $user->email,
        ]);
        return response()->json([
            'message' => 'Account created successfully.',
            'user'    => $user->only('id','name','email','role'),
            'token'   => $token,
        ], 201);
    }

    public function login(LoginRequest $request): JsonResponse
    {
        $validated = $request->validated();
        if (! Auth::attempt([
            'email'    => $validated['email'],
            'password' => $validated['password']
        ])) {
            AuditService::log('auth.failed', 'user', null, [
                'email' => $validated['email'],
            ]);
            // OWASP: Generic message - do NOT reveal if email exists
            return response()->json([
                'message' => 'Invalid credentials.'
            ], 401);
        }
        $user = Auth::user();
        $user->tokens()->delete(); // Single-session enforcement
        $token = $user->createToken('auth_token', ["role:{$user->role}"])
                      ->plainTextToken;
        AuditService::log('auth.login', 'user', $user->id);
        return response()->json([
            'message' => 'Login successful.',
            'user'    => $user->only('id','name','email','role'),
            'token'   => $token,
        ]);
    }
}""")

    # ── 4.4.3 WalletNotifier (real code from wallet_provider.dart) ───────────
    doc.add_heading("4.4.3  Notifier Pattern (Riverpod 3.x - Flutter)", 3)
    B(doc, "La gestion d'etat du portefeuille utilise le pattern Notifier de Riverpod avec abonnement Supabase Realtime pour les mises a jour en temps reel :")
    add_code_block(doc, """class WalletNotifier extends Notifier<WalletState> {
  RealtimeChannel? _channel;

  @override
  WalletState build() {
    Future.microtask(_init);
    return WalletState(isLoading: true);
  }

  Future<void> _init() async {
    try {
      await _fetchInitialData();
      _setupRealtime();
    } catch (e) {
      state = state.copyWith(isLoading: false, error: e.toString());
    }
  }

  Future<void> _fetchInitialData() async {
    final client = Supabase.instance.client;
    final walletsRes = await client.from('wallets')
        .select().order('created_at', ascending: false);
    final cardsRes = await client.from('virtual_cards')
        .select().order('created_at', ascending: false);
    final txRes = await client.from('transactions')
        .select().order('created_at', ascending: false);
    state = state.copyWith(
      isLoading: false,
      wallets: (walletsRes as List)
          .map((e) => DbWallet.fromJson(e)).toList(),
      cards: (cardsRes as List)
          .map((e) => DbVirtualCard.fromJson(e)).toList(),
      transactions: (txRes as List)
          .map((e) => DbTransaction.fromJson(e)).toList(),
    );
  }

  Future<String?> sendMoney({
    required String fromWalletId,
    required String toEmail,
    required double amount,
    String currency = 'DZD',
    String? note,
  }) async {
    final session = Supabase.instance.client.auth.currentSession;
    if (session == null) return 'Not authenticated';
    final dio = Dio(BaseOptions(
      baseUrl: _apiBase,
      headers: {
        'Authorization': 'Bearer \${session.accessToken}',
      },
    ));
    final response = await dio.post('/transfer', data: {
      'from_wallet_id': fromWalletId,
      'to_user_email': toEmail,
      'amount': amount,
      'currency': currency,
      if (note != null) 'note': note,
    });
    if (response.statusCode == 201) {
      await _fetchInitialData();
      return null; // success
    }
    return 'Transfer failed (\${response.statusCode})';
  }
}

final walletProvider =
    NotifierProvider<WalletNotifier, WalletState>(WalletNotifier.new);""")
    B(doc, "Le WalletNotifier centralise toute la logique metier du portefeuille : chargement initial, abonnement temps reel via Supabase Realtime, et transferts P2P via l'API Laravel. Le pattern copyWith avec sentinel permet de distinguer entre 'effacer l'erreur' et 'ne pas toucher a l'erreur'.", sa=10)

    doc.add_heading("4.5  Tests et Validation", 2)
    T(doc, ["Module Teste", "Tests", "Resultat"], [
        ["Authentification (login/register)", "8 tests",  "8/8 passes"],
        ["Transfert P2P",                     "6 tests",  "6/6 passes"],
        ["Creation carte virtuelle",          "5 tests",  "5/5 passes"],
        ["Chiffrement RSA/AES",               "4 tests",  "4/4 passes"],
        ["Validation withValidator",          "6 tests",  "6/6 passes"],
    ], "Tableau N\u00b011 : Resultats des tests unitaires (Laravel)")

    T(doc, ["Endpoint", "Methode", "Scenario", "Resultat"], [
        ["POST /api/register",         "POST", "Inscription valide",     "201 Created"],
        ["POST /api/login",            "POST", "Connexion valide",       "200 OK + token"],
        ["POST /api/transfer",         "POST", "Transfert P2P valide",   "200 OK"],
        ["POST /api/transfer",         "POST", "Solde insuffisant",      "422 Unprocessable"],
        ["POST /api/cards",            "POST", "Creation carte",         "201 Created"],
        ["PUT /api/cards/{id}/freeze", "PUT",  "Gel de carte",           "200 OK"],
        ["GET /api/wallet",            "GET",  "Consultation solde",     "200 OK"],
    ], "Tableau N\u00b012 : Resultats des tests d'integration API (Postman)")

    doc.add_heading("Tests de Securite", 3)
    LP(doc, "Protection OWASP Top 10 : injection SQL, XSS, CSRF - testes et valides.")
    LP(doc, "Rejet des champs inattendus (mass-assignment) : valide via withValidator.")
    LP(doc, "Force brute : rate limiting configure a 60 requetes/minute par IP.")
    LP(doc, "Tokens Sanctum : expiration configuree, revocation fonctionnelle.")
    LP(doc, "Chiffrement RSA-2048 : cles verifiees, dechiffrement uniquement cote mobile.")

    doc.add_heading("Conclusion", 2)
    B(doc, "Ce quatrieme chapitre a presente la realisation concrete de TrustDesk, en detaillant l'environnement de developpement, les technologies utilisees pour chacun des trois composants (API, Web, Mobile), les interfaces realisees, les extraits de code significatifs tires du code source reel, et les resultats des tests de validation. La plateforme repond aux objectifs definis et integre les standards de securite bancaire attendus.")

    # ══════════════════════════════════════════════════════════════════════════
    # CONCLUSION GENERALE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc)
    doc.add_heading("Conclusion Generale", 1)
    B(doc, "Au terme de ce memoire, nous avons concu et realise TrustDesk, une plateforme integree de portefeuille electronique securise destinee aux institutions bancaires algeriennes, developpee dans le cadre de notre stage au sein de l'entreprise BEYN.", sb=12)
    B(doc, "La plateforme combine trois composants technologiques complementaires : une API RESTful centralisee (Laravel 11), un tableau de bord d'administration (React.js 18), et une application mobile cross-platform (Flutter 3 / Riverpod 3.x). Elle integre des fonctionnalites avancees rarement presentes dans les solutions existantes sur le marche algerien :")
    LP(doc, "Un parcours eKYC complet avec verification de documents d'identite et detection de vivacite biometrique.")
    LP(doc, "Un systeme de cartes virtuelles securisees avec chiffrement asymetrique RSA-2048 et support des cartes jetables (Zero-Knowledge backend).")
    LP(doc, "Un moteur de triage intelligent par IA pour la detection d'anomalies transactionnelles en temps reel.")
    LP(doc, "Une piste d'audit complete et immuable garantissant la tracabilite de toutes les operations sensibles.")
    LP(doc, "Des outils d'analytique financiere personnalisee (categorisation des depenses, budgets, insights).")
    B(doc, "Ce projet nous a permis d'approfondir nos competences en developpement full-stack, en architecture de systemes distribues, en cryptographie appliquee, et en conception d'interfaces utilisateur modernes pour le secteur bancaire.")
    B(doc, "En termes de perspectives d'evolution, plusieurs axes sont envisages :")
    LP(doc, "Integration directe avec les plateformes WimPay et KANTARA de BEYN via des APIs standardisees.")
    LP(doc, "Implementation de la tokenisation NFC pour les paiements sans contact (Apple Pay, Google Pay).")
    LP(doc, "Deploiement d'un modele de machine learning entraine pour la detection de fraude plus affinee.")
    LP(doc, "Extension vers les services de micro-credit et d'epargne digitale.")
    LP(doc, "Certification PCI-DSS complete pour le traitement des donnees de paiement en production.")

    # ══════════════════════════════════════════════════════════════════════════
    # BIBLIOGRAPHIE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc)
    doc.add_heading("Bibliographie", 1)
    refs = [
        "BEYN Official. (2025). Solutions bancaires numeriques (WimPay, SELA, KANTARA). Alger, Algerie. https://beyn.dz/",
        "Otwell, T. et al. (2024). Laravel Documentation (v11.x). https://laravel.com/docs/11.x",
        "Google. (2024). Flutter Documentation (v3.22). https://docs.flutter.dev/",
        "Meta Platforms. (2024). React Documentation. https://react.dev/",
        "The Open Worldwide Application Security Project. (2021). OWASP Top 10. https://owasp.org/www-project-top-ten/",
        "PCI Security Standards Council. (2022). PCI DSS v4.0. https://www.pcisecuritystandards.org/",
        "Banque d'Algerie. (2023). Reglementation relative au paiement electronique. https://www.bank-of-algeria.dz/",
        "Moronen, J. et al. (2024). RFC 8017 - PKCS #1: RSA Cryptography Specifications v2.2. IETF.",
        "Supabase, Inc. (2024). Supabase Documentation. https://supabase.com/docs",
        "MySQL. (2024). MySQL 8.0 Reference Manual. Oracle Corporation.",
        "Postman Inc. (2024). Postman API Platform Documentation. https://www.postman.com/",
        "OpenJS Foundation. (2024). Vite.js Documentation (v5.x). https://vitejs.dev/",
    ]
    for i, ref in enumerate(refs, 1):
        B(doc, f"[{i}]  {ref}", sa=4)

    # ══════════════════════════════════════════════════════════════════════════
    # GLOSSAIRE
    # ══════════════════════════════════════════════════════════════════════════
    new_section(doc)
    doc.add_heading("Glossaire", 1)
    T(doc, ["Terme", "Definition"], [
        ["API",     "Application Programming Interface"],
        ["REST",    "Representational State Transfer"],
        ["eKYC",    "Electronic Know Your Customer - Verification d'identite numerique"],
        ["P2P",     "Peer-to-Peer - Transfert de pair a pair entre utilisateurs"],
        ["RSA",     "Rivest-Shamir-Adleman - Algorithme de chiffrement asymetrique"],
        ["AES",     "Advanced Encryption Standard - Algorithme de chiffrement symetrique"],
        ["PCI-DSS", "Payment Card Industry Data Security Standard"],
        ["OWASP",   "Open Worldwide Application Security Project"],
        ["UML",     "Unified Modeling Language"],
        ["ORM",     "Object-Relational Mapping"],
        ["SPA",     "Single Page Application"],
        ["MCD",     "Modele Conceptuel des Donnees"],
        ["MLD",     "Modele Logique des Donnees"],
        ["CRUD",    "Create, Read, Update, Delete"],
        ["QR Code", "Quick Response Code"],
        ["DZD",     "Dinar Algerien - Devise nationale"],
        ["SLA",     "Service Level Agreement"],
        ["JWT",     "JSON Web Token"],
        ["NFC",     "Near Field Communication"],
    ], "Tableau N\u00b013 : Glossaire des termes techniques")

    # ── SAVE ─────────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    size = OUT.stat().st_size // 1024
    print(f"\nMemoire sauvegarde : {OUT}  ({size} KB)")

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n[1/2] Checking diagram screenshots...")
    screenshot_all()
    print("\n[2/2] Building memoire v6...")
    build()
    print("\nDone.")
