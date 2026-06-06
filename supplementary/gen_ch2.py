# -*- coding: utf-8 -*-
"""Append Chapter 2 to the memoire"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v4.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b; r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa); p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.bold=b; r.font.italic=it
def E(): doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr(); numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}))
    pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(12); r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1); p.alignment=a
    for r in p.runs: r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.name='Times New Roman'
def new_section(header_text=""):
    doc.add_section()
    sec = doc.sections[-1]
    sec.top_margin=Cm(2.36); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2); sec.right_margin=Cm(1.25)
    if header_text:
        hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(header_text)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.font.italic=True

# ════════════════════════════════════════════════════════════════
# CHAPITRE II - GUARD PAGE
# ════════════════════════════════════════════════════════════════
new_section()
for _ in range(8): E()
C('CHAPITRE II :', 18, True, 6)
C('Analyse et Sp\u00e9cification des Besoins', 16, False, 8)
for _ in range(8): E()
doc.add_page_break()

# CHAPITRE II CONTENT
new_section("Analyse et Sp\u00e9cification des Besoins")

H1("Les concepts th\u00e9oriques :")
H2("Les d\u00e9finitions g\u00e9n\u00e9rales :")

B("Application web :", b=True)
B("Une application web est un logiciel applicatif h\u00e9berg\u00e9 sur un serveur et accessible via un navigateur web \u00e0 travers le protocole HTTP/HTTPS. Contrairement aux logiciels de bureau traditionnels qui n\u00e9cessitent une installation locale sur chaque poste, les applications web sont accessibles depuis n\u2019importe quel appareil disposant d\u2019une connexion Internet et d\u2019un navigateur compatible. Cette approche offre plusieurs avantages majeurs : d\u00e9ploiement centralis\u00e9, mises \u00e0 jour instantan\u00e9es sans intervention de l\u2019utilisateur, et accessibilit\u00e9 universelle ind\u00e9pendante du syst\u00e8me d\u2019exploitation.")
B("Dans le cas de TrustDesk, l\u2019application web constitue le module de supervision destin\u00e9 aux administrateurs. Elle a \u00e9t\u00e9 d\u00e9velopp\u00e9e sous forme de Single Page Application (SPA) avec React.js 18 et TypeScript, offrant une exp\u00e9rience utilisateur fluide et r\u00e9active sans rechargement complet de la page.")

B("Application mobile :", b=True)
B("Une application mobile est un logiciel sp\u00e9cifiquement con\u00e7u pour fonctionner sur des appareils mobiles tels que les smartphones et les tablettes. Les applications mobiles peuvent \u00eatre de trois types : natives (d\u00e9velopp\u00e9es sp\u00e9cifiquement pour un syst\u00e8me d\u2019exploitation), hybrides (utilisant des technologies web encapsul\u00e9es), ou cross-platform (partageant un code source unique pour plusieurs plateformes).")
B("Pour TrustDesk, nous avons opt\u00e9 pour une approche cross-platform avec Flutter 3 et le langage Dart, permettant de g\u00e9n\u00e9rer des applications natives performantes pour Android et iOS \u00e0 partir d\u2019un code source unique. Ce choix r\u00e9duit consid\u00e9rablement le temps de d\u00e9veloppement tout en garantissant des performances proches du natif gr\u00e2ce \u00e0 la compilation AOT (Ahead-Of-Time).")

B("API REST (Representational State Transfer) :", b=True)
B("Une API RESTful est une interface de programmation applicative qui respecte les contraintes architecturales du style REST, d\u00e9fini par Roy Fielding dans sa th\u00e8se doctorale en 2000. Elle permet \u00e0 diff\u00e9rentes applications de communiquer entre elles via le protocole HTTP en utilisant des m\u00e9thodes standard : GET (lecture), POST (cr\u00e9ation), PUT/PATCH (modification), et DELETE (suppression).")
B("Les principes fondamentaux d\u2019une API REST incluent : l\u2019architecture client-serveur, l\u2019absence d\u2019\u00e9tat (stateless), la mise en cache, l\u2019interface uniforme, et le syst\u00e8me en couches. Dans TrustDesk, l\u2019API REST d\u00e9velopp\u00e9e avec Laravel 11 constitue le point central de communication entre l\u2019application web et l\u2019application mobile.")

B("SOC (Security Operations Center) :", b=True)
B("Un Centre d\u2019Op\u00e9rations de S\u00e9curit\u00e9 (SOC) est une structure centralis\u00e9e d\u00e9di\u00e9e \u00e0 la surveillance continue, la d\u00e9tection, l\u2019analyse et la r\u00e9ponse aux incidents de s\u00e9curit\u00e9. Le concept de SOC, traditionnellement utilis\u00e9 dans le domaine de la cybers\u00e9curit\u00e9, a \u00e9t\u00e9 adapt\u00e9 dans TrustDesk pour couvrir la s\u00e9curit\u00e9 physique des campus universitaires.")

B("Architecture trois tiers (Three-Tier Architecture) :", b=True)
B("L\u2019architecture trois tiers est un mod\u00e8le d\u2019architecture logicielle qui s\u00e9pare l\u2019application en trois couches logiques distinctes : la couche de pr\u00e9sentation (interface utilisateur), la couche m\u00e9tier (logique applicative), et la couche de donn\u00e9es (stockage et persistance). Cette s\u00e9paration favorise la modularit\u00e9, la maintenabilit\u00e9 et l\u2019\u00e9volutivit\u00e9 du syst\u00e8me.")
B("Dans TrustDesk, cette architecture se traduit par : React.js/Flutter pour la couche de pr\u00e9sentation, Laravel 11 pour la couche m\u00e9tier, et MySQL 8.x pour la couche de donn\u00e9es.")

B("UML (Unified Modeling Language) :", b=True)
B("UML est un langage de mod\u00e9lisation graphique standardis\u00e9 par l\u2019Object Management Group (OMG), utilis\u00e9 pour sp\u00e9cifier, visualiser, construire et documenter les artefacts d\u2019un syst\u00e8me logiciel. Il fournit un ensemble de diagrammes permettant de repr\u00e9senter les diff\u00e9rents aspects d\u2019un syst\u00e8me : structurels (diagramme de classes, de composants) et comportementaux (diagramme de cas d\u2019utilisation, de s\u00e9quence, d\u2019activit\u00e9).")

B("RBAC (Role-Based Access Control) :", b=True)
B("Le contr\u00f4le d\u2019acc\u00e8s bas\u00e9 sur les r\u00f4les est un mod\u00e8le de s\u00e9curit\u00e9 dans lequel les autorisations d\u2019acc\u00e8s sont attribu\u00e9es aux utilisateurs en fonction de leur r\u00f4le au sein de l\u2019organisation. Dans TrustDesk, trois r\u00f4les principaux sont d\u00e9finis : administrateur (acc\u00e8s complet au dashboard web), agent de s\u00e9curit\u00e9 (acc\u00e8s terrain via mobile), et utilisateur standard (\u00e9tudiant/personnel avec signalement d\u2019incidents).")

H2("Identification des acteurs :")
B("Un acteur repr\u00e9sente une entit\u00e9 externe qui interagit avec le syst\u00e8me. Dans le cadre de TrustDesk, nous avons identifi\u00e9 les acteurs suivants :")
T(['Acteur','Description','Interface'],[
    ['Administrateur','Supervise les incidents, g\u00e8re les utilisateurs, consulte les statistiques','Application Web'],
    ['Agent de s\u00e9curit\u00e9','Re\u00e7oit les alertes, intervient sur le terrain, d\u00e9clenche le bouton panique','Application Mobile'],
    ['Utilisateur (Etudiant)','Signale des incidents, utilise le radar communautaire, bouton panique','Application Mobile'],
    ['Syst\u00e8me (API)','Traite les requ\u00eates, calcule les priorit\u00e9s, envoie les notifications','Backend Laravel'],
])
C("Tableau N\u00b06 : Identification des acteurs du syst\u00e8me", 12, True, 8)

H2("Besoins fonctionnels :")
B("Les besoins fonctionnels d\u00e9crivent les services et les fonctionnalit\u00e9s que le syst\u00e8me doit fournir. Nous les avons organis\u00e9s par module :")

H3("Module Authentification :")
LP("Inscription avec validation des donn\u00e9es (nom, email, mot de passe complexe, r\u00f4le).")
LP("Connexion s\u00e9curis\u00e9e avec g\u00e9n\u00e9ration de token Sanctum.")
LP("D\u00e9connexion avec r\u00e9vocation du token d\u2019acc\u00e8s.")
LP("Gestion des sessions avec expiration automatique.")

H3("Module Gestion des Incidents :")
LP("Cr\u00e9ation d\u2019un incident avec titre, description, cat\u00e9gorie, localisation et niveau d\u2019urgence.")
LP("Calcul automatique du score de priorit\u00e9 bas\u00e9 sur la s\u00e9v\u00e9rit\u00e9, l\u2019impact et l\u2019urgence.")
LP("Assignation des incidents aux agents de s\u00e9curit\u00e9 comp\u00e9tents.")
LP("Suivi du statut en temps r\u00e9el : ouvert, en cours, r\u00e9solu, cl\u00f4tur\u00e9.")
LP("Historique complet des actions effectu\u00e9es sur chaque incident.")

H3("Module Bouton Panique :")
LP("D\u00e9clenchement d\u2019une alerte d\u2019urgence en un seul geste depuis l\u2019application mobile.")
LP("Capture automatique des coordonn\u00e9es GPS au moment du d\u00e9clenchement.")
LP("Notification imm\u00e9diate de tous les agents de s\u00e9curit\u00e9 connect\u00e9s.")
LP("Activation automatique des restrictions d\u2019acc\u00e8s au campus si n\u00e9cessaire.")
LP("R\u00e9solution de l\u2019alerte panique avec rapport d\u2019intervention.")

H3("Module Radar Communautaire :")
LP("Signalement d\u2019activit\u00e9s suspectes avec description et g\u00e9olocalisation.")
LP("Affichage des signalements sur une carte interactive.")
LP("Validation et mod\u00e9ration des signalements par les administrateurs.")

H3("Module Portefeuille \u00c9lectronique :")
LP("Consultation du solde en temps r\u00e9el.")
LP("Historique des transactions (cr\u00e9dits et d\u00e9bits).")
LP("Gel et d\u00e9gel de fonds par les administrateurs.")
LP("Suivi des op\u00e9rations financi\u00e8res li\u00e9es aux services de s\u00e9curit\u00e9.")

H2("Besoins non fonctionnels :")
B("Les besoins non fonctionnels d\u00e9finissent les contraintes de qualit\u00e9 du syst\u00e8me :")
LP("Performance : Temps de r\u00e9ponse API inf\u00e9rieur \u00e0 500ms pour 95% des requ\u00eates.")
LP("S\u00e9curit\u00e9 : Authentification par token Sanctum, validation stricte des entr\u00e9es, rate limiting, protection OWASP Top 10.")
LP("Disponibilit\u00e9 : Le syst\u00e8me doit \u00eatre disponible 24h/24 et 7j/7.")
LP("Ergonomie : Interface intuitive respectant les principes de Material Design.")
LP("Scalabilit\u00e9 : L\u2019architecture doit supporter une mont\u00e9e en charge progressive.")
LP("Maintenabilit\u00e9 : Code source document\u00e9, architecture modulaire, tests automatis\u00e9s.")
LP("Compatibilit\u00e9 : Support multi-navigateurs et multi-plateformes mobiles.")

H2("Diagramme de cas d\u2019utilisation :")
B("Le diagramme de cas d\u2019utilisation permet de repr\u00e9senter les interactions entre les acteurs et le syst\u00e8me. Il identifie les fonctionnalit\u00e9s principales offertes par la plateforme TrustDesk \u00e0 chacun de ses utilisateurs.")
B("[>>> INS\u00c9RER ICI : Figure N\u00b02 \u2013 Diagramme de cas d\u2019utilisation g\u00e9n\u00e9ral de TrustDesk <<<]", it=True)
C("Figure N\u00b02 : Diagramme de cas d\u2019utilisation g\u00e9n\u00e9ral", 12, True, 8)
B("Le diagramme ci-dessus illustre les principaux cas d\u2019utilisation du syst\u00e8me TrustDesk, r\u00e9partis entre les trois acteurs identifi\u00e9s :")
LP("L\u2019administrateur peut : g\u00e9rer les utilisateurs, superviser les incidents, consulter les statistiques, configurer le syst\u00e8me, et g\u00e9rer le portefeuille \u00e9lectronique.")
LP("L\u2019agent de s\u00e9curit\u00e9 peut : recevoir les alertes, d\u00e9clencher le bouton panique, mettre \u00e0 jour le statut des incidents, et contribuer au radar communautaire.")
LP("L\u2019utilisateur (\u00e9tudiant) peut : signaler un incident, utiliser le bouton panique, et consulter le radar communautaire.")

H2("Diagrammes de s\u00e9quence :")
B("Les diagrammes de s\u00e9quence repr\u00e9sentent les interactions entre les objets du syst\u00e8me dans un ordre chronologique. Ils permettent de visualiser le flux de messages \u00e9chang\u00e9s lors de l\u2019ex\u00e9cution d\u2019un sc\u00e9nario sp\u00e9cifique.")

H3("Diagramme de s\u00e9quence : Authentification")
B("Ce diagramme illustre le processus d\u2019authentification d\u2019un utilisateur sur la plateforme TrustDesk. L\u2019utilisateur saisit ses identifiants (email et mot de passe), le client (web ou mobile) envoie une requ\u00eate POST \u00e0 l\u2019API, qui valide les donn\u00e9es, v\u00e9rifie les identifiants dans la base de donn\u00e9es, et retourne un token Sanctum en cas de succ\u00e8s.")
B("[>>> INS\u00c9RER ICI : Figure N\u00b03 \u2013 Diagramme de s\u00e9quence \u2013 Authentification <<<]", it=True)
C("Figure N\u00b03 : Diagramme de s\u00e9quence \u2013 Authentification", 12, True, 8)

H3("Diagramme de s\u00e9quence : D\u00e9clenchement du bouton panique")
B("Ce diagramme d\u00e9taille le flux d\u2019ex\u00e9cution lors du d\u00e9clenchement du bouton panique depuis l\u2019application mobile. L\u2019agent appuie sur le bouton panique, l\u2019application capture automatiquement les coordonn\u00e9es GPS, envoie une requ\u00eate POST \u00e0 l\u2019API avec la localisation, l\u2019API cr\u00e9e l\u2019alerte, active les restrictions campus si n\u00e9cessaire, et notifie tous les agents connect\u00e9s.")
B("[>>> INS\u00c9RER ICI : Figure N\u00b04 \u2013 Diagramme de s\u00e9quence \u2013 Bouton Panique <<<]", it=True)
C("Figure N\u00b04 : Diagramme de s\u00e9quence \u2013 Bouton Panique", 12, True, 8)

H3("Diagramme de s\u00e9quence : Signalement d\u2019incident")
B("Ce diagramme repr\u00e9sente le processus de signalement d\u2019un incident par un utilisateur. L\u2019utilisateur remplit le formulaire de signalement (titre, description, cat\u00e9gorie, localisation), le client envoie la requ\u00eate \u00e0 l\u2019API, qui valide les donn\u00e9es, calcule le score de priorit\u00e9, enregistre l\u2019incident dans la base de donn\u00e9es, et notifie les agents concern\u00e9s.")
B("[>>> INS\u00c9RER ICI : Figure N\u00b05 \u2013 Diagramme de s\u00e9quence \u2013 Signalement d\u2019incident <<<]", it=True)
C("Figure N\u00b05 : Diagramme de s\u00e9quence \u2013 Signalement d\u2019incident", 12, True, 8)

H2("Outils et technologies utilis\u00e9s :")
B("Le tableau suivant pr\u00e9sente l\u2019ensemble des technologies et outils utilis\u00e9s pour le d\u00e9veloppement de la plateforme TrustDesk :")
T(['Composante','Technologie','Version','R\u00f4le'],[
    ['Backend','Laravel (PHP)','11.x','Framework API REST'],
    ['Authentification','Laravel Sanctum','4.x','Tokens API s\u00e9curis\u00e9s'],
    ['Base de donn\u00e9es','MySQL','8.x','Stockage relationnel'],
    ['Frontend Web','React.js','18.x','Interface SPA'],
    ['Typage','TypeScript','5.x','Typage statique JavaScript'],
    ['Bundler','Vite','5.x','Build & serveur de d\u00e9veloppement'],
    ['Routing Web','React Router','6.x','Navigation SPA'],
    ['Mobile','Flutter','3.x','Application cross-platform'],
    ['Langage Mobile','Dart','3.x','Langage de programmation Flutter'],
    ['\u00c9tat Mobile','Riverpod','2.x','Gestion d\u2019\u00e9tat r\u00e9active'],
    ['Navigation','GoRouter','14.x','Navigation d\u00e9clarative Flutter'],
    ['HTTP Client','Dio','5.x','Requ\u00eates HTTP avanc\u00e9es'],
    ['G\u00e9olocalisation','Geolocator','13.x','Capture coordonn\u00e9es GPS'],
    ['IDE','VS Code','Latest','\u00c9diteur de code principal'],
    ['VCS','Git + GitHub','Latest','Contr\u00f4le de version'],
    ['Serveur local','Laragon','Latest','Environnement de d\u00e9veloppement'],
])
C("Tableau N\u00b07 : Technologies et outils utilis\u00e9s", 12, True, 8)

H2("Conclusion :")
B("Ce deuxi\u00e8me chapitre nous a permis de d\u00e9finir les concepts th\u00e9oriques fondamentaux sous-jacents au projet TrustDesk, d\u2019identifier les acteurs du syst\u00e8me et de formaliser les besoins fonctionnels et non fonctionnels de la plateforme.")
B("La mod\u00e9lisation UML, \u00e0 travers les diagrammes de cas d\u2019utilisation et de s\u00e9quence, a permis de repr\u00e9senter de mani\u00e8re formelle les interactions entre les diff\u00e9rents acteurs et le syst\u00e8me, ainsi que les flux d\u2019ex\u00e9cution des principaux sc\u00e9narios d\u2019utilisation.")
B("Le chapitre suivant sera consacr\u00e9 \u00e0 la conception d\u00e9taill\u00e9e de l\u2019application, o\u00f9 nous pr\u00e9senterons l\u2019architecture globale du syst\u00e8me, le mod\u00e8le de donn\u00e9es, et les choix de conception technique retenus.")

doc.add_page_break()

# SAVE
out = r'D:\TrustDesk\Memoire_TrustDesk_FINAL_v4.docx'
doc.save(out)
print(f'[OK] Chapter 2 added: {out} ({os.path.getsize(out)/1024:.1f} KB)')
