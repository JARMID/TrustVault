# -*- coding: utf-8 -*-
"""Chapter 3 - Conception de l'Application (Wallet)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx')

def C(t,sz=12,b=False,sa=0):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(t);r.font.size=Pt(sz);r.font.bold=b;r.font.name='Times New Roman'
def B(t,b=False,it=False,sa=6):
    p=doc.add_paragraph(style='Body Text');p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(sa);p.paragraph_format.line_spacing=1.5
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.bold=b;r.font.italic=it
def E():doc.add_paragraph(style='Body Text')
def LP(t):
    p=doc.add_paragraph(style='List Paragraph');p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12)
    pPr=p._p.get_or_add_pPr();numPr=pPr.makeelement(qn('w:numPr'),{})
    numPr.append(pPr.makeelement(qn('w:ilvl'),{qn('w:val'):'0'}))
    numPr.append(pPr.makeelement(qn('w:numId'),{qn('w:val'):'1'}));pPr.append(numPr)
def T(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h);r.font.bold=True;r.font.size=Pt(12);r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci];c.text='';p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val));r.font.size=Pt(12);r.font.name='Times New Roman'
def H1(t,a=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_heading(t,level=1);p.alignment=a
    for r in p.runs:r.font.name='Times New Roman'
def H2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs:r.font.name='Times New Roman'
def H3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs:r.font.name='Times New Roman'

# ═══════ CHAPTER COVER ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)
for _ in range(8):E()
C('CHAPITRE III :',18,True,6)
C('Conception de l\u2019Application',16,False,8)
for _ in range(8):E()
doc.add_page_break()

# ═══════ CONTENT ═══════
doc.add_section()
s=doc.sections[-1];s.top_margin=Cm(2.36);s.bottom_margin=Cm(2);s.left_margin=Cm(2);s.right_margin=Cm(1.25)

H1("Introduction :")
B("Ce chapitre pr\u00e9sente la conception d\u00e9taill\u00e9e de la plateforme TrustDesk. Nous y d\u00e9crirons l\u2019architecture globale du syst\u00e8me, le mod\u00e8le de donn\u00e9es relationnel, les diagrammes de classes, ainsi que les diagrammes de s\u00e9quence d\u00e9taill\u00e9s illustrant les interactions entre les diff\u00e9rents composants.")

# ─── ARCHITECTURE ───
H1("Architecture globale du syst\u00e8me :")
B("La plateforme TrustDesk adopte une architecture trois tiers (three-tier architecture) compos\u00e9e de trois couches distinctes, communiquant via des APIs RESTful s\u00e9curis\u00e9es :")

H2("Couche Pr\u00e9sentation (Frontend) :")
LP("Application Web (React.js 18 + TypeScript + Vite) : Tableau de bord d\u2019administration destin\u00e9 aux administrateurs bancaires.")
LP("Application Mobile (Flutter 3 / Dart) : Portefeuille \u00e9lectronique destin\u00e9 aux clients bancaires.")
B("Ces deux interfaces communiquent exclusivement avec la couche m\u00e9tier via des appels API REST authentifi\u00e9s par tokens Sanctum.")

H2("Couche M\u00e9tier (Backend API) :")
LP("API RESTful d\u00e9velopp\u00e9e avec Laravel 11 (PHP 8.2).")
LP("Authentification par Laravel Sanctum (tokens bearer).")
LP("Modules m\u00e9tier : gestion des portefeuilles, transactions, cartes virtuelles, eKYC, triage IA, audit.")
LP("Middleware de s\u00e9curit\u00e9 : validation stricte des entr\u00e9es (withValidator), rate limiting, CORS.")
LP("Chiffrement RSA/AES pour les donn\u00e9es sensibles (num\u00e9ros de carte, CVV).")

H2("Couche Donn\u00e9es (Base de donn\u00e9es) :")
LP("MySQL 8.x pour le stockage relationnel.")
LP("Migrations Laravel pour le versionnement du sch\u00e9ma.")
LP("Indexation optimis\u00e9e pour les requ\u00eates transactionnelles \u00e0 haute fr\u00e9quence.")

B("[>>> INS\u00c9RER ICI : Figure \u2013 Architecture globale trois tiers de TrustDesk <<<]", it=True)
C("Figure N\u00b06 : Architecture globale du syst\u00e8me TrustDesk", 12, True, 8)

# ─── ARCHITECTURE DETAILLEE ───
H2("Architecture d\u00e9taill\u00e9e de l\u2019API :")
B("L\u2019API Laravel suit le pattern MVC (Model-View-Controller) enrichi de couches suppl\u00e9mentaires pour la s\u00e9curit\u00e9 et la modularit\u00e9 :")

T(['Couche','Composant Laravel','R\u00f4le'],[
    ['Routes','routes/api.php','D\u00e9finition des endpoints REST'],
    ['Middleware','app/Http/Middleware','Authentification, rate limiting, CORS'],
    ['Form Requests','app/Http/Requests','Validation stricte + rejet des champs inattendus'],
    ['Controllers','app/Http/Controllers','Logique de coordination'],
    ['Services','app/Services','Logique m\u00e9tier (portefeuille, cartes, triage)'],
    ['Models','app/Models','Entit\u00e9s Eloquent ORM + relations'],
    ['Encryption','app/Services/EncryptionService','Chiffrement RSA/AES des donn\u00e9es sensibles'],
    ['Events/Listeners','app/Events, app/Listeners','Journalisation audit, notifications'],
])
C("Tableau N\u00b013 : Couches de l\u2019architecture API Laravel", 12, True, 8)

# ─── DIAGRAMME DE CLASSES ───
H1("Diagramme de classes :")
B("Le diagramme de classes mod\u00e9lise les entit\u00e9s principales du syst\u00e8me TrustDesk et leurs relations. Les classes principales sont :")

H2("Classe User :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['name','string','Nom complet'],
    ['email','string','Adresse email (unique)'],
    ['phone','string','Num\u00e9ro de t\u00e9l\u00e9phone'],
    ['password','string','Mot de passe hash\u00e9 (bcrypt)'],
    ['role','enum','admin | client'],
    ['avatar','string','URL de l\u2019avatar'],
    ['kyc_status','enum','pending | verified | rejected'],
    ['is_active','boolean','Statut du compte'],
])
C("Tableau N\u00b014 : Attributs de la classe User", 12, True, 8)

H2("Classe Wallet :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['user_id','bigint','R\u00e9f\u00e9rence utilisateur (FK)'],
    ['balance','decimal(15,2)','Solde courant en DZD'],
    ['currency','string','Devise (DZD par d\u00e9faut)'],
    ['status','enum','active | frozen | suspended'],
    ['frozen_at','timestamp','Date de gel (nullable)'],
    ['daily_limit','decimal','Plafond journalier'],
])
C("Tableau N\u00b015 : Attributs de la classe Wallet", 12, True, 8)

H2("Classe Transaction :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['wallet_id','bigint','Portefeuille source (FK)'],
    ['type','enum','p2p | topup | payment | bill | qr'],
    ['amount','decimal(15,2)','Montant de la transaction'],
    ['recipient_id','bigint','Destinataire (FK, nullable)'],
    ['description','text','Libell\u00e9 de la transaction'],
    ['status','enum','pending | completed | failed | reversed'],
    ['reference','string','R\u00e9f\u00e9rence unique (UUID)'],
    ['category','string','Cat\u00e9gorie (alimentation, transport, etc.)'],
])
C("Tableau N\u00b016 : Attributs de la classe Transaction", 12, True, 8)

H2("Classe VirtualCard :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['user_id','bigint','Propri\u00e9taire (FK)'],
    ['card_number_enc','text','Num\u00e9ro chiffr\u00e9 RSA'],
    ['cvv_enc','text','CVV chiffr\u00e9 RSA'],
    ['expiry','date','Date d\u2019expiration'],
    ['status','enum','active | frozen | expired | revoked'],
    ['is_disposable','boolean','Carte jetable (usage unique)'],
    ['public_key','text','Cl\u00e9 publique associ\u00e9e'],
])
C("Tableau N\u00b017 : Attributs de la classe VirtualCard", 12, True, 8)

H2("Classe KycDocument :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['user_id','bigint','Client concern\u00e9 (FK)'],
    ['document_type','enum','cni | passport | permis'],
    ['document_url','string','URL du document stock\u00e9'],
    ['selfie_url','string','URL du selfie'],
    ['liveness_score','float','Score de d\u00e9tection de vivacit\u00e9'],
    ['status','enum','pending | approved | rejected'],
])
C("Tableau N\u00b018 : Attributs de la classe KycDocument", 12, True, 8)

H2("Classe Incident :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['user_id','bigint','Client concern\u00e9 (FK)'],
    ['type','enum','fraud | suspicious | complaint | system'],
    ['description','text','Description de l\u2019incident'],
    ['priority','enum','low | medium | high | critical'],
    ['status','enum','open | in_progress | resolved | closed'],
    ['assigned_to','bigint','Admin assign\u00e9 (FK, nullable)'],
    ['ai_triage_result','json','R\u00e9sultat du triage IA'],
])
C("Tableau N\u00b019 : Attributs de la classe Incident", 12, True, 8)

H2("Classe AuditLog :")
T(['Attribut','Type','Description'],[
    ['id','bigint','Identifiant unique (PK)'],
    ['user_id','bigint','Auteur de l\u2019action (FK)'],
    ['action','string','Type d\u2019action (login, freeze, transfer, etc.)'],
    ['entity_type','string','Type d\u2019entit\u00e9 concern\u00e9e'],
    ['entity_id','bigint','ID de l\u2019entit\u00e9'],
    ['ip_address','string','Adresse IP source'],
    ['metadata','json','D\u00e9tails suppl\u00e9mentaires'],
])
C("Tableau N\u00b020 : Attributs de la classe AuditLog", 12, True, 8)

B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de classes UML complet <<<]", it=True)
C("Figure N\u00b07 : Diagramme de classes du syst\u00e8me TrustDesk", 12, True, 8)

# ─── RELATIONS ───
H2("Relations entre les classes :")
LP("User (1) \u2194 (1) Wallet : Chaque utilisateur poss\u00e8de un portefeuille unique.")
LP("User (1) \u2194 (0..*) VirtualCard : Un utilisateur peut avoir plusieurs cartes virtuelles.")
LP("User (1) \u2194 (0..*) KycDocument : Un utilisateur soumet un ou plusieurs documents eKYC.")
LP("Wallet (1) \u2194 (0..*) Transaction : Un portefeuille est associ\u00e9 \u00e0 plusieurs transactions.")
LP("User (1) \u2194 (0..*) Incident : Un utilisateur peut \u00eatre li\u00e9 \u00e0 plusieurs incidents.")
LP("User (1) \u2194 (0..*) AuditLog : Chaque action sensible g\u00e9n\u00e8re un enregistrement d\u2019audit.")
LP("User (1) \u2194 (0..*) Notification : Un utilisateur re\u00e7oit des notifications.")

# ─── MCD ───
H1("Mod\u00e8le Conceptuel de Donn\u00e9es (MCD) :")
B("Le MCD repr\u00e9sente les entit\u00e9s du syst\u00e8me et leurs associations dans un formalisme Entit\u00e9-Association. Les cardinalit\u00e9s refl\u00e8tent les r\u00e8gles m\u00e9tier du portefeuille \u00e9lectronique bancaire.")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Mod\u00e8le Conceptuel de Donn\u00e9es (MCD) <<<]", it=True)
C("Figure N\u00b08 : Mod\u00e8le Conceptuel de Donn\u00e9es", 12, True, 8)

# ─── MLD ───
H1("Mod\u00e8le Logique de Donn\u00e9es (MLD) :")
B("La transformation du MCD en MLD produit les tables relationnelles suivantes :")
B("users (id, name, email, phone, password, role, avatar, kyc_status, is_active, created_at, updated_at)", b=True)
B("wallets (id, #user_id, balance, currency, status, frozen_at, daily_limit, created_at, updated_at)", b=True)
B("transactions (id, #wallet_id, type, amount, #recipient_id, description, status, reference, category, created_at)", b=True)
B("virtual_cards (id, #user_id, card_number_enc, cvv_enc, expiry, status, is_disposable, public_key, created_at)", b=True)
B("kyc_documents (id, #user_id, document_type, document_url, selfie_url, liveness_score, status, created_at)", b=True)
B("incidents (id, #user_id, type, description, priority, status, #assigned_to, ai_triage_result, created_at)", b=True)
B("audit_logs (id, #user_id, action, entity_type, entity_id, ip_address, metadata, created_at)", b=True)
B("notifications (id, #user_id, title, body, type, is_read, created_at)", b=True)
B("Note : Le symbole # indique une cl\u00e9 \u00e9trang\u00e8re (FK).", it=True)

# ─── SEQUENCE DIAGRAMS DETAILLES ───
H1("Diagrammes de s\u00e9quence d\u00e9taill\u00e9s :")

H2("S\u00e9quence : Authentification et cr\u00e9ation de session :")
B("1. Le client saisit ses identifiants (email + mot de passe) dans l\u2019application mobile.")
B("2. L\u2019application envoie une requ\u00eate POST /api/login au backend Laravel.")
B("3. Le LoginRequest valide les champs et rejette les entr\u00e9es inattendues (withValidator).")
B("4. Le AuthController v\u00e9rifie les identifiants via Auth::attempt().")
B("5. En cas de succ\u00e8s, Sanctum g\u00e9n\u00e8re un token bearer retourn\u00e9 au client.")
B("6. L\u2019\u00e9v\u00e9nement est journalis\u00e9 dans audit_logs (action: \u2018login\u2019).")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de s\u00e9quence \u2014 Authentification <<<]", it=True)
C("Figure N\u00b09 : Diagramme de s\u00e9quence \u2014 Authentification", 12, True, 8)

H2("S\u00e9quence : Cr\u00e9ation d\u2019une carte virtuelle chiffr\u00e9e :")
B("1. Le client authentifi\u00e9 demande la cr\u00e9ation d\u2019une carte virtuelle.")
B("2. L\u2019API v\u00e9rifie le statut KYC (doit \u00eatre \u2018verified\u2019).")
B("3. Le VirtualCardService g\u00e9n\u00e8re un num\u00e9ro de carte et un CVV al\u00e9atoires.")
B("4. L\u2019EncryptionService chiffre le num\u00e9ro et le CVV avec la cl\u00e9 publique RSA du client.")
B("5. La carte chiffr\u00e9e est persist\u00e9e en base de donn\u00e9es.")
B("6. Le client re\u00e7oit les d\u00e9tails d\u00e9chiffr\u00e9s uniquement c\u00f4t\u00e9 mobile (cl\u00e9 priv\u00e9e locale).")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Diagramme de s\u00e9quence \u2014 Cr\u00e9ation carte virtuelle <<<]", it=True)
C("Figure N\u00b010 : Diagramme de s\u00e9quence \u2014 Cr\u00e9ation carte virtuelle", 12, True, 8)

# ─── CHIFFREMENT ───
H1("Architecture de chiffrement :")
B("La s\u00e9curit\u00e9 des donn\u00e9es sensibles de la plateforme TrustDesk repose sur un mod\u00e8le de chiffrement hybride RSA/AES :")
H2("Chiffrement asym\u00e9trique (RSA-2048) :")
LP("Utilis\u00e9 pour le chiffrement des donn\u00e9es de cartes virtuelles (num\u00e9ro, CVV).")
LP("Chaque client poss\u00e8de une paire de cl\u00e9s (publique/priv\u00e9e).")
LP("La cl\u00e9 publique est stock\u00e9e c\u00f4t\u00e9 serveur pour le chiffrement.")
LP("La cl\u00e9 priv\u00e9e reste exclusivement sur l\u2019appareil mobile du client.")
H2("Chiffrement sym\u00e9trique (AES-256) :")
LP("Utilis\u00e9 pour le chiffrement des payloads volumineux en transit.")
LP("La cl\u00e9 AES est chiffr\u00e9e par RSA pour le transport s\u00e9curis\u00e9 (enveloppe hybride).")
B("[>>> INS\u00c9RER ICI : Figure \u2013 Sch\u00e9ma du mod\u00e8le de chiffrement hybride RSA/AES <<<]", it=True)
C("Figure N\u00b011 : Mod\u00e8le de chiffrement hybride RSA/AES", 12, True, 8)

# ─── CONCLUSION ───
H2("Conclusion :")
B("Ce troisi\u00e8me chapitre a permis de d\u00e9tailler la conception compl\u00e8te de la plateforme TrustDesk. Nous avons pr\u00e9sent\u00e9 l\u2019architecture trois tiers, le mod\u00e8le de donn\u00e9es (MCD et MLD), les diagrammes de classes avec leurs attributs d\u00e9taill\u00e9s, les diagrammes de s\u00e9quence illustrant les flux cl\u00e9s, ainsi que l\u2019architecture de chiffrement hybride RSA/AES garantissant la s\u00e9curit\u00e9 des donn\u00e9es bancaires.")
B("Le chapitre suivant sera consacr\u00e9 \u00e0 la r\u00e9alisation et l\u2019impl\u00e9mentation, o\u00f9 nous pr\u00e9senterons les outils utilis\u00e9s, les interfaces d\u00e9velopp\u00e9es et les r\u00e9sultats des tests de validation.")

doc.add_page_break()
out=r'D:\TrustDesk\Memoire_TrustDesk_Wallet.docx'
doc.save(out)
print(f'[OK] Ch1-3 saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
