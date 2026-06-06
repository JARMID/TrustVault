"""Pass 2: Fix names + rewrite Ch1 content for TrustDesk security platform"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx')
ns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# === Fix ALL text everywhere ===
replacements = {
    'M.Machat Wassim': 'M.Khalef Abdelmadjid',
    'Machat Wassim': 'Khalef Abdelmadjid',
    'M.Benbehaz Houssam Elddine': 'M.Dai Khaled Wassim',
    'Benbehaz Houssam Elddine': 'Dai Khaled Wassim',
    'Benbehaz Houssam': 'Dai Khaled Wassim',
    'Wassim M.': 'Khalef A.',
    'Houssam Elddine B.': 'Dai K.W.',
    'Wassim M. , Houssam Elddine B.': 'Khalef A. , Dai K.W.',
    'caisses r\u00e9gionales': 'incidents de s\u00e9curit\u00e9',
    'caisse r\u00e9gionale': 'incident de s\u00e9curit\u00e9',
    'Caisse R\u00e9gionale': 'Incident de S\u00e9curit\u00e9',
    'op\u00e9rations financi\u00e8res': 'incidents de s\u00e9curit\u00e9',
}

for p in doc.paragraphs:
    for r in p.runs:
        if r.text:
            for old, new in replacements.items():
                if old in r.text:
                    r.text = r.text.replace(old, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text:
                        for old, new in replacements.items():
                            if old in r.text:
                                r.text = r.text.replace(old, new)

# === Rewrite Introduction content (P118-P132) ===
intro_texts = [
    "Depuis toujours, les \u00e9tablissements d\u2019enseignement sup\u00e9rieur cherchent \u00e0 am\u00e9liorer leurs m\u00e9thodes de gestion et de suivi de la s\u00e9curit\u00e9 sur leurs campus. La diversit\u00e9 des acteurs, l\u2019\u00e9tendue g\u00e9ographique des sites universitaires et la nature vari\u00e9e des incidents pouvant survenir repr\u00e9sentent des d\u00e9fis majeurs pour les responsables de la s\u00e9curit\u00e9. Les m\u00e9thodes traditionnelles de signalement, bas\u00e9es sur des registres papier, des appels t\u00e9l\u00e9phoniques ou des d\u00e9placements physiques, montrent leurs limites face aux exigences croissantes de r\u00e9activit\u00e9 et de tra\u00e7abilit\u00e9.",
    "Avec l\u2019essor des technologies de l\u2019information et de la communication, les \u00e9tablissements d\u2019enseignement sup\u00e9rieur disposent aujourd\u2019hui d\u2019outils puissants pour transformer leurs processus de gestion de la s\u00e9curit\u00e9. Les frameworks de d\u00e9veloppement modernes, les architectures API-first et les applications mobiles offrent des possibilit\u00e9s de centralisation et de r\u00e9activit\u00e9 sans pr\u00e9c\u00e9dent, permettant une coordination efficace entre les diff\u00e9rents intervenants.",
    "Par ailleurs, la g\u00e9n\u00e9ralisation de l\u2019Internet et le d\u00e9veloppement des technologies mobiles ont profond\u00e9ment modifi\u00e9 les attentes des usagers en mati\u00e8re de r\u00e9activit\u00e9 et d\u2019accessibilit\u00e9 des services. Un \u00e9tudiant confront\u00e9 \u00e0 une situation d\u2019urgence doit pouvoir signaler un incident instantan\u00e9ment depuis son smartphone, tandis que les agents de s\u00e9curit\u00e9 doivent pouvoir recevoir et traiter ces alertes en temps r\u00e9el.",
    "Dans ce contexte, la gestion des incidents de s\u00e9curit\u00e9 constitue un aspect particuli\u00e8rement sensible dans les universit\u00e9s. Les processus actuels, souvent fragment\u00e9s et non num\u00e9ris\u00e9s, ne permettent pas d\u2019assurer une tra\u00e7abilit\u00e9 ad\u00e9quate ni une coordination efficace entre les diff\u00e9rents services concern\u00e9s (administration, s\u00e9curit\u00e9, maintenance).",
    "C\u2019est dans cette perspective que s\u2019inscrit le projet TrustDesk, qui vise \u00e0 concevoir et r\u00e9aliser une plateforme int\u00e9gr\u00e9e de type Security Operations Center (SOC), compos\u00e9e d\u2019une application web de supervision, d\u2019une application mobile de terrain et d\u2019une API RESTful centralis\u00e9e, d\u00e9di\u00e9e \u00e0 la gestion compl\u00e8te du cycle de vie des incidents de s\u00e9curit\u00e9 en milieu universitaire.",
    "Afin de r\u00e9pondre \u00e0 ce besoin, notre projet consiste \u00e0 concevoir et r\u00e9aliser une plateforme web et mobile qui permettra de centraliser le signalement, le triage, l\u2019assignation, le suivi et la cl\u00f4ture des incidents, tout en offrant un m\u00e9canisme de bouton panique avec g\u00e9olocalisation et un radar communautaire collaboratif.",
    "Ainsi, la r\u00e9alisation de ce projet nous am\u00e8ne \u00e0 nous poser la probl\u00e9matique suivante :",
    "\u00ab Comment concevoir et r\u00e9aliser une plateforme web et mobile centralis\u00e9e permettant d\u2019assurer une gestion compl\u00e8te, s\u00e9curis\u00e9e et tra\u00e7able du cycle de vie des incidents de s\u00e9curit\u00e9 en milieu universitaire ? \u00bb",
    "Pour r\u00e9pondre \u00e0 cette probl\u00e9matique, nous avons structur\u00e9 notre m\u00e9moire en quatre chapitres.",
    "Le premier chapitre, intitul\u00e9 \u00ab \u00c9tude pr\u00e9alable \u00bb, pr\u00e9sente le contexte du projet, l\u2019organisme d\u2019accueil, la probl\u00e9matique identifi\u00e9e, les objectifs vis\u00e9s ainsi que la m\u00e9thodologie de travail adopt\u00e9e.",
    "Le deuxi\u00e8me chapitre, intitul\u00e9 \u00ab Analyse et sp\u00e9cification des besoins \u00bb, est consacr\u00e9 \u00e0 l\u2019identification et \u00e0 la formalisation des exigences fonctionnelles et non fonctionnelles du syst\u00e8me \u00e0 travers les cas d\u2019utilisation UML.",
    None,  # skip P130
    "Le troisi\u00e8me chapitre, intitul\u00e9 \u00ab Conception de l\u2019application \u00bb, pr\u00e9sente la conception globale du syst\u00e8me \u00e0 travers les diagrammes UML, l\u2019architecture logicielle et le mod\u00e8le de donn\u00e9es.",
    "Le quatri\u00e8me chapitre, intitul\u00e9 \u00ab R\u00e9alisation et impl\u00e9mentation \u00bb, d\u00e9crit les diff\u00e9rentes \u00e9tapes du d\u00e9veloppement, les interfaces r\u00e9alis\u00e9es et les r\u00e9sultats des tests de validation.",
    "Enfin, nous terminerons ce m\u00e9moire par une conclusion g\u00e9n\u00e9rale, dans laquelle nous pr\u00e9senterons le bilan du travail r\u00e9alis\u00e9 et les perspectives d\u2019am\u00e9lioration envisag\u00e9es.",
]

# Replace intro paragraphs (P118-P132)
intro_start = 118
for idx, new_text in enumerate(intro_texts):
    pi = intro_start + idx
    if pi < len(doc.paragraphs) and new_text is not None:
        p = doc.paragraphs[pi]
        # Clear existing runs and add new text
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = new_text
            # Make the problematique bold
            if new_text.startswith('\u00ab'):
                p.runs[0].font.bold = True

# === Rewrite Chapter 1 content ===
# P154: Heading "Presentation de TrustDesk :"
# P155: Body text about TrustDesk
ch1_rewrites = {
    154: ("Pr\u00e9sentation du projet :", True),
    155: ("TrustDesk est une plateforme int\u00e9gr\u00e9e de type Security Operations Center (SOC), sp\u00e9cialement con\u00e7ue pour la gestion des incidents de s\u00e9curit\u00e9 en milieu universitaire. Elle se compose de trois modules compl\u00e9mentaires : une API RESTful backend d\u00e9velopp\u00e9e avec Laravel 11, une application web de supervision construite avec React.js 18 et TypeScript, et une application mobile de terrain r\u00e9alis\u00e9e avec Flutter 3. Cette solution multi-plateforme permet de centraliser l\u2019ensemble du cycle de vie des incidents, depuis le signalement initial jusqu\u2019\u00e0 la cl\u00f4ture, en passant par le triage, l\u2019assignation et le suivi en temps r\u00e9el. TrustDesk int\u00e8gre \u00e9galement un m\u00e9canisme innovant de bouton panique avec g\u00e9olocalisation GPS, ainsi qu\u2019un radar communautaire collaboratif permettant aux usagers du campus de signaler les activit\u00e9s suspectes.", False),
    156: ("Fiche Technique du projet :", True),
}

for pi, (text, is_heading) in ch1_rewrites.items():
    if pi < len(doc.paragraphs):
        p = doc.paragraphs[pi]
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = text

# Rewrite body paragraphs for Chapter 1 (P241 onwards - Presentation du sujet)
ch1_body = {
    241: "Pr\u00e9sentation du sujet :",
    242: "Le syst\u00e8me de gestion des incidents de s\u00e9curit\u00e9 TrustDesk consiste en une plateforme int\u00e9gr\u00e9e compos\u00e9e d\u2019une application web, d\u2019une application mobile et d\u2019une API RESTful centralis\u00e9e. Cette solution est destin\u00e9e aux \u00e9tablissements d\u2019enseignement sup\u00e9rieur souhaitant moderniser et optimiser leur gestion de la s\u00e9curit\u00e9 sur campus.",
    243: "L\u2019application web constitue le c\u0153ur du syst\u00e8me de supervision et permet aux administrateurs de g\u00e9rer l\u2019ensemble des incidents signal\u00e9s, de visualiser les statistiques en temps r\u00e9el via un tableau de bord interactif, de g\u00e9rer les utilisateurs et leurs r\u00f4les, et de configurer les param\u00e8tres du syst\u00e8me. Le dashboard int\u00e8gre des indicateurs cl\u00e9s de performance (KPIs) tels que le nombre total d\u2019incidents, le temps moyen de r\u00e9solution, le taux de r\u00e9solution et la r\u00e9partition par cat\u00e9gorie.",
    244: "En compl\u00e9ment, l\u2019application mobile permet aux agents de s\u00e9curit\u00e9 et aux usagers du campus d\u2019acc\u00e9der aux fonctionnalit\u00e9s essentielles directement depuis leur terminal mobile. Elle offre notamment un bouton panique permettant de d\u00e9clencher une alerte d\u2019urgence en un seul geste avec capture automatique des coordonn\u00e9es GPS, un radar communautaire pour signaler les activit\u00e9s suspectes avec g\u00e9olocalisation, et un syst\u00e8me de notifications en temps r\u00e9el pour les alertes entrantes.",
    245: "Gr\u00e2ce \u00e0 cette solution centralis\u00e9e et multi-plateforme, les \u00e9tablissements universitaires disposent d\u2019un outil performant et s\u00e9curis\u00e9 pour la gestion compl\u00e8te du cycle de vie des incidents de s\u00e9curit\u00e9, depuis le signalement initial jusqu\u2019\u00e0 la cl\u00f4ture, en passant par le triage, l\u2019assignation, le suivi et la r\u00e9solution.",
    247: "Probl\u00e9matique :",
    248: "Face aux difficult\u00e9s de gestion de la s\u00e9curit\u00e9 auxquelles les campus universitaires alg\u00e9riens sont confront\u00e9s, plusieurs insuffisances majeures ont \u00e9t\u00e9 identifi\u00e9es : l\u2019absence de centralisation des signalements qui transitent par des canaux h\u00e9t\u00e9rog\u00e8nes (appels t\u00e9l\u00e9phoniques, d\u00e9placements physiques, e-mails), le temps de r\u00e9ponse \u00e9lev\u00e9 en l\u2019absence d\u2019outils d\u2019alerte instantan\u00e9e, le manque de tra\u00e7abilit\u00e9 des incidents non document\u00e9s num\u00e9riquement, et le cloisonnement entre les diff\u00e9rents acteurs (\u00e9tudiants, personnel de s\u00e9curit\u00e9, administration).",
    249: "Ces contraintes engendrent des retards dans le traitement des situations d\u2019urgence, des pertes d\u2019informations critiques et une incapacit\u00e9 \u00e0 analyser statistiquement les tendances en mati\u00e8re de s\u00e9curit\u00e9 pour am\u00e9liorer la pr\u00e9vention.",
    250: "D\u2019o\u00f9 notre probl\u00e9matique :",
    251: "\u00ab Comment concevoir et r\u00e9aliser une plateforme web et mobile int\u00e9gr\u00e9e, s\u00e9curis\u00e9e et performante, permettant la gestion compl\u00e8te du cycle de vie des incidents de s\u00e9curit\u00e9 dans un environnement universitaire ? \u00bb",
    252: "Les objectifs :",
    253: "Objectifs g\u00e9n\u00e9raux :",
    254: "Nous devons concevoir et d\u00e9velopper une plateforme (web et mobile) pour la gestion centralis\u00e9e des incidents de s\u00e9curit\u00e9 en milieu universitaire, int\u00e9grant les fonctionnalit\u00e9s suivantes :",
    255: "Mettre en \u0153uvre un syst\u00e8me de gestion centralis\u00e9e des incidents (signalement, triage, assignation, suivi, r\u00e9solution, cl\u00f4ture).",
    256: "Impl\u00e9menter un m\u00e9canisme de bouton panique instantan\u00e9 avec g\u00e9olocalisation GPS depuis l\u2019application mobile.",
    257: "Mettre en place un syst\u00e8me d\u2019authentification s\u00e9curis\u00e9 bas\u00e9 sur les r\u00f4les (administrateur, agent de s\u00e9curit\u00e9, utilisateur) avec Laravel Sanctum.",
    258: "Assurer les notifications en temps r\u00e9el pour alerter les agents de s\u00e9curit\u00e9 concern\u00e9s lors de chaque nouveau signalement ou alerte panique.",
    259: "Mettre en place un radar communautaire permettant aux usagers de signaler les activit\u00e9s suspectes avec g\u00e9olocalisation.",
    260: "Objectifs Espace Administrateur :",
    261: "Superviser l\u2019ensemble des incidents d\u00e9clar\u00e9s sur la plateforme via un tableau de bord interactif.",
    262: "G\u00e9rer les utilisateurs, les r\u00f4les et les permissions d\u2019acc\u00e8s.",
    263: "Consulter les statistiques globales et les indicateurs cl\u00e9s de performance (KPIs).",
    264: "Configurer les cat\u00e9gories d\u2019incidents, les niveaux de priorit\u00e9 et les workflows de triage.",
    265: "Objectifs Espace Agent de terrain (Mobile) :",
    266: "Recevoir les alertes d\u2019incidents en temps r\u00e9el avec g\u00e9olocalisation.",
    267: "D\u00e9clencher un bouton de panique en situation d\u2019urgence avec capture GPS automatique.",
    268: "Les Contraintes :",
    269: "L\u2019application doit \u00eatre enti\u00e8rement responsive, s\u2019adaptant aux diff\u00e9rents types d\u2019\u00e9crans (desktop, tablette, mobile).",
    270: "Mise en place d\u2019un syst\u00e8me d\u2019authentification s\u00e9curis\u00e9 (Laravel Sanctum) et de chiffrement des donn\u00e9es sensibles.",
    271: "Temps de r\u00e9ponse API inf\u00e9rieur \u00e0 500ms pour 95% des requ\u00eates.",
    272: "L\u2019application doit \u00eatre compatible avec les navigateurs modernes (Chrome, Firefox, Edge, Safari) et les terminaux mobiles Android 8.0+ et iOS 12+.",
    273: "Le public vis\u00e9 :",
    274: "Les administrateurs universitaires, les agents de s\u00e9curit\u00e9, le personnel enseignant et administratif, ainsi que les \u00e9tudiants de l\u2019\u00e9tablissement.",
}

for pi, text in ch1_body.items():
    if pi < len(doc.paragraphs):
        p = doc.paragraphs[pi]
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = text
            if text.startswith('\u00ab'):
                p.runs[0].font.bold = True

# Save
out = r'D:\TrustDesk\Memoire_TrustDesk_CH1.docx'
try:
    doc.save(out)
except PermissionError:
    out = r'D:\TrustDesk\Memoire_TrustDesk_CH1_v2.docx'
    doc.save(out)
print(f'[OK] Saved: {out} ({os.path.getsize(out)/1024:.1f} KB)')
