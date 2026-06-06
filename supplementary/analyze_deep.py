"""
Strategy: Use template.docx as base, strip its content, replace with TrustDesk content.
This preserves all text boxes, dark backgrounds, shapes, and exact formatting.
"""
from docx import Document
from lxml import etree
import os, copy

ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'v': 'urn:schemas-microsoft-com:vml',
}

doc = Document(r'D:\TrustDesk\template.docx')
body = doc.element.body

# Find all text in AlternateContent boxes to understand what they contain
alts = body.findall('.//mc:AlternateContent', ns)
print(f"Total text boxes: {len(alts)}")

# Map text box contents
for ai, alt in enumerate(alts):
    texts = alt.findall('.//w:t', ns)
    combined = ''.join(t.text or '' for t in texts)
    safe = combined.encode('ascii', 'replace').decode()
    print(f"  Box {ai}: '{safe[:100]}'")

# Count paragraphs
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# Find what sections exist
sectPrs = body.findall('.//w:sectPr', ns)
print(f"Sections: {len(sectPrs)}")

# Find section break positions (pPr/sectPr means end of section in the para)
print("\n=== SECTION BREAKS (in paragraphs) ===")
for i, p in enumerate(doc.paragraphs):
    pPr = p._p.find('w:pPr', ns)
    if pPr is not None:
        sect = pPr.find('w:sectPr', ns)
        if sect is not None:
            print(f"  Section break at P{i}: '{p.text[:50]}'")

# Find page breaks
print("\n=== PAGE BREAKS ===")
for i, p in enumerate(doc.paragraphs[:50]):
    runs = p._p.findall('.//w:br', ns)
    for br in runs:
        btype = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
        if btype == 'page':
            print(f"  Page break at P{i}: '{p.text[:50]}'")

# Check text box 0 (page de garde title box) - find what text it has
print("\n=== BOX 0 (Page de Garde title) ===")
box0_texts = alts[0].findall('.//w:t', ns)
for t in box0_texts:
    safe = (t.text or '').encode('ascii','replace').decode()
    print(f"  '{safe}'")
